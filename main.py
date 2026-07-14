import telebot
from telebot import types
import json, os, uuid, time, threading, logging, sys, requests, re, gc
from datetime import datetime, timedelta
import psycopg2
from psycopg2 import pool as pg_pool

def run_with_timeout(fn, args=(), kwargs=None, timeout=90):
    """يشغّل fn بخيط منفصل بمهلة أقصى (ثواني)، ويرجع (ok, result_or_exception).
    يحمي أي عملية تحويل/دمج ثقيلة (EPUB/PDF كبيرة) من التعليق الصامت اللي يخلي
    المستخدم يشوف 'جاري التحويل...' ومايصير أي شي بعدها للأبد — بدل هذا، لو
    العملية أخذت وقت أطول من المعقول، نوقف الانتظار ونبلغ المستخدم بوضوح."""
    kwargs = kwargs or {}
    result = {}
    def _target():
        try:
            result["value"] = fn(*args, **kwargs)
        except Exception as e:
            result["error"] = e
    t = threading.Thread(target=_target, daemon=True)
    t.start()
    t.join(timeout)
    if t.is_alive():
        return False, TimeoutError(f"العملية استغرقت أكثر من {timeout} ثانية ولم تكتمل")
    if "error" in result:
        return False, result["error"]
    return True, result.get("value")

# مكتبات تحويل صيغ الملفات (PDF/DOCX/DOC/EPUB) — تحميل كسول (lazy) واختياري حتى
# ما يوقف تشغيل البوت لو ما كانت مثبتة بعد. المطور لازم يضيفها لملف requirements.txt:
# pip install python-docx pypdf reportlab ebooklib arabic-reshaper python-bidi
_CONVERT_LIBS_READY = False
_CONVERT_IMPORT_ERROR = None
def _ensure_convert_libs():
    """يحاول استيراد مكتبات التحويل عند أول استخدام فعلي فقط، ويرجع رسالة خطأ واضحة
    للمطور لو كانت ناقصة، بدل ما يكسر تشغيل البوت بالكامل من البداية."""
    global _CONVERT_LIBS_READY, _CONVERT_IMPORT_ERROR
    if _CONVERT_LIBS_READY: return True
    try:
        global docx, pypdf, canvas, A4, cm, pdfmetrics, TTFont, epub, arabic_reshaper, bidi_get_display
        import docx
        import pypdf
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from ebooklib import epub
        import arabic_reshaper
        from bidi.algorithm import get_display as bidi_get_display
        _CONVERT_LIBS_READY = True
        return True
    except Exception as e:
        _CONVERT_IMPORT_ERROR = str(e)
        return False

# ── مكتبة الترجمة المحلية (deep-translator + langdetect) — استيراد كسول بنفس
# فلسفة مكتبات التحويل أعلاه: لا تُحمَّل إلا عند أول استخدام فعلي للمحرك
# المحلي (translation_settings.engine == "local")، حتى لا نثقّل الذاكرة على
# البوت طول الوقت لو المطور مستخدم الذكاء الاصطناعي فقط (الوضع الافتراضي).
# المطور لازم يضيفها لملف requirements.txt: pip install deep-translator langdetect
_TRANSLATE_LIBS_READY = False
_TRANSLATE_IMPORT_ERROR = None
def _ensure_translation_libs():
    global _TRANSLATE_LIBS_READY, _TRANSLATE_IMPORT_ERROR
    if _TRANSLATE_LIBS_READY: return True
    try:
        global GoogleTranslator, detect_lang
        from deep_translator import GoogleTranslator
        from langdetect import detect as detect_lang
        _TRANSLATE_LIBS_READY = True
        return True
    except Exception as e:
        _TRANSLATE_IMPORT_ERROR = str(e)
        return False

# ── مكتبة python-magic (كشف نوع الملف الحقيقي من محتواه، بغض النظر عن
# الامتداد المكتوب بالاسم) — استيراد كسول لأغراض أمنية: تمنع رفع ملف خبيث
# باسم "chapter.pdf" وهو فعليًا ملف تنفيذي مثلاً. المطور لازم يضيفها لملف
# requirements.txt: pip install python-magic (وعلى لينكس تحتاج أيضًا حزمة
# النظام libmagic1: apt-get install libmagic1)
_SECURITY_LIBS_READY = False
_SECURITY_IMPORT_ERROR = None
def _ensure_security_libs():
    global _SECURITY_LIBS_READY, _SECURITY_IMPORT_ERROR
    if _SECURITY_LIBS_READY: return True
    try:
        global magic
        import magic
        _SECURITY_LIBS_READY = True
        return True
    except Exception as e:
        _SECURITY_IMPORT_ERROR = str(e)
        return False

def verify_file_type(file_path, expected_exts):
    """يتحقق أن نوع الملف الحقيقي (من محتواه الثنائي عبر python-magic) يطابق
    فعليًا أحد الامتدادات المتوقعة، بغض النظر عن اسم/امتداد الملف المُرسَل.
    يحمي من ملفات خبيثة أو تالفة تتنكر بامتداد بريء (مثل .exe باسم .pdf).
    يرجّع (True, mime_type) لو الفحص نجح أو المكتبة غير متاحة (fail-open حتى
    لا نعطّل الميزة بالكامل لو المكتبة/حزمة libmagic غير مثبتة على الاستضافة)،
    أو (False, mime_type) لو النوع الحقيقي لا يطابق أي امتداد متوقع صراحة."""
    if not _ensure_security_libs():
        logger.warning(f"⚠️ python-magic غير متاح ({_SECURITY_IMPORT_ERROR})، تخطي التحقق الأمني من نوع الملف.")
        return True, None
    try:
        mime_type = magic.Magic(mime=True).from_file(file_path)
    except Exception as e:
        logger.warning(f"⚠️ فشل فحص نوع الملف بـ python-magic: {e}")
        return True, None
    mime_map = {
        "pdf": ["application/pdf"],
        "docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/zip"],
        "doc": ["application/msword"],
        "epub": ["application/epub+zip", "application/zip"],
        "txt": ["text/plain"],
        "html": ["text/html", "text/plain"],
        "htm": ["text/html", "text/plain"],
    }
    allowed_mimes = set()
    for ext in expected_exts:
        allowed_mimes.update(mime_map.get(ext.lower().lstrip("."), []))
    if not allowed_mimes:
        return True, mime_type  # امتداد غير مسجَّل بالخريطة، ما نحكم عليه (fail-open)
    return (mime_type in allowed_mimes), mime_type



# ── مكتبة السحب/scraping فقط — هذي مستخدمة مباشرة عند الإقلاع (خارج مسار
# التحويل)، فتبقى استيراد فوري. مكتبات التحويل الثقيلة (docx/pypdf/reportlab/
# ebooklib/arabic_reshaper/bidi) أُزيلت من هنا لأنها كانت تُستورد مرتين —
# مرة هنا فورًا عند الإقلاع، ومرة تانية بشكل كسول جوا _ensure_convert_libs()
# (اللي فعليًا هي الوحيدة المستخدمة من convert_file). هذا الازدواج كان يحمّل
# كل هالمكتبات الثقيلة بالذاكرة فور تشغيل البوت حتى لو ما حد استخدم ميزة
# التحويل أبدًا — وعلى استضافة محدودة الرام هذا سبب مباشر لاستهلاك ذاكرة
# زائد وربما "نفاد الذاكرة" (OOM). الحين تُحمّل فقط أول مرة فعليًا يُستخدم
# التحويل، عبر _ensure_convert_libs().
try:
    from bs4 import BeautifulSoup
except Exception:
    BeautifulSoup = None

# مكتبة wikipedia تُستخدم فقط من SmartSearchEngine (نظام المحادثة الذكية AI Talk)،
# استيراد اختياري بنفس فلسفة BeautifulSoup أعلاه: لو غير مثبتة، الميزة تتعطل بأمان
# بدل ما توقف تشغيل البوت بالكامل. لتفعيلها أضف "wikipedia" لـ requirements.txt.
try:
    import wikipedia
except Exception:
    wikipedia = None
import random

# ==============================================================================
# 1. CORE CONFIGURATION
# ==============================================================================
# مهم جداً: هذي القيم تُقرأ من متغيرات البيئة (Environment Variables) في
# الاستضافة فقط. تم حذف القيم الاحتياطية (fallback) اللي كانت مكتوبة هنا
# مباشرة بالكود — كانت تحتوي أسرار حقيقية مكشوفة (توكن البوت، كلمة مرور
# قاعدة بيانات Supabase القديمة، مفتاح OpenRouter)، وأي شخص يشوف الملف
# (رفعه لمكان عام، مشاركته، إلخ) يصير عنده وصول كامل لها. الحين لازم تكون
# كل هذي القيم مضبوطة كمتغيرات بيئة بلوحة تحكم الاستضافة قبل ما يشتغل البوت.
# قاعدة البيانات الحالية: Neon (عبر BOT_DATABASE_URL) بدل Supabase القديمة.
API_TOKEN           = os.environ.get("BOT_API_TOKEN")
OWNER_ID            = int(os.environ.get("BOT_OWNER_ID", "0"))
DATABASE_URL        = os.environ.get("BOT_DATABASE_URL")
OR_KEY  = os.environ.get("BOT_OPENROUTER_KEY", "")
# ── مفاتيح خدمات التحويل السحابية (كلها اختيارية) ──────────────────────────
# convert_file يجرب هذي الخدمات بالترتيب التالي كل ما كان مفتاحها مضبوطًا:
# 1) CloudConvert  2) Convertio  3) التحويل المحلي (يشتغل دايمًا بدون أي مفتاح)
# لو فشلت كل الخدمات السحابية أو مفاتيحها فاضية، يرجع تلقائيًا للمحلي — البوت
# لا يتعطل أبدًا بسبب غياب أو انقطاع أي خدمة خارجية.
CLOUDCONVERT_KEY = os.environ.get("CLOUDCONVERT_API_KEY", "")
CONVERTIO_KEY    = os.environ.get("CONVERTIO_API_KEY", "")
# ConvertAPI: خدمة سحابية ثالثة اختيارية (منفصلة عن CloudConvert/Convertio)،
# تُستخدم فقط لو حددت CONVERTAPI_SECRET كمتغير بيئة بلوحة الاستضافة (غير
# موجود افتراضيًا). سلسلة التحويل تجربها كخيار ثالث قبل المحلي.
CONVERTAPI_SECRET = os.environ.get("CONVERTAPI_SECRET", "")
# ── Upstash Redis (اختياري) — يُستخدم كتخزين خفيف وسريع لحالة "مهام التقطيع
# والاستئناف" (chunk jobs) لتحويل/دمج الملفات الكبيرة، بدل تحميل الرام أو حتى
# PostgreSQL بتحديثات متكررة كل جزء. لو المتغيرات غير مضبوطة أو الاتصال فشل،
# النظام يرجع تلقائيًا لتخزين نفس البيانات بجدول PostgreSQL الموجود أصلاً —
# البوت لا يتعطل أبدًا بسبب غياب أو انقطاع Redis.
UPSTASH_REDIS_URL   = os.environ.get("UPSTASH_REDIS_REST_URL", "")
UPSTASH_REDIS_TOKEN = os.environ.get("UPSTASH_REDIS_REST_TOKEN", "")
# HTML2PDF_API_KEY وOCF_API_KEY مقروءة ومتاحة كمتغيرات جاهزة للاستخدام، لكن
# بدون تكامل فعلي مبرمج بعد لأنها ليست جزءًا من سلسلة convert_file الحالية —
# لو حبيت ربطهما بميزة معيّنة مستقبلاً (مثلاً HTML2PDF لتحويل صفحات ويب PDF
# مباشرة)، وضّح لي الاستخدام المطلوب بالضبط وأربطهما.
HTML2PDF_KEY     = os.environ.get("HTML2PDF_API_KEY", "")
OCF_KEY          = os.environ.get("OCF_API_KEY", "")
# CHANGETHISFILE_API_KEY: نفس الحالة، مقروء ومتاح، بدون تكامل فعلي بعد.
CHANGETHISFILE_KEY = os.environ.get("CHANGETHISFILE_API_KEY", "")
A_MDL            = "openrouter/free" # الراوتر التلقائي المجاني من OpenRouter: يختار أقوى موديل مجاني متاح حاليًا (يشمل DeepSeek R1 لو متوفر) تلقائيًا، ويتجنب مشكلة الموديل يوقف يشتغل فجأة لو انسحب من قائمة المجاني

if not API_TOKEN or not DATABASE_URL:
    # نوقف التشغيل فورًا بدل ما نكمل بقيم فاضية ونكتشف المشكلة متأخر (مثلاً
    # بعد ما يحاول يتصل بقاعدة بيانات مالها عنوان أصلاً).
    missing = [name for name, val in (("BOT_API_TOKEN", API_TOKEN), ("BOT_DATABASE_URL", DATABASE_URL)) if not val]
    print(f"❌ متغيرات بيئة ناقصة وأساسية للتشغيل: {', '.join(missing)}. أضفها بإعدادات الاستضافة قبل التشغيل.", file=sys.stderr)
    sys.exit(1)

bot = telebot.TeleBot(API_TOKEN, threaded=True, num_threads=8)
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s',
                    handlers=[logging.StreamHandler(sys.stdout)])
logger = logging.getLogger("TRAIKA")

# ==============================================================================
# 2. DEFAULT DATABASE
# ==============================================================================
DEFAULT_MENU_BUTTONS = [
    {"id": "btn_browse",   "label": "📚 تصفح الأقسام", "action": "nav_root",      "visible": True, "order": 1, "type": "main_action", "sub_buttons": []},
    {"id": "btn_search",   "label": "🔍 بحث مفلتر",      "action": "search_filter", "visible": True, "order": 2, "type": "main_action", "sub_buttons": []},
    {"id": "btn_favs",     "label": "⭐ المفضلة",        "action": "view_favs",     "visible": True, "order": 3, "type": "main_action", "sub_buttons": []},
    {"id": "btn_notify",   "label": "🔔 التنبيهات",      "action": "view_notifications", "visible": True, "order": 4, "type": "main_action", "sub_buttons": []},
    {"id": "btn_redeem",   "label": "🎁 شحن نقاط",       "action": "redeem_code",   "visible": True, "order": 5, "type": "main_action", "sub_buttons": []},
    {"id": "btn_account",  "label": "👤 حسابي",          "action": "my_account",    "visible": True, "order": 6, "type": "main_action", "sub_buttons": []},
    {"id": "btn_sub",      "label": "💎 الاشتراك",       "action": "my_subscription","visible": True, "order": 7, "type": "main_action", "sub_buttons": []},
    {"id": "btn_support",  "label": "📞 الدعم",          "action": "contact_us",    "visible": True, "order": 8, "type": "main_action", "sub_buttons": []},
    {"id": "btn_merge",     "label": "🧩 دمج ملفات",       "action": "filemerge_start", "visible": True, "order": 9, "type": "main_action", "sub_buttons": []},
]

DEFAULT_DB = {
    "config": {
        "owner_id": OWNER_ID,
        "admins": {},
        "maintenance": False,
        "welcome_msg": "💎 *أهلاً بك في نظام ترايكا V5*\n\nاختر من القائمة أدناه 👇",
        "welcome_rich": None,  # {"text": str|None, "photo_id": str|None, "video_id": str|None} — إذا موجود يُستخدم بدل welcome_msg، ويُرسل كرسالة مستقلة قبل رسالة القائمة
        "mandatory_channels": [],
        "announce_channel": {
            "chat_id": None,           # آيدي/يوزرنيم قناة الإعلانات (يختلف عن القنوات الإجبارية)
            "enabled": True,           # تشغيل/تعطيل الإعلان التلقائي بالكامل
            "mode": "selected",        # "all" = كل الأعمال، "selected" = بس الأعمال المختارة
            "selected_novels": [],     # قائمة آيدي الروايات المسموح لها تُعلن (لو mode=selected)
            "message_template": "🆕 تم تحديث العمل: *{title}*\n\nآخر فصل: {chapter}",
            "button_label": "📖 اقرأ الآن"
        },
        "content_protection": {"copy": True, "save": True},
        "archive_channels": {
            "novels": None,   # آيدي/يوزرنيم قناة أرشيف الروايات (نسخة كاملة من كل فصل يُرفع)
            "manga": None,    # آيدي/يوزرنيم قناة أرشيف المانجا/المانهوا/الكوميكس
            "series": None,   # آيدي/يوزرنيم قناة أرشيف المسلسلات والأفلام
        },
        "archive_upload_allowed_users": [],  # آيدي المستخدمين المسموح لهم يرفعوا مباشرة بقناة الأرشيف ويربطوا الملف بفصل (غير المطور)
        "news_settings": {
            "publish_channel": None,      # آيدي/يوزرنيم القناة اللي تُنشر فيها الأخبار
            "source_button_label": "🔗 المصدر",  # نص زر الانتقال لمصدر الخبر
            "novel_button_label": "📖 الدخول للرواية",  # نص زر الدخول للرواية المرتبطة
            # ── السحب التلقائي والتلخيص الذكي (إضافة متطورة فوق النظام الأساسي) ──
            "auto_scrape_enabled": False,   # تشغيل/تعطيل السحب التلقائي من المصادر بالكامل
            "sources": [],                  # روابط RSS أو صفحات أخبار يُسحب منها تلقائيًا
            "scrape_interval": 6,           # كل كم ساعة يُعاد فحص المصادر
            "max_news_per_day": 5,          # حد أقصى لعدد الأخبار المسحوبة بكل دورة
            "auto_publish": False,          # نشر تلقائي فور السحب (بدل ترك الأخبار كمسودات للمراجعة)
            "publish_time": "06:00",        # وقت تشغيل دورة السحب (لو auto_publish أو auto_scrape مفعّلين)
            "last_scrape": None,            # آخر وقت تم فيه فحص المصادر (ISO datetime)
            "scraped_news": [],             # روابط أخبار سُحبت سابقًا (لتجنب التكرار، آخر 100)
        },
        "scraper_settings": {
            "enabled": False,               # تشغيل/تعطيل ميزة السحب بالكامل
            "allowed_users": [],            # آيدي المستخدمين المسموح لهم يستخدمون السحب (غير المطور)
            "public": False,                # لو True تظهر الميزة لأي شخص، لو False فقط المطور والمسموح لهم
            "batch_sizes": [25, 35, 45, 55],  # أحجام الدفعات المدموجة اللي تُبنى تلقائيًا بعد كل عدد فصول
            "auto_send_to_archive": True,   # هل يُرسل كل شيء لقناة الأرشيف تلقائيًا بعد السحب
            "default_translate_to": None,   # لغة ترجمة تلقائية بعد السحب (None = بدون ترجمة، أو "ar"/"en"/"ko"/"zh")
            "site_profiles": {},            # "profile_id": {...} — إعدادات كل موقع (selectors)، تُدار من لوحة التحكم
            # site_profiles[id] = {
            #   "name": str, "domain": str,
            #   "title_selector": str, "description_selector": str, "poster_selector": str,
            #   "chapter_content_selector": str, "chapter_title_selector": str,
            #   "next_page_selector": str,  # CSS selector لرابط "الفصل التالي" أو "الصفحة التالية"
            # }
            # ── قائمة انتظار وجدولة تلقائية (إضافة متطورة فوق النظام الأساسي) ──
            "auto_schedule_enabled": False, # تشغيل/تعطيل السحب التلقائي المجدول من قائمة الانتظار
            "schedule_time": "06:00",       # وقت تشغيل دورة السحب التلقائي يوميًا
            "novels_per_schedule": 1,       # كم رابط من قائمة الانتظار يُسحب بكل دورة جدولة
            "max_chapters_per_run": 50,     # حد أقصى للفصول المسحوبة بجلسة سحب واحدة (يحمي من التعليق الطويل)
            "pending_links": [],            # [{"url","profile_id","work_type","translate_to","added_at"}]
            "processed_links": [],          # روابط اكتملت (نفس بنية pending_links + "job_id")
            "failed_links": [],             # روابط فشلت (نفس البنية + "error")
            "last_schedule_run": None,      # آخر وقت شُغّلت فيه دورة الجدولة (لمنع التكرار بنفس الدقيقة)
        },
        "new_member_settings": {
            "enabled": False,             # تشغيل/تعطيل الميزة بالكامل
            "watched_groups": [],         # قائمة آيدي/يوزرنيم المجموعات اللي يراقبها البوت لدخول أعضاء جدد
            "dm_enabled": True,           # هل يرسل رسالة خاصة للعضو الجديد
            "dm_message": "👋 أهلاً فيك! نورت المجموعة.",  # نص الرسالة الخاصة (المطور يقدر يغيّرها)
            "dm_button_label": None,      # نص زر اختياري تحت الرسالة الخاصة (None = بدون زر)
            "dm_button_url": None,        # رابط الزر الاختياري
            "auto_add_channel": None,     # آيدي/يوزرنيم قناة يُنشأ لها رابط دعوة تلقائي يُرسل للعضو الجديد (None = تعطيل)
            "auto_add_channel_message": "📢 انضم لقناتنا أيضًا:",  # النص المرافق لرابط قناة الإضافة التلقائية
            "dm_failed_users": [],        # آيدي المستخدمين اللي فشلت رسالتهم الخاصة (خصوصية/حظر للبوت) — يسجلهم المطور لمتابعتهم يدويًا
        },
        "points_system": {"max_free_download": 5, "max_sub_download": 25},
        "merge_limits": {
            "free_daily": 10,       # حد التنزيلات المدمجة اليومي للمستخدم العادي
            "sub_daily": None,      # حد المشترك — None يعني غير محدود
            "user_overrides": {}    # {user_id: عدد مخصص لهذا المستخدم بالذات، يتجاوز الحدين فوق}
        },
        "convert_allowed_users": [],  # قائمة آيدي مستخدمين مسموح لهم استخدام تحويل الصيغة، غير المطور
        "translation_settings": {
            "public": False,          # لو True تظهر الترجمة لأي شخص، لو False فقط المطور والمسموح لهم
            "allowed_users": [],      # قائمة آيدي مستخدمين مسموح لهم استخدام الترجمة تحديدًا (منفصلة عن تحويل الصيغة)
            # محرك الترجمة الفعلي: "ai" (افتراضي، دقيق وأدبي عبر q_tr/OpenRouter، لكنه
            # يستهلك رصيد الذكاء الاصطناعي) أو "local" (عبر deep-translator/Google، فوري
            # ومجاني بدون أي رصيد، لكن ترجمة أقرب للحرفية بلا صقل أدبي). يبدّله المطور
            # من لوحة التحكم متى ما احتاج (تنظيم استهلاك، أو تعطيل الذكاء الاصطناعي مؤقتًا).
            "engine": "ai",
        },
        "public_merge_tool": {
            "enabled": True,          # تشغيل/تعطيل أداة الدمج الذاتي للمستخدمين العاديين بالكامل
            "max_files_free": 5,      # حد أقصى لعدد الملفات يقدر المستخدم العادي يدمجها بعملية وحدة (المطور والمسموح لهم دائمًا غير محدود)
        },
        "always_remove_phrases": [],  # عبارات/روابط تُحذف تلقائيًا من أي ملف يُرفع لأداة حذف النص، بدون ما تكتبها كل مرة
        "translation_phrasing_rules": {},  # {"الكلمة_الأصلية": "الصياغة المفضّلة بالترجمة"} — قواعد ثابتة تُطبّق دايمًا وقت الترجمة
        "sub_settings": {"base_price": 5.0, "total_sub_revenue": 0.0, "currency": "$"},
        "sub_plans": {"7": 3.0, "30": 5.0, "90": 12.0},
        "stats": {
            "total_users": 0, "total_downloads": 0, "total_searches": 0,
            "novel_downloads": {},   # {novel_id: count}
            "item_downloads": {},    # {item_id: count}
            "novel_views": {},       # {novel_id: count} — فتح صفحة الرواية
            "item_views": {},        # {item_id: count}
            "search_queries": {},    # {query_text_or_tag: count} — لمعرفة أكثر الطلبات
            "daily_new_users": {},   # {"YYYY-MM-DD": count}
            "daily_downloads": {},   # {"YYYY-MM-DD": count}
        },
        "stats_visible_to_users": False,  # هل تظهر إحصائيات مبسطة للمستخدمين العاديين أيضًا
        "menu_buttons": DEFAULT_MENU_BUTTONS,
        "custom_features": {},
        "feature_nesting_map": {}, # Stores parent-child relationships for features
        "custom_feature_template": {"label": "", "action": "", "visible": True, "active": True, "dev_only": False, "sub_buttons": []}, # Template for custom features with sub-buttons
        "banned_users": [],
        "exempt_users": [],
        "tags": ["شونين", "سينين", "رومنسي", "دراما", "أكشن", "خيال"],
        "tags_layout": "grid2",
        "novel_watch_label": "👁️ شاهد الآن",
        "category_layouts": {},  # {cat_id: "vertical"/"horizontal"/"grid2"/"grid3"} لعرض الأقسام الفرعية
        "list_sections": [
            {"id": "sec_watching", "name": "👁️ أشاهدها حالياً"},
            {"id": "sec_done", "name": "✅ تمت المشاهدة"},
            {"id": "sec_favs", "name": "⭐ المفضلة"},
            {"id": "sec_wish", "name": "📌 أرغب بمشاهدتها"},
        ],
        "progress_marker": "🔴",  # العلامة اللي تظهر جنب رقم آخر فصل وصله المستخدم بقائمته
        # نظام المسابقات المتطور — يدعم القنوات والمجموعات معًا، وضعا إجابة (أزرار/تعليق)،
        # أسئلة ذكية من الذكاء الاصطناعي أو من بيانات البوت نفسه (روايات/مانجا/مسلسلات) مع
        # منع تكرار الأسئلة، جدولة مرنة بالوقت والأيام، وجوائز قابلة للتخصيص بالكامل.
        "contests": {
            "enabled": False,
            "mode": "manual",              # manual / auto
            "schedule_time": "18:00",
            "schedule_days": ["sat", "sun", "mon", "tue", "wed", "thu", "fri"],
            "target_chat_id": None,        # القناة/المجموعة المستهدفة؛ افتراضيًا announce_channel لو فاضية

            "question_type": "mixed",      # mixed / novel / manga / series / anime / general
            "question_count": 5,
            "choices_count": 2,            # 2 أو 4 خيارات لكل سؤال
            "answer_mode": "button",       # button (أزرار) / comment (كتابة في التعليق)
            "question_source": "ai",       # ai / database / mixed

            "prizes": {"points": 0, "vip_days": 0, "role": ""},

            "time_between_questions": 5,   # ثواني بين إرسال كل سؤال (وضع الأزرار)
            "auto_end_minutes": 15,
            "allow_duplicate_answers": False,

            "active_contests": {},
            "past_contests": [],
            "used_questions": [],          # نصوص الأسئلة المستخدمة سابقًا (لمنع التكرار)
            "last_run": None,
            "contest_counter": 0
        },
        # الدعوة التلقائية عند التفاعل — ترسل رسالة خاصة مع رابط دعوة لقناة/مجموعة مستهدفة
        # لأي مستخدم يتفاعل مع البوت (ضغط زر/رسالة/انضمام)، باعتبار التفاعل موافقة مسبقة.
        "auto_invite": {
            "enabled": False,
            "target_chat_id": None,
            "invite_link": None,
            "confirmation_text": "✅ لقد وافقت على الانضمام إلى قناتنا المميزة.\nاضغط الزر أدناه للتأكيد والدخول:",
            "button_label": "📢 أوافق وانضم الآن",
            "trigger_on": ["callback", "message", "new_member"],
            "pending_users": {}
        },
        # نظام تقطيع الملفات الكبيرة والاستئناف — معطّل افتراضيًا. لو فعّلته،
        # أي PDF أكبر من pages_per_chunk صفحة يُقسَّم ويُحوَّل جزءًا-جزءًا مع
        # مراقبة رام واستئناف حقيقي بدل تحويله دفعة واحدة بالكامل.
        "chunking_settings": {
            "enabled": False,
            "pages_per_chunk": 50,       # كم صفحة بكل جزء PDF
            "pause_between_chunks": 3,   # ثواني انتظار بين كل جزء ومعالجة التالي
            "max_ram_mb": 200,           # الحد الأقصى للرام قبل ما ننتظر لتنخفض
            "job_retention_days": 7,     # كم يوم تُحفظ المهام المكتملة/الفاشلة قبل التنظيف التلقائي
        },
        # نظام "المحادثة الذكية" — منشورات تلقائية/يدوية عن أعمال (أفلام/أنمي/مانجا/روايات)
        # مدعومة بمعلومات حقيقية من مصادر خارجية (ويكيبيديا وغيرها) عبر SmartSearchEngine.
        "ai_talk": {
            "enabled": True,
            "auto_post": False,
            "post_channel": None,
            "post_time": "18:00",
            "post_count": 1,
            "content_types": {
                "movies": True, "series": True, "anime": True, "cartoon": True,
                "manga": True, "manhwa": True, "novels": True, "light_novels": True
            },
            "talk_about_news": True,
            "language": "arabic",
            "personality": "friendly",   # friendly/professional/sarcastic/enthusiastic/calm/mysterious
            "tone": "casual",
            "approved_topics": [],
            "blocked_topics": [],
            "history": [],
            "posts_history": []
        },
        # إعدادات محرك البحث الذكي المساعد لنظام المحادثة الذكية (ويكيبيديا/IMDB/MAL/DuckDuckGo).
        "ai_search": {
            "enabled": True,
            "search_engines": ["wikipedia", "myanimelist", "imdb"],
            "search_depth": "medium",
            "max_results": 5,
            "cache_duration": 3600,
            "search_cache": {},
            "search_history": []
        }
    },
    "users": {},
    "categories": {"root": {"id": "root", "name": "🏠 القائمة الرئيسية", "parent": None, "children": [], "items": [], "novels": [], "series": [], "manga": []}},
    "items": {}, "codes": {}, "contact_inquiries": [], "pending_actions": {},
    "novels": {},
    "series": {},  # مسلسلات وأفلام — نفس بنية novels لكن الحلقات فيديو بجودات متعددة
    "manga": {},   # مانهوا ومانجا — نفس بنية novels لكن الفصول PDF/EPUB (صور فصول)
    "content_features": {  # تشغيل/تعطيل كل نوع محتوى بالكامل للمستخدمين (زي أي ميزة ثانية بالبوت)
        "novels": {"active": True, "label": "📚 الروايات"},
        "series": {"active": True, "label": "🎬 المسلسلات والأفلام"},
        "manga": {"active": True, "label": "🎨 المانهوا والمانجا"},
    },
    "comments": {},  # {"novel_<nid>": [...], "novel_<nid>_ch_<num>": [...]} — تعليقات على عمل كامل أو فصل معيّن
    "archive_index": {},
    "news_items": {},
    "scrape_jobs": {}
    # news_items[news_id] = {
    #   "id": news_id, "title": str, "description": str, "poster_url": str|None,
    #   "poster_file_id": str|None, "category": str|None, "source_url": str,
    #   "linked_type": "novel"/"manga"/"series"/None, "linked_id": str|None,
    #   "status": "draft"/"published", "created_by": uid, "created_at": iso_str,
    #   "published_msg_id": int|None, "published_chat": str|None
    # }
    # scrape_jobs[job_id] = {
    #   "id": job_id, "profile_id": str, "start_url": str,
    #   "work_type": "novel"/"manga"/"series", "work_id": str|None (يُملأ بعد إنشاء العمل),
    #   "status": "running"/"paused"/"done"/"failed", "chapters_done": int,
    #   "last_chapter_url": str|None, "translate_to": str|None,
    #   "created_by": uid, "created_at": iso_str, "error": str|None,
    #   "scraped_texts": {"12": {"title":..., "text":...}, ...}  # مؤقت، يُنظّف بعد كل دفعة مدموجة
    # }
    # archive_index["novels"/"manga"/"series"][archive_msg_id] = {
    #   "work_id": nid/mid/sid, "chapter": ch_num, "fmt": "pdf"/"epub"/quality, "file_id": ...
    # } — فهرس يربط كل رسالة بقناة الأرشيف بالفصل/الحلقة اللي تخصه، يُستخدم لإعادة
    # بناء الربط تلقائيًا لو انمسحت بيانات البوت بالكامل وبقيت قناة الأرشيف موجودة.
    # novels[nid] = {
    #   "id": nid, "title": str, "story": str, "tag": str, "poster_file_id": str,
    #   "category": cat_id, "linked_chat": None or "@channel"/chat_id,
    #   "chapters": {"1": {"pdf": file_id|None, "epub": file_id|None}, ...},
    #   "merged": {"1-25": {"pdf": file_id|None, "epub": file_id|None}, ...},
    #   "created_at": iso_str
    # }
    # users[uid]["novel_progress"] = {nid: last_chapter_number}
}


db = {}

# ==============================================================================
# 3. DATABASE ENGINE (Connection Pool + Debounced Sync)
# ==============================================================================
# بدل ما نفتح اتصال جديد بقاعدة البيانات مع كل عملية حفظ (بطيء ومكلف خصوصًا
# مع عدة مستخدمين بنفس الوقت)، نستخدم "بركة اتصالات" (Connection Pool) جاهزة
# ومُعاد استخدامها. كذلك نؤجل الحفظ الفعلي جزء من ثانية (debounce) عشان لو
# صار كذا ضغطة زر بنفس اللحظة (من نفس المستخدم أو عدة مستخدمين)، تنكتب
# دفعة وحدة بدل عدة كتابات متكررة على نفس البيانات.
_db_pool = None
_sync_lock = threading.Lock()
_sync_pending = False
_sync_timer = None
DEBOUNCE_SECONDS = 2.5
_db_load_failed = False  # يصير True إذا فشل تحميل القاعدة عند الإقلاع؛ يمنع الكتابة فوق البيانات الحقيقية بنسخة فاضية

def init_pool():
    global _db_pool
    # حماية إضافية: لو فيه بركة اتصال سابقة مفتوحة (نادرًا، لو init_db انسبت تنّفذ
    # أكثر من مرة لأي سبب)، نقفلها أول قبل ما نفتح وحدة جديدة، حتى ما تتراكم اتصالات.
    if _db_pool is not None:
        try: _db_pool.closeall()
        except Exception: pass
    try:
        _db_pool = pg_pool.ThreadedConnectionPool(minconn=1, maxconn=5, dsn=DATABASE_URL, sslmode='require')
        logger.info("✅ DB connection pool ready.")
    except Exception as e:
        logger.error(f"❌ Pool init error: {e}")
        _db_pool = None

def get_conn():
    """يرجع اتصال من البركة إذا جاهزة، وإلا يفتح اتصال مباشر كحل احتياطي."""
    if _db_pool:
        return _db_pool.getconn()
    return psycopg2.connect(DATABASE_URL, sslmode='require')

def release_conn(conn):
    if _db_pool and conn:
        try: _db_pool.putconn(conn)
        except Exception: pass
    elif conn:
        try: conn.close()
        except Exception: pass

def init_db():
    global db, _db_load_failed
    last_error = None
    # محاولات متعددة عند الإقلاع: بعض المستضيفات تحتاج لحظات حتى تجهز الشبكة/DNS
    # بعد إعادة التشغيل مباشرة، فمحاولة وحدة فورية ممكن تفشل رغم إن القاعدة سليمة.
    for attempt in range(1, 6):
        init_pool()
        conn = None
        try:
            conn = get_conn()
            cur = conn.cursor()
            cur.execute('''CREATE TABLE IF NOT EXISTS bot_core_data (id TEXT PRIMARY KEY, content JSONB)''')
            conn.commit()
            cur.execute("SELECT content FROM bot_core_data WHERE id='main_db'")
            row = cur.fetchone()
            if row:
                db = row[0]
                for key in DEFAULT_DB:
                    if key not in db: db[key] = DEFAULT_DB[key]
                if "tags" not in db["config"]: db["config"]["tags"] = DEFAULT_DB["config"]["tags"]
                if "tags_layout" not in db["config"]: db["config"]["tags_layout"] = "grid2"
                if "novel_watch_label" not in db["config"]: db["config"]["novel_watch_label"] = "👁️ شاهد الآن"
                if "category_layouts" not in db["config"]: db["config"]["category_layouts"] = {}
                if "novels" not in db: db["novels"] = {}
                if "manga" not in db: db["manga"] = {}
                if "series" not in db: db["series"] = {}
                if "content_features" not in db["config"]: db["config"]["content_features"] = DEFAULT_DB["config"].get("content_features", {})
                if "archive_channels" not in db["config"]: db["config"]["archive_channels"] = DEFAULT_DB["config"].get("archive_channels", {})
                if "archive_upload_allowed_users" not in db["config"]: db["config"]["archive_upload_allowed_users"] = []
                if "archive_index" not in db: db["archive_index"] = {}
                if "always_remove_phrases" not in db["config"]: db["config"]["always_remove_phrases"] = []
                if "translation_phrasing_rules" not in db["config"]: db["config"]["translation_phrasing_rules"] = {}
                if "public_merge_tool" not in db["config"]: db["config"]["public_merge_tool"] = DEFAULT_DB["config"].get("public_merge_tool", {"enabled": True, "max_files_free": 5})
                if "new_member_settings" not in db["config"]: db["config"]["new_member_settings"] = DEFAULT_DB["config"]["new_member_settings"]
                if "news_settings" not in db["config"]: db["config"]["news_settings"] = DEFAULT_DB["config"]["news_settings"]
                else:
                    for _k, _v in DEFAULT_DB["config"]["news_settings"].items():
                        if _k not in db["config"]["news_settings"]: db["config"]["news_settings"][_k] = _v
                if "news_items" not in db: db["news_items"] = {}
                if "scraper_settings" not in db["config"]: db["config"]["scraper_settings"] = DEFAULT_DB["config"]["scraper_settings"]
                else:
                    for _k, _v in DEFAULT_DB["config"]["scraper_settings"].items():
                        if _k not in db["config"]["scraper_settings"]: db["config"]["scraper_settings"][_k] = _v
                if "scrape_jobs" not in db: db["scrape_jobs"] = {}
                if "translation_settings" not in db["config"]: db["config"]["translation_settings"] = DEFAULT_DB["config"]["translation_settings"]
                else:
                    for _k, _v in DEFAULT_DB["config"]["translation_settings"].items():
                        if _k not in db["config"]["translation_settings"]: db["config"]["translation_settings"][_k] = _v
                if "contests" not in db["config"]: db["config"]["contests"] = DEFAULT_DB["config"]["contests"]
                else:
                    for _k, _v in DEFAULT_DB["config"]["contests"].items():
                        if _k not in db["config"]["contests"]: db["config"]["contests"][_k] = _v
                    # ترحيل من البنية القديمة (settings متداخلة) للبنية الجديدة المسطّحة، إن وُجدت
                    _old_settings = db["config"]["contests"].pop("settings", None)
                    if _old_settings:
                        _map = {
                            "question_type": "question_type", "questions_per_contest": "question_count",
                            "time_between_contests": None, "allowed_days": "schedule_days",
                            "start_time": None, "end_time": None,
                        }
                        for _ok, _nk in _map.items():
                            if _nk and _ok in _old_settings and _nk not in db["config"]["contests"]:
                                db["config"]["contests"][_nk] = _old_settings[_ok]
                    if "prizes" not in db["config"]["contests"]:
                        db["config"]["contests"]["prizes"] = DEFAULT_DB["config"]["contests"]["prizes"]
                if "auto_invite" not in db["config"]: db["config"]["auto_invite"] = DEFAULT_DB["config"]["auto_invite"]
                else:
                    for _k, _v in DEFAULT_DB["config"]["auto_invite"].items():
                        if _k not in db["config"]["auto_invite"]: db["config"]["auto_invite"][_k] = _v
                if "chunking_settings" not in db["config"]: db["config"]["chunking_settings"] = DEFAULT_DB["config"]["chunking_settings"]
                else:
                    for _k, _v in DEFAULT_DB["config"]["chunking_settings"].items():
                        if _k not in db["config"]["chunking_settings"]: db["config"]["chunking_settings"][_k] = _v
                if "chunk_jobs" not in db: db["chunk_jobs"] = {}
                if "ai_talk" not in db["config"]: db["config"]["ai_talk"] = DEFAULT_DB["config"]["ai_talk"]
                else:
                    for _k, _v in DEFAULT_DB["config"]["ai_talk"].items():
                        if _k not in db["config"]["ai_talk"]: db["config"]["ai_talk"][_k] = _v
                if "ai_search" not in db["config"]: db["config"]["ai_search"] = DEFAULT_DB["config"]["ai_search"]
                else:
                    for _k, _v in DEFAULT_DB["config"]["ai_search"].items():
                        if _k not in db["config"]["ai_search"]: db["config"]["ai_search"][_k] = _v
                for cat in db.get("categories", {}).values():
                    if "novels" not in cat: cat["novels"] = []
                    if "manga" not in cat: cat["manga"] = []
                    if "series" not in cat: cat["series"] = []
                for u in db.get("users", {}).values():
                    if "novel_progress" not in u: u["novel_progress"] = {}
                    if "manga_progress" not in u: u["manga_progress"] = {}
                    if "series_progress" not in u: u["series_progress"] = {}
                # ترحيل بيانات قديمة: تحويل "tag" (نص واحد) إلى "tags" (قائمة) لكل الروايات والملفات،
                # حتى تشتغل خاصية التصنيفات المتعددة مع المحتوى المُضاف قبل التحديث بدون فقدان بياناته.
                for nv in db.get("novels", {}).values():
                    if "tags" not in nv:
                        old = nv.pop("tag", None)
                        nv["tags"] = [old] if old else ["عام"]
                for it in db.get("items", {}).values():
                    if "tags" not in it:
                        old = it.pop("tag", None)
                        it["tags"] = [old] if old else ["عام"]
                logger.info(f"✅ DB loaded (محاولة رقم {attempt}).")
            else:
                cur.execute("INSERT INTO bot_core_data (id,content) VALUES ('main_db',%s)", (json.dumps(DEFAULT_DB),))
                conn.commit()
                db = DEFAULT_DB.copy()
                logger.info("✅ DB initialized (أول تشغيل — قاعدة جديدة فعلاً فاضية).")
            cur.close()
            _db_load_failed = False
            return  # نجح الاتصال — نخرج فورًا بدون تجربة محاولات إضافية
        except Exception as e:
            last_error = e
            logger.error(f"❌ DB init error (محاولة {attempt}/5): {e}")
            if attempt < 5:
                time.sleep(2 * attempt)  # تأخير متزايد قبل المحاولة التالية
        finally:
            if conn: release_conn(conn)
    # كل المحاولات فشلت فعلاً — الآن فقط نعتبرها فشل حقيقي ونحمي البيانات من الاستبدال
    logger.error(f"🚨 فشلت كل محاولات الاتصال بقاعدة البيانات ({last_error}). "
                 "البوت راح يشتغل بذاكرة فاضية مؤقتة وممنوع من الحفظ نهائيًا حتى تُحل المشكلة، "
                 "حماية للبيانات الأصلية من الاستبدال.")
    db = DEFAULT_DB.copy()
    _db_load_failed = True

def create_snapshot(label=""):
    """يأخذ نسخة كاملة من قاعدة البيانات الحالية (كل شي: روايات، ملفات، مستخدمين،
    ميزات، إعدادات...) ويخزّنها بصف منفصل بجدول قاعدة البيانات نفسه (السحابة)،
    مربوطة بكود استرجاع قصير. يرجع الكود لو نجح، أو None لو فشل."""
    conn = None
    try:
        code = uuid.uuid4().hex[:8].upper()
        conn = get_conn()
        cur = conn.cursor()
        cur.execute('''CREATE TABLE IF NOT EXISTS bot_snapshots (code TEXT PRIMARY KEY, content JSONB, label TEXT, created_at TEXT)''')
        cur.execute("INSERT INTO bot_snapshots(code, content, label, created_at) VALUES(%s, %s, %s, %s)",
                    (code, json.dumps(db), label, str(datetime.now())))
        conn.commit(); cur.close()
        return code
    except Exception as e:
        logger.error(f"❌ فشل أخذ نسخة احتياطية: {e}")
        return None
    finally:
        if conn: release_conn(conn)

def list_snapshots():
    conn = None
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("CREATE TABLE IF NOT EXISTS bot_snapshots (code TEXT PRIMARY KEY, content JSONB, label TEXT, created_at TEXT)")
        conn.commit()
        cur.execute("SELECT code, label, created_at FROM bot_snapshots ORDER BY created_at DESC LIMIT 20")
        rows = cur.fetchall(); cur.close()
        return rows
    except Exception as e:
        logger.error(f"❌ فشل جلب النسخ الاحتياطية: {e}")
        return []
    finally:
        if conn: release_conn(conn)

def restore_snapshot(code):
    """يرجّع البوت بالكامل لنفس حالته وقت أخذ هذي النسخة بالضبط — كل الروايات،
    الملفات، الميزات، الإعدادات، المستخدمين — كأن شي ما تغيّر من ذاك الوقت."""
    global db
    conn = None
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("SELECT content FROM bot_snapshots WHERE code=%s", (code,))
        row = cur.fetchone()
        cur.close()
        if not row:
            return False
        db = row[0]
        sync_db_now()  # نحفظ فورًا حتى تنعكس الاستعادة بقاعدة البيانات الرئيسية مباشرة
        return True
    except Exception as e:
        logger.error(f"❌ فشل استرجاع النسخة الاحتياطية: {e}")
        return False
    finally:
        if conn: release_conn(conn)

def _write_db_now():
    global _sync_pending
    if _db_load_failed:
        # حماية حرجة: لا نكتب أي شي فوق القاعدة الحقيقية إذا فشل تحميلها بالأساس،
        # لأن db الحالية بهذي الحالة هي DEFAULT_DB فاضية، وكتابتها ستمسح بيانات المستخدمين الحقيقية.
        logger.error("🚨 تم تجاهل محاولة حفظ لأن تحميل القاعدة فشل سابقًا (حماية من مسح البيانات).")
        with _sync_lock:
            _sync_pending = False
        return
    conn = None
    try:
        conn = get_conn()
        cur = conn.cursor()
        # حارس أمان: قبل أي كتابة، نقارن حجم البيانات الحالية بالمخزّنة فعليًا بالقاعدة.
        # لو الكتابة الجديدة راح تمسح جزء كبير من الروايات/الملفات/المستخدمين فجأة (أكبر من 50%)،
        # نوقف الكتابة وننبّه المالك بدل ما نسمح بمسح صامت — هذا يحمي من أي DATABASE_URL غلط
        # أو كود فيه خلل يفرّغ البيانات بالغلط.
        cur.execute("SELECT content FROM bot_core_data WHERE id='main_db'")
        existing_row = cur.fetchone()
        if existing_row and existing_row[0]:
            old = existing_row[0]
            for key in ("novels", "items", "users", "manga", "series"):
                old_count = len(old.get(key, {}) or {})
                new_count = len(db.get(key, {}) or {})
                if old_count >= 5 and new_count < old_count * 0.5:
                    logger.error(f"🚨 تم رفض الحفظ: {key} كان فيه {old_count} وصار {new_count} — انخفاض مفاجئ يشبه مسح بيانات. راجع الاتصال بقاعدة البيانات فورًا.")
                    try:
                        bot.send_message(OWNER_ID,
                            f"🚨 *تنبيه حرج جدًا:* تم إيقاف عملية حفظ لأنها كانت ستمسح جزء كبير من بيانات "
                            f"«{key}» ({old_count} → {new_count}). لم يُحفظ شي حماية لبياناتك. "
                            f"تأكد من DATABASE_URL والاتصال بقاعدة البيانات.",
                            parse_mode="Markdown")
                    except Exception: pass
                    cur.close()
                    with _sync_lock: _sync_pending = False
                    return
            # نسخة احتياطية دوّارة: نحفظ آخر نسخة سليمة قبل الكتابة فوقها، حتى لو صار خطأ لاحقًا
            # نقدر نرجع لها يدويًا من قاعدة البيانات مباشرة (صف id='main_db_backup').
            cur.execute("INSERT INTO bot_core_data(id,content) VALUES('main_db_backup',%s) "
                        "ON CONFLICT(id) DO UPDATE SET content=EXCLUDED.content", (json.dumps(old),))
        cur.execute("INSERT INTO bot_core_data(id,content) VALUES('main_db',%s) "
                    "ON CONFLICT(id) DO UPDATE SET content=EXCLUDED.content", (json.dumps(db),))
        conn.commit(); cur.close()
    except Exception as e:
        logger.error(f"❌ Sync: {e}")
    finally:
        if conn: release_conn(conn)
        with _sync_lock:
            _sync_pending = False

def sync_db():
    """يجدول كتابة للقاعدة خلال DEBOUNCE_SECONDS. لو فيه كتابة مجدولة أصلاً،
    ما يفتح مؤقت ثاني — الكتابة القادمة راح تاخذ آخر نسخة من db تلقائيًا."""
    global _sync_pending, _sync_timer
    with _sync_lock:
        if _sync_pending:
            return
        _sync_pending = True
        _sync_timer = threading.Timer(DEBOUNCE_SECONDS, _write_db_now)
        _sync_timer.daemon = True
        _sync_timer.start()

def sync_db_now():
    """كتابة فورية بدون تأجيل، تُستخدم عند إغلاق البوت أو لحظات حرجة."""
    global _sync_pending
    with _sync_lock:
        if _sync_timer: _sync_timer.cancel()
        _sync_pending = False
    _write_db_now()

init_db()

# ==============================================================================
# 4. ANTI-SLEEP & KEEP-ALIVE SYSTEM
# ==============================================================================
# ملاحظة: Bot-Hosting.net يوفر 24/7 uptime بدون نوم أصلاً على الخطة المجانية،
# فهذا النظام غير ضروري هناك ويستهلك موارد إضافية (كل 60 ثانية اتصال DB).
# تم تحويله لنظام قابل للتفعيل/التعطيل عبر متغير بيئة بدل حذفه، حتى لو رجعت
# لاستضافة تنام (مثل Render) تقدر تفعّله بدون تعديل الكود.
# للتفعيل: أضف متغير بيئة ANTI_SLEEP=1 في إعدادات الاستضافة. افتراضيًا معطّل.
ANTI_SLEEP_ENABLED = os.environ.get("ANTI_SLEEP", "0") == "1"

def anti_sleep_ping():
    while True:
        try:
            conn = get_conn(); cur = conn.cursor(); cur.execute("SELECT 1"); cur.fetchone(); cur.close(); conn.close()
            logger.info("🕒 Anti-Sleep: Database and Bot are active.")
        except Exception as e: logger.error(f"⚠️ Anti-Sleep Error: {e}")
        time.sleep(60)

def db_health_watchdog():
    """يراقب صحة الاتصال بقاعدة البيانات باستمرار (بغض النظر عن تفعيل Anti-Sleep)،
    ولو رجع الاتصال يشتغل بعد فشل سابق، يعيد تحميل البيانات الحقيقية تلقائيًا
    بدل ما يضل البوت عالق بذاكرة فاضية وممنوع من الحفظ للأبد."""
    while True:
        time.sleep(30)
        if not _db_load_failed:
            continue
        try:
            conn = get_conn(); cur = conn.cursor(); cur.execute("SELECT 1"); cur.fetchone(); cur.close(); release_conn(conn)
            logger.info("🔄 الاتصال بالقاعدة رجع يشتغل — إعادة تحميل البيانات الحقيقية تلقائيًا...")
            init_db()
            if not _db_load_failed:
                try:
                    bot.send_message(OWNER_ID, "✅ رجع الاتصال بقاعدة البيانات وتم تحميل بياناتك الحقيقية بنجاح. الحفظ يشتغل عادي الآن.")
                except Exception: pass
        except Exception as e:
            logger.error(f"⚠️ DB Watchdog: لسا ما رجع الاتصال ({e})")

threading.Thread(target=db_health_watchdog, daemon=True).start()

SNAPSHOT_MIN_INTERVAL = 23 * 60 * 60  # أقل فاصل مسموح بين رسالتين تلقائيتين للمطور (يحمي من السبام لو البوت يعيد التشغيل كثير)

def daily_snapshot_worker():
    """يأخذ نسخة احتياطية كاملة تلقائيًا كل 24 ساعة تقريبًا، ويرسل كود الاسترجاع
    للمطور بتلغرام — لكن فقط لو مر وقت كافٍ (23 ساعة) من آخر إشعار فعليًا، حتى لو
    البوت يعيد التشغيل عدة مرات بفترة قصيرة (كراش/نشر متكرر) ما يكرر الإزعاج.
    وقت آخر إشعار يُخزَّن بقاعدة البيانات نفسها فيضل يتذكره حتى بعد أي ريستارت."""
    while True:
        if _db_load_failed:
            time.sleep(60)
            continue  # لا نأخذ نسخة من بيانات فاضية بحالة فشل الاتصال

        last_ts = db["config"].get("last_auto_snapshot_at")
        now = datetime.now()
        due = True
        if last_ts:
            try:
                elapsed = (now - datetime.fromisoformat(last_ts)).total_seconds()
                due = elapsed >= SNAPSHOT_MIN_INTERVAL
            except Exception:
                due = True

        if not due:
            # ننام لحد وقت الاستحقاق التالي بدل ما نلف كل شوي بالفاضي
            try:
                remaining = SNAPSHOT_MIN_INTERVAL - (now - datetime.fromisoformat(last_ts)).total_seconds()
            except Exception:
                remaining = SNAPSHOT_MIN_INTERVAL
            time.sleep(max(60, min(remaining, 24 * 60 * 60)))
            continue

        code = create_snapshot(label="تلقائية يومية")
        if code:
            logger.info(f"✅ نسخة احتياطية تلقائية: {code}")
            db["config"]["last_auto_snapshot_at"] = str(now)
            sync_db()
            try:
                bot.send_message(OWNER_ID,
                    f"💾 *نسخة احتياطية تلقائية*\n\n"
                    f"تم أخذ نسخة كاملة من كل بيانات البوت (روايات، مانجا، مسلسلات، "
                    f"مستخدمين، إعدادات، ميزات) الآن.\n\n"
                    f"كود الاسترجاع:\n`{code}`\n\n"
                    f"احتفظ بهذا الكود. لو احتجت ترجع البوت لهذي اللحظة بالضبط بأي وقت، "
                    f"استخدم «⚙️ لوحة التحكم» → «💾 نسخ احتياطية واسترجاع» → «🔑 استرجاع بكود».",
                    parse_mode="Markdown")
            except Exception as e:
                logger.error(f"❌ فشل إرسال كود النسخة الاحتياطية للمطور: {e}")
        else:
            time.sleep(300)  # فشل أخذ النسخة (مشكلة اتصال مؤقتة) — نحاول مرة ثانية بعد 5 دقايق بدل ما ننتظر يوم كامل

threading.Thread(target=daily_snapshot_worker, daemon=True).start()

if ANTI_SLEEP_ENABLED:
    threading.Thread(target=anti_sleep_ping, daemon=True).start()
    logger.info("🕒 Anti-Sleep system: ENABLED (ANTI_SLEEP=1)")
else:
    logger.info("🕒 Anti-Sleep system: DISABLED (not needed on 24/7 hosts like Bot-Hosting.net)")

# ==============================================================================
# 5. HELPERS
# ==============================================================================
def md_safe(text):
    """يهرب الرموز اللي تكسر تنسيق Markdown القديم بتليجرام (* _ ` [) لو كانت موجودة
    بنص كتبه المستخدم (اسم رواية، قصة، إلخ) قبل حقنه داخل نص منسّق. بدون هذا، أي
    عنوان فيه '_' أو '*' يفشّل الرسالة كاملة بصمت (can't parse entities)."""
    if not text: return text
    for ch in ("_", "*", "`", "["):
        text = text.replace(ch, f"\\{ch}")
    return text

def is_owner(uid):  return int(uid) == OWNER_ID

def comment_key_novel(nid): return f"novel_{nid}"
def comment_key_chapter(nid, ch): return f"novel_{nid}_ch_{ch}"

def add_comment(key, uid, text):
    db["comments"].setdefault(key, [])
    entry = {
        "id": str(uuid.uuid4())[:8], "uid": str(uid),
        "name": db["users"].get(str(uid), {}).get("first_name", "مستخدم"),
        "text": text, "at": str(datetime.now())
    }
    db["comments"][key].append(entry)
    if len(db["comments"][key]) > 200:  # حد أقصى 200 تعليق لكل عمل/فصل، يحتفظ بالأحدث
        db["comments"][key] = db["comments"][key][-200:]
    sync_db()
    return entry

def get_comments(key):
    return db["comments"].get(key, [])

def delete_comment(key, comment_id):
    if key in db["comments"]:
        db["comments"][key] = [c for c in db["comments"][key] if c["id"] != comment_id]
        sync_db()

def comments_view_kb(key, back_cb, uid):
    m = types.InlineKeyboardMarkup(row_width=1)
    comments = get_comments(key)
    for c in comments[-15:]:  # آخر 15 تعليق فقط بالعرض حتى ما تطول القائمة
        label = f"💬 {c['name']}: {c['text'][:35]}"
        if is_admin(uid):
            m.add(types.InlineKeyboardButton(label[:60], callback_data="noop_"),
                  types.InlineKeyboardButton("🗑️", callback_data=f"delcomment_{key}_{c['id']}"))
        else:
            m.add(types.InlineKeyboardButton(label[:64], callback_data="noop_"))
    m.add(types.InlineKeyboardButton("✍️ أضف تعليق", callback_data=f"addcomment_{key}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=back_cb))
    return m

def push_notification_all(text, target_novel=None):
    """يضيف تنبيه بصندوق كل مستخدم (زر 🔔 التنبيهات) عن عمل جديد أو فصل جديد.
    لا يرسل رسالة تلغرام مباشرة لكل مستخدم (تفاديًا للإزعاج/الحظر الجماعي)،
    بس يخزّن التنبيه ليشوفه المستخدم لما يفتح الجرس بنفسه."""
    entry = {"text": text, "at": str(datetime.now()), "read": False, "target": target_novel}
    for u in db["users"].values():
        u.setdefault("notifications", []).append(entry.copy())
        if len(u["notifications"]) > 50:  # حد أقصى 50 تنبيه محفوظ لكل مستخدم
            u["notifications"] = u["notifications"][-50:]
    sync_db()

def notifications_kb(ustr):
    m = types.InlineKeyboardMarkup(row_width=1)
    notifs = db["users"].get(ustr, {}).get("notifications", [])
    for n in reversed(notifs[-20:]):
        label = f"{'🔵' if not n.get('read') else '⚪'} {n['text'][:50]}"
        cb = f"opennotif_{n['target']}" if n.get("target") else "noop_"
        m.add(types.InlineKeyboardButton(label[:64], callback_data=cb))
    if notifs:
        m.add(types.InlineKeyboardButton("🗑️ مسح الكل", callback_data="clear_notifications"))
    m.add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
    return m

def mirror_to_archive(content_type, work_id, chapter_label, fmt_or_quality, file_id, file_kind="document"):
    """يرسل نسخة من الملف المرفوع لقناة الأرشيف المخصصة لنوع المحتوى (روايات/مانجا/مسلسلات)،
    ويسجل بفهرس الأرشيف أي رسالة تخص أي عمل/فصل — حتى لو انمسحت بيانات البوت بالكامل،
    نقدر نعيد بناء الربط من قناة الأرشيف نفسها لأن كل ملف فيها موسوم بوضوح."""
    channel = db["config"].get("archive_channels", {}).get(content_type)
    if not channel:
        return  # ما فيه قناة أرشيف مربوطة لهذا النوع، نتجاهل بصمت
    work = None
    if content_type == "novels": work = get_novel(work_id)
    elif content_type == "manga": work = get_manga(work_id)
    elif content_type == "series": work = get_series(work_id)
    if not work: return
    caption = f"📦 *أرشيف تلقائي*\n\nالعمل: {work['title']}\nالمعرّف: `{work_id}`\nالفصل/الجزء: {chapter_label}\nالصيغة/الجودة: {fmt_or_quality}"
    try:
        if file_kind == "video":
            sent = bot.send_video(channel, file_id, caption=caption, parse_mode="Markdown")
        else:
            sent = bot.send_document(channel, file_id, caption=caption, parse_mode="Markdown")
        db["archive_index"].setdefault(content_type, {})
        db["archive_index"][content_type][str(sent.message_id)] = {
            "work_id": work_id, "chapter": chapter_label, "fmt": fmt_or_quality, "file_id": file_id
        }
        sync_db()
    except Exception as e:
        logger.error(f"❌ فشل إرسال نسخة أرشيف ({content_type}/{work_id}): {e}")

def announce_new_chapter(nid, chapter_num):
    """ينشر إعلان بقناة الإعلانات (لو مفعّلة ومربوطة) كل ما ينضاف فصل جديد لعمل، مع
    زر 'اقرأ الآن' يودّي المستخدم للبوت مباشرة على نفس العمل عبر Deep Link. يحترم
    إعداد 'الأعمال المختارة فقط' لو المطور مفعّله بدل نشر كل الأعمال تلقائيًا."""
    ac = db["config"].get("announce_channel", {})
    if not ac.get("enabled") or not ac.get("chat_id"):
        return
    if ac.get("mode") == "selected" and nid not in ac.get("selected_novels", []):
        return
    novel = get_novel(nid)
    if not novel: return
    try:
        bot_username = bot.get_me().username
    except Exception:
        bot_username = None
    deep_link = f"https://t.me/{bot_username}?start=novel_{nid}" if bot_username else None
    text = ac.get("message_template", "🆕 تم تحديث العمل: *{title}*\n\nآخر فصل: {chapter}").format(
        title=novel["title"], chapter=chapter_num)
    m = types.InlineKeyboardMarkup()
    if deep_link:
        m.add(types.InlineKeyboardButton(ac.get("button_label", "📖 اقرأ الآن"), url=deep_link))
    try:
        if novel.get("poster_file_id"):
            bot.send_photo(ac["chat_id"], novel["poster_file_id"], caption=text, parse_mode="Markdown", reply_markup=m)
        else:
            bot.send_message(ac["chat_id"], text, parse_mode="Markdown", reply_markup=m)
    except Exception as e:
        logger.error(f"❌ فشل نشر إعلان الفصل الجديد بالقناة: {e}")

def should_protect_content(uid):
    """يتحقق هل لازم نمنع النسخ/إعادة التوجيه/الحفظ لهذا المستخدم. المطور والمشرفين
    دايمًا مستثنون (يقدرون يحفظون وينسخون بحرية للإدارة)، والإعداد قابل للتغيير من لوحة التحكم."""
    if is_admin(uid): return False
    cp = db["config"].get("content_protection", {})
    return bool(cp.get("copy") or cp.get("save"))

def _trim_daily_stats(daily_dict, keep_days=90):
    """يحتفظ بآخر 90 يوم بس من أي إحصائية يومية، حتى ما تكبر قاعدة البيانات للأبد
    مع مرور الوقت (سنة كاملة = 365 مفتاح، غير ضروري نحتفظ فيها كلها)."""
    if len(daily_dict) <= keep_days: return
    for old_day in sorted(daily_dict.keys())[:-keep_days]:
        daily_dict.pop(old_day, None)

def track_event(kind, key=None, _skip_sync=False):
    """يسجل حدث إحصائي: تنزيل، مشاهدة، أو بحث. يُستخدم بكل نقطة تنزيل/عرض/بحث بالبوت
    حتى تكون الإحصائيات حقيقية بدل أرقام ثابتة على صفر."""
    s = db["config"]["stats"]
    today = str(datetime.now().date())
    if kind == "download":
        s["total_downloads"] = s.get("total_downloads", 0) + 1
        s.setdefault("daily_downloads", {})
        s["daily_downloads"][today] = s["daily_downloads"].get(today, 0) + 1
        _trim_daily_stats(s["daily_downloads"])
    elif kind in ("novel_download", "manga_download", "series_download"):
        bucket = {"novel_download": "novel_downloads", "manga_download": "manga_downloads", "series_download": "series_downloads"}[kind]
        s.setdefault(bucket, {})
        s[bucket][key] = s[bucket].get(key, 0) + 1
        track_event("download", _skip_sync=True)  # نحدّث العداد العام بدون كتابة مزدوجة للقاعدة
    elif kind == "item_download":
        s.setdefault("item_downloads", {})
        s["item_downloads"][key] = s["item_downloads"].get(key, 0) + 1
        track_event("download", _skip_sync=True)
    elif kind in ("novel_view", "manga_view", "series_view"):
        bucket = {"novel_view": "novel_views", "manga_view": "manga_views", "series_view": "series_views"}[kind]
        s.setdefault(bucket, {})
        s[bucket][key] = s[bucket].get(key, 0) + 1
    elif kind == "item_view":
        s.setdefault("item_views", {})
        s["item_views"][key] = s["item_views"].get(key, 0) + 1
    elif kind == "search":
        s["total_searches"] = s.get("total_searches", 0) + 1
        s.setdefault("search_queries", {})
        s["search_queries"][key] = s["search_queries"].get(key, 0) + 1
    if not _skip_sync:
        sync_db()

def is_admin(uid):
    """المطور دايمًا مشرف كامل. أي مشرف ثاني نتأكد أولًا إنه ما انتهت صلاحيته المؤقتة
    (لو محدد له مدة)، وإلا نعتبره غير مشرف تلقائيًا بدون أي إشعار له بالسبب."""
    if is_owner(uid): return True
    ustr = str(uid)
    admins = db["config"]["admins"]
    if ustr not in admins: return False
    rec = admins[ustr]
    if isinstance(rec, str):
        return True  # توافق مع بيانات قديمة كانت تخزن مشرف كنص بسيط فقط
    expiry = rec.get("expires_at")
    if expiry:
        try:
            if datetime.now() >= datetime.fromisoformat(expiry):
                return False  # انتهت مدته — يفقد صلاحياته تلقائيًا دون أي تنبيه له
        except Exception:
            pass
    return True

def get_admin_record(uid):
    """يرجع بيانات المشرف الكاملة (لقب، صلاحيات، مدة) أو None لو مو مشرف حقيقي."""
    ustr = str(uid)
    if is_owner(uid):
        return {"title": "المطور", "permissions": {"all": True}, "expires_at": None}
    admins = db["config"]["admins"]
    rec = admins.get(ustr)
    if not rec: return None
    if isinstance(rec, str):
        return {"title": rec, "permissions": {"all": True}, "expires_at": None}
    return rec

def has_permission(uid, perm):
    """يتحقق هل المشرف عنده صلاحية معيّنة (مثلاً 'add_novel'). المالك وأي مشرف بصلاحية
    'all' يعدّون مصرّح لهم بكل شي تلقائيًا."""
    if is_owner(uid): return True
    rec = get_admin_record(uid)
    if not rec: return False
    perms = rec.get("permissions", {})
    return bool(perms.get("all") or perms.get(perm))

def is_banned(uid): return str(uid) in db["config"]["banned_users"]

def ensure_user(msg):
    uid = str(msg.from_user.id)
    if uid not in db["users"]:
        db["users"][uid] = {
            "first_name": msg.from_user.first_name or "",
            "username": msg.from_user.username or "",
            "joined_at": str(datetime.now()),
            "sub_expiry": None, "sub_months": 0,
            "downloads": 0, "points": 0,
            "favs": [], "history": [],
            "novel_progress": {},  # {novel_id: last_chapter_number}
            "list_items": {},      # {section_id: [{"type": "novel"|"item", "id": ...}]}
            "seen_start_prompt": False,  # هل شاف زر "ابدأ" مرة، عشان ما نكرره على كل نص عشوائي
            "notifications": []  # [{"text":.., "at":.., "read": bool, "target": novel_id|None}]
        }
        db["config"]["stats"]["total_users"] += 1
        today = str(datetime.now().date())
        db["config"]["stats"].setdefault("daily_new_users", {})
        db["config"]["stats"]["daily_new_users"][today] = db["config"]["stats"]["daily_new_users"].get(today, 0) + 1
        _trim_daily_stats(db["config"]["stats"]["daily_new_users"])
        sync_db()
    if "novel_progress" not in db["users"][uid]:
        db["users"][uid]["novel_progress"] = {}
    if "list_items" not in db["users"][uid]:
        db["users"][uid]["list_items"] = {}
        # ترحيل المفضلة القديمة (favs) لقسم "المفضلة" الافتراضي بالقائمة الجديدة، حتى ما يخسر
        # المستخدم مفضلته القديمة بعد إضافة نظام الأقسام.
        old_favs = db["users"][uid].get("favs", [])
        if old_favs:
            db["users"][uid]["list_items"]["sec_favs"] = [{"type": "item", "id": iid} for iid in old_favs]
        sync_db()
    return db["users"][uid]

def check_sub(uid):
    u = db["users"].get(str(uid))
    if not u: return False
    if u.get("sub_expiry"):
        try: return datetime.now() < datetime.fromisoformat(u["sub_expiry"])
        except: return False
    return False

def get_merge_limit(uid):
    """يرجع الحد المسموح للمستخدم بالتنزيلات المدمجة اليوم. None = غير محدود."""
    ustr = str(uid)
    ml = db["config"].get("merge_limits", {})
    override = ml.get("user_overrides", {}).get(ustr)
    if override is not None:
        return None if override == -1 else override  # -1 يعني غير محدود لهذا المستخدم تحديدًا
    if check_sub(uid):
        return ml.get("sub_daily")  # None يعني غير محدود للمشتركين
    return ml.get("free_daily", 10)

def get_merge_usage_today(ustr):
    u = db["users"].get(ustr, {})
    today = str(datetime.now().date())
    return u.get("merge_downloads", {}).get(today, 0)

def can_download_merge(uid):
    limit = get_merge_limit(uid)
    if limit is None: return True, None, None
    used = get_merge_usage_today(str(uid))
    return used < limit, used, limit

def record_merge_download(ustr):
    today = str(datetime.now().date())
    u = db["users"].setdefault(ustr, {})
    u.setdefault("merge_downloads", {})
    u["merge_downloads"][today] = u["merge_downloads"].get(today, 0) + 1
    sync_db()

def check_channels(uid):
    if is_admin(uid): return []  # المطور والمشرفين معفيين من الاشتراك الإجباري بالكامل
    out = []
    for ch in db["config"].get("mandatory_channels", []):
        try:
            m = bot.get_chat_member(ch, uid)
            if m.status in ["left","kicked"]: out.append(ch)
        except: out.append(ch)
    return out

def send_channels_msg(chat_id, chs):
    m = types.InlineKeyboardMarkup()
    for ch in chs:
        m.add(types.InlineKeyboardButton(f"📢 {ch}", url=f"https://t.me/{ch.lstrip('@')}"))
    m.add(types.InlineKeyboardButton("✅ تحققت", callback_data="check_joined"))
    bot.send_message(chat_id, "⚠️ *اشترك في القنوات أولاً:*", reply_markup=m, parse_mode="Markdown")

def do_redeem(chat_id, ustr, code):
    code = code.upper().strip()
    c = db.get("codes", {}).get(code)
    if not c: bot.send_message(chat_id, "❌ الكود غير صحيح."); return
    if c.get("used"): bot.send_message(chat_id, "❌ الكود مستخدم مسبقاً."); return
    if c.get("expires_at"):
        try:
            if datetime.fromisoformat(c["expires_at"]) < datetime.now():
                bot.send_message(chat_id, "❌ انتهت صلاحية هذا الكود."); return
        except: pass
    if ustr not in db["users"]: db["users"][ustr] = {"points": 0}
    pts = c["points"]
    db["users"][ustr]["points"] = db["users"][ustr].get("points", 0) + pts
    db["codes"][code]["used"] = True; sync_db()
    bot.send_message(chat_id, f"🎉 تم استبدال الكود!\n*+{pts}* نقطة أضيفت لحسابك.", parse_mode="Markdown")

# ==============================================================================
# 6. AI (OPENROUTER - DEEPSEEK)
# ==============================================================================
def q_ai(prompt, system_prompt=None, image_url=None):
    # فحص مبكر: إذا المفتاح لسا القيمة الافتراضية (placeholder) أو فاضي، لا نرسل
    # الطلب أصلاً (توفير وقت + رسالة خطأ واضحة للمطور بدل خطأ عام غامض).
    if not OR_KEY or OR_KEY.strip() == "":
        logger.error("❌ AI error: OPENROUTER_API_KEY لم يُضبط بعد. ضع مفتاح OpenRouter حقيقي كمتغير بيئة BOT_OPENROUTER_KEY.")
        return {"__ai_error__": "missing_key"}
    if not system_prompt:
        system_prompt = (
            "أنت مساعد ذكاء اصطناعي متطور لبوت تليجرام. "
            "أعد استجابتك بتنسيق JSON فقط مع الحقول المطلوبة. "
            "بدون أي نص إضافي أو علامات اقتباس خلفية (backticks). "
            "يجب أن تحتوي استجابتك على حقل `suggested_code` (كود Python مقترح) وحقل `explanation` (شرح للتعديل). "
            "إذا كان التعديل معقداً جداً أو يتطلب معلومات إضافية، يمكنك طلب توضيح."
        )
    try:
        r = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {OR_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://manus.ai",
                "X-Title": "Traika Bot"
            },
            json={
                "model": A_MDL,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": [
                        {"type": "text", "text": prompt},
                        *([{"type": "image_url", "image_url": {"url": image_url}}] if image_url else [])
                    ]}
                ],
                "temperature": 0.7, "max_tokens": 1000
            }, timeout=45
        )
        data = r.json()
        if "choices" in data:
            content = data["choices"][0]["message"]["content"].strip()
            if "```" in content:
                content = content.split("```")[1]
                if content.startswith("json"): content = content[4:]
            return json.loads(content.strip())
        else:
            logger.error(f"AI API Error: {data}")
            return {"__ai_error__": "api_error", "__ai_error_detail__": str(data)[:300]}
    except Exception as e:
        logger.error(f"AI error: {e}")
        return {"__ai_error__": "exception", "__ai_error_detail__": str(e)[:300]}

# ==============================================================================
# 6ب. SMART SEARCH ENGINE (يدعم نظام المحادثة الذكية بمعلومات حقيقية)
# ==============================================================================
class SmartSearchEngine:
    """محرك بحث خارجي بسيط يجمع معلومات عن عمل معيّن (فيلم/أنمي/رواية/إلخ) من عدة
    مصادر مجانية (ويكيبيديا، MyAnimeList، IMDB، DuckDuckGo) عشان تكون منشورات
    "المحادثة الذكية" مبنية على معلومات حقيقية بدل ما الذكاء الاصطناعي يختلقها بالكامل.
    كل نتيجة تُخزَّن مؤقتًا (cache) بقاعدة البيانات لمدة ai_search.cache_duration
    ثانية حتى ما نكرر نفس الطلبات الخارجية لنفس الموضوع بكل مرة."""

    def __init__(self):
        ai_search_cfg = db["config"].get("ai_search", {})
        self.cache = ai_search_cfg.get("search_cache", {})
        self.cache_duration = ai_search_cfg.get("cache_duration", 3600)

    def search(self, query, search_type="general"):
        cache_key = f"{query}_{search_type}"
        cached = self.cache.get(cache_key)
        if cached and (time.time() - cached.get("timestamp", 0) < self.cache_duration):
            return cached["data"]

        results = {}
        wiki_data = self.search_wikipedia(query)
        if wiki_data:
            results["wikipedia"] = wiki_data

        if search_type in ("anime", "manga", "general"):
            mal_data = self.search_myanimelist(query)
            if mal_data:
                results["myanimelist"] = mal_data

        if search_type in ("movie", "series", "general"):
            imdb_data = self.search_imdb(query)
            if imdb_data:
                results["imdb"] = imdb_data

        if not results or search_type == "general":
            ddg_data = self.search_duckduckgo(query)
            if ddg_data:
                results["duckduckgo"] = ddg_data

        combined_text = self.combine_search_results(results, query)
        db["config"].setdefault("ai_search", {}).setdefault("search_cache", {})[cache_key] = {
            "timestamp": time.time(), "data": combined_text
        }
        sync_db()
        return combined_text

    def search_wikipedia(self, query):
        if not wikipedia:
            return None
        try:
            wikipedia.set_lang("ar")
            try:
                page = wikipedia.page(query, auto_suggest=True)
                return {"title": page.title, "summary": page.summary[:1000], "url": page.url}
            except Exception:
                wikipedia.set_lang("en")
                page = wikipedia.page(query, auto_suggest=True)
                return {"title": page.title, "summary": page.summary[:1000], "url": page.url}
        except Exception:
            return None

    def search_myanimelist(self, query):
        if not BeautifulSoup:
            return None
        try:
            url = f"https://myanimelist.net/search/all?q={query.replace(' ', '+')}"
            headers = {"User-Agent": "Mozilla/5.0"}
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code != 200:
                return None
            soup = BeautifulSoup(response.text, "html.parser")
            result = {"source": "MyAnimeList"}
            first_result = soup.select_one(".js-search-result")
            if first_result:
                title = first_result.select_one(".h2_anime_title a")
                if title:
                    result["title"] = title.text.strip()
            return result if result.get("title") else None
        except Exception:
            return None

    def search_imdb(self, query):
        if not BeautifulSoup:
            return None
        try:
            url = f"https://www.imdb.com/find?q={query.replace(' ', '+')}"
            headers = {"User-Agent": "Mozilla/5.0"}
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code != 200:
                return None
            soup = BeautifulSoup(response.text, "html.parser")
            results = {"source": "IMDB", "results": []}
            for item in soup.select(".find-result-item")[:3]:
                title = item.select_one(".result_text a")
                if title:
                    results["results"].append({"title": title.text.strip(), "link": "https://www.imdb.com" + title.get("href", "")})
            return results if results["results"] else None
        except Exception:
            return None

    def search_duckduckgo(self, query):
        try:
            url = f"https://api.duckduckgo.com/?q={query.replace(' ', '+')}&format=json"
            response = requests.get(url, timeout=10)
            if response.status_code != 200:
                return None
            data = response.json()
            if not data.get("Abstract"):
                return None
            return {"source": "DuckDuckGo", "summary": data.get("Abstract", ""), "url": data.get("AbstractURL", "")}
        except Exception:
            return None

    def combine_search_results(self, results, query):
        combined_text = f"📝 معلومات عن: {query}\n\n"
        if "wikipedia" in results:
            wiki = results["wikipedia"]
            combined_text += f"📚 من ويكيبيديا:\n{wiki['summary'][:500]}...\n\n"
        if "myanimelist" in results and results["myanimelist"].get("title"):
            combined_text += f"🎌 من MyAnimeList:\n{results['myanimelist']['title']}\n\n"
        if "imdb" in results and results["imdb"].get("results"):
            combined_text += "🎬 من IMDB:\n"
            for item in results["imdb"]["results"][:3]:
                combined_text += f"• {item['title']}\n"
            combined_text += "\n"
        if "duckduckgo" in results and results["duckduckgo"].get("summary"):
            combined_text += f"🔍 معلومات عامة:\n{results['duckduckgo']['summary'][:300]}...\n\n"
        if combined_text.strip() == f"📝 معلومات عن: {query}":
            combined_text += "لا توجد معلومات إضافية متاحة من المصادر الخارجية حاليًا.\n"
        return combined_text

# ==============================================================================
# 6ج. نظام المسابقات المتطور (Contests) — أسئلة تلقائية بجوائز نقاط/VIP اختيارية
# ==============================================================================
def validate_questions(questions, choices_count=None):
    """يتحقق من صحة قائمة أسئلة ويصلحها إذا لزم الأمر: يضمن وجود خيارات كافية،
    عدم تكرارها، وأن مؤشر الإجابة الصحيحة صالح ضمن نطاق الخيارات."""
    if not questions:
        return []
    if choices_count is None:
        choices_count = db["config"]["contests"].get("choices_count", 2)
    validated = []
    for q in questions:
        if not q.get("question"):
            continue
        if not q.get("options") or len(q["options"]) < 2:
            q["options"] = ["نعم", "لا"] if choices_count == 2 else ["أ", "ب", "ج", "د"]
        q["options"] = list(dict.fromkeys(q["options"]))  # إزالة تكرار الخيارات مع حفظ الترتيب
        if "correct" not in q or q["correct"] >= len(q["options"]):
            q["correct"] = 0
        validated.append(q)
    return validated

def generate_questions_ai(question_type, count=5, choices=2):
    """يستخدم الذكاء الاصطناعي (q_ai) لتوليد أسئلة جديدة غير مكررة، مع تتبع نصوص
    الأسئلة المستخدمة سابقًا (آخر 100) وتمريرها للنموذج صراحة حتى يتجنبها."""
    if not OR_KEY or OR_KEY.strip() == "":
        return None
    topics = {
        "novel": "الروايات العربية والعالمية", "manga": "المانجا والمانهوا",
        "series": "المسلسلات والأفلام", "anime": "الأنمي", "general": "المعرفة العامة",
        "mixed": "الأنمي، المانجا، الروايات، المسلسلات",
    }
    topic = topics.get(question_type, "المعرفة العامة")
    used = db["config"]["contests"].get("used_questions", [])
    used_text = "\n".join(used[-20:]) if used else "لا يوجد"
    prompt = (
        f"أنت خبير في {topic}. قم بإنشاء {count} أسئلة اختيار من متعدد، كل سؤال له {choices} خيارات.\n"
        f"تجنب تمامًا الأسئلة التالية (لا تكررها بأي شكل):\n{used_text}\n\n"
        'الصيغة المطلوبة (JSON): [{"question":"نص السؤال","options":["خيار1","خيار2"],'
        '"correct":0,"explanation":"شرح مختصر اختياري"}]\n'
        "تأكد من تنوع الأسئلة وصحة المعلومات (الأسماء والتواريخ والحقائق). "
        "أجب فقط بصيغة JSON قابلة للتحليل، بدون أي نص إضافي."
    )
    for attempt in range(3):
        try:
            resp = q_ai(prompt, system_prompt="أنت مساعد متخصص في توليد الأسئلة. أجب فقط بصيغة JSON صحيحة.")
            if isinstance(resp, dict) and "__ai_error__" in resp:
                continue
            if isinstance(resp, list):
                questions = resp
            elif isinstance(resp, dict) and "questions" in resp:
                questions = resp["questions"]
            else:
                questions = []
                if isinstance(resp, dict):
                    for val in resp.values():
                        if isinstance(val, list) and val:
                            questions = val
                            break
            if not questions:
                continue
            validated = []
            for q in questions[:count]:
                if isinstance(q, dict) and "question" in q:
                    if "options" not in q or not q["options"]:
                        q["options"] = ["نعم", "لا"] if choices == 2 else ["أ", "ب", "ج", "د"][:choices]
                    if "correct" not in q or q["correct"] >= len(q["options"]):
                        q["correct"] = 0
                    q_text = q["question"][:50]
                    if q_text not in used:
                        validated.append(q)
                        used.append(q_text)
            if validated:
                db["config"]["contests"]["used_questions"] = used[-100:]
                sync_db()
                return validated
        except Exception as e:
            logger.error(f"❌ محاولة توليد أسئلة {attempt+1} فشلت: {e}")
    return None

def generate_questions_from_db(question_type, count=5, choices=2):
    """يولّد أسئلة من بيانات البوت الفعلية (روايات/مانجا/مسلسلات مخزّنة): يبني سؤال
    'ما هو عنوان العمل الذي نصه كذا؟' مع خيارات مموّهة من عناوين أعمال أخرى حقيقية."""
    pool = []
    if question_type in ("novel", "mixed"):
        for nid, novel in db.get("novels", {}).items():
            if novel.get("title") and novel.get("story"):
                pool.append({"type": "novel", "title": novel["title"], "story": novel["story"][:500]})
    if question_type in ("manga", "mixed"):
        for mid, manga in db.get("manga", {}).items():
            if manga.get("title") and manga.get("story"):
                pool.append({"type": "manga", "title": manga["title"], "story": manga["story"][:500]})
    if question_type in ("series", "mixed"):
        for sid, series in db.get("series", {}).items():
            if series.get("title") and series.get("story"):
                pool.append({"type": "series", "title": series["title"], "story": series["story"][:500]})
    if not pool:
        return None
    selected = random.sample(pool, min(count, len(pool)))
    questions = []
    for item in selected:
        title, story = item["title"], item["story"]
        question_text = f'ما هو عنوان العمل الذي نصه: "{story[:100]}..."؟'
        options = [title]
        other_titles = [p["title"] for p in pool if p["title"] != title]
        if other_titles:
            wrong = random.sample(other_titles, min(choices - 1, len(other_titles)))
            options.extend(wrong)
        random.shuffle(options)
        correct_index = options.index(title)
        questions.append({"question": question_text, "options": options, "correct": correct_index, "source": item["type"]})
    return questions[:count] if questions else None

def generate_fallback_questions(question_type, count=5, choices=2):
    """أسئلة احتياطية جاهزة تُستخدم فقط لو فشل كل من الذكاء الاصطناعي وقاعدة البيانات،
    حتى تبقى المسابقة قابلة للتشغيل دومًا بدون انقطاع."""
    fallback_pool = {
        "novel": [
            {"question": "من هو مؤلف رواية 1984؟", "options": ["جورج أورويل", "ألدوس هكسلي"], "correct": 0},
            {"question": "ما هي عاصمة مصر؟", "options": ["القاهرة", "الإسكندرية"], "correct": 0},
        ],
        "manga": [{"question": "من هو مؤلف ون بيس؟", "options": ["إيتشيرو أودا", "أوياما"], "correct": 0}],
        "series": [{"question": "في أي مسلسل يظهر شخصية 'والتر وايت'؟", "options": ["صراع العروش", "اختلال ضال"], "correct": 1}],
        "anime": [{"question": "من هو بطل أنمي ناروتو؟", "options": ["ناروتو", "ساسكي"], "correct": 0}],
        "general": [{"question": "ما هو أكبر محيط في العالم؟", "options": ["المحيط الهادئ", "المحيط الأطلسي"], "correct": 0}],
    }
    pool = fallback_pool.get(question_type, fallback_pool["general"])
    if question_type == "mixed":
        pool = fallback_pool["anime"] + fallback_pool["manga"] + fallback_pool["novel"]
    if len(pool) < count:
        pool = pool * (count // len(pool) + 1)
    selected = random.sample(pool, min(count, len(pool)))
    for q in selected:
        q = dict(q)  # نسخة حتى لا نعدّل القاموس الأصلي بالـ pool المشترك
        if len(q["options"]) != choices:
            if choices == 2 and len(q["options"]) > 2:
                q["options"] = q["options"][:2]
                if q["correct"] >= 2:
                    q["correct"] = 0
            elif choices == 4 and len(q["options"]) < 4:
                while len(q["options"]) < 4:
                    q["options"].append(f"خيار {len(q['options'])+1}")
    return selected

def generate_questions_from_data(question_type, count=5, choices=2):
    """نقطة الدخول الموحّدة لتوليد الأسئلة: يجرّب المصدر المفضّل بالإعدادات (ai/database/mixed)
    ثم يتدرّج للاحتياطي التالي تلقائيًا حتى لا تفشل المسابقة أبدًا."""
    source = db["config"]["contests"].get("question_source", "ai")
    questions = None
    if source != "database" and OR_KEY and OR_KEY.strip():
        questions = generate_questions_ai(question_type, count, choices)
    if not questions:
        questions = generate_questions_from_db(question_type, count, choices)
    if not questions:
        questions = generate_fallback_questions(question_type, count, choices)
    return validate_questions(questions, choices)

def add_points(user_id, points):
    """يضيف نقاط لمستخدم موجود بالفعل. يُستخدم من نظام المسابقات كجائزة."""
    ustr = str(user_id)
    if ustr in db["users"]:
        db["users"][ustr]["points"] = db["users"][ustr].get("points", 0) + points
        sync_db()
        return True
    return False

def activate_vip(user_id, days):
    """اسم مختصر لتفعيل اشتراك VIP كجائزة مسابقة — يستخدم نفس منطق activate_subscription."""
    return activate_subscription(str(user_id), days)

def _format_prizes(prizes):
    """يبني سطر نصي مختصر يلخّص جوائز المسابقة (نقاط/VIP/رتبة) لعرضه بلوحة الإعدادات."""
    if not prizes:
        return "🚫 بدون جوائز"
    text = ""
    if prizes.get("points"): text += f"💎 {prizes['points']} نقطة "
    if prizes.get("vip_days"): text += f"👑 {prizes['vip_days']} يوم "
    if prizes.get("role"): text += f"🏅 {prizes['role']}"
    return text or "🚫 بدون جوائز"

def prepare_contest_data():
    """يجهّز بيانات مسابقة جديدة كاملة: أسئلة مولّدة ومُتحقَّق منها + جوائز مضبوطة."""
    cfg = db["config"]["contests"]
    questions = generate_questions_from_data(cfg.get("question_type", "mixed"), cfg.get("question_count", 5), cfg.get("choices_count", 2))
    if not questions:
        logger.error("❌ فشل توليد أسئلة المسابقة من كل المصادر")
        return None
    prizes = cfg.get("prizes", {})
    has_prizes = any(prizes.get(k) for k in ("points", "vip_days", "role"))
    return {"name": f"مسابقة {cfg.get('question_type', 'عامة')}", "questions": questions, "prizes": prizes if has_prizes else {}}

def send_contest_to_chat(target, contest):
    """يرسل المسابقة إلى القناة/المجموعة المستهدفة، بوضع الأزرار أو التعليق حسب الإعداد."""
    questions = contest["questions"]
    prizes = contest.get("prizes", {})
    answer_mode = contest.get("answer_mode", "button")

    text = f"🎮 *{contest['name']}*\n\n"
    for i, q in enumerate(questions, 1):
        if answer_mode == "button":
            opts = "\n".join([f"{chr(65+j)}. {opt}" for j, opt in enumerate(q["options"])])
            text += f"*{i}. {q['question']}*\n{opts}\n\n"
        else:
            text += f"*{i}. {q['question']}*\n📝 *اكتب إجابتك في التعليق*\n\n"

    if prizes:
        ptext = "\n🎁 *الجوائز:*\n"
        if prizes.get("points"): ptext += f"💎 {prizes['points']} نقطة\n"
        if prizes.get("vip_days"): ptext += f"👑 اشتراك VIP {prizes['vip_days']} أيام\n"
        if prizes.get("role"): ptext += f"🏅 رتبة '{prizes['role']}'\n"
        text += ptext

    text += f"\n⏰ تنتهي المسابقة بعد {db['config']['contests'].get('auto_end_minutes', 15)} دقيقة"
    bot.send_message(target, text, parse_mode="Markdown")

    if answer_mode == "button":
        for i, q in enumerate(questions):
            markup = types.InlineKeyboardMarkup(row_width=2)
            for j, opt in enumerate(q["options"]):
                markup.add(types.InlineKeyboardButton(chr(65 + j), callback_data=f"contest_answer_{contest['id']}_{i}_{j}"))
            bot.send_message(target, f"❓ *السؤال {i+1}:*\n{q['question']}\nاختر إجابتك:", reply_markup=markup, parse_mode="Markdown")
            time.sleep(db["config"]["contests"].get("time_between_questions", 5))
    else:
        bot.send_message(
            target,
            "📝 *للإجابة:* أرسل رقم السؤال ثم إجابتك في رسالة واحدة\nمثال: `1 ج` أو `2 ب`\n\nيمكنك الإجابة على كل الأسئلة في رسائل منفصلة.",
            parse_mode="Markdown",
        )

def start_contest(contest_data=None):
    """يبدأ مسابقة جديدة في القناة/المجموعة المستهدفة (target_chat_id أو announce_channel احتياطيًا)."""
    target = db["config"]["contests"].get("target_chat_id") or db["config"]["announce_channel"].get("chat_id")
    if not target:
        logger.error("❌ لا توجد قناة/مجموعة مستهدفة للمسابقة")
        return None
    if contest_data is None:
        contest_data = prepare_contest_data()
    if not contest_data:
        return None
    contest_id = str(uuid.uuid4())[:8]
    answer_mode = db["config"]["contests"].get("answer_mode", "button")
    contest = {
        "id": contest_id,
        "name": contest_data.get("name", "مسابقة"),
        "questions": contest_data["questions"],
        "participants": {},
        "start_time": datetime.now().isoformat(),
        "end_time": (datetime.now() + timedelta(minutes=db["config"]["contests"].get("auto_end_minutes", 15))).isoformat(),
        "prizes": contest_data.get("prizes", {}),
        "answer_mode": answer_mode,
        "status": "active",
    }
    db["config"]["contests"]["active_contests"][contest_id] = contest
    db["config"]["contests"]["contest_counter"] = db["config"]["contests"].get("contest_counter", 0) + 1
    sync_db()
    try:
        send_contest_to_chat(target, contest)
    except Exception as e:
        logger.error(f"❌ فشل إرسال مسابقة: {e}")
    threading.Timer(db["config"]["contests"].get("auto_end_minutes", 15) * 60, auto_end_contest, args=[contest_id]).start()
    return contest_id

def handle_comment_answer(msg, contest_id, user_id):
    """يعالج إجابة مكتوبة في التعليق (وضع comment) بصيغة 'رقم_السؤال حرف_الإجابة'،
    ويدعم عدة صيغ كتابة عربية/إنجليزية شائعة."""
    active = db["config"]["contests"].get("active_contests", {})
    contest = active.get(contest_id)
    if not contest:
        return
    text = (msg.text or "").strip()
    if not text:
        return
    patterns = [
        r'^(\d+)\s*[\.\-:]\s*([أ-يa-zA-Z])$',
        r'^(\d+)\s*([أ-يa-zA-Z])$',
        r'^السؤال\s*(\d+)\s*[\.\-:]\s*([أ-يa-zA-Z])$',
        r'^س\s*(\d+)\s*[\.\-:]\s*([أ-يa-zA-Z])$',
    ]
    ustr = str(user_id)
    for pattern in patterns:
        match = re.match(pattern, text)
        if not match:
            continue
        q_num = int(match.group(1))
        answer_char = match.group(2).upper()
        arabic_map = {"أ": 0, "ب": 1, "ج": 2, "د": 3}
        if match.group(2) in arabic_map:
            answer_idx = arabic_map[match.group(2)]
        else:
            answer_idx = ord(answer_char) - ord('A')
        if 0 <= q_num - 1 < len(contest["questions"]):
            participant = contest["participants"].setdefault(ustr, {"answers": [], "score": 0})
            if len(participant["answers"]) <= q_num - 1:
                participant["answers"].append(answer_idx)
                sync_db()
                try:
                    bot.reply_to(msg, f"✅ تم تسجيل إجابتك للسؤال {q_num}")
                except Exception:
                    pass
            else:
                try:
                    bot.reply_to(msg, f"⚠️ لقد أجبت على السؤال {q_num} مسبقًا")
                except Exception:
                    pass
        return
    try:
        bot.reply_to(msg, "⚠️ صيغة غير صحيحة. استخدم: `رقم_السؤال الحرف`\nمثال: `1 أ` أو `1 A`", parse_mode="Markdown")
    except Exception:
        pass

def auto_end_contest(contest_id):
    """ينهي مسابقة، يحسب النتائج (لكل من وضعي الأزرار والتعليق)، ويعلن الفائزين ويوزّع الجوائز."""
    active = db["config"]["contests"]["active_contests"]
    contest = active.pop(contest_id, None)
    if not contest:
        return
    contest["status"] = "ended"
    contest["end_time"] = datetime.now().isoformat()

    results = []
    for uid_key, data in contest["participants"].items():
        score = 0
        for i, answer in enumerate(data.get("answers", [])):
            if i < len(contest["questions"]) and answer == contest["questions"][i]["correct"]:
                score += 1
        data["score"] = score
        results.append((uid_key, score))
    results.sort(key=lambda x: x[1], reverse=True)

    target = db["config"]["contests"].get("target_chat_id") or db["config"]["announce_channel"].get("chat_id")
    if target:
        text = f"🏆 *نتائج المسابقة*\n\n*{contest['name']}*\n"
        if results:
            for i, (uid_key, score) in enumerate(results[:5], 1):
                try:
                    user = bot.get_chat(int(uid_key))
                    name = user.first_name or user.username or uid_key
                except Exception:
                    name = uid_key
                text += f"{i}. {name} - {score} نقطة\n"
        else:
            text += "لا توجد مشاركات"
        if contest.get("answer_mode") == "comment":
            text += "\n📝 *الإجابات الصحيحة:*\n"
            for i, q in enumerate(contest["questions"], 1):
                correct_idx = q.get("correct", 0)
                correct_opt = q["options"][correct_idx] if correct_idx < len(q["options"]) else "غير معروف"
                text += f"{i}. {q['question'][:50]}... → {correct_opt}\n"
        try:
            bot.send_message(target, text, parse_mode="Markdown")
        except Exception as e:
            logger.error(f"❌ فشل إرسال نتائج المسابقة: {e}")

        prizes = contest.get("prizes", {})
        if prizes and results:
            for rank, (uid_key, score) in enumerate(results[:3], 1):
                if score == 0 and rank > 1:
                    continue
                pts = prizes.get("points", 0) // (rank * 2) if rank > 1 else prizes.get("points", 0)
                vdays = prizes.get("vip_days", 0) // (rank * 2) if rank > 1 else prizes.get("vip_days", 0)
                role = f"{prizes.get('role', 'بطل')} #{rank}" if rank > 1 else prizes.get("role", "بطل")
                if pts: add_points(uid_key, pts)
                if vdays: activate_vip(uid_key, vdays)
                try:
                    wmsg = f"🎉 *تهانينا!*\nفزت بـ {contest['name']}!\nمركزك: #{rank}\n"
                    if pts: wmsg += f"💎 +{pts} نقطة\n"
                    if vdays: wmsg += f"👑 اشتراك VIP {vdays} أيام\n"
                    if role: wmsg += f"🏅 رتبة: {role}\n"
                    bot.send_message(int(uid_key), wmsg, parse_mode="Markdown")
                except Exception:
                    pass

    past = db["config"]["contests"]["past_contests"]
    past.append(contest)
    if len(past) > 50:
        db["config"]["contests"]["past_contests"] = past[-50:]
    sync_db()

def record_contest_answer(contest_id, uid, q_index, option_index):
    """يسجّل إجابة مستخدم على سؤال معيّن (وضع الأزرار). يمنع الإجابة المكررة لنفس السؤال."""
    active = db["config"]["contests"].get("active_contests", {})
    contest = active.get(contest_id)
    if not contest:
        return False, "❌ هذي المسابقة انتهت."
    ustr = str(uid)
    if q_index >= len(contest["questions"]):
        return False, "❌ سؤال غير صالح."
    participant = contest["participants"].setdefault(ustr, {"answers": [], "score": 0})
    if len(participant["answers"]) <= q_index:
        participant["answers"].append(option_index)
        sync_db()
        is_correct = (option_index == contest["questions"][q_index].get("correct"))
        return True, (f"✅ تم تسجيل إجابتك للسؤال {q_index+1}" if not db["config"]["contests"].get("allow_duplicate_answers") else ("✅ إجابة صحيحة!" if is_correct else "❌ إجابة خاطئة."))
    return False, "⚠️ لقد أجبت على هذا السؤال مسبقًا"

def contests_scheduler_loop():
    """يعمل في الخلفية لبدء المسابقات التلقائية وفق الوضع (auto) والوقت والأيام المحددة."""
    while True:
        try:
            cfg = db["config"]["contests"]
            if cfg.get("enabled") and cfg.get("mode") == "auto":
                now = datetime.now()
                today = now.strftime("%a").lower()
                schedule_time = cfg.get("schedule_time", "18:00")
                allowed_days = cfg.get("schedule_days", ["sat", "sun", "mon", "tue", "wed", "thu", "fri"])
                if today in allowed_days and now.strftime("%H:%M") == schedule_time:
                    if not cfg.get("active_contests"):
                        today_contests = [c for c in cfg.get("past_contests", []) if c.get("start_time", "").startswith(now.date().isoformat())]
                        if len(today_contests) < 1:
                            logger.info("🚀 بدء مسابقة تلقائية")
                            start_contest()
                            time.sleep(60)
            time.sleep(30)
        except Exception as e:
            logger.error(f"❌ خطأ في جدولة المسابقات: {e}")
            time.sleep(60)

def _contests_panel_content():
    """يبني نص وأزرار لوحة تحكم المسابقات المتطورة للوحة الإدارة."""
    cfg = db["config"]["contests"]
    status = "🟢 مفعّل" if cfg.get("enabled") else "🔴 معطّل"
    mode = "🔄 تلقائي" if cfg.get("mode") == "auto" else "✋ يدوي"
    active = len(cfg.get("active_contests", {}))
    past = len(cfg.get("past_contests", []))
    answer_mode = "🔘 أزرار" if cfg.get("answer_mode") == "button" else "📝 تعليق"
    target = cfg.get("target_chat_id") or db["config"]["announce_channel"].get("chat_id") or "غير محدد"
    prize_text = _format_prizes(cfg.get("prizes", {}))
    text = (f"🎮 *المسابقات المتطورة*\n\nالحالة: {status}\nالوضع: {mode}\nطريقة الإجابة: {answer_mode}\n"
            f"الهدف: `{target}`\n\nالأسئلة: {cfg.get('question_count', 5)} سؤال، {cfg.get('choices_count', 2)} خيار\n"
            f"النوع: {cfg.get('question_type', 'mixed')}\nالمصدر: {cfg.get('question_source', 'ai')}\nالجوائز: {prize_text}\n\n"
            f"⏰ الجدولة: {cfg.get('schedule_time', '18:00')}\nأيام: {', '.join(cfg.get('schedule_days', []))}\n\n"
            f"المسابقات النشطة: {active}\nالمسابقات السابقة: {past}")
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if cfg.get('enabled') else '🟢 تفعيل'}", callback_data="contest_toggle_enabled"))
    m.add(types.InlineKeyboardButton(f"{'⏸️ إيقاف الجدولة' if cfg.get('mode')=='auto' else '▶️ تشغيل الجدولة'}", callback_data="contest_toggle_mode"))
    m.add(types.InlineKeyboardButton("▶️ بدء مسابقة الآن", callback_data="contest_start_now"))
    m.add(types.InlineKeyboardButton("🔘 تبديل طريقة الإجابة", callback_data="contest_toggle_answer"))
    m.add(types.InlineKeyboardButton("🎯 تعيين الهدف", callback_data="contest_set_target"))
    m.add(types.InlineKeyboardButton("⚙️ الإعدادات", callback_data="contest_settings"))
    m.add(types.InlineKeyboardButton("🏆 النتائج السابقة", callback_data="contest_past_results"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return text, m

def _contests_settings_panel_content():
    """يبني نص وأزرار لوحة الإعدادات التفصيلية للمسابقات (نوع الأسئلة، الجدولة، الجوائز، إلخ)."""
    cfg = db["config"]["contests"]
    text = (f"⚙️ *إعدادات المسابقات*\n\nنوع الأسئلة: {cfg.get('question_type', 'mixed')}\n"
            f"عدد الأسئلة: {cfg.get('question_count', 5)}\nعدد الخيارات: {cfg.get('choices_count', 2)}\n"
            f"مصدر الأسئلة: {cfg.get('question_source', 'ai')}\n\nوقت الجدولة: {cfg.get('schedule_time', '18:00')}\n"
            f"أيام الجدولة: {', '.join(cfg.get('schedule_days', []))}\n\nالجوائز:\n{_format_prizes(cfg.get('prizes', {}))}")
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(types.InlineKeyboardButton("📝 نوع الأسئلة", callback_data="contest_set_type"))
    m.add(types.InlineKeyboardButton("🔢 عدد الأسئلة", callback_data="contest_set_count"))
    m.add(types.InlineKeyboardButton("🔢 عدد الخيارات", callback_data="contest_set_choices"))
    m.add(types.InlineKeyboardButton("📚 مصدر الأسئلة", callback_data="contest_set_source"))
    m.add(types.InlineKeyboardButton("⏰ وقت الجدولة", callback_data="contest_set_time"))
    m.add(types.InlineKeyboardButton("📅 أيام الجدولة", callback_data="contest_set_days"))
    m.add(types.InlineKeyboardButton("🎁 تعديل الجوائز", callback_data="contest_set_prizes"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_contests"))
    return text, m

# ==============================================================================
# 6هـ. الدعوة التلقائية عند التفاعل (Auto Invite) — دعوة لقناة/مجموعة أخرى تُرسل
# خاصة لأي مستخدم يتفاعل مع البوت (زر/رسالة/انضمام)، باعتبار التفاعل موافقة مسبقة.
# ==============================================================================
def get_or_create_target_invite_link():
    """ينشئ رابط دعوة للقناة/المجموعة المستهدفة ويخزّنه (بدون إعادة إنشائه في كل مرة)."""
    cfg = db["config"]["auto_invite"]
    if not cfg.get("enabled") or not cfg.get("target_chat_id"):
        return None
    if cfg.get("invite_link"):
        return cfg["invite_link"]
    try:
        link = bot.create_chat_invite_link(cfg["target_chat_id"], name="دعوة تلقائية", creates_join_request=False)
        cfg["invite_link"] = link.invite_link
        sync_db()
        return link.invite_link
    except Exception as e:
        logger.error(f"❌ فشل إنشاء رابط الدعوة: {e}")
        return None

def send_invite_to_user(user_id, trigger_type="تفاعل"):
    """يرسل رسالة تأكيد خاصة للمستخدم مع رابط دعوة القناة/المجموعة المستهدفة."""
    link = get_or_create_target_invite_link()
    if not link:
        return
    cfg = db["config"]["auto_invite"]
    text = cfg["confirmation_text"].format(link=link)
    m = types.InlineKeyboardMarkup()
    m.add(types.InlineKeyboardButton(cfg["button_label"], url=link))
    try:
        bot.send_message(user_id, text, reply_markup=m)
        logger.info(f"✅ تم إرسال رابط الدعوة للمستخدم {user_id} (سبب: {trigger_type})")
    except Exception as e:
        logger.warning(f"⚠️ فشل إرسال الدعوة للمستخدم {user_id}: {e}")

def _auto_invite_panel_content():
    """يبني نص وأزرار لوحة تحكم الدعوة التلقائية للوحة الإدارة."""
    cfg = db["config"]["auto_invite"]
    status = "🟢 مفعّل" if cfg.get("enabled") else "🔴 معطّل"
    target = cfg.get("target_chat_id") or "غير محدد"
    triggers = "، ".join(cfg.get("trigger_on", []))
    text = f"🔗 *الدعوة التلقائية (بموافقة المستخدم)*\n\nالحالة: {status}\nالوجهة: `{target}`\nمسببات التفاعل: {triggers}"
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if cfg.get('enabled') else '🟢 تفعيل'}", callback_data="invite_toggle_enabled"))
    m.add(types.InlineKeyboardButton("🎯 تعيين الوجهة", callback_data="invite_set_target"))
    m.add(types.InlineKeyboardButton("📝 تعديل نص الموافقة", callback_data="invite_edit_text"))
    m.add(types.InlineKeyboardButton("🔘 تعديل زر الموافقة", callback_data="invite_edit_button"))
    m.add(types.InlineKeyboardButton("⚙️ اختيار أنواع التفاعل", callback_data="invite_triggers"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return text, m

# ==============================================================================
# 6د. نظام "المحادثة الذكية" (AI Talk) — منشورات تلقائية/يدوية مبنية على SmartSearchEngine
# ==============================================================================
def select_random_topic(talk_config):
    """يختار نوع محتوى عشوائي من الأنواع المفعّلة بالإعدادات، ثم عملاً عشوائيًا مرتبطًا به."""
    type_names = {
        "movies": {"name": "أفلام", "emoji": "🎬"},
        "series": {"name": "مسلسلات", "emoji": "📺"},
        "anime": {"name": "أنمي", "emoji": "🎌"},
        "cartoon": {"name": "كرتون", "emoji": "🖍️"},
        "manga": {"name": "مانجا", "emoji": "📖"},
        "manhwa": {"name": "مانهوا", "emoji": "🎨"},
        "novels": {"name": "روايات", "emoji": "📚"},
        "light_novels": {"name": "روايات خفيفة", "emoji": "📖"},
    }
    active_types = [type_names[t] for t, enabled in talk_config.get("content_types", {}).items() if enabled and t in type_names]
    if not active_types:
        active_types = list(type_names.values())
    selected = random.choice(active_types)
    works = {
        "أفلام": ["Inception", "The Dark Knight", "Interstellar", "The Godfather", "Pulp Fiction"],
        "أنمي": ["Attack on Titan", "Naruto", "One Piece", "Demon Slayer", "Jujutsu Kaisen", "Spy x Family"],
        "مانجا": ["One Piece", "Berserk", "Naruto", "Death Note", "Fullmetal Alchemist"],
        "روايات": ["1984", "The Great Gatsby", "Pride and Prejudice", "The Catcher in the Rye"],
    }
    work = random.choice(works.get(selected["name"], ["عمل مشهور"]))
    type_id = next((k for k, v in type_names.items() if v["name"] == selected["name"]), "general")
    return {"type": selected["name"], "emoji": selected["emoji"], "subtopic": "عام", "work": work, "type_id": type_id}

def create_smart_talk_post(talk_config, topic=None, search_query=None, news_item=None):
    """ينشئ نص منشور نقاش عن عمل معيّن (أو خبر) بأسلوب الشخصية المضبوطة، مستعينًا
    بمعلومات حقيقية من SmartSearchEngine. لو فشل الذكاء الاصطناعي، يبني منشورًا
    احتياطيًا من نتائج البحث مباشرة حتى لا تتعطل الميزة بالكامل."""
    if not topic:
        topic = select_random_topic(talk_config)
    query = search_query or topic.get("work", topic["type"])
    search_engine = SmartSearchEngine()
    search_results = search_engine.search(query, topic.get("type_id", "general"))
    personality_map = {
        "friendly": "ودود وحماسي", "professional": "محترف", "sarcastic": "ساخر",
        "enthusiastic": "حماسي", "calm": "هادئ", "mysterious": "غامض",
    }
    personality = personality_map.get(talk_config.get("personality"), "ودود")
    if news_item:
        prompt = (f"الخبر: {news_item.get('title')}\n{news_item.get('description')}\n"
                  f"المعلومات: {search_results}\nتحدث عن الخبر بأسلوب {personality} (300-500 كلمة). "
                  'أجب بصيغة JSON: {"title":"عنوان قصير","text":"نص المنشور","hashtags":["وسم1","وسم2"]}')
    else:
        prompt = (f"تحدث عن {topic['type']}: {topic['work']}\nالمعلومات: {search_results}\n"
                  f"بأسلوب {personality} (300-500 كلمة). "
                  'أجب بصيغة JSON: {"title":"عنوان قصير","text":"نص المنشور","hashtags":["وسم1","وسم2"]}')
    try:
        response = q_ai(prompt, system_prompt="أجب فقط بصيغة JSON صالحة للتحليل بالحقول المطلوبة، بدون أي نص إضافي.")
        if isinstance(response, dict) and "__ai_error__" not in response:
            post_text = response.get("text", "")
            title = response.get("title", f"💬 {topic['type']}")
            hashtags = response.get("hashtags", [])
        else:
            post_text, title, hashtags = f"{search_results}\n\nما رأيك في هذا الموضوع؟", f"💬 {topic['type']}", ["نقاش"]
    except Exception:
        post_text, title, hashtags = f"{search_results}\n\nما رأيك في هذا الموضوع؟", f"💬 {topic['type']}", ["نقاش"]
    hashtags_line = " ".join([f"#{tag}" for tag in hashtags]) if hashtags else ""
    return f"*{title}*\n\n{post_text}\n\n{hashtags_line}\n\n📌 *بحثت عن:* {query}"

def _ai_talk_panel_content():
    """يبني نص وأزرار لوحة تحكم المحادثة الذكية للوحة الإدارة."""
    talk_config = db["config"]["ai_talk"]
    status = "🟢 مفعّل" if talk_config["enabled"] else "🔴 معطّل"
    auto = "🔄 تلقائي" if talk_config["auto_post"] else "✋ يدوي"
    text = f"🤖 *المحادثة الذكية*\n\nالحالة: {status}\nالنشر: {auto}\nالشخصية: {talk_config['personality']}\nوقت النشر التلقائي: {talk_config['post_time']}"
    markup = types.InlineKeyboardMarkup(row_width=2)
    markup.add(
        types.InlineKeyboardButton("🔄 تبديل التشغيل", callback_data="talk_toggle"),
        types.InlineKeyboardButton("📤 تبديل النشر التلقائي", callback_data="talk_toggle_auto"),
    )
    markup.add(
        types.InlineKeyboardButton("🎭 تغيير الشخصية", callback_data="talk_personality"),
        types.InlineKeyboardButton("📝 إنشاء منشور الآن", callback_data="talk_create_now"),
    )
    markup.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return text, markup

def ai_talk_scheduler_loop():
    """يفحص كل 30 ثانية هل حان وقت النشر التلقائي المضبوط، وينشر منشورًا جديدًا بحدود post_count يوميًا."""
    while True:
        try:
            now = datetime.now()
            cfg = db["config"]["ai_talk"]
            if cfg.get("enabled") and cfg.get("auto_post"):
                if now.strftime("%H:%M") == cfg.get("post_time"):
                    today_posts = [p for p in cfg.get("posts_history", []) if p.get("timestamp", "").startswith(now.date().isoformat())]
                    if len(today_posts) < cfg.get("post_count", 1):
                        post = create_smart_talk_post(cfg)
                        channel = cfg.get("post_channel") or OWNER_ID
                        bot.send_message(channel, post, parse_mode="Markdown")
                        cfg.setdefault("posts_history", []).append({"timestamp": now.isoformat(), "text": post})
                        sync_db()
        except Exception as e:
            logger.error(f"AI talk scheduler error: {e}")
        time.sleep(30)

# ==============================================================================
# 7. MENU BUILDERS
# ==============================================================================
def build_main_menu(uid, parent_feature_id=None):
    m = types.InlineKeyboardMarkup(row_width=2)
    
    current_buttons = []
    back_button_callback = "go_home"

    if parent_feature_id:
        # If a parent_feature_id is provided, we are building a submenu
        parent_btn = next((b for b in db["config"]["menu_buttons"] if b["id"] == parent_feature_id), None)
        if not parent_btn:
            parent_btn = db["config"]["custom_features"].get(parent_feature_id)

        if parent_btn and parent_btn.get("sub_buttons"):
            current_buttons = parent_btn["sub_buttons"]
            # Determine the back button for submenus
            if parent_btn.get("parent_id"): # If this sub-feature itself has a parent
                pid = parent_btn['parent_id']
                back_button_callback = f"open_feature_{pid}"
            else:
                back_button_callback = "go_home"
        else:
            current_buttons = []
            back_button_callback = "go_home"
    else:
        # Building the main menu
        current_buttons = sorted(db["config"].get("menu_buttons", DEFAULT_MENU_BUTTONS),
                                 key=lambda x: x.get("order", 99))
        # Add top-level custom features that are not sub-buttons of any other feature
        for fid, f in db["config"].get("custom_features", {}).items():
            is_sub_button = False
            # Check if this custom feature is a sub-button of any main menu button
            for main_btn in db["config"]["menu_buttons"]:
                if any(sub_btn.get("id") == fid for sub_btn in main_btn.get("sub_buttons", [])):
                    is_sub_button = True
                    break
            # Check if this custom feature is a sub-button of any other custom feature
            if not is_sub_button:
                for custom_feat_id, custom_feat_data in db["config"]["custom_features"].items():
                    if custom_feat_id != fid and any(sub_btn.get("id") == fid for sub_btn in custom_feat_data.get("sub_buttons", [])):
                        is_sub_button = True
                        break
            
            if not is_sub_button and f.get("visible", True) and f.get("active", True):
                if f.get("dev_only", False) and not is_admin(uid): continue
                current_buttons.append({"id": fid, "label": f["label"], "action": f"open_feature_{fid}", "type": "custom_feature"})

    row = []
    for btn in current_buttons:
        if not btn.get("visible", True): continue
        if not btn.get("active", True) and not is_admin(uid): continue  # الميزات المعطّلة تختفي عن غير المطور

        callback_data = btn["action"]
        bid_ = btn.get("id", "")
        if btn.get("type") == "custom_feature":
            callback_data = f"open_feature_{bid_}"
        elif btn.get("type") == "main_action" and btn.get("sub_buttons") and not btn.get("action"):
            # فقط الأزرار الرئيسية اللي ما لها أمر (action) فعلي حقيقي تتحول لقائمة فرعية.
            # الأزرار اللي عندها أمر أصلي مثل "بحث مفلتر" أو "الاشتراك" لازم تشتغل بمنطقها
            # الحقيقي دائمًا، حتى لو انضاف لها أزرار فرعية بالغلط أو بالذكاء الاصطناعي.
            callback_data = f"open_feature_{bid_}"
        
        row.append(types.InlineKeyboardButton(btn["label"], callback_data=callback_data))
        if len(row) == 2:
            m.add(*row); row = []
    if row: m.add(*row)

    # Add back button if not in main menu
    if parent_feature_id:
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=back_button_callback))
    
    if is_admin(uid) and not parent_feature_id:
        m.add(types.InlineKeyboardButton("⚙️ لوحة التحكم", callback_data="admin_panel"))

    if not parent_feature_id and db["config"].get("stats_visible_to_users") and not is_admin(uid):
        m.add(types.InlineKeyboardButton("📊 إحصائيات البوت", callback_data="view_public_stats"))

    m.add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
    return m

def send_welcome_and_menu(chat_id, uid, menu_intro="اختر من القائمة أدناه 👇"):
    """يرسل الترحيب كرسالة مستقلة (نص/صورة/فيديو أو أي دمج بينها) ثم رسالة القائمة
    منفصلة تمامًا بعدها — عشان الترحيب ما يأثر على شكل الأزرار ولا يخليها تنكمش،
    ويبقى شكل الاثنين طبيعي متتابع وكأنهم رسالة وحدة من نظر المستخدم."""
    rich = db["config"].get("welcome_rich")
    sent_something = False
    if rich:
        text = rich.get("text")
        photo_id = rich.get("photo_id")
        video_id = rich.get("video_id")
        if photo_id:
            bot.send_photo(chat_id, photo_id, caption=text, parse_mode="Markdown")
            sent_something = True
        if video_id:
            bot.send_video(chat_id, video_id, caption=text if not photo_id else None, parse_mode="Markdown")
            sent_something = True
        if text and not photo_id and not video_id:
            bot.send_message(chat_id, text, parse_mode="Markdown")
            sent_something = True
    if not sent_something:
        # توافق مع الإعداد القديم (نص فقط) لو ما فيه welcome_rich بعد
        legacy = db["config"].get("welcome_msg")
        if legacy:
            bot.send_message(chat_id, legacy, parse_mode="Markdown")
    bot.send_message(chat_id, menu_intro, reply_markup=build_main_menu(uid))

def build_category_menu(cat_id, uid):
    cat = db["categories"].get(cat_id)
    if not cat: return types.InlineKeyboardMarkup()

    layout = db["config"].get("category_layouts", {}).get(cat_id, "grid2")

    # أزرار الأقسام الفرعية (تُبنى حسب شكل العرض المختار)
    child_buttons = []
    for cid in cat.get("children", []):
        c = db["categories"].get(cid)
        if c: child_buttons.append(types.InlineKeyboardButton(f"📁 {c['name']}", callback_data=f"nav_{cid}"))

    m = apply_layout(None, child_buttons, layout)

    for iid in cat.get("items", []):
        item = db["items"].get(iid)
        if item:
            icon = "📄"
            if item.get("type") == "photo": icon = "🖼️"
            elif item.get("type") == "video": icon = "🎬"
            elif item.get("type") == "audio": icon = "🎵"
            m.add(types.InlineKeyboardButton(f"{icon} {item['title']}", callback_data=f"item_{iid}"))

    novels_active = db["config"].get("content_features", {}).get("novels", {}).get("active", True)
    if novels_active or is_admin(uid):
        for nid in cat.get("novels", []):
            novel = db["novels"].get(nid)
            if novel: m.add(types.InlineKeyboardButton(f"📖 {novel['title']}", callback_data=f"open_novel_{nid}"))

    manga_active = db["config"].get("content_features", {}).get("manga", {}).get("active", True)
    if manga_active or is_admin(uid):
        for mid in cat.get("manga", []):
            manga = db["manga"].get(mid)
            if manga: m.add(types.InlineKeyboardButton(f"🎨 {manga['title']}", callback_data=f"open_manga_{mid}"))

    series_active = db["config"].get("content_features", {}).get("series", {}).get("active", True)
    if series_active or is_admin(uid):
        for sid in cat.get("series", []):
            series = db["series"].get(sid)
            if series: m.add(types.InlineKeyboardButton(f"🎬 {series['title']}", callback_data=f"open_series_{sid}"))

    if is_admin(uid):
        m.add(types.InlineKeyboardButton("➕ قسم فرعي", callback_data=f"add_cat_{cat_id}"),
              types.InlineKeyboardButton("📥 إضافة ملف", callback_data=f"add_item_{cat_id}"))
        m.add(types.InlineKeyboardButton("📖 إضافة رواية", callback_data=f"add_novel_{cat_id}"))
        m.add(types.InlineKeyboardButton("🎨 إضافة مانهوا/مانجا", callback_data=f"add_manga_{cat_id}"))
        m.add(types.InlineKeyboardButton("🎬 إضافة مسلسل/فيلم", callback_data=f"add_series_{cat_id}"))
        if child_buttons:
            m.add(types.InlineKeyboardButton("📐 شكل عرض الأقسام الفرعية", callback_data=f"catlayout_{cat_id}"))
        if cat_id != "root":
            m.add(types.InlineKeyboardButton("🗑️ حذف القسم", callback_data=f"del_cat_{cat_id}"))
    if cat_id != "root":
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"nav_{cat.get('parent','root')}"))
    m.add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
    return m

def category_layout_picker_kb(cat_id):
    m = types.InlineKeyboardMarkup(row_width=1)
    for key, label in LAYOUT_OPTIONS.items():
        m.add(types.InlineKeyboardButton(label, callback_data=f"setcatlayout_{cat_id}_{key}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"nav_{cat_id}"))
    return m

def tags_admin_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    tags = db["config"].get("tags", [])
    for i, t in enumerate(tags):
        name = _tag_name(t); active = _tag_active(t)
        act_icon = "🟢" if active else "🔴"
        m.add(types.InlineKeyboardButton(f"{act_icon} {name}", callback_data=f"tag_edit_{i}"))
    m.add(types.InlineKeyboardButton("➕ إضافة تصنيف جديد", callback_data="tag_add"))
    m.add(types.InlineKeyboardButton("📐 شكل عرض التصنيفات", callback_data="tags_layout"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

def tags_layout_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("↕️ عمودي (واحد تحت الثاني)", callback_data="settagslayout_vertical"))
    m.add(types.InlineKeyboardButton("↔️ أفقي (كلهم بصف واحد)", callback_data="settagslayout_horizontal"))
    m.add(types.InlineKeyboardButton("▦ شبكة 2 بكل صف", callback_data="settagslayout_grid2"))
    m.add(types.InlineKeyboardButton("▦ شبكة 3 بكل صف", callback_data="settagslayout_grid3"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_tags"))
    return m

def tag_editor_kb(idx):
    m = types.InlineKeyboardMarkup(row_width=2)
    tags = db["config"].get("tags", [])
    if idx < 0 or idx >= len(tags): return m
    t = tags[idx]
    active = _tag_active(t)
    act_label = "🔴 تعطيل" if active else "🟢 تفعيل"
    m.add(
        types.InlineKeyboardButton("✏️ تغيير الاسم", callback_data=f"tag_rename_{idx}"),
        types.InlineKeyboardButton(act_label,        callback_data=f"tag_toggle_{idx}")
    )
    m.add(
        types.InlineKeyboardButton("⬆️ تحريك لأعلى", callback_data=f"tag_up_{idx}"),
        types.InlineKeyboardButton("⬇️ تحريك لأسفل", callback_data=f"tag_down_{idx}")
    )
    m.add(types.InlineKeyboardButton("🗑️ حذف التصنيف", callback_data=f"tag_del_{idx}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_tags"))
    return m

def _tag_name(t):
    return t.get("name", "") if isinstance(t, dict) else t

def _tag_active(t):
    return t.get("active", True) if isinstance(t, dict) else True

def _item_tags(obj):
    """يرجع قائمة تصنيفات العنصر (رواية أو ملف) بشكل موحّد، مع التوافق مع البيانات
    القديمة اللي كانت تخزن تصنيف واحد فقط كنص (tag) بدل قائمة (tags)."""
    if not obj: return []
    if "tags" in obj and isinstance(obj["tags"], list):
        return obj["tags"]
    old = obj.get("tag")
    return [old] if old else []

def build_tag_search_menu(selected=None, mode="search"):
    """قائمة اختيار تصنيفات متعددة (بدون حد أقصى). selected: set من أسماء التصنيفات
    المختارة حاليًا (تظهر بعلامة ✅). mode: 'search' (بحث فعلي) أو 'assign' (اختيار
    تصنيفات لعنصر/رواية جديدة قيد الإضافة)."""
    selected = selected or []
    layout = db["config"].get("tags_layout", "grid2")
    tags = db["config"].get("tags", ["شونين", "سينين", "رومنسي", "دراما", "أكشن", "خيال"])
    buttons = []
    for i, t in enumerate(tags):
        if not _tag_active(t): continue
        name = _tag_name(t)
        mark = "✅ " if name in selected else ""
        buttons.append(types.InlineKeyboardButton(f"{mark}{name}", callback_data=f"searchtag_{i}"))
    m = apply_layout(None, buttons, layout)
    if mode == "search":
        m.add(types.InlineKeyboardButton("🔎 بحث بالاسم", callback_data="search_by_name"))
        if selected:
            m.add(types.InlineKeyboardButton(f"✅ عرض النتائج ({len(selected)} تصنيف)", callback_data="run_tag_search"))
    else:
        m.add(types.InlineKeyboardButton("✅ تم الاختيار — متابعة", callback_data="confirm_tags_assign"))
    m.add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
    return m

# ==============================================================================
# 7B. FEATURE NESTING (dمج ميزة داخل ميزة) + LAYOUT SYSTEM
# ==============================================================================
def get_feature_owner_type(fid):
    """يرجع ('main', obj) أو ('custom', obj) حسب مكان الميزة."""
    for b in db["config"]["menu_buttons"]:
        if b["id"] == fid: return "main", b
    if fid in db["config"]["custom_features"]:
        return "custom", db["config"]["custom_features"][fid]
    return None, None

def nest_feature(parent_fid, child_fid):
    """يدمج ميزة (child) كزر فرعي داخل ميزة أخرى (parent). يشتغل فعليًا: بعد الدمج
    الزر الفرعي في القائمة يفتح منطق الميزة الأصلي (child) كامل، مو مجرد شكل."""
    if parent_fid == child_fid: return False
    _, parent_obj = get_feature_owner_type(parent_fid)
    child_type, child_obj = get_feature_owner_type(child_fid)
    if not parent_obj or not child_obj: return False

    parent_obj.setdefault("sub_buttons", [])
    # تفادي التكرار
    if any(sb.get("nested_fid") == child_fid for sb in parent_obj["sub_buttons"]):
        return True

    nested_label = child_obj.get("label", child_fid)
    # الزر الفرعي يخزن مرجع (nested_fid + nested_type) بدل نسخ منطق الميزة،
    # بذا أي تحديث على الميزة الأصلية ينعكس تلقائيًا هنا (دمج حقيقي لا شكلي)
    parent_obj["sub_buttons"].append({
        "id": str(uuid.uuid4())[:8],
        "label": f"🔗 {nested_label}",
        "nested_fid": child_fid,
        "nested_type": child_type
    })
    db["config"]["feature_nesting_map"].setdefault(parent_fid, [])
    if child_fid not in db["config"]["feature_nesting_map"][parent_fid]:
        db["config"]["feature_nesting_map"][parent_fid].append(child_fid)

    # منع تعارض: إذا الميزة المدمجة كانت زر رئيسي في القائمة الرئيسية، نخفيها من هناك
    # حتى ما تتكرر (تبقى موجودة وتشتغل، بس فقط داخل الميزة الأم الجديدة)
    if child_type == "main":
        child_obj["visible"] = False
    return True

def nest_feature_picker_kb(parent_fid):
    m = types.InlineKeyboardMarkup(row_width=1)
    for b in db["config"]["menu_buttons"]:
        if b["id"] != parent_fid:
            m.add(types.InlineKeyboardButton(f"⚙️ {b['label']}", callback_data=f"donest_{parent_fid}_{b['id']}"))
    for fid, f in db["config"]["custom_features"].items():
        if fid != parent_fid:
            m.add(types.InlineKeyboardButton(f"✨ {f['label']}", callback_data=f"donest_{parent_fid}_{fid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"feat_subs_{parent_fid}"))
    return m

LAYOUT_OPTIONS = {
    "vertical": "▤ عمودي (زر تحت زر)",
    "horizontal": "▬ أفقي (كل الأزرار بصف واحد)",
    "grid2": "▦ شبكة 2×2",
    "grid3": "▦ شبكة 3×3",
}

def layout_picker_kb(fid):
    m = types.InlineKeyboardMarkup(row_width=1)
    for key, label in LAYOUT_OPTIONS.items():
        m.add(types.InlineKeyboardButton(label, callback_data=f"setlayout_{fid}_{key}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"feat_subs_{fid}"))
    return m

def apply_layout(markup_rows_builder, buttons, layout):
    """يبني صفوف الأزرار حسب شكل العرض المختار. buttons: list of InlineKeyboardButton."""
    m = types.InlineKeyboardMarkup()
    if layout == "vertical":
        for b in buttons: m.add(b)
    elif layout == "horizontal":
        if buttons: m.add(*buttons)
    elif layout == "grid3":
        row = []
        for b in buttons:
            row.append(b)
            if len(row) == 3: m.add(*row); row = []
        if row: m.add(*row)
    else:  # grid2 (default)
        row = []
        for b in buttons:
            row.append(b)
            if len(row) == 2: m.add(*row); row = []
        if row: m.add(*row)
    return m

# ==============================================================================
# 7C. SUBSCRIPTION SYSTEM (يدوي + عبر كود + خطط أسعار)
# ==============================================================================
def activate_subscription(ustr, days):
    ustr = str(ustr)
    if ustr not in db["users"]:
        db["users"][ustr] = {"points": 0, "downloads": 0, "favs": [], "history": []}
    u = db["users"][ustr]
    now = datetime.now()
    current_expiry = None
    if u.get("sub_expiry"):
        try: current_expiry = datetime.fromisoformat(u["sub_expiry"])
        except: current_expiry = None
    base = current_expiry if (current_expiry and current_expiry > now) else now
    new_expiry = base + timedelta(days=days)
    u["sub_expiry"] = new_expiry.isoformat()
    u["sub_months"] = u.get("sub_months", 0) + 1
    db["config"]["sub_settings"]["total_sub_revenue"] = db["config"]["sub_settings"].get("total_sub_revenue", 0.0)
    sync_db()
    return new_expiry

def do_sub_redeem(chat_id, ustr, code):
    code = code.upper().strip()
    c = db.get("codes", {}).get(code)
    if not c or "sub_days" not in c:
        bot.send_message(chat_id, "❌ كود اشتراك غير صحيح."); return
    if c.get("used"):
        bot.send_message(chat_id, "❌ الكود مستخدم مسبقاً."); return
    days = c["sub_days"]
    new_expiry = activate_subscription(ustr, days)
    db["codes"][code]["used"] = True; sync_db()
    bot.send_message(chat_id, f"🎉 تم تفعيل اشتراكك ({days} يوم)!\nينتهي في: {new_expiry.strftime('%Y-%m-%d')}", parse_mode="Markdown")

def sub_status_text(ustr):
    u = db["users"].get(ustr, {})
    if check_sub(ustr):
        exp = u.get("sub_expiry", "")[:10]
        return f"💎 *اشتراكك نشط*\n\nينتهي في: {exp}\nعدد مرات التجديد: {u.get('sub_months', 0)}"
    return "❌ *لا يوجد اشتراك نشط حاليًا.*\n\nاشترك للحصول على تنزيلات إضافية ومزايا حصرية."

def sub_plans_text():
    plans = db["config"].get("sub_plans", {})
    if not plans: return "لا توجد خطط أسعار محددة بعد."
    return "\n".join([f"• {d} يوم — {p}$" for d, p in sorted(plans.items(), key=lambda x: int(x[0]))])

def sub_menu_kb(uid):
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(types.InlineKeyboardButton("💳 خطط الاشتراك", callback_data="sub_plans"),
          types.InlineKeyboardButton("🎟️ لدي كود", callback_data="sub_redeem"))
    m.add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
    return m

def sub_plans_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    plans = db["config"].get("sub_plans", {})
    for d, p in sorted(plans.items(), key=lambda x: int(x[0])):
        m.add(types.InlineKeyboardButton(f"💎 اشتراك {d} يوم — {p}$", callback_data=f"sub_request_{d}"))
    m.add(types.InlineKeyboardButton("🎟️ عندي كود اشتراك", callback_data="sub_redeem"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="my_subscription"))
    return m

def admin_subs_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    plans = db["config"].get("sub_plans", {})
    for d, p in sorted(plans.items(), key=lambda x: int(x[0])):
        m.add(types.InlineKeyboardButton(f"✏️ {d} يوم — {p}$ (تعديل/حذف)", callback_data=f"editplan_{d}"))
    m.add(types.InlineKeyboardButton("➕ إضافة خطة بمدة وسعر مخصصين", callback_data="adm_sub_addplan"))
    m.add(types.InlineKeyboardButton("🎟️ إنشاء كود اشتراك", callback_data="adm_sub_gencode"))
    m.add(types.InlineKeyboardButton("✍️ تفعيل يدوي لمستخدم", callback_data="adm_sub_manual"))
    m.add(types.InlineKeyboardButton("📋 قائمة المشتركين", callback_data="adm_sub_list"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

# ==============================================================================
# 7D. AI STUDIO — يبني/يعدّل ميزات فعليًا (نص + صورة، محادثة تفاعلية)
# ==============================================================================
def get_telegram_file_url(file_id):
    file_info = bot.get_file(file_id)
    return f"https://api.telegram.org/file/bot{API_TOKEN}/{file_info.file_path}"

AI_STUDIO_SYSTEM_PROMPT = (
    "أنت مساعد ذكاء اصطناعي داخل استوديو بناء ميزات لبوت تليجرام (Python + pyTelegramBotAPI). "
    "دورك اثنان معًا حسب طلب المستخدم: "
    "(1) إنشاء ميزة/زر جديد بالكامل في البوت من وصف نصي أو صورة يرسلها المستخدم، "
    "(2) تعديل منطق ميزة موجودة فعليًا (تغيير سلوكها الحقيقي، لا مجرد اقتراح نص). "
    "إذا كان طلب المستخدم غامضًا، اطلب توضيحًا محددًا بدل التخمين. "
    "أعد ردك بصيغة JSON فقط بدون أي نص أو backticks خارج الـ JSON، بالحقول التالية: "
    "`needs_clarification` (true/false)، "
    "`clarification_question` (سؤال توضيحي إن احتجت، وإلا فارغ)، "
    "`action_type` ('create_feature' أو 'modify_feature' أو 'none')، "
    "`feature_label` (اسم الميزة المقترح إن كان إنشاء)، "
    "`response_text` (النص الذي يرد به الزر عند الضغط عليه — يجب أن يكون محتوى فعلي جاهز للاستخدام)، "
    "`explanation` (شرح موجز بالعربية لما تم فعله)."
)

def q_tr(text, target_lang, phrasing_rules=None):
    """يترجم نص (فصل رواية كامل مثلاً) عبر الذكاء الاصطناعي بترجمة دقيقة ولغوية طبيعية،
    مو حرفية كلمة-بكلمة. يقسّم النصوص الطويلة لأجزاء (chunks) حتى ما يتجاوز حد الموديل،
    ويطبّق قواعد صياغة ثابتة حددها المطور (مثل: إله → حاكم) بكل الأجزاء بثبات."""
    if not OR_KEY or OR_KEY.strip() == "":
        return {"__ai_error__": "missing_key"}
    lang_names = {"ar": "العربية", "en": "English", "ko": "한국어 (Korean)", "zh": "中文 (Chinese)"}
    target_name = lang_names.get(target_lang, target_lang)
    rules = phrasing_rules or db["config"].get("translation_phrasing_rules", {})
    rules_text = ""
    if rules:
        rules_lines = "\n".join(f'- "{k}" تُترجم دائمًا كـ "{v}"' for k, v in rules.items())
        rules_text = f"\n\nقواعد صياغة ثابتة يجب اتباعها دائمًا بدون استثناء:\n{rules_lines}"
    system_prompt = (
        f"أنت مترجم أدبي محترف متخصص بترجمة الروايات والقصص. مهمتك ترجمة النص التالي "
        f"إلى {target_name} بدقة عالية وبأسلوب أدبي طبيعي وسلس، كأنه مكتوب أصلًا بهذي اللغة "
        f"من مترجم بشري محترف، مو ترجمة حرفية آلية. حافظ على المعنى والسياق والنبرة العاطفية "
        f"للنص الأصلي بدقة. لا تحذف ولا تضف أي معلومة. لا تشرح ولا تعلّق — أعد النص المترجم فقط "
        f"بدون أي مقدمة أو خاتمة أو علامات markdown.{rules_text}"
    )
    # تقسيم النص الطويل لأجزاء (chunks) حتى ما نتجاوز حد الموديل — كل جزء يترجم لحاله
    # مع الحفاظ على السياق العام عبر إرسال النظام نفسه لكل جزء
    chunk_size = 3000  # حروف تقريبًا لكل جزء، حد آمن لمعظم الموديلات المجانية
    paragraphs = text.split("\n")
    chunks = []
    current = ""
    for p in paragraphs:
        if len(current) + len(p) > chunk_size and current:
            chunks.append(current)
            current = p
        else:
            current = current + "\n" + p if current else p
    if current: chunks.append(current)

    translated_parts = []
    for i, chunk in enumerate(chunks):
        try:
            r = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {OR_KEY}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://manus.ai",
                    "X-Title": "Traika Bot"
                },
                json={
                    "model": A_MDL,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": chunk}
                    ],
                    "temperature": 0.3, "max_tokens": 4000  # حرارة منخفضة لدقة أعلى وثبات بالمصطلحات
                }, timeout=90
            )
            data = r.json()
            if "choices" in data:
                translated_parts.append(data["choices"][0]["message"]["content"].strip())
            else:
                return {"__ai_error__": "api_error", "__ai_error_detail__": str(data)[:300]}
        except Exception as e:
            return {"__ai_error__": "exception", "__ai_error_detail__": str(e)[:300]}
    return {"translated_text": "\n\n".join(translated_parts)}

def translate_via_local_engine(text, target_lang):
    """يترجم نصًا عبر deep-translator (خدمة Google الترجمة، بدون حاجة لمفتاح
    API). أسرع وأوفر من الذكاء الاصطناعي، لكن ترجمة أقرب للحرفية بلا صقل أدبي
    ولا تطبيق قواعد صياغة ثابتة. يقسّم النص لأجزاء آمنة (Google يحدد ~5000 حرف)
    ويرجع بنفس صيغة q_tr ({"translated_text": ...} أو {"__ai_error__": ...})
    حتى يبقى المستدعي (send/handle) موحّدًا بغض النظر عن المحرك المختار."""
    if not _ensure_translation_libs():
        return {"__ai_error__": "missing_local_translator", "__ai_error_detail__": _TRANSLATE_IMPORT_ERROR}
    try:
        chunk_size = 4500
        paragraphs = text.split("\n")
        chunks, current = [], ""
        for p in paragraphs:
            if len(current) + len(p) > chunk_size and current:
                chunks.append(current); current = p
            else:
                current = current + "\n" + p if current else p
        if current: chunks.append(current)

        translator = GoogleTranslator(source="auto", target=target_lang)
        translated_parts = [translator.translate(chunk) for chunk in chunks if chunk.strip()]
        return {"translated_text": "\n\n".join(translated_parts)}
    except Exception as e:
        return {"__ai_error__": "exception", "__ai_error_detail__": str(e)[:300]}

def translate_text(text, target_lang, phrasing_rules=None):
    """نقطة الدخول الموحّدة للترجمة بكل أنحاء البوت: تختار المحرك الفعلي حسب
    إعداد المطور (translation_settings.engine) — "ai" عبر q_tr (دقيق وأدبي،
    يستهلك رصيد OpenRouter) أو "local" عبر deep-translator (فوري ومجاني، بدون
    صقل أدبي ولا قواعد صياغة مخصصة). أي كود جديد يحتاج ترجمة يستدعي هذي الدالة
    بدل q_tr مباشرة، حتى يحترم اختيار المطور دومًا."""
    engine = db["config"].get("translation_settings", {}).get("engine", "ai")
    if engine == "local":
        return translate_via_local_engine(text, target_lang)
    return q_tr(text, target_lang, phrasing_rules)

def ai_error_message(res, uid):
    """يبني رسالة خطأ الذكاء الاصطناعي: تفاصيل تقنية للمالك، ورسالة عامة لباقي المستخدمين."""
    reason = res.get("__ai_error__") if isinstance(res, dict) else None
    detail = res.get("__ai_error_detail__", "") if isinstance(res, dict) else ""
    if is_owner(uid):
        if reason == "missing_key":
            return ("❌ *الذكاء الاصطناعي غير مفعّل بعد.*\n\n"
                    "لازم تضبط مفتاح OpenRouter حقيقي كمتغير بيئة `BOT_OPENROUTER_KEY` بإعدادات الاستضافة "
                    "(احصل عليه من openrouter.ai) — المتغير حاليًا فاضي أو غير مضبوط.")
        elif reason == "api_error":
            return f"❌ *خطأ من خادم الذكاء الاصطناعي (OpenRouter):*\n\n`{detail}`\n\nتأكد أن المفتاح صحيح وله رصيد كافٍ."
        elif reason == "exception":
            return f"❌ *خطأ تقني أثناء الاتصال بالذكاء الاصطناعي:*\n\n`{detail}`"
    return "❌ حدث خطأ أثناء التواصل مع الذكاء الاصطناعي. حاول مرة أخرى لاحقاً."

def h_studio(msg, pending, txt):
    ustr = str(msg.from_user.id)
    image_url = None
    if msg.content_type == "photo":
        image_url = get_telegram_file_url(msg.photo[-1].file_id)

    history = pending.get("history", [])
    user_turn = txt or "(صورة بدون نص)"
    history.append({"role": "user", "content": user_turn})

    bot.send_message(msg.chat.id, "🧠 جاري التحليل...")
    convo_context = "\n".join([f"{h['role']}: {h['content']}" for h in history[-6:]])
    res = q_ai(f"سياق المحادثة:\n{convo_context}\n\nطلب المستخدم الحالي: {user_turn}",
                 AI_STUDIO_SYSTEM_PROMPT, image_url)

    if not res or (isinstance(res, dict) and "__ai_error__" in res):
        bot.send_message(msg.chat.id, ai_error_message(res, msg.from_user.id), parse_mode="Markdown")
        db["pending_actions"][ustr] = pending; sync_db()
        return

    if res.get("needs_clarification"):
        q = res.get("clarification_question", "ممكن توضح أكثر؟")
        history.append({"role": "assistant", "content": q})
        db["pending_actions"][ustr] = {"action": "ai_studio_chat", "mode": pending.get("mode", "new"), "history": history}
        sync_db()
        bot.send_message(msg.chat.id, f"🤔 {q}")
        return

    action_type = res.get("action_type", "none")
    explanation = res.get("explanation", "")
    label = res.get("feature_label", "ميزة جديدة")
    response_text = res.get("response_text", "")

    if action_type == "create_feature" and response_text:
        fid = str(uuid.uuid4())[:8]
        db["config"]["custom_features"][fid] = {
            "label": label, "response_text": response_text,
            "active": True, "visible": True, "sub_buttons": [],
            "created_by_ai": True
        }
        sync_db()
        bot.send_message(msg.chat.id, f"✅ *تم إنشاء الميزة فعليًا وهي تعمل الآن:*\n\n*{label}*\n{explanation}", parse_mode="Markdown")
    elif action_type == "modify_feature" and pending.get("target_fid"):
        f = db["config"]["custom_features"].get(pending["target_fid"])
        if f and response_text:
            f["response_text"] = response_text
            sync_db()
            bot.send_message(msg.chat.id, f"✅ *تم تعديل منطق الميزة فعليًا:*\n\n{explanation}", parse_mode="Markdown")
        else:
            bot.send_message(msg.chat.id, "❌ لم أجد الميزة المطلوب تعديلها.")
    else:
        bot.send_message(msg.chat.id, f"ℹ️ {explanation or 'لم أتمكن من تنفيذ تغيير فعلي بعد، وضّح أكثر.'}")

    # يبقي المحادثة مفتوحة لمزيد من التوضيح أو طلبات جديدة
    db["pending_actions"][ustr] = {"action": "ai_studio_chat", "mode": pending.get("mode", "new"),
                                     "history": history, "target_fid": pending.get("target_fid")}
    sync_db()

def handle_ai_logic_request(msg, pending, txt):
    """تعديل منطق ميزة موجودة عبر AI (نص + صورة اختيارية) — من قائمة 'تعديل منطق الميزات'."""
    ustr = str(msg.from_user.id)
    image_url = None
    if msg.content_type == "photo":
        image_url = get_telegram_file_url(msg.photo[-1].file_id)
        bot.send_message(msg.chat.id, "🧠 جاري تحليل طلبك والصورة...")
    else:
        bot.send_message(msg.chat.id, "🧠 جاري تحليل طلبك...")

    btn_id = pending.get("btn_id", "")
    res = q_ai(f"طلب المستخدم: {txt}\nالميزة المراد تعديلها: {btn_id}\nيرجى تعديل رد/سلوك هذه الميزة فعليًا.",
                 AI_STUDIO_SYSTEM_PROMPT, image_url)
    if not res or (isinstance(res, dict) and "__ai_error__" in res):
        bot.send_message(msg.chat.id, ai_error_message(res, msg.from_user.id), parse_mode="Markdown")
        return

    response_text = res.get("response_text", "")
    explanation = res.get("explanation", "")
    f = db["config"]["custom_features"].get(btn_id)
    if f and response_text:
        f["response_text"] = response_text
        sync_db()
        bot.send_message(msg.chat.id, f"✅ *تم تعديل منطق الميزة فعليًا:*\n\n{explanation}", parse_mode="Markdown")
    elif response_text:
        bot.send_message(msg.chat.id, f"ℹ️ الزر '{btn_id}' من الأزرار الأساسية للبوت ولا يمكن تعديل رده تلقائيًا هنا، لكن هذا المقترح:\n\n{response_text}\n\n{explanation}")
    else:
        bot.send_message(msg.chat.id, f"ℹ️ {explanation or 'لم أفهم طلب التعديل، جرّب توضيحه أكثر.'}")

# ==============================================================================
# 7E. NOVEL / CHAPTER SYSTEM (نظام الروايات والفصول)
# ==============================================================================
# فكرة النظام:
#   - كل رواية (novel) تُخزن منفصلة عن items العادية.
#   - لها بوستر + اسم + قصة + تصنيف، وتحتها زر "👁️ شاهد الآن".
#   - داخل "شاهد الآن": الفصل الأول | آخر فصل وصله المستخدم (يُحفظ تلقائيًا لكل مستخدم).
#   - تصفح فصل بفصل بأزرار ◀️ / ▶️.
#   - فصول مدمجة (1-25 / 26-50 / ...): يرفعها الأدمن يدويًا كملف PDF وملف EPUB لكل مجموعة.
#   - عند اختيار أي فصل مفرد أو مجموعة مدمجة: يُعرض للمستخدم خيار PDF أو EPUB.

ADMIN_PERMISSIONS = [
    ("all", "🔑 كل الصلاحيات (مشرف كامل)"),
    ("add_novel", "📚 إضافة رواية جديدة (وإدارة رواياته فقط)"),
    ("delete_chapters", "🗑️ حذف فصول من رواياته الخاصة"),
    ("add_item", "📄 إضافة ملفات عادية"),
    ("manage_categories", "📁 إدارة الأقسام"),
    ("broadcast", "📢 إرسال إذاعة للمستخدمين"),
    ("manage_codes", "🎟️ إدارة أكواد النقاط"),
    ("ban_users", "🚫 حظر/رفع حظر المستخدمين"),
]

def admin_expiry_edit_kb(target):
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("♾️ اجعله غير محدود", callback_data=f"setexp_none_{target}"))
    m.add(types.InlineKeyboardButton("📅 غيّر المدة (أيام جديدة)", callback_data=f"setexp_custom_{target}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_users"))
    return m

def admin_perm_picker_kb(selected):
    m = types.InlineKeyboardMarkup(row_width=1)
    for key, label in ADMIN_PERMISSIONS:
        mark = "✅ " if key in selected else ""
        m.add(types.InlineKeyboardButton(f"{mark}{label}", callback_data=f"newadm_perm_{key}"))
    m.add(types.InlineKeyboardButton("➡️ متابعة (تحديد المدة)", callback_data="newadm_perms_done"))
    return m

def admin_expiry_picker_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("♾️ غير محدود (حتى تشيله بنفسك)", callback_data="newadm_expiry_none"))
    m.add(types.InlineKeyboardButton("📅 مدة محددة (تحدد عدد الأيام)", callback_data="newadm_expiry_custom"))
    return m

def can_use_convert(uid):
    """المطور دايمًا مسموح له. أي مستخدم ثاني لازم يكون بقائمة المسموح لهم صراحة،
    حتى لو كان مشرف عادي — الوصول لهذي الميزة يُمنح بالاسم/الآيدي وليس بمنصب الإشراف."""
    if is_owner(uid): return True
    return str(uid) in db["config"].get("convert_allowed_users", [])

def can_use_scraper(uid):
    """صلاحية استخدام السحب التلقائي — المطور دايمًا، وباقي الحالات حسب إعدادات
    scraper_settings (عام للكل، أو قائمة مستخدمين محددة بالآيدي فقط)."""
    if is_owner(uid): return True
    ss = db["config"].get("scraper_settings", {})
    if not ss.get("enabled"): return False
    if ss.get("public"): return True
    return str(uid) in ss.get("allowed_users", [])

def can_use_translate(uid):
    """صلاحية استخدام الترجمة — منفصلة تمامًا عن صلاحية تحويل الصيغة. المطور دايمًا
    مسموح له، وباقي الحالات حسب translation_settings: إما عامة للكل، أو محصورة
    بقائمة مستخدمين يحددهم المطور بالآيدي، أو مقفولة بالكامل (يظهر للمطور فقط)."""
    if is_owner(uid): return True
    ts = db["config"].get("translation_settings", {})
    if ts.get("public"): return True
    return str(uid) in ts.get("allowed_users", [])

def fetch_page(url, retries=2, retry_delay=2):
    """يجيب HTML صفحة واحدة، مع User-Agent متصفح حقيقي عشان أغلب المواقع ما تحظر
    الطلب. يعيد المحاولة تلقائيًا (retries مرة، بفاصل retry_delay ثانية) عند
    أخطاء شبكة مؤقتة (timeout/connection error) قبل ما يستسلم — يفرّق بين فشل
    عابر (نعيد المحاولة) وفشل دائم مثل 404 (نرجع None فورًا بدون إعادة محاولة
    عديمة الفائدة). يرجّع BeautifulSoup جاهز أو None لو فشل نهائيًا."""
    if not BeautifulSoup:
        return None
    last_error = None
    for attempt in range(retries + 1):
        try:
            r = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0 (compatible; TraikaBot/1.0)"})
            r.raise_for_status()
            return BeautifulSoup(r.text, "html.parser")
        except requests.exceptions.HTTPError as e:
            # خطأ HTTP صريح (404/403/500...) — مو مفيد نعيد المحاولة، الصفحة فعلاً
            # مو موجودة أو ممنوعة، غالبًا يعني نهاية الفصول أو رابط خاطئ.
            logger.error(f"❌ فشل جلب الصفحة {url} (HTTP): {e}")
            return None
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            last_error = e
            if attempt < retries:
                time.sleep(retry_delay)
        except Exception as e:
            logger.error(f"❌ فشل جلب الصفحة {url}: {e}")
            return None
    logger.error(f"❌ فشل جلب الصفحة {url} بعد {retries+1} محاولات (شبكة): {last_error}")
    return None

def scrape_chapter(url, profile):
    """يسحب فصل واحد بالاعتماد على selectors محددة بملف تعريف الموقع (profile).
    يرجّع dict فيه العنوان والنص ورابط الفصل التالي (أو None لو ما لقى شيء)."""
    soup = fetch_page(url)
    if not soup:
        return None
    content_sel = profile.get("chapter_content_selector", "")
    title_sel = profile.get("chapter_title_selector", "")
    next_sel = profile.get("next_page_selector", "")

    content_el = soup.select_one(content_sel) if content_sel else None
    if not content_el:
        return None
    # نجمع النص فقرة فقرة عشان نحافظ على فواصل الأسطر بدل ما يصير النص كتلة وحدة
    paragraphs = [p.get_text(strip=True) for p in content_el.find_all(["p", "br", "div"]) if p.get_text(strip=True)]
    text = "\n\n".join(paragraphs) if paragraphs else content_el.get_text("\n", strip=True)

    title_el = soup.select_one(title_sel) if title_sel else None
    chapter_title = title_el.get_text(strip=True) if title_el else None

    next_el = soup.select_one(next_sel) if next_sel else None
    next_url = next_el.get("href") if next_el else None
    if next_url and next_url.startswith("/"):
        from urllib.parse import urljoin
        next_url = urljoin(url, next_url)

    return {"title": chapter_title, "text": text, "next_url": next_url}

def scrape_work_metadata(url, profile):
    """يسحب بيانات العمل الأساسية (العنوان، الوصف، البوستر) من صفحة الرواية
    الرئيسية — يُستخدم مرة وحدة بداية كل مهمة سحب قبل البدء بسحب الفصول."""
    soup = fetch_page(url)
    if not soup:
        return None
    title_sel = profile.get("title_selector", "")
    desc_sel = profile.get("description_selector", "")
    poster_sel = profile.get("poster_selector", "")

    title_el = soup.select_one(title_sel) if title_sel else None
    desc_el = soup.select_one(desc_sel) if desc_sel else None
    poster_el = soup.select_one(poster_sel) if poster_sel else None

    poster_url = None
    if poster_el:
        poster_url = poster_el.get("src") or poster_el.get("content") or poster_el.get("data-src")
        if poster_url and poster_url.startswith("/"):
            from urllib.parse import urljoin
            poster_url = urljoin(url, poster_url)

    return {
        "title": title_el.get_text(strip=True) if title_el else "بدون عنوان",
        "description": desc_el.get_text(strip=True) if desc_el else "",
        "poster_url": poster_url,
    }

def create_scrape_job(profile_id, start_url, work_type, created_by, translate_to=None):
    job_id = str(uuid.uuid4())[:8]
    db["scrape_jobs"][job_id] = {
        "id": job_id, "profile_id": profile_id, "start_url": start_url,
        "work_type": work_type, "work_id": None,
        "status": "running", "chapters_done": 0,
        "last_chapter_url": start_url, "translate_to": translate_to,
        "created_by": str(created_by), "created_at": datetime.now().isoformat(),
        "error": None, "scraped_texts": {},
    }
    sync_db()
    return job_id

def get_scrape_job(job_id):
    return db["scrape_jobs"].get(job_id)

def _create_work_by_type(work_type, title, description, created_by):
    """يستدعي دالة الإنشاء الصحيحة حسب نوع العمل — بدون أي فحوصات globals()
    هشة؛ الدوال الثلاث مؤكدة موجودة بنفس التوقيع بالملف."""
    if work_type == "novel":
        return create_novel(title, description, ["عام"], None, None, created_by=created_by)
    elif work_type == "manga":
        return create_manga(title, description, ["عام"], None, None, created_by=created_by)
    elif work_type == "series":
        return create_series(title, description, ["عام"], None, None, created_by=created_by)
    return None

def _get_work_by_type(work_type, work_id):
    if work_type == "novel": return get_novel(work_id)
    elif work_type == "manga": return get_manga(work_id)
    elif work_type == "series": return get_series(work_id)
    return None

def run_scrape_job_step(job_id, max_chapters=5):
    """يسحب دفعة من الفصول (افتراضيًا 5 كل استدعاء) بدل ما يحاول يسحب كل الرواية
    دفعة وحدة — هذا يحمي من انقطاع الاتصال بمنتصف عمل طويل، ويسمح للمهمة تكمل
    من حيث وقفت عبر استدعاءات متكررة (من جدولة الخلفية أو من ضغطة المطور).
    نص كل فصل يُحفظ مؤقتًا بسجل المهمة (scraped_texts) لحين اكتمال أول دفعة
    مدموجة تحويه، وبعدها يُحذف من هناك عشان قاعدة البيانات ما تتضخم.

    تحسينات الاستقرار:
    - كل ملف مؤقت يُسجَّل بـ temp_file_manager ويُحذف فورًا بعد الرفع (بدل ما
      يتراكم بـ /tmp للأبد).
    - حماية من الحلقة اللانهائية: لو next_url رجع نفس الرابط الحالي (موقع
      معطوب)، نوقف المهمة بدل ما نعلق بحلقة أبدية.
    - fetch_page الآن عندها retry داخلي، فـ scrape_chapter يفشل فقط عند فشل
      حقيقي (مو انقطاع شبكة عابر) — نميّز هذا بحالة 'paused' القابلة للاستئناف
      بدل 'failed' النهائية، إلا لو فشل من أول فصل بالمهمة كلها."""
    job = get_scrape_job(job_id)
    if not job or job["status"] != "running":
        return job

    profile = db["config"]["scraper_settings"]["site_profiles"].get(job["profile_id"])
    if not profile:
        job["status"] = "failed"; job["error"] = "ملف تعريف الموقع غير موجود"; sync_db()
        return job

    # أول خطوة بالمهمة: نسحب بيانات العمل الأساسية وننشئه بالبوت لو ما انسوى قبل
    if not job["work_id"]:
        meta = scrape_work_metadata(job["start_url"], profile)
        if not meta:
            job["status"] = "failed"; job["error"] = "فشل سحب بيانات العمل الأساسية"; sync_db()
            return job
        work_id = _create_work_by_type(job["work_type"], meta["title"], meta["description"], job["created_by"])
        if not work_id:
            job["status"] = "failed"; job["error"] = "فشل إنشاء العمل بالبوت"; sync_db()
            return job
        job["work_id"] = work_id
        job.setdefault("scraped_texts", {})
        sync_db()

    url = job["last_chapter_url"]
    done_this_step = 0
    ss = db["config"]["scraper_settings"]
    batch_sizes = ss.get("batch_sizes", [25, 35, 45, 55])
    scraped_texts = job.setdefault("scraped_texts", {})  # {"12": {"title":..., "text":...}, ...}

    while url and done_this_step < max_chapters:
        memory_monitor.wait_if_needed()
        result = scrape_chapter(url, profile)
        if not result or not result.get("text"):
            # فشل السحب: لو كان أول فصل بالمهمة كلها، غالبًا profile خاطئ من
            # الأساس (فشل دائم). لو صار بعد ما سحبنا فصول سابقة بنجاح، الأرجح
            # انقطاع مؤقت أو نهاية فصول الموقع — نخليها 'paused' قابلة للاستئناف
            # بضغطة زر بدل 'failed' نهائية تحتاج بدء مهمة جديدة من الصفر.
            job["status"] = "failed" if job["chapters_done"] == 0 else "paused"
            job["error"] = f"توقف السحب عند: {url}"
            sync_db()
            break
        job["chapters_done"] += 1
        ch_num = job["chapters_done"]
        ch_title = result.get("title") or f"الفصل {ch_num}"
        ch_text = result["text"]

        if job.get("translate_to"):
            translated = translate_text(ch_text, job["translate_to"])
            if translated: ch_text = translated

        # نحتفظ بنص الفصل مؤقتًا بسجل المهمة عشان نقدر نبني الملف المدموج لاحقًا
        # بدون إعادة سحبه من الموقع، حتى لو صار السحب على أكثر من استدعاء منفصل.
        scraped_texts[str(ch_num)] = {"title": ch_title, "text": ch_text}

        # نبني ملف PDF وEPUB للفصل المنفرد باستخدام نفس دوال البناء المستخدمة بأداة الدمج
        # (شكل الأقسام هنا لازم يكون tuples (label, text) بالضبط زي merge_files_to_one)
        sections = [(ch_title, ch_text)]
        tmp_dir = f"/tmp/scrape_{job_id}_{ch_num}"
        os.makedirs(tmp_dir, exist_ok=True)
        pdf_path = temp_file_manager.register(f"{tmp_dir}/ch{ch_num}.pdf")
        epub_path = temp_file_manager.register(f"{tmp_dir}/ch{ch_num}.epub")
        try:
            build_pdf_from_sections(sections, pdf_path)
            build_epub_from_sections(sections, epub_path, title=ch_title)
            save_scraped_chapter_files(job["work_type"], job["work_id"], ch_num, pdf_path, epub_path, ss)
        except Exception as e:
            logger.error(f"❌ فشل بناء ملفات الفصل {ch_num} بمهمة {job_id}: {e}")
        finally:
            temp_file_manager.safe_delete(pdf_path)
            temp_file_manager.safe_delete(epub_path)
            try: os.rmdir(tmp_dir)
            except Exception: pass
            gc.collect()

        # كل ما وصلنا حجم دفعة معرّف، نبني ملف مدموج تلقائي من آخر نقطة توقف
        for size in batch_sizes:
            if ch_num % size == 0:
                build_auto_merged_batch(job["work_type"], job["work_id"], ch_num - size + 1, ch_num, ss, scraped_texts)

        done_this_step += 1
        next_url = result.get("next_url")
        if next_url and next_url == url:
            # حماية من الحلقة اللانهائية: الموقع رجّع نفس الرابط الحالي كـ
            # "الفصل التالي" — هذا خطأ بملف التعريف (selector خاطئ) مو نهاية
            # طبيعية، نوقف المهمة بحالة قابلة للمراجعة بدل التعليق للأبد.
            job["status"] = "paused"
            job["error"] = f"رابط الفصل التالي مطابق للحالي (احتمال selector خاطئ بملف التعريف): {url}"
            sync_db()
            break
        url = next_url
        job["last_chapter_url"] = url
        sync_db()

    if not url and job["status"] == "running":
        job["status"] = "done"
        sync_db()
    return job

def save_scraped_chapter_files(work_type, work_id, ch_num, pdf_path, epub_path, scraper_settings):
    """يرفع ملفات الفصل المسحوب لقناة الأرشيف (لو مفعّلة) ويربطها بالفصل بالبوت،
    بنفس طريقة أي فصل يُرفع يدويًا — عشان المستخدم يقدر يفتحه عادي من داخل البوت."""
    work = _get_work_by_type(work_type, work_id)
    if not work:
        return
    file_ids = {"pdf": None, "epub": None}
    archive_key = "novels" if work_type == "novel" else work_type
    archive_ch = db["config"]["archive_channels"].get(archive_key)
    for fmt, path in (("pdf", pdf_path), ("epub", epub_path)):
        try:
            with open(path, "rb") as f:
                if scraper_settings.get("auto_send_to_archive") and archive_ch:
                    sent = bot.send_document(archive_ch, f, caption=f"{work.get('title','')} — الفصل {ch_num} ({fmt.upper()})")
                    file_ids[fmt] = sent.document.file_id
                else:
                    sent = bot.send_document(OWNER_ID, f, caption=f"{work.get('title','')} — الفصل {ch_num} ({fmt.upper()})")
                    file_ids[fmt] = sent.document.file_id
        except Exception as e:
            logger.error(f"❌ فشل رفع ملف {fmt} للفصل {ch_num}: {e}")
    work.setdefault("chapters", {})[str(ch_num)] = file_ids
    sync_db()

def build_auto_merged_batch(work_type, work_id, ch_from, ch_to, scraper_settings, scraped_texts):
    """يبني ملف مدموج تلقائي (PDF/EPUB) لمجموعة فصول متسلسلة بعد ما يوصلها السحب،
    بنفس منطق أداة الدمج اليدوية — كل فصل بعنوانه بمكانه الصحيح داخل الملف.
    scraped_texts هو نفس القاموس المؤقت اللي يحتفظ به run_scrape_job_step بنص كل
    فصل لحين اكتمال أول دفعة تحويه، وبعدها يُنظّف عشان قاعدة البيانات ما تتضخم."""
    work = _get_work_by_type(work_type, work_id)
    if not work:
        return
    # شكل الأقسام هنا لازم يطابق تمامًا (label, text) tuples — نفس ما تتوقعه merge_files_to_one
    sections = []
    for ch in range(ch_from, ch_to + 1):
        entry = scraped_texts.get(str(ch))
        if not entry:
            continue  # نص الفصل مو متوفر (حالة نادرة، غالبًا بسبب استئناف مهمة قديمة) — نتجاوزه بدل ما نوقف الدفعة كلها
        sections.append((entry["title"], entry["text"]))

    if not sections:
        logger.error(f"⚠️ ما فيه نصوص متوفرة لبناء الدفعة {ch_from}-{ch_to} للعمل {work_id}")
        return

    range_key = f"{ch_from}-{ch_to}"
    tmp_dir = f"/tmp/merge_{work_id}_{range_key}"
    os.makedirs(tmp_dir, exist_ok=True)
    archive_key = "novels" if work_type == "novel" else work_type
    generated_paths = []
    for ext in ("pdf", "epub"):
        out_path = temp_file_manager.register(f"{tmp_dir}/merged.{ext}")
        generated_paths.append(out_path)
        try:
            merge_files_to_one(sections, ext, out_path, title=work.get("title", "Document"))
            archive_ch = db["config"]["archive_channels"].get(archive_key)
            target = archive_ch if (scraper_settings.get("auto_send_to_archive") and archive_ch) else OWNER_ID
            with open(out_path, "rb") as f:
                sent = bot.send_document(target, f, caption=f"{work.get('title','')} — فصول {range_key} ({ext.upper()})")
            work.setdefault("merged", {})[range_key] = work.get("merged", {}).get(range_key, {})
            work["merged"][range_key][ext] = sent.document.file_id
        except Exception as e:
            logger.error(f"❌ فشل بناء الدفعة المدموجة {range_key}: {e}")

    for p in generated_paths:
        temp_file_manager.safe_delete(p)
    try: os.rmdir(tmp_dir)
    except Exception: pass
    gc.collect()

    # نظّف نصوص الفصول اللي دخلت بالدفعة المكتملة من سجل المهمة عشان ما تتراكم
    # بقاعدة البيانات بلا داعي — الملفات النهائية محفوظة أصلًا كـ file_id بتيليجرام.
    for ch in range(ch_from, ch_to + 1):
        scraped_texts.pop(str(ch), None)
    sync_db()

def get_matching_profile_id(url):
    """يبحث بملفات تعريف المواقع المسجّلة (site_profiles) عن أول بروفايل يطابق
    دومين الرابط المُعطى تلقائيًا، حتى لا يحتاج المستخدم يختار profile_id يدويًا
    في كل مرة بنظام الجدولة التلقائية. يرجّع None لو ما فيه تطابق."""
    try:
        domain = re.search(r"https?://([^/]+)", url).group(1).lower().lstrip("www.")
    except Exception:
        return None
    profiles = db["config"]["scraper_settings"].get("site_profiles", {})
    for pid, profile in profiles.items():
        pdomain = (profile.get("domain") or "").lower().lstrip("www.")
        if pdomain and pdomain in domain:
            return pid
    return None

def add_pending_scrape_link(url, profile_id, work_type, translate_to=None):
    """يضيف رابطًا لقائمة الانتظار ليُسحب لاحقًا (يدويًا أو بالجدولة التلقائية)."""
    ss = db["config"]["scraper_settings"]
    link_data = {
        "url": url, "profile_id": profile_id, "work_type": work_type,
        "translate_to": translate_to, "added_at": datetime.now().isoformat(),
    }
    ss.setdefault("pending_links", []).append(link_data)
    sync_db()
    return link_data

def run_scheduled_scrape_job(link_data):
    """ينشئ مهمة سحب من بيانات رابط بقائمة الانتظار ويشغّلها بخطوات متتالية حتى
    الانتهاء أو الفشل أو بلوغ الحد الأقصى للفصول بجلسة واحدة (max_chapters_per_run)
    — تحمي من تعليق الخيط لفترة طويلة جدًا على رواية ضخمة. تُبنى بالكامل فوق
    create_scrape_job وrun_scrape_job_step الموجودتين، بدون تكرار منطقهما."""
    ss = db["config"]["scraper_settings"]
    profile_id = link_data.get("profile_id") or get_matching_profile_id(link_data["url"])
    if not profile_id or profile_id not in ss.get("site_profiles", {}):
        return None, "لا يوجد ملف تعريف مطابق لهذا الموقع (سجّله يدويًا أولاً من ⚙️ ملفات التعريف)."
    job_id = create_scrape_job(profile_id, link_data["url"], link_data.get("work_type", "novel"),
                                OWNER_ID, translate_to=link_data.get("translate_to") or ss.get("default_translate_to"))
    max_total = ss.get("max_chapters_per_run", 50)
    job = get_scrape_job(job_id)
    last_notified_at = 0
    while job and job["status"] == "running" and job["chapters_done"] < max_total:
        remaining = max_total - job["chapters_done"]
        job = run_scrape_job_step(job_id, max_chapters=min(5, remaining))
        if not job:
            break
        # تنبيه تقدم للمطور كل 20 فصل (بدل رسالة لكل دفعة صغيرة، تجنبًا للإزعاج)
        if job["chapters_done"] - last_notified_at >= 20:
            last_notified_at = job["chapters_done"]
            try:
                bot.send_message(OWNER_ID, f"🕷️ تقدم سحب `{job_id}`: {job['chapters_done']}/{max_total} فصل حتى الآن.", parse_mode="Markdown")
            except Exception:
                pass
    return job_id, None

def run_auto_scrape_schedule():
    """ينفّذ دورة جدولة كاملة: يأخذ أول novels_per_schedule رابط من قائمة الانتظار،
    يسحب كل واحد منها بالكامل (أو حتى الحد الأقصى للفصول)، وينقله لقائمة المعالَجة
    أو الفاشلة حسب النتيجة."""
    ss = db["config"]["scraper_settings"]
    if not ss.get("enabled") or not ss.get("auto_schedule_enabled"):
        return
    pending = ss.get("pending_links", [])
    if not pending:
        return
    count = min(ss.get("novels_per_schedule", 1), len(pending))
    batch = pending[:count]
    logger.info(f"🕷️ بدء دورة سحب تلقائية: {len(batch)} رابط")
    for link_data in list(batch):
        try:
            ss["pending_links"].remove(link_data)
        except ValueError:
            continue
        job_id, err = run_scheduled_scrape_job(link_data)
        if err:
            link_data["error"] = err
            ss.setdefault("failed_links", []).append(link_data)
            logger.error(f"❌ فشل سحب {link_data['url']}: {err}")
        else:
            job = get_scrape_job(job_id)
            link_data["job_id"] = job_id
            if job and job.get("status") in ("done", "paused"):
                ss.setdefault("processed_links", []).append(link_data)
                logger.info(f"✅ تم سحب {link_data['url']} — {job.get('chapters_done', 0)} فصل")
            else:
                link_data["error"] = (job or {}).get("error", "غير معروف")
                ss.setdefault("failed_links", []).append(link_data)
                logger.error(f"❌ توقفت مهمة سحب {link_data['url']} بحالة غير مكتملة")
        sync_db()
    ss["processed_links"] = ss.get("processed_links", [])[-100:]
    ss["failed_links"] = ss.get("failed_links", [])[-100:]
    sync_db()

def scraper_scheduler_loop():
    """يعمل بالخلفية لتشغيل دورة السحب التلقائي مرة يوميًا بالوقت المضبوط
    (schedule_time)، طالما auto_schedule_enabled مفعّل."""
    while True:
        try:
            ss = db["config"]["scraper_settings"]
            if ss.get("enabled") and ss.get("auto_schedule_enabled"):
                now = datetime.now()
                if now.strftime("%H:%M") == ss.get("schedule_time", "06:00"):
                    last_run = ss.get("last_schedule_run")
                    already_ran_today = last_run and last_run.startswith(now.date().isoformat())
                    if not already_ran_today:
                        run_auto_scrape_schedule()
                        ss["last_schedule_run"] = now.isoformat()
                        sync_db()
            time.sleep(60)
        except Exception as e:
            logger.error(f"❌ خطأ في جدولة السحب التلقائي: {e}")
            time.sleep(60)

def _scraper_advanced_panel_content():
    """يبني نص وأزرار لوحة إدارة قائمة الانتظار والجدولة التلقائية للسحب."""
    ss = db["config"]["scraper_settings"]
    status = "🟢 مفعّل" if ss.get("auto_schedule_enabled") else "🔴 معطّل"
    pending = len(ss.get("pending_links", []))
    processed = len(ss.get("processed_links", []))
    failed = len(ss.get("failed_links", []))
    text = (f"🕷️ *السحب التلقائي (قائمة الانتظار)*\n\nالجدولة: {status}\n"
            f"وقت التشغيل اليومي: {ss.get('schedule_time', '06:00')}\n"
            f"روابط لكل دورة: {ss.get('novels_per_schedule', 1)}\n"
            f"حد الفصول بالجلسة: {ss.get('max_chapters_per_run', 50)}\n\n"
            f"⏳ بالانتظار: {pending}\n✅ مكتملة: {processed}\n❌ فاشلة: {failed}")
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(types.InlineKeyboardButton(f"{'⏸️ إيقاف الجدولة' if ss.get('auto_schedule_enabled') else '▶️ تشغيل الجدولة'}", callback_data="scr_toggle_auto"))
    m.add(types.InlineKeyboardButton("➕ إضافة رابط للانتظار", callback_data="scr_add_link"))
    m.add(types.InlineKeyboardButton("📋 قائمة الانتظار", callback_data="scr_view_pending"),
          types.InlineKeyboardButton("📊 السجل", callback_data="scr_view_history"))
    m.add(types.InlineKeyboardButton("⏰ تعديل الوقت", callback_data="scr_edit_time"))
    m.add(types.InlineKeyboardButton("🔢 روابط/دورة", callback_data="scr_edit_count"))
    m.add(types.InlineKeyboardButton("▶️ تشغيل دورة الآن", callback_data="scr_run_now"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_scraper"))
    return text, m

def can_manage_novel(uid, novel):
    """يتحقق هل المستخدم يقدر يدير رواية معيّنة: المطور ومشرف صلاحية 'all' يقدرون بكل الأعمال،
    أما مشرف بصلاحية 'add_novel' بس (محدودة) فيقدر يدير فقط الروايات اللي أضافها هو بنفسه."""
    if is_owner(uid): return True
    rec = get_admin_record(uid)
    if not rec: return False
    perms = rec.get("permissions", {})
    if perms.get("all"): return True
    if perms.get("add_novel") and novel and novel.get("created_by") == str(uid):
        return True
    return False

def can_delete_chapter(uid, novel):
    """صلاحية حذف فصل منفصلة عن صلاحية الإضافة — حتى مشرف يقدر يضيف روايته الخاصة،
    ما يقدر يحذف فصول منها إلا لو عنده صلاحية 'all' أو صلاحية حذف صريحة."""
    if is_owner(uid): return True
    rec = get_admin_record(uid)
    if not rec: return False
    perms = rec.get("permissions", {})
    if perms.get("all"): return True
    if perms.get("delete_chapters") and novel and novel.get("created_by") == str(uid):
        return True
    return False

def extract_url_metadata(url):
    """يجيب عنوان/وصف/صورة أي رابط عبر وسوم Open Graph القياسية (يدعمها أغلب
    المواقع)، مع رجوع لوسوم meta العادية لو ما وجد og:*. ما يعدّل ولا يحفظ أي شيء —
    بس يرجّع البيانات الخام عشان المطور يراجعها ويعدّلها قبل النشر."""
    if not BeautifulSoup:
        return {"error": "مكتبة BeautifulSoup غير مثبتة على الاستضافة"}
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0 (compatible; TraikaBot/1.0)"})
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")

        def meta(prop_name, attr="property"):
            tag = soup.find("meta", attrs={attr: prop_name})
            return tag.get("content", "").strip() if tag and tag.get("content") else None

        title = meta("og:title") or meta("twitter:title") or (soup.title.string.strip() if soup.title and soup.title.string else None)
        description = meta("og:description") or meta("twitter:description") or meta("description", attr="name")
        image = meta("og:image") or meta("twitter:image")
        return {
            "title": title or "بدون عنوان",
            "description": description or "",
            "poster_url": image,
            "source_url": url,
        }
    except Exception as e:
        return {"error": str(e)[:300]}

def create_news_item(data, created_by):
    news_id = str(uuid.uuid4())[:8]
    db["news_items"][news_id] = {
        "id": news_id,
        "title": data.get("title", "بدون عنوان"),
        "description": data.get("description", ""),
        "poster_url": data.get("poster_url"),
        "poster_file_id": None,
        "category": data.get("category"),
        "source_url": data.get("source_url"),
        "linked_type": None,
        "linked_id": None,
        "status": "draft",
        "created_by": str(created_by),
        "created_at": datetime.now().isoformat(),
        "published_msg_id": None,
        "published_chat": None,
    }
    sync_db()
    return news_id

def get_news_item(news_id):
    return db["news_items"].get(news_id)

def publish_news_item(news_id):
    """ينشر خبرًا (مسودة) بقناة النشر المضبوطة. يُستخدم من زر '✅ نشر الآن'
    ومن النشر التلقائي (auto_publish) في news_scheduler_loop على حد سواء،
    بحيث يبقى منطق النشر بمكان واحد. يرجّع (True, None) عند النجاح أو
    (False, رسالة الخطأ) عند الفشل."""
    n = get_news_item(news_id)
    ns = db["config"]["news_settings"]
    channel = ns.get("publish_channel")
    if not n:
        return False, "الخبر غير موجود."
    if not channel:
        return False, "لا توجد قناة نشر مضبوطة."
    caption = f"*{n['title']}*\n\n{n['description']}"
    if n.get("category"): caption += f"\n\n🏷️ {n['category']}"
    try:
        if n.get("poster_file_id"):
            sent = bot.send_photo(channel, n["poster_file_id"], caption=caption, parse_mode="Markdown", reply_markup=news_publish_kb(news_id))
        elif n.get("poster_url"):
            sent = bot.send_photo(channel, n["poster_url"], caption=caption, parse_mode="Markdown", reply_markup=news_publish_kb(news_id))
        else:
            sent = bot.send_message(channel, caption, parse_mode="Markdown", reply_markup=news_publish_kb(news_id))
        n["status"] = "published"
        n["published_msg_id"] = sent.message_id
        n["published_chat"] = str(channel)
        sync_db()
        return True, None
    except Exception as e:
        return False, str(e)[:150]

def extract_best_image(soup, url):
    """يستخرج أفضل صورة من الصفحة (og:image، twitter:image، ثم أول صورة كبيرة
    داخل المحتوى الرئيسي، ثم أي صورة معقولة الحجم بالصفحة كخيار أخير)."""
    og_image = soup.find("meta", property="og:image")
    if og_image and og_image.get("content"):
        return og_image["content"]
    tw_image = soup.find("meta", attrs={"name": "twitter:image"})
    if tw_image and tw_image.get("content"):
        return tw_image["content"]
    for tag in ["article", "main", "div.content", "div.entry-content"]:
        container = soup.select_one(tag)
        if container:
            img = container.find("img")
            if img and img.get("src"):
                src = img["src"]
                if src.startswith("//"):
                    src = "https:" + src
                elif src.startswith("/"):
                    from urllib.parse import urljoin
                    src = urljoin(url, src)
                return src
    for img in soup.find_all("img"):
        src = img.get("src")
        if not src:
            continue
        width, height = img.get("width"), img.get("height")
        if width and height:
            try:
                if int(width) > 100 and int(height) > 100:
                    return src
            except ValueError:
                pass
        if src.startswith("http") and "logo" not in src.lower():
            return src
    return None

def extract_main_content(soup):
    """يحاول استخراج النص الرئيسي من الصفحة (يتجاهل الإعلانات/الهوامش/التنقل)."""
    for selector in ["article", "main", ".content", ".entry-content", ".post-content", ".article-body"]:
        elem = soup.select_one(selector)
        if elem:
            for unwanted in elem.find_all(["script", "style", "iframe", "ins", "nav", "header", "footer", "aside"]):
                unwanted.decompose()
            text = elem.get_text(separator="\n", strip=True)
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            return "\n".join(lines)
    body = soup.find("body")
    if body:
        for unwanted in body.find_all(["script", "style", "iframe", "ins", "nav", "header", "footer", "aside"]):
            unwanted.decompose()
        text = body.get_text(separator="\n", strip=True)
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        return "\n".join(lines[:50])
    return None

def extract_category(soup, url):
    """يستخرج التصنيف من الصفحة (عنصر category/tags) أو من مسار الرابط نفسه."""
    for selector in [".category", ".tags", ".post-category", "a[rel='category']", ".entry-category"]:
        elem = soup.select_one(selector)
        if elem:
            cat = elem.get_text(strip=True)
            if cat:
                return cat
    match = re.search(r"/(category|tag|topic)/([^/]+)", url)
    if match:
        return match.group(2).replace("-", " ").title()
    return None

def summarize_with_ai(text, max_length=300):
    """يلخص نصًا طويلاً عبر q_ai. لو الذكاء الاصطناعي غير متاح أو فشل، يرجع
    اقتطاعًا بسيطًا للنص الأصلي حتى يبقى الخبر قابلاً للعرض دومًا."""
    if not text:
        return ""
    if not OR_KEY or OR_KEY.strip() == "":
        return text[:max_length] + "..." if len(text) > max_length else text
    try:
        prompt = f"لخص النص التالي في {max_length} حرفًا أو أقل، مع الحفاظ على الجوهر:\n\n{text[:3000]}"
        resp = q_ai(prompt, system_prompt="أنت مساعد تلخيص محترف. أجب فقط بصيغة JSON: {\"summary\":\"...\"}")
        if isinstance(resp, dict) and "__ai_error__" not in resp:
            summary = resp.get("summary") or resp.get("text") or resp.get("response")
            if summary:
                return summary
    except Exception as e:
        logger.error(f"❌ فشل تلخيص الخبر بالذكاء الاصطناعي: {e}")
    return text[:max_length] + "..." if len(text) > max_length else text

def suggest_category_ai(text):
    """يقترح تصنيفًا واحدًا مناسبًا للنص عبر q_ai (يُستخدم فقط لو الاستخراج
    المباشر من الصفحة فشل بإيجاد تصنيف)."""
    if not text or not OR_KEY or OR_KEY.strip() == "":
        return None
    try:
        prompt = f"اقترح تصنيفًا واحدًا مناسبًا للنص التالي (مثل: رياضة، تقنية، صحة، سياسة، فن، اقتصاد):\n\n{text[:500]}"
        resp = q_ai(prompt, system_prompt="أجب فقط بصيغة JSON: {\"category\":\"كلمة واحدة\"}")
        if isinstance(resp, dict) and "__ai_error__" not in resp:
            cat = resp.get("category") or resp.get("text") or resp.get("response")
            if cat:
                return str(cat).strip()
    except Exception as e:
        logger.error(f"❌ فشل اقتراح تصنيف الخبر: {e}")
    return None

def scrape_news_metadata_enhanced(url):
    """نسخة محسّنة لاستخراج بيانات خبر من رابط: عنوان، وصف، صورة، تصنيف،
    والمحتوى الكامل (يُستخدم لاحقًا للتلخيص). تُبنى فوق fetch_page الموجودة."""
    soup = fetch_page(url)
    if not soup:
        return None
    title = None
    for sel in ["meta[property='og:title']", "meta[name='twitter:title']", "title", "h1"]:
        elem = soup.select_one(sel)
        if elem:
            title = elem.get("content") if elem.name == "meta" else elem.get_text(strip=True)
            if title:
                break
    description = None
    for sel in ["meta[property='og:description']", "meta[name='twitter:description']", "meta[name='description']"]:
        elem = soup.select_one(sel)
        if elem and elem.get("content"):
            description = elem["content"]
            break
    full_content = extract_main_content(soup)
    if not description and full_content:
        description = full_content[:500] + "..." if len(full_content) > 500 else full_content
    poster_url = extract_best_image(soup, url)
    category = extract_category(soup, url)
    return {
        "title": title or "بدون عنوان",
        "description": description or "",
        "poster_url": poster_url,
        "category": category or "عام",
        "full_content": full_content,
        "source_url": url,
    }

def create_news_item_enhanced(data, created_by):
    """ينشئ خبرًا (مسودة) مع تلخيص تلقائي للوصف الطويل واقتراح تصنيف بالذكاء
    الاصطناعي عند غيابه. يُبنى فوق create_news_item الحالية بدون تعديلها."""
    description = data.get("description", "")
    if len(description) > 500:
        description = summarize_with_ai(description, 300)
    category = data.get("category")
    if (not category or category == "عام") and data.get("full_content"):
        suggested = suggest_category_ai(data["full_content"])
        if suggested:
            category = suggested
    news_id = create_news_item({
        "title": data.get("title", "بدون عنوان"),
        "description": description,
        "poster_url": data.get("poster_url"),
        "category": category or "عام",
        "source_url": data.get("source_url"),
    }, created_by)
    return news_id

def fetch_news_from_source(url):
    """يسحب روابط أخبار من مصدر واحد: يدعم RSS (item/link) وصفحات HTML عادية
    (يبحث عن روابط تشبه مقالات أخبار عبر أنماط مسار شائعة)."""
    if not BeautifulSoup:
        return []
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0 (compatible; TraikaBot/1.0)"})
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        news_links = []
        content_type = r.headers.get("Content-Type", "").lower()
        if "rss" in content_type or url.endswith(".rss") or url.endswith(".xml"):
            for item in soup.find_all("item"):
                link = item.find("link")
                if link and link.text:
                    news_links.append(link.text.strip())
                elif item.find("guid"):
                    news_links.append(item.find("guid").text.strip())
            return news_links
        from urllib.parse import urljoin
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if any(seg in href for seg in ("/news/", "/article/", "/post/", "/story/")):
                if not href.startswith("http"):
                    href = urljoin(url, href)
                news_links.append(href)
        return news_links[:20]
    except Exception as e:
        logger.error(f"❌ فشل سحب الأخبار من {url}: {e}")
        return []

def auto_scrape_news():
    """يسحب أخبارًا جديدة من كل المصادر المضبوطة، ينشئ مسودة لكل خبر جديد
    (غير مكرر)، ويحمّل صورته إن وُجدت، وينشرها تلقائيًا لو auto_publish مفعّل."""
    settings = db["config"]["news_settings"]
    if not settings.get("auto_scrape_enabled"):
        return
    sources = settings.get("sources", [])
    if not sources:
        return
    last_scrape = settings.get("last_scrape")
    if last_scrape:
        try:
            last_time = datetime.fromisoformat(last_scrape)
            interval_hours = settings.get("scrape_interval", 6)
            if (datetime.now() - last_time).total_seconds() < interval_hours * 3600:
                return
        except Exception:
            pass
    all_links = []
    for src in sources:
        all_links.extend(fetch_news_from_source(src))
    all_links = list(dict.fromkeys(all_links))  # إزالة تكرار مع حفظ الترتيب
    scraped = settings.get("scraped_news", [])
    new_links = [link for link in all_links if link not in scraped][: settings.get("max_news_per_day", 5)]

    for link in new_links:
        meta = scrape_news_metadata_enhanced(link)
        if not meta or not meta.get("title"):
            continue
        news_id = create_news_item_enhanced(meta, OWNER_ID)
        if meta.get("poster_url"):
            try:
                img_data = requests.get(meta["poster_url"], timeout=10)
                if img_data.status_code == 200:
                    tmp = f"/tmp/news_{news_id}.jpg"
                    with open(tmp, "wb") as f:
                        f.write(img_data.content)
                    with open(tmp, "rb") as f:
                        sent = bot.send_photo(OWNER_ID, f, caption="بوستر الخبر (نسخة احتياطية)")
                        db["news_items"][news_id]["poster_file_id"] = sent.photo[-1].file_id
                        sync_db()
                    os.remove(tmp)
            except Exception as e:
                logger.error(f"❌ فشل تحميل بوستر الخبر: {e}")
        scraped.append(link)
        logger.info(f"✅ تم سحب خبر جديد: {meta['title']}")

    settings["last_scrape"] = datetime.now().isoformat()
    settings["scraped_news"] = scraped[-100:]
    sync_db()

    if settings.get("auto_publish"):
        for news_id, item in list(db["news_items"].items()):
            if item.get("status") == "draft" and item.get("created_by") == str(OWNER_ID):
                ok, err = publish_news_item(news_id)
                if ok:
                    logger.info(f"✅ تم نشر الخبر تلقائيًا: {item['title']}")
                else:
                    logger.error(f"❌ فشل نشر الخبر تلقائيًا ({item['title']}): {err}")

def news_scheduler_loop():
    """يعمل بالخلفية لتشغيل دورة سحب الأخبار تلقائيًا حسب الإعدادات: إما بوقت
    نشر محدد (لو auto_publish مفعّل) أو بفحص دوري كل 5 دقائق (سحب بدون نشر)."""
    while True:
        try:
            settings = db["config"]["news_settings"]
            if settings.get("auto_scrape_enabled"):
                now = datetime.now()
                if settings.get("auto_publish"):
                    if now.strftime("%H:%M") == settings.get("publish_time", "06:00"):
                        auto_scrape_news()
                        time.sleep(60)  # تجنب تكرار التشغيل بنفس الدقيقة
                else:
                    auto_scrape_news()
            time.sleep(300)
        except Exception as e:
            logger.error(f"❌ خطأ في جدولة الأخبار: {e}")
            time.sleep(300)

def _news_advanced_panel_content():
    """يبني نص وأزرار لوحة إعدادات السحب التلقائي للأخبار (فوق لوحة adm_news الحالية)."""
    settings = db["config"]["news_settings"]
    drafts = [n for n in db["news_items"].values() if n["status"] == "draft"]
    status = "🟢 مفعّل" if settings.get("auto_scrape_enabled") else "🔴 معطّل"
    auto_pub = "🟢 مفعّل" if settings.get("auto_publish") else "🔴 معطّل"
    text = (f"📰 *السحب التلقائي للأخبار*\n\nالحالة: {status}\nالنشر التلقائي: {auto_pub}\n"
            f"المصادر: {len(settings.get('sources', []))}\nالمسودات الحالية: {len(drafts)}\n"
            f"كل: {settings.get('scrape_interval', 6)} ساعة، حتى {settings.get('max_news_per_day', 5)} خبر/دورة\n"
            f"وقت النشر التلقائي: {settings.get('publish_time', '06:00')}")
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(types.InlineKeyboardButton("🔄 تبديل السحب", callback_data="news_toggle_scrape"),
          types.InlineKeyboardButton("📤 تبديل النشر", callback_data="news_toggle_autopub"))
    m.add(types.InlineKeyboardButton("➕ إضافة مصدر", callback_data="news_add_source"),
          types.InlineKeyboardButton("📋 المصادر", callback_data="news_list_sources"))
    m.add(types.InlineKeyboardButton("⏰ تعديل التوقيت", callback_data="news_set_time"))
    m.add(types.InlineKeyboardButton("🔢 عدد الأخبار", callback_data="news_set_max"))
    m.add(types.InlineKeyboardButton("▶️ سحب الآن", callback_data="news_scrape_now"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_news"))
    return text, m

def news_review_text(news_id):
    n = get_news_item(news_id)
    if not n: return "❌ الخبر غير موجود."
    link_info = "غير مربوط بأي عمل" if not n.get("linked_id") else f"مربوط بـ {n['linked_type']}: {n['linked_id']}"
    return (f"📰 *مراجعة الخبر*\n\n"
            f"*العنوان:* {n['title']}\n\n"
            f"*الوصف:*\n{n['description'] or '—'}\n\n"
            f"*التصنيف:* {n.get('category') or '—'}\n"
            f"*الربط:* {link_info}\n"
            f"*المصدر:* `{n['source_url']}`\n"
            f"*الصورة:* {'موجودة' if (n.get('poster_url') or n.get('poster_file_id')) else 'بدون صورة'}")

def news_review_kb(news_id):
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(types.InlineKeyboardButton("✏️ العنوان", callback_data=f"news_edit_title_{news_id}"),
          types.InlineKeyboardButton("📝 الوصف", callback_data=f"news_edit_desc_{news_id}"))
    m.add(types.InlineKeyboardButton("🏷️ التصنيف", callback_data=f"news_edit_cat_{news_id}"),
          types.InlineKeyboardButton("🖼️ الصورة", callback_data=f"news_edit_poster_{news_id}"))
    m.add(types.InlineKeyboardButton("🔗 ربط برواية/عمل", callback_data=f"news_link_pick_type_{news_id}"))
    n = get_news_item(news_id)
    if n and n.get("linked_id"):
        m.add(types.InlineKeyboardButton("❌ إلغاء الربط", callback_data=f"news_unlink_{news_id}"))
    m.add(types.InlineKeyboardButton("👁️ معاينة", callback_data=f"news_preview_{news_id}"))
    m.add(types.InlineKeyboardButton("✅ نشر الآن", callback_data=f"news_publish_{news_id}"))
    m.add(types.InlineKeyboardButton("🗑️ حذف المسودة", callback_data=f"news_delete_{news_id}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_news"))
    return m

def news_publish_kb(news_id):
    n = get_news_item(news_id)
    ns = db["config"]["news_settings"]
    m = types.InlineKeyboardMarkup(row_width=1)
    if n.get("linked_id"):
        m.add(types.InlineKeyboardButton(ns.get("novel_button_label", "📖 الدخول للرواية"), callback_data=f"open_{n['linked_type']}_{n['linked_id']}"))
    if n.get("source_url"):
        m.add(types.InlineKeyboardButton(ns.get("source_button_label", "🔗 المصدر"), url=n["source_url"]))
    return m

def create_novel(title, story, tags, poster_file_id, category_id, created_by=None):
    nid = str(uuid.uuid4())[:8]
    tags_list = tags if isinstance(tags, list) else ([tags] if tags else ["عام"])
    db["novels"][nid] = {
        "id": nid, "title": title, "story": story, "tags": tags_list,
        "poster_file_id": poster_file_id, "category": category_id,
        "linked_chat": None, "created_by": str(created_by) if created_by else None,
        "chapters": {},   # {"1": {"pdf": file_id|None, "epub": file_id|None}}
        "merged": {},     # {"1-25": {"pdf": file_id|None, "epub": file_id|None}}
        "created_at": str(datetime.now())
    }
    if category_id in db["categories"]:
        db["categories"][category_id].setdefault("novels", []).append(nid)
    sync_db()
    return nid

def get_novel(nid):
    return db["novels"].get(nid)

def novel_chapter_numbers(novel):
    return sorted([int(k) for k in novel.get("chapters", {}).keys()])

def novel_merge_groups(novel):
    """يرجع أسماء المجموعات المدمجة المتاحة فعليًا مرتبة (مثلاً ['1-25','26-50'])."""
    return sorted(novel.get("merged", {}).keys(), key=lambda g: int(g.split("-")[0]))

def merge_existing_chapters(nid, chapter_list, cid, progress_msg_id=None):
    """يدمج فصول مرفوعة أصلاً بالرواية (مو رفع ملف جاهز يدويًا): يحمّل ملف كل فصل
    من تلغرام، يستخرج نصه الخام، ثم يبني ملف PDF وEPUB واحد يحوي كل الفصول المختارة
    مرتبة، بالضبط بنفس منطق أداة الدمج القديمة لكن من مصدر حقيقي (الفصول المنفردة)
    بدل ملف يرفعه الأدمن يدويًا. يرجع (ok: bool, message: str)."""
    novel = get_novel(nid)
    if not novel:
        return False, "❌ الرواية غير موجودة."

    chapter_list = sorted(set(chapter_list))
    if not chapter_list:
        return False, "❌ ما فيه فصول محددة."

    tmp_dir = f"/tmp/manual_merge_{nid}_{uuid.uuid4().hex[:6]}"
    os.makedirs(tmp_dir, exist_ok=True)
    sections = []
    missing = []
    try:
        for ch_num in chapter_list:
            ch = novel.get("chapters", {}).get(str(ch_num))
            src_file_id = None
            src_ext = None
            if ch and ch.get("pdf"):
                src_file_id, src_ext = ch["pdf"], "pdf"
            elif ch and ch.get("epub"):
                src_file_id, src_ext = ch["epub"], "epub"
            if not src_file_id:
                missing.append(ch_num)
                continue
            try:
                file_info = bot.get_file(src_file_id)
                downloaded = bot.download_file(file_info.file_path)
                src_path = f"{tmp_dir}/ch_{ch_num}.{src_ext}"
                with open(src_path, "wb") as f:
                    f.write(downloaded)
                text = extract_text_from_file(src_path, src_ext)
                sections.append((f"الفصل {ch_num}", text))
            except Exception as e:
                logger.error(f"❌ فشل تحميل/استخراج نص الفصل {ch_num} للدمج: {e}")
                missing.append(ch_num)

        if not sections:
            return False, "❌ ما قدرت أجهز أي فصل من المختارة (تحقق إن الفصول عندها ملفات مرفوعة)."

        range_key = f"{chapter_list[0]}-{chapter_list[-1]}" if chapter_list == list(range(chapter_list[0], chapter_list[-1] + 1)) \
            else "+".join(str(c) for c in chapter_list)

        built = {}
        for ext in ("pdf", "epub"):
            out_path = f"{tmp_dir}/merged.{ext}"
            # ما فيه قتل قسري بمهلة ثابتة هنا — بناء PDF/EPUB لعدد كبير من الفصول
            # (خصوصًا مع تشكيل عربي لكل سطر) قد ياخذ وقت متفاوت، وهذي الدالة أصلًا
            # تُستدعى من خيط خلفية منفصل (شوف مواقع النداء)، فما فيه داعي لقتلها
            # قسريًا — بس نسجّل الوقت المستغرق لأي تشخيص مستقبلي.
            t_build = time.time()
            try:
                merge_files_to_one(sections, ext, out_path, title=novel.get("title", "Document"))
                logger.info(f"[merge] بناء {ext} لـ {len(sections)} فصل: {time.time()-t_build:.1f}ث")
            except Exception as e:
                logger.error(f"❌ فشل بناء ملف {ext} للدمج اليدوي بعد {time.time()-t_build:.1f}ث: {e}")
                continue
            try:
                with open(out_path, "rb") as f:
                    sent = bot.send_document(cid, f, caption=f"📦 *{md_safe(novel['title'])}* — تجهيز الفصول {range_key} ({ext.upper()})", parse_mode="Markdown")
                built[ext] = sent.document.file_id
            except Exception as e:
                logger.error(f"❌ فشل إرسال ملف {ext} المدمج: {e}")

        if not built:
            return False, "❌ فشل بناء أي ملف مدمج (قد يكون بسبب فصل كبير جدًا أو خطأ بالتحويل). جرب بعدد فصول أقل بكل مرة."

        novel.setdefault("merged", {})[range_key] = novel.get("merged", {}).get(range_key, {"pdf": None, "epub": None})
        novel["merged"][range_key].update(built)
        sync_db()

        msg = f"✅ تم دمج {len(sections)} فصل بنجاح تحت اسم المجموعة «{range_key}»."
        if missing:
            msg += f"\n⚠️ تعذّر تجهيز الفصول التالية (ملف مفقود): {', '.join(str(m) for m in missing)}"
        return True, msg
    finally:
        try:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass

def next_merge_range(novel):
    """يحسب المدى التالي المتاح للدمج بناءً على أكبر فصل مرفوع (25 فصل لكل مجموعة كحد أقصى)."""
    chs = novel_chapter_numbers(novel)
    if not chs: return None
    max_ch = max(chs)
    existing_groups = novel_merge_groups(novel)
    covered_end = 0
    for g in existing_groups:
        end = int(g.split("-")[1])
        covered_end = max(covered_end, end)
    if covered_end >= max_ch:
        return None  # لا يوجد فصول جديدة كافية بعد
    start = covered_end + 1
    end = min(start + 24, max_ch)
    return f"{start}-{end}"

def get_user_novel_progress(ustr, nid):
    u = db["users"].get(ustr, {})
    return u.get("novel_progress", {}).get(nid)

def set_user_novel_progress(ustr, nid, chapter_num):
    if ustr not in db["users"]: return
    db["users"][ustr].setdefault("novel_progress", {})
    db["users"][ustr]["novel_progress"][nid] = chapter_num
    sync_db()

# ── Keyboards ────────────────────────────────────────────────────────────────

def novel_card_kb(nid, uid):
    """الكيبورد اللي يظهر تحت بوستر الرواية مباشرة (زر شاهد الآن + نقل للقائمة)."""
    m = types.InlineKeyboardMarkup(row_width=1)
    watch_label = db["config"].get("novel_watch_label", "👁️ شاهد الآن")
    m.add(types.InlineKeyboardButton(watch_label, callback_data=f"novel_watch_{nid}"))
    m.add(types.InlineKeyboardButton("➕ نقل إلى القائمة", callback_data=f"addlist_novel_{nid}"))
    comments_count = len(get_comments(comment_key_novel(nid)))
    m.add(types.InlineKeyboardButton(f"💬 التعليقات ({comments_count})", callback_data=f"comments_{comment_key_novel(nid)}"))
    novel = get_novel(nid)
    if is_admin(uid):
        m.add(types.InlineKeyboardButton("⚙️ إدارة الرواية", callback_data=f"novel_admin_{nid}"))
    cat_id = novel.get("category", "root") if novel else "root"
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"nav_{cat_id}"))
    return m

def novel_watch_kb(nid, ustr):
    """قائمة 'شاهد الآن': أول فصل انضاف / آخر فصل انضاف (حسب تاريخ الإضافة، مو رقمياً
    ومو حسب آخر فصل قرأه المستخدم)، وزر الفصول المدمجة."""
    novel = get_novel(nid)
    m = types.InlineKeyboardMarkup(row_width=2)
    chs = novel_chapter_numbers(novel)
    if not chs:
        m.add(types.InlineKeyboardButton("❌ لا توجد فصول مرفوعة بعد", callback_data="noop_"))
    else:
        first_added_ch = chs[0]
        last_added_ch = chs[-1]
        m.add(
            types.InlineKeyboardButton(f"📖 الفصل {first_added_ch}", callback_data=f"novel_ch_{nid}_{first_added_ch}"),
            types.InlineKeyboardButton(f"🆕 الفصل {last_added_ch}", callback_data=f"novel_ch_{nid}_{last_added_ch}")
        )
    if novel.get("merged"):
        m.add(types.InlineKeyboardButton("📦 الفصول المدمجة", callback_data=f"novel_merged_{nid}"))
    if len(chs) > 4:
        m.add(types.InlineKeyboardButton("🔢 كل الفصول (شبكة أرقام)", callback_data=f"chgrid_{nid}_0"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"open_novel_{nid}"))
    return m

def novel_chapter_nav_kb(nid, chapter_num):
    """كيبورد تصفح فصل بفصل: ◀️ فصل - رقم - فصل ▶️."""
    novel = get_novel(nid)
    chs = novel_chapter_numbers(novel)
    m = types.InlineKeyboardMarkup(row_width=3)
    idx = chs.index(chapter_num) if chapter_num in chs else 0
    prev_ch = chs[idx - 1] if idx > 0 else None
    next_ch = chs[idx + 1] if idx < len(chs) - 1 else None
    row = []
    row.append(types.InlineKeyboardButton("◀️ السابق" if prev_ch else "▪️", callback_data=f"novel_ch_{nid}_{prev_ch}" if prev_ch else "noop_"))
    row.append(types.InlineKeyboardButton(f"فصل {chapter_num}", callback_data="noop_"))
    row.append(types.InlineKeyboardButton("التالي ▶️" if next_ch else "▪️", callback_data=f"novel_ch_{nid}_{next_ch}" if next_ch else "noop_"))
    m.add(*row)
    m.add(
        types.InlineKeyboardButton("📄 PDF", callback_data=f"novel_get_{nid}_ch_{chapter_num}_pdf"),
        types.InlineKeyboardButton("📱 EPUB", callback_data=f"novel_get_{nid}_ch_{chapter_num}_epub")
    )
    m.add(types.InlineKeyboardButton("🔢 كل الفصول", callback_data=f"chgrid_{nid}_0"))
    ch_key = comment_key_chapter(nid, chapter_num)
    m.add(types.InlineKeyboardButton(f"💬 تعليقات الفصل ({len(get_comments(ch_key))})", callback_data=f"comments_{ch_key}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"novel_watch_{nid}"))
    return m

CHAPTER_GRID_PAGE_SIZE = 20  # عدد أرقام الفصول المعروضة بكل صفحة شبكة

def novel_chapter_grid_kb(nid, page=0):
    """شبكة أرقام الفصول (زي: 44 45 46 47...) بحيث المستخدم يقدر يضغط أي رقم فصل
    مباشرة بدل ما يمرر فصل فصل بزر التالي — مفيدة خصوصًا بالروايات الطويلة (+500 فصل)."""
    novel = get_novel(nid)
    chs = novel_chapter_numbers(novel)
    m = types.InlineKeyboardMarkup(row_width=5)
    start = page * CHAPTER_GRID_PAGE_SIZE
    end = start + CHAPTER_GRID_PAGE_SIZE
    page_chs = chs[start:end]
    row = []
    for ch in page_chs:
        row.append(types.InlineKeyboardButton(str(ch), callback_data=f"novel_ch_{nid}_{ch}"))
        if len(row) == 5:
            m.add(*row); row = []
    if row: m.add(*row)
    nav_row = []
    if start > 0:
        nav_row.append(types.InlineKeyboardButton("⏮️ السابقة", callback_data=f"chgrid_{nid}_{page-1}"))
    if end < len(chs):
        nav_row.append(types.InlineKeyboardButton("التالية ⏭️", callback_data=f"chgrid_{nid}_{page+1}"))
    if nav_row: m.add(*nav_row)
    total_pages = max(1, (len(chs) + CHAPTER_GRID_PAGE_SIZE - 1) // CHAPTER_GRID_PAGE_SIZE)
    m.add(types.InlineKeyboardButton(f"📄 صفحة {page+1}/{total_pages}", callback_data="noop_"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"novel_watch_{nid}"))
    return m

def novel_merged_list_kb(nid):
    novel = get_novel(nid)
    m = types.InlineKeyboardMarkup(row_width=1)
    for g in novel_merge_groups(novel):
        m.add(types.InlineKeyboardButton(f"📦 الفصول {g}", callback_data=f"novel_mopen_{nid}_{g}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"novel_watch_{nid}"))
    return m

def novel_merged_format_kb(nid, group):
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(
        types.InlineKeyboardButton("📄 PDF", callback_data=f"novel_get_{nid}_mg_{group}_pdf"),
        types.InlineKeyboardButton("📱 EPUB", callback_data=f"novel_get_{nid}_mg_{group}_epub")
    )
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"novel_merged_{nid}"))
    return m

def novel_merge_pick_kb(nid):
    """الخطوة الأولى بأداة دمج الفصول الجديدة: يختار الأدمن الطريقة — مدى متسلسل
    (1-25) أو تحديد فصول معينة يدويًا (مو بالضرورة متسلسلة)."""
    novel = get_novel(nid)
    chs = novel_chapter_numbers(novel)
    m = types.InlineKeyboardMarkup(row_width=1)
    if not chs:
        m.add(types.InlineKeyboardButton("❌ لا يوجد فصول مرفوعة بعد بهذي الرواية", callback_data="noop_"))
    else:
        m.add(types.InlineKeyboardButton(f"🔢 مدى متسلسل (مثال: 1-25) — متوفر {chs[0]}-{chs[-1]}", callback_data=f"novel_mergerange_{nid}"))
        m.add(types.InlineKeyboardButton("✅ اختيار فصول محددة يدويًا", callback_data=f"novel_mergemanual_{nid}_0"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"novel_admin_{nid}"))
    return m

MERGE_PICK_PAGE_SIZE = 20

def build_novel_merge_manual_kb(nid, selected, page=0):
    """شبكة أرقام فصول قابلة للتحديد المتعدد (تشيك/إلغاء تشيك) للدمج اليدوي —
    الفصول المختارة تُخزّن مؤقتًا بـ pending_actions لحد ما يضغط 'تم، ابدأ الدمج'."""
    novel = get_novel(nid)
    chs = novel_chapter_numbers(novel)
    m = types.InlineKeyboardMarkup(row_width=5)
    start = page * MERGE_PICK_PAGE_SIZE
    end = start + MERGE_PICK_PAGE_SIZE
    page_chs = chs[start:end]
    row = []
    for ch in page_chs:
        mark = "✅" if ch in selected else str(ch)
        row.append(types.InlineKeyboardButton(mark, callback_data=f"novel_mergetoggle_{nid}_{page}_{ch}"))
        if len(row) == 5:
            m.add(*row); row = []
    if row: m.add(*row)
    nav_row = []
    if start > 0:
        nav_row.append(types.InlineKeyboardButton("⏮️ السابقة", callback_data=f"novel_mergemanual_{nid}_{page-1}"))
    if end < len(chs):
        nav_row.append(types.InlineKeyboardButton("التالية ⏭️", callback_data=f"novel_mergemanual_{nid}_{page+1}"))
    if nav_row: m.add(*nav_row)
    total_pages = max(1, (len(chs) + MERGE_PICK_PAGE_SIZE - 1) // MERGE_PICK_PAGE_SIZE)
    m.add(types.InlineKeyboardButton(f"📄 صفحة {page+1}/{total_pages} — محدد: {len(selected)}", callback_data="noop_"))
    if selected:
        m.add(types.InlineKeyboardButton(f"✅ تم، ابدأ الدمج ({len(selected)} فصل)", callback_data=f"novel_mergeconfirm_{nid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"novel_mergepick_{nid}"))
    return m

# ── Admin: novel management ──────────────────────────────────────────────────

def novel_admin_kb(nid):
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("📥 رفع فصل جديد", callback_data=f"novel_addch_{nid}"))
    m.add(types.InlineKeyboardButton("🗑️ حذف فصل", callback_data=f"novel_delch_{nid}"))
    m.add(types.InlineKeyboardButton("📦 دمج فصول موجودة", callback_data=f"novel_mergepick_{nid}"))
    m.add(types.InlineKeyboardButton("🔗 ربط بقناة/مجموعة", callback_data=f"novel_link_{nid}"))
    m.add(types.InlineKeyboardButton("🗑️ حذف الرواية", callback_data=f"novel_del_{nid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"open_novel_{nid}"))
    return m

def set_user_item_view(ustr, iid):
    """يسجل وقت آخر مرة شاف/حمّل فيها المستخدم ملف معين، لعرضه له لاحقًا كـ 'آخر مرة شاهدته'."""
    if ustr not in db["users"]: return
    db["users"][ustr].setdefault("item_views", {})
    db["users"][ustr]["item_views"][iid] = str(datetime.now())
    sync_db()

def get_user_item_view(ustr, iid):
    return db["users"].get(ustr, {}).get("item_views", {}).get(iid)

def open_novel(uid, cid, mid, nid, as_new_message=False):
    """يعرض بوستر الرواية (صورة) + الاسم + القصة + التصنيف + زر شاهد الآن."""
    novel = get_novel(nid)
    if not novel:
        bot.send_message(cid, "❌ الرواية غير موجودة.")
        return
    if not db["config"].get("content_features", {}).get("novels", {}).get("active", True) and not is_admin(uid):
        bot.send_message(cid, "⚠️ هذي الميزة معطّلة حاليًا.")
        return
    track_event("novel_view", nid)
    tags_display = "، ".join(_item_tags(novel)) or "عام"
    ustr = str(uid)
    progress = get_user_novel_progress(ustr, nid)
    marker = db["config"].get("progress_marker", "🔴")
    progress_line = f"\n\n{marker} آخر فصل قرأته: *{progress}*" if progress else ""
    caption = f"*{md_safe(novel['title'])}*\n\n{md_safe(novel.get('story',''))}\n\n🏷️ {md_safe(tags_display)}{progress_line}"
    kb = novel_card_kb(nid, uid)
    try:
        if novel.get("poster_file_id"):
            bot.send_photo(cid, novel["poster_file_id"], caption=caption, parse_mode="Markdown", reply_markup=kb, protect_content=should_protect_content(uid))
        else:
            bot.send_message(cid, caption, parse_mode="Markdown", reply_markup=kb, protect_content=should_protect_content(uid))
    except Exception as e:
        bot.send_message(cid, f"❌ خطأ بعرض الرواية: {e}")

# ==============================================================================
# MANGA / MANHWA SYSTEM (نفس بنية الروايات تمامًا: فصول PDF/EPUB)
# ==============================================================================

def create_manga(title, story, tags, poster_file_id, category_id, created_by=None):
    mid = str(uuid.uuid4())[:8]
    tags_list = tags if isinstance(tags, list) else ([tags] if tags else ["عام"])
    db["manga"][mid] = {
        "id": mid, "title": title, "story": story, "tags": tags_list,
        "poster_file_id": poster_file_id, "category": category_id,
        "linked_chat": None, "created_by": str(created_by) if created_by else None,
        "chapters": {}, "merged": {}, "created_at": str(datetime.now())
    }
    if category_id in db["categories"]:
        db["categories"][category_id].setdefault("manga", []).append(mid)
    sync_db()
    return mid

def get_manga(mid): return db["manga"].get(mid)
def manga_chapter_numbers(manga): return sorted([int(k) for k in manga.get("chapters", {}).keys()])
def manga_merge_groups(manga): return sorted(manga.get("merged", {}).keys(), key=lambda g: int(g.split("-")[0]))

def manga_next_merge_range(manga):
    chs = manga_chapter_numbers(manga)
    if not chs: return None
    max_ch = max(chs)
    covered_end = 0
    for g in manga_merge_groups(manga):
        covered_end = max(covered_end, int(g.split("-")[1]))
    if covered_end >= max_ch: return None
    start = covered_end + 1
    return f"{start}-{min(start + 24, max_ch)}"

def get_user_manga_progress(ustr, mid): return db["users"].get(ustr, {}).get("manga_progress", {}).get(mid)

def set_user_manga_progress(ustr, mid, chapter_num):
    if ustr not in db["users"]: return
    db["users"][ustr].setdefault("manga_progress", {})
    db["users"][ustr]["manga_progress"][mid] = chapter_num
    sync_db()

def manga_card_kb(mid, uid):
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("👁️ شاهد الآن", callback_data=f"manga_watch_{mid}"))
    m.add(types.InlineKeyboardButton("➕ نقل إلى القائمة", callback_data=f"addlist_manga_{mid}"))
    m.add(types.InlineKeyboardButton(f"💬 التعليقات ({len(get_comments(f'manga_{mid}'))})", callback_data=f"comments_manga_{mid}"))
    manga = get_manga(mid)
    if is_admin(uid):
        m.add(types.InlineKeyboardButton("⚙️ إدارة العمل", callback_data=f"manga_admin_{mid}"))
    cat_id = manga.get("category", "root") if manga else "root"
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"nav_{cat_id}"))
    return m

def manga_watch_kb(mid, ustr):
    manga = get_manga(mid)
    m = types.InlineKeyboardMarkup(row_width=2)
    chs = manga_chapter_numbers(manga)
    if not chs:
        m.add(types.InlineKeyboardButton("❌ لا توجد فصول مرفوعة بعد", callback_data="noop_"))
    else:
        m.add(
            types.InlineKeyboardButton(f"📖 الفصل {chs[0]}", callback_data=f"manga_ch_{mid}_{chs[0]}"),
            types.InlineKeyboardButton(f"🆕 الفصل {chs[-1]}", callback_data=f"manga_ch_{mid}_{chs[-1]}")
        )
    if manga.get("merged"):
        m.add(types.InlineKeyboardButton("📦 الفصول المدمجة", callback_data=f"manga_merged_{mid}"))
    if len(chs) > 4:
        m.add(types.InlineKeyboardButton("🔢 كل الفصول (شبكة أرقام)", callback_data=f"mangagrid_{mid}_0"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"open_manga_{mid}"))
    return m

def manga_chapter_nav_kb(mid, chapter_num):
    manga = get_manga(mid)
    chs = manga_chapter_numbers(manga)
    m = types.InlineKeyboardMarkup(row_width=3)
    idx = chs.index(chapter_num) if chapter_num in chs else 0
    prev_ch = chs[idx - 1] if idx > 0 else None
    next_ch = chs[idx + 1] if idx < len(chs) - 1 else None
    m.add(
        types.InlineKeyboardButton("◀️ السابق" if prev_ch else "▪️", callback_data=f"manga_ch_{mid}_{prev_ch}" if prev_ch else "noop_"),
        types.InlineKeyboardButton(f"فصل {chapter_num}", callback_data="noop_"),
        types.InlineKeyboardButton("التالي ▶️" if next_ch else "▪️", callback_data=f"manga_ch_{mid}_{next_ch}" if next_ch else "noop_")
    )
    m.add(
        types.InlineKeyboardButton("📄 PDF", callback_data=f"manga_get_{mid}_ch_{chapter_num}_pdf"),
        types.InlineKeyboardButton("📱 EPUB", callback_data=f"manga_get_{mid}_ch_{chapter_num}_epub")
    )
    m.add(types.InlineKeyboardButton("🔢 كل الفصول", callback_data=f"mangagrid_{mid}_0"))
    m.add(types.InlineKeyboardButton(f"💬 تعليقات الفصل ({len(get_comments(f'manga_{mid}_ch_{chapter_num}'))})", callback_data=f"comments_manga_{mid}_ch_{chapter_num}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"manga_watch_{mid}"))
    return m

MANGA_GRID_PAGE_SIZE = 20

def manga_chapter_grid_kb(mid, page=0):
    manga = get_manga(mid)
    chs = manga_chapter_numbers(manga)
    m = types.InlineKeyboardMarkup(row_width=5)
    start = page * MANGA_GRID_PAGE_SIZE
    page_chs = chs[start:start + MANGA_GRID_PAGE_SIZE]
    row = []
    for ch in page_chs:
        row.append(types.InlineKeyboardButton(str(ch), callback_data=f"manga_ch_{mid}_{ch}"))
        if len(row) == 5: m.add(*row); row = []
    if row: m.add(*row)
    nav_row = []
    if start > 0: nav_row.append(types.InlineKeyboardButton("⏮️ السابقة", callback_data=f"mangagrid_{mid}_{page-1}"))
    if start + MANGA_GRID_PAGE_SIZE < len(chs): nav_row.append(types.InlineKeyboardButton("التالية ⏭️", callback_data=f"mangagrid_{mid}_{page+1}"))
    if nav_row: m.add(*nav_row)
    total_pages = max(1, (len(chs) + MANGA_GRID_PAGE_SIZE - 1) // MANGA_GRID_PAGE_SIZE)
    m.add(types.InlineKeyboardButton(f"📄 صفحة {page+1}/{total_pages}", callback_data="noop_"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"manga_watch_{mid}"))
    return m

def manga_admin_kb(mid):
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("📥 رفع فصل جديد", callback_data=f"manga_addch_{mid}"))
    m.add(types.InlineKeyboardButton("🗑️ حذف فصل", callback_data=f"manga_delch_{mid}"))
    m.add(types.InlineKeyboardButton("📦 رفع مجموعة مدمجة", callback_data=f"manga_addmerge_{mid}"))
    m.add(types.InlineKeyboardButton("🔗 ربط بقناة/مجموعة", callback_data=f"manga_link_{mid}"))
    m.add(types.InlineKeyboardButton("🗑️ حذف العمل", callback_data=f"manga_del_{mid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"open_manga_{mid}"))
    return m

def open_manga(uid, cid, mid_msg, mid, as_new_message=False):
    manga = get_manga(mid)
    if not manga:
        bot.send_message(cid, "❌ العمل غير موجود.")
        return
    if not db["config"].get("content_features", {}).get("manga", {}).get("active", True) and not is_admin(uid):
        bot.send_message(cid, "⚠️ هذي الميزة معطّلة حاليًا.")
        return
    track_event("manga_view", mid)
    tags_display = "، ".join(_item_tags(manga)) or "عام"
    ustr = str(uid)
    progress = get_user_manga_progress(ustr, mid)
    marker = db["config"].get("progress_marker", "🔴")
    progress_line = f"\n\n{marker} آخر فصل قرأته: *{progress}*" if progress else ""
    caption = f"*{md_safe(manga['title'])}*\n\n{md_safe(manga.get('story',''))}\n\n🏷️ {md_safe(tags_display)}{progress_line}"
    kb = manga_card_kb(mid, uid)
    try:
        if manga.get("poster_file_id"):
            bot.send_photo(cid, manga["poster_file_id"], caption=caption, parse_mode="Markdown", reply_markup=kb, protect_content=should_protect_content(uid))
        else:
            bot.send_message(cid, caption, parse_mode="Markdown", reply_markup=kb, protect_content=should_protect_content(uid))
    except Exception as e:
        bot.send_message(cid, f"❌ خطأ بعرض العمل: {e}")

def send_manga_chapter_file(cid, mid, chapter_num, fmt):
    manga = get_manga(mid)
    if not manga: return False
    ch = manga.get("chapters", {}).get(str(chapter_num))
    if not ch or not ch.get(fmt):
        bot.send_message(cid, f"❌ لا يوجد ملف بصيغة {fmt.upper()} لهذا الفصل بعد.")
        return False
    try:
        bot.send_document(cid, ch[fmt], caption=f"🎨 *{md_safe(manga['title'])}* — الفصل {chapter_num} ({fmt.upper()})", parse_mode="Markdown", protect_content=should_protect_content(cid))
        track_event("manga_download", mid)
        return True
    except Exception as e:
        bot.send_message(cid, f"❌ خطأ بإرسال الملف: {e}")
        return False

def send_manga_merged_file(cid, mid, group, fmt):
    manga = get_manga(mid)
    if not manga: return False
    mg = manga.get("merged", {}).get(group)
    if not mg or not mg.get(fmt):
        bot.send_message(cid, f"❌ لا يوجد ملف مدمج بصيغة {fmt.upper()} لهذي المجموعة بعد.")
        return False
    try:
        bot.send_document(cid, mg[fmt], caption=f"📦 *{md_safe(manga['title'])}* — الفصول {group} ({fmt.upper()})", parse_mode="Markdown", protect_content=should_protect_content(cid))
        track_event("manga_download", mid)
        return True
    except Exception as e:
        bot.send_message(cid, f"❌ خطأ بإرسال الملف: {e}")
        return False

# ==============================================================================
# SERIES / MOVIES SYSTEM (حلقات فيديو بجودات متعددة: 480/720/1080/2K/4K)
# ==============================================================================

SERIES_QUALITIES = ["480p", "720p", "1080p", "2K", "4K"]

def create_series(title, story, tags, poster_file_id, category_id, created_by=None):
    sid = str(uuid.uuid4())[:8]
    tags_list = tags if isinstance(tags, list) else ([tags] if tags else ["عام"])
    db["series"][sid] = {
        "id": sid, "title": title, "story": story, "tags": tags_list,
        "poster_file_id": poster_file_id, "category": category_id,
        "linked_chat": None, "created_by": str(created_by) if created_by else None,
        "episodes": {},  # {"1": {"480p": {"file_id":.., "sub_only": bool}, "720p": {...}, ...}}
        "created_at": str(datetime.now())
    }
    if category_id in db["categories"]:
        db["categories"][category_id].setdefault("series", []).append(sid)
    sync_db()
    return sid

def get_series(sid): return db["series"].get(sid)
def series_episode_numbers(series): return sorted([int(k) for k in series.get("episodes", {}).keys()])
def series_episode_qualities(series, ep_num):
    """يرجع الجودات المتوفرة فعليًا لحلقة معيّنة (اللي المطور رفعها فقط)."""
    ep = series.get("episodes", {}).get(str(ep_num), {})
    return [q for q in SERIES_QUALITIES if q in ep]

def get_user_series_progress(ustr, sid): return db["users"].get(ustr, {}).get("series_progress", {}).get(sid)

def set_user_series_progress(ustr, sid, ep_num):
    if ustr not in db["users"]: return
    db["users"][ustr].setdefault("series_progress", {})
    db["users"][ustr]["series_progress"][sid] = ep_num
    sync_db()

def series_card_kb(sid, uid):
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("👁️ شاهد الآن", callback_data=f"series_watch_{sid}"))
    m.add(types.InlineKeyboardButton("➕ نقل إلى القائمة", callback_data=f"addlist_series_{sid}"))
    m.add(types.InlineKeyboardButton(f"💬 التعليقات ({len(get_comments(f'series_{sid}'))})", callback_data=f"comments_series_{sid}"))
    series = get_series(sid)
    if is_admin(uid):
        m.add(types.InlineKeyboardButton("⚙️ إدارة العمل", callback_data=f"series_admin_{sid}"))
    cat_id = series.get("category", "root") if series else "root"
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"nav_{cat_id}"))
    return m

def series_watch_kb(sid, ustr):
    series = get_series(sid)
    m = types.InlineKeyboardMarkup(row_width=2)
    eps = series_episode_numbers(series)
    if not eps:
        m.add(types.InlineKeyboardButton("❌ لا توجد حلقات مرفوعة بعد", callback_data="noop_"))
    else:
        m.add(
            types.InlineKeyboardButton(f"🎬 الحلقة {eps[0]}", callback_data=f"series_ep_{sid}_{eps[0]}"),
            types.InlineKeyboardButton(f"🆕 الحلقة {eps[-1]}", callback_data=f"series_ep_{sid}_{eps[-1]}")
        )
    if len(eps) > 4:
        m.add(types.InlineKeyboardButton("🔢 كل الحلقات (شبكة أرقام)", callback_data=f"seriesgrid_{sid}_0"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"open_series_{sid}"))
    return m

def series_quality_kb(sid, ep_num, uid):
    """يعرض الجودات المتوفرة لهذي الحلقة، مع قفل 🔒 على أي جودة مقصورة على المشتركين
    لو المستخدم مو مشترك، وفتح مباشر للجودات المجانية أو للمشرفين."""
    series = get_series(sid)
    ep = series.get("episodes", {}).get(str(ep_num), {})
    m = types.InlineKeyboardMarkup(row_width=2)
    eps = series_episode_numbers(series)
    idx = eps.index(ep_num) if ep_num in eps else 0
    prev_ep = eps[idx - 1] if idx > 0 else None
    next_ep = eps[idx + 1] if idx < len(eps) - 1 else None
    m.add(
        types.InlineKeyboardButton("◀️ السابقة" if prev_ep else "▪️", callback_data=f"series_ep_{sid}_{prev_ep}" if prev_ep else "noop_"),
        types.InlineKeyboardButton(f"حلقة {ep_num}", callback_data="noop_"),
        types.InlineKeyboardButton("التالية ▶️" if next_ep else "▪️", callback_data=f"series_ep_{sid}_{next_ep}" if next_ep else "noop_")
    )
    is_sub = check_sub(uid)
    for q in SERIES_QUALITIES:
        if q not in ep: continue
        sub_only = ep[q].get("sub_only", False)
        locked = sub_only and not is_sub and not is_admin(uid)
        label = f"🔒 {q} (اشتراك)" if locked else f"▶️ {q}"
        m.add(types.InlineKeyboardButton(label, callback_data=f"series_get_{sid}_{ep_num}_{q}"))
    m.add(types.InlineKeyboardButton("🔢 كل الحلقات", callback_data=f"seriesgrid_{sid}_0"))
    m.add(types.InlineKeyboardButton(f"💬 تعليقات الحلقة ({len(get_comments(f'series_{sid}_ep_{ep_num}'))})", callback_data=f"comments_series_{sid}_ep_{ep_num}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"series_watch_{sid}"))
    return m

SERIES_GRID_PAGE_SIZE = 20

def series_episode_grid_kb(sid, page=0):
    series = get_series(sid)
    eps = series_episode_numbers(series)
    m = types.InlineKeyboardMarkup(row_width=5)
    start = page * SERIES_GRID_PAGE_SIZE
    page_eps = eps[start:start + SERIES_GRID_PAGE_SIZE]
    row = []
    for ep in page_eps:
        row.append(types.InlineKeyboardButton(str(ep), callback_data=f"series_ep_{sid}_{ep}"))
        if len(row) == 5: m.add(*row); row = []
    if row: m.add(*row)
    nav_row = []
    if start > 0: nav_row.append(types.InlineKeyboardButton("⏮️ السابقة", callback_data=f"seriesgrid_{sid}_{page-1}"))
    if start + SERIES_GRID_PAGE_SIZE < len(eps): nav_row.append(types.InlineKeyboardButton("التالية ⏭️", callback_data=f"seriesgrid_{sid}_{page+1}"))
    if nav_row: m.add(*nav_row)
    total_pages = max(1, (len(eps) + SERIES_GRID_PAGE_SIZE - 1) // SERIES_GRID_PAGE_SIZE)
    m.add(types.InlineKeyboardButton(f"📄 صفحة {page+1}/{total_pages}", callback_data="noop_"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"series_watch_{sid}"))
    return m

def series_admin_kb(sid):
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("📥 رفع حلقة جديدة", callback_data=f"series_addep_{sid}"))
    m.add(types.InlineKeyboardButton("🗑️ حذف حلقة", callback_data=f"series_delep_{sid}"))
    m.add(types.InlineKeyboardButton("🔗 ربط بقناة/مجموعة", callback_data=f"series_link_{sid}"))
    m.add(types.InlineKeyboardButton("🗑️ حذف العمل", callback_data=f"series_del_{sid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"open_series_{sid}"))
    return m

def open_series(uid, cid, mid_msg, sid, as_new_message=False):
    series = get_series(sid)
    if not series:
        bot.send_message(cid, "❌ العمل غير موجود.")
        return
    if not db["config"].get("content_features", {}).get("series", {}).get("active", True) and not is_admin(uid):
        bot.send_message(cid, "⚠️ هذي الميزة معطّلة حاليًا.")
        return
    track_event("series_view", sid)
    tags_display = "، ".join(_item_tags(series)) or "عام"
    ustr = str(uid)
    progress = get_user_series_progress(ustr, sid)
    marker = db["config"].get("progress_marker", "🔴")
    progress_line = f"\n\n{marker} آخر حلقة شاهدتها: *{progress}*" if progress else ""
    caption = f"*{md_safe(series['title'])}*\n\n{md_safe(series.get('story',''))}\n\n🏷️ {md_safe(tags_display)}{progress_line}"
    kb = series_card_kb(sid, uid)
    try:
        if series.get("poster_file_id"):
            bot.send_photo(cid, series["poster_file_id"], caption=caption, parse_mode="Markdown", reply_markup=kb, protect_content=should_protect_content(uid))
        else:
            bot.send_message(cid, caption, parse_mode="Markdown", reply_markup=kb, protect_content=should_protect_content(uid))
    except Exception as e:
        bot.send_message(cid, f"❌ خطأ بعرض العمل: {e}")

def send_series_episode(cid, sid, ep_num, quality, uid):
    """يرسل الحلقة بالجودة المطلوبة، مع التحقق من قفل الاشتراك لو الجودة مقصورة."""
    series = get_series(sid)
    if not series: return False
    ep = series.get("episodes", {}).get(str(ep_num), {})
    q_data = ep.get(quality)
    if not q_data or not q_data.get("file_id"):
        bot.send_message(cid, f"❌ لا توجد نسخة بجودة {quality} لهذي الحلقة بعد.")
        return False
    if q_data.get("sub_only") and not check_sub(uid) and not is_admin(uid):
        bot.send_message(cid, f"🔒 جودة {quality} مقصورة على المشتركين. اشترك من «💎 الاشتراك» للوصول لها.")
        return False
    try:
        bot.send_video(cid, q_data["file_id"], caption=f"🎬 *{md_safe(series['title'])}* — الحلقة {ep_num} ({quality})", parse_mode="Markdown", protect_content=should_protect_content(cid))
        track_event("series_download", sid)
        return True
    except Exception as e:
        bot.send_message(cid, f"❌ خطأ بإرسال الحلقة: {e}")
        return False

FONT_CACHE_DIR = "/tmp/bot_fonts"
ARABIC_FONT_PATH = os.path.join(FONT_CACHE_DIR, "Amiri-Regular.ttf")
ARABIC_FONT_URL = "https://github.com/google/fonts/raw/main/ofl/amiri/Amiri-Regular.ttf"

def _ensure_arabic_font():
    """يحمّل خط عربي مجاني (Amiri) مرة وحدة ويخزّنه محليًا، يحتاجه reportlab لعرض
    العربي صح بدل مربعات فاضية أو حروف مفكوكة."""
    if os.path.exists(ARABIC_FONT_PATH):
        return ARABIC_FONT_PATH
    t0 = time.time()
    try:
        os.makedirs(FONT_CACHE_DIR, exist_ok=True)
        r = requests.get(ARABIC_FONT_URL, timeout=20)
        r.raise_for_status()
        with open(ARABIC_FONT_PATH, "wb") as f:
            f.write(r.content)
        logger.info(f"[font] تحميل الخط العربي من الشبكة: {time.time()-t0:.1f}ث")
        return ARABIC_FONT_PATH
    except Exception as e:
        logger.error(f"❌ فشل تحميل الخط العربي بعد {time.time()-t0:.1f}ث: {e}")
        return None

def _shape_arabic_line(line):
    """يعالج أي سطر فيه عربي: يوصل الحروف ببعض ويرتبها من اليمين لليسار صح،
    بدون ما يغيّر أي كلمة أو حرف من النص نفسه — بس شكل العرض التقني."""
    try:
        reshaped = arabic_reshaper.reshape(line)
        return bidi_get_display(reshaped)
    except Exception:
        return line  # لو صار خطأ بالمعالجة، نرجع النص الأصلي بدل ما نفقده

def extract_text_from_file(path, ext):
    """يستخرج النص الخام من أي صيغة مدعومة بدون تغيير كلمة واحدة."""
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        reader = pypdf.PdfReader(path)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext in ("docx", "doc"):
        d = docx.Document(path)
        return "\n".join(p.text for p in d.paragraphs)
    elif ext == "epub":
        book = epub.read_epub(path)
        import re
        parts = []
        for item in book.get_items():
            if item.get_type() == 9:  # ITEM_DOCUMENT
                html = item.get_content().decode("utf-8", errors="ignore")
                text = re.sub(r"<[^>]+>", "\n", html)  # إزالة وسوم HTML فقط، بدون لمس النص
                parts.append(text.strip())
        return "\n\n".join(parts)
    elif ext == "txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext in ("html", "htm"):
        import re
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            html = f.read()
        return re.sub(r"<[^>]+>", "\n", html).strip()
    else:
        raise ValueError(f"صيغة غير مدعومة للاستخراج: {ext}")

def build_txt_from_text(text, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)

def build_html_from_text(text, out_path, title="Document"):
    paragraphs = "".join(f"<p>{p}</p>" for p in text.split("\n") if p.strip())
    html = f"<html dir='rtl' lang='ar'><head><meta charset='utf-8'><title>{title}</title></head><body>{paragraphs}</body></html>"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html)

def remove_text_from_content(text, phrases):
    """يحذف كل ظهور لأي عبارة من قائمة العبارات (نص أو رابط) من المحتوى كامل، مهما
    تكررت، بالإضافة لأي عبارات دائمة الحذف مسجّلة بإعدادات المطور، ثم ينظّف الفراغات
    الزايدة اللي تنتج عن الحذف (سطور فاضية متتالية، مسافات مضاعفة) حتى ما يضل أثر
    فاضي مكان النص المحذوف."""
    import re
    all_phrases = list(phrases) + db["config"].get("always_remove_phrases", [])
    for phrase in all_phrases:
        if not phrase.strip(): continue
        text = text.replace(phrase, "")
    # تنظيف الفراغات الناتجة عن الحذف: مسافات مضاعفة بنفس السطر، وأسطر فاضية متتالية
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()

def build_pdf_from_sections(sections, out_path):
    """يبني PDF من عدة أقسام (كل قسم = ملف مدموج)، كل قسم يبدأ بصفحة جديدة وعنوان
    واضح (رقم/اسم الفصل)، بالضبط زي الفصول بالروايات العادية بس كلهم بملف وحد."""
    font_path = _ensure_arabic_font()
    font_name = "Helvetica"
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("Amiri", font_path))
            font_name = "Amiri"
        except Exception as e:
            logger.error(f"❌ فشل تسجيل الخط: {e}")
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    margin = 2 * cm
    max_width = width - 2 * margin
    line_height = 16
    font_size = 12
    for label, text in sections:
        c.showPage()  # كل قسم/فصل يبدأ بصفحة جديدة تمامًا
        y = height - margin
        c.setFont(font_name, 16)
        title_line = _shape_arabic_line(label) if any("\u0600" <= ch <= "\u06FF" for ch in label) else label
        try:
            c.drawRightToLeft(width - margin, y, title_line) if hasattr(c, "drawRightToLeft") else c.drawString(margin, y, title_line)
        except Exception:
            c.drawString(margin, y, label)
        y -= line_height * 2
        c.setFont(font_name, font_size)
        for raw_line in text.split("\n"):
            wrapped = _wrap_line_for_pdf(c, raw_line, font_name, font_size, max_width)
            for raw_sub in wrapped:
                if y < margin:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = height - margin
                display_line = _shape_arabic_line(raw_sub) if raw_sub.strip() else raw_sub
                try:
                    c.drawRightToLeft(width - margin, y, display_line) if hasattr(c, "drawRightToLeft") else c.drawString(margin, y, display_line)
                except Exception:
                    c.drawString(margin, y, raw_sub)
                y -= line_height
    c.save()

def build_docx_from_sections(sections, out_path):
    """يبني DOCX من عدة أقسام، كل قسم يبدأ بعنوان (Heading) وصفحة جديدة."""
    d = docx.Document()
    for i, (label, text) in enumerate(sections):
        if i > 0:
            d.add_page_break()
        heading = d.add_heading(label, level=1)
        if any("\u0600" <= ch <= "\u06FF" for ch in label):
            heading.alignment = 2
        for para in text.split("\n"):
            p = d.add_paragraph(para)
            if any("\u0600" <= ch <= "\u06FF" for ch in para):
                p.alignment = 2
                pPr = p._p.get_or_add_pPr()
                bidi = pPr.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", {})
                pPr.append(bidi)
    d.save(out_path)

def build_epub_from_sections(sections, out_path, title="Document"):
    """يبني EPUB من عدة أقسام، كل قسم يصير فصل (chapter) منفصل بجدول المحتويات،
    بالضبط زي الفصل 1 والفصل 2 بمنهج القراءة العادي."""
    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(title)
    book.set_language("ar")
    chapters = []
    for i, (label, text) in enumerate(sections):
        paragraphs = "".join(f"<p>{p}</p>" for p in text.split("\n") if p.strip())
        ch = epub.EpubHtml(title=label, file_name=f"chapter_{i+1}.xhtml", lang="ar")
        ch.content = f"<html dir='rtl'><body><h1>{label}</h1>{paragraphs}</body></html>"
        book.add_item(ch)
        chapters.append(ch)
    book.toc = tuple(epub.Link(ch.file_name, ch.title, f"ch{i}") for i, ch in enumerate(chapters))
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + chapters
    epub.write_epub(out_path, book, {})

def build_txt_from_sections(sections, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        for i, (label, text) in enumerate(sections):
            if i > 0: f.write("\n\n" + "="*40 + "\n\n")
            f.write(f"{label}\n\n{text}")

def build_html_from_sections(sections, out_path, title="Document"):
    body_parts = []
    for label, text in sections:
        paragraphs = "".join(f"<p>{p}</p>" for p in text.split("\n") if p.strip())
        body_parts.append(f"<div class='section'><h1>{label}</h1>{paragraphs}</div><hr/>")
    html = f"<html dir='rtl' lang='ar'><head><meta charset='utf-8'><title>{title}</title></head><body>{''.join(body_parts)}</body></html>"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html)

def merge_files_to_one(sections, target_ext, out_path, title="Document"):
    """يدمج قائمة (label, text) لملف وحد بالصيغة المطلوبة، كل قسم يظهر بمكانه الخاص
    (فصل 1، فصل 2، إلخ) بالضبط زي الفصول العادية بالروايات."""
    if not _ensure_convert_libs():
        raise RuntimeError(f"مكتبات التحويل غير مثبتة: {_CONVERT_IMPORT_ERROR}")
    target_ext = target_ext.lower().lstrip(".")
    if target_ext == "pdf": build_pdf_from_sections(sections, out_path)
    elif target_ext in ("docx", "doc"): build_docx_from_sections(sections, out_path)
    elif target_ext == "epub": build_epub_from_sections(sections, out_path, title)
    elif target_ext == "txt": build_txt_from_sections(sections, out_path)
    elif target_ext in ("html", "htm"): build_html_from_sections(sections, out_path, title)
    else: raise ValueError(f"صيغة هدف غير مدعومة: {target_ext}")

def _wrap_line_for_pdf(c, raw_line, font_name, font_size, max_width):
    """يلف أي سطر طويل (شائع بالنصوص المستخرجة من EPUB اللي ما فيها أسطر طبيعية)
    لعدة أسطر تناسب عرض الصفحة، بدل ما ينرسم خارج حدود الصفحة ويضيع/يتقطع بصمت."""
    if not raw_line.strip():
        return [raw_line]
    words = raw_line.split(" ")
    lines = []
    current = ""
    for w in words:
        candidate = f"{current} {w}".strip()
        try:
            width = c.stringWidth(candidate, font_name, font_size)
        except Exception:
            width = len(candidate) * (font_size * 0.6)  # تقدير احتياطي لو فشل قياس العرض
        if width <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines or [raw_line]

def build_pdf_from_text(text, out_path, title="Document"):
    """يبني ملف PDF من نص خام، مع دعم كامل للعربي (تشكيل + اتجاه صحيح) بدون تغيير الكلمات."""
    font_path = _ensure_arabic_font()
    font_name = "Helvetica"
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("Amiri", font_path))
            font_name = "Amiri"
        except Exception as e:
            logger.error(f"❌ فشل تسجيل الخط بـ reportlab: {e}")
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    margin = 2 * cm
    max_width = width - 2 * margin
    y = height - margin
    line_height = 16
    font_size = 12
    c.setFont(font_name, font_size)
    for raw_line in text.split("\n"):
        wrapped = _wrap_line_for_pdf(c, raw_line, font_name, font_size, max_width)
        for raw_sub in wrapped:
            if y < margin:
                c.showPage()
                c.setFont(font_name, font_size)
                y = height - margin
            display_line = _shape_arabic_line(raw_sub) if raw_sub.strip() else raw_sub
            try:
                c.drawRightToLeft(width - margin, y, display_line) if hasattr(c, "drawRightToLeft") else c.drawString(margin, y, display_line)
            except Exception:
                c.drawString(margin, y, raw_sub)  # آخر حل احتياطي: نص عادي بدون تشكيل لو صار خطأ عرض
            y -= line_height
    c.save()

def build_docx_from_text(text, out_path):
    """يبني ملف DOCX/DOC من نص خام، مع ضبط اتجاه الفقرة تلقائيًا لليمين لليسار للعربي."""
    d = docx.Document()
    for para in text.split("\n"):
        p = d.add_paragraph(para)
        # لو الفقرة فيها عربي، نضبط اتجاهها لليمين لليسار حتى تنعرض صح بمحرر النصوص
        if any("\u0600" <= ch <= "\u06FF" for ch in para):
            p.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
            pPr = p._p.get_or_add_pPr()
            bidi = pPr.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", {})
            pPr.append(bidi)
    d.save(out_path)

def build_epub_from_text(text, out_path, title="Document"):
    """يبني ملف EPUB صحيح من نص خام، بدون تغيير أي كلمة."""
    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(title)
    book.set_language("ar")
    chapter = epub.EpubHtml(title=title, file_name="content.xhtml", lang="ar")
    paragraphs = "".join(f"<p>{p}</p>" for p in text.split("\n") if p.strip())
    chapter.content = f"<html dir='rtl'><body>{paragraphs}</body></html>"
    book.add_item(chapter)
    book.toc = (epub.Link("content.xhtml", title, "content"),)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]
    epub.write_epub(out_path, book, {})

def convert_locally(src_path, src_ext, target_ext, out_path, title="Document"):
    """المحول المحلي: يستخرج النص من الصيغة المصدر ويبنيه بالصيغة الهدف، بدون
    تغيير كلمة واحدة من المحتوى الفعلي — بس تغيير الحاوية/الصيغة نفسها.
    هذا هو خط الدفاع الأخير (fallback) اللي يشتغل دايمًا بدون أي اتصال خارجي."""
    t0 = time.time()
    if not _ensure_convert_libs():
        raise RuntimeError(f"مكتبات التحويل غير مثبتة على الاستضافة: {_CONVERT_IMPORT_ERROR}")
    logger.info(f"[convert] تحميل المكتبات: {time.time()-t0:.1f}ث")
    t1 = time.time()
    text = extract_text_from_file(src_path, src_ext)
    logger.info(f"[convert] استخراج النص ({len(text)} حرف): {time.time()-t1:.1f}ث")
    target_ext = target_ext.lower().lstrip(".")
    t2 = time.time()
    if target_ext == "pdf":
        build_pdf_from_text(text, out_path, title)
    elif target_ext in ("docx", "doc"):
        build_docx_from_text(text, out_path)
    elif target_ext == "epub":
        build_epub_from_text(text, out_path, title)
    elif target_ext == "txt":
        build_txt_from_text(text, out_path)
    elif target_ext in ("html", "htm"):
        build_html_from_text(text, out_path, title)
    else:
        raise ValueError(f"صيغة هدف غير مدعومة: {target_ext}")
    logger.info(f"[convert] بناء {target_ext}: {time.time()-t2:.1f}ث | الإجمالي: {time.time()-t0:.1f}ث")

def convert_via_cloudconvert(src_path, src_ext, target_ext, out_path):
    """يحوّل الملف عبر CloudConvert (تحويل حقيقي على مستوى الملف نفسه، وليس
    استخراج نص — أدق للملفات ذات التنسيق المعقد مثل الجداول والصور المدمجة).
    يرمي استثناء عند أي فشل حتى يلتقطه convert_file ويرجع للمحلي فورًا."""
    import cloudconvert
    cloudconvert.configure(api_key=CLOUDCONVERT_KEY, sandbox=False)
    target_ext = target_ext.lower().lstrip(".")
    job = cloudconvert.Job.create(payload={
        "tasks": {
            "import-file": {"operation": "import/upload"},
            "convert-file": {
                "operation": "convert",
                "input": "import-file",
                "output_format": target_ext,
            },
            "export-file": {"operation": "export/url", "input": "convert-file"},
        }
    })
    import_task = next(t for t in job["tasks"] if t["name"] == "import-file")
    upload_task = cloudconvert.Task.find(id=import_task["id"])
    with open(src_path, "rb") as f:
        cloudconvert.Task.upload(file_name=os.path.basename(src_path), task=upload_task, file=f)
    job = cloudconvert.Job.wait(id=job["id"])
    export_task = next(t for t in job["tasks"] if t["name"] == "export-file" and t["status"] == "finished")
    file_info = export_task["result"]["files"][0]
    download_url = file_info["url"]
    r = requests.get(download_url, timeout=60)
    r.raise_for_status()
    with open(out_path, "wb") as f:
        f.write(r.content)

def convert_via_convertio(src_path, src_ext, target_ext, out_path):
    """يحوّل الملف عبر Convertio (REST API بسيط: رفع base64 → استعلام دوري عن
    الحالة → تنزيل الناتج). يُستخدم كخيار ثانٍ بعد CloudConvert، قبل الرجوع
    للتحويل المحلي. يرمي استثناء عند أي فشل حتى يلتقطه convert_file."""
    import base64
    target_ext = target_ext.lower().lstrip(".")
    with open(src_path, "rb") as f:
        content_b64 = base64.b64encode(f.read()).decode()
    start_resp = requests.post(
        "https://api.convertio.co/convert",
        json={
            "apikey": CONVERTIO_KEY,
            "input": "base64",
            "file": content_b64,
            "filename": os.path.basename(src_path),
            "outputformat": target_ext,
        },
        timeout=60,
    )
    start_resp.raise_for_status()
    start_data = start_resp.json()
    if start_data.get("code") != 200:
        raise RuntimeError(f"Convertio رفض الطلب: {start_data.get('error', start_data)}")
    conv_id = start_data["data"]["id"]

    # نستعلم عن حالة التحويل كل ثانيتين لحد ما يخلص أو يفشل أو نتجاوز مهلة معقولة
    for _ in range(60):  # حتى 120 ثانية انتظار كحد أقصى
        time.sleep(2)
        status_resp = requests.get(f"https://api.convertio.co/convert/{conv_id}/status", timeout=30)
        status_resp.raise_for_status()
        status_data = status_resp.json()
        step = status_data.get("data", {}).get("step")
        if step == "finish":
            download_url = status_data["data"]["output"]["url"]
            r = requests.get(download_url, timeout=60)
            r.raise_for_status()
            with open(out_path, "wb") as f:
                f.write(r.content)
            return
        if step == "error":
            raise RuntimeError(f"Convertio فشل بالتحويل: {status_data.get('data', {}).get('error', 'غير معروف')}")
    raise TimeoutError("Convertio تجاوز مهلة الانتظار (120 ثانية) بدون اكتمال.")

def convert_via_convertapi(src_path, src_ext, target_ext, out_path):
    """يحوّل الملف عبر ConvertAPI (REST API: رفع مباشر multipart → تنزيل الناتج
    من رابط النتيجة). خيار سحابي ثالث، يُستخدم فقط لو CONVERTAPI_SECRET مضبوط
    (غير موجود افتراضيًا). يرمي استثناء عند أي فشل حتى يلتقطه convert_file."""
    target_ext = target_ext.lower().lstrip(".")
    src_ext_clean = src_ext.lower().lstrip(".")
    with open(src_path, "rb") as f:
        r = requests.post(
            f"https://v2.convertapi.com/convert/{src_ext_clean}/to/{target_ext}"
            f"?Secret={CONVERTAPI_SECRET}",
            files={"file": (os.path.basename(src_path), f)},
            timeout=90,
        )
    r.raise_for_status()
    data = r.json()
    files_out = data.get("Files")
    if not files_out:
        raise RuntimeError(f"ConvertAPI لم يرجع أي ملف ناتج: {data}")
    download_url = files_out[0]["Url"]
    dl = requests.get(download_url, timeout=60)
    dl.raise_for_status()
    with open(out_path, "wb") as f:
        f.write(dl.content)

def convert_file(src_path, src_ext, target_ext, out_path, title="Document"):
    """المحول الرئيسي: يجرب الخدمات السحابية بالترتيب — CloudConvert أولاً
    (لو المفتاح مضبوط)، ثم Convertio، ثم ConvertAPI، وعند فشل الكل أو غياب
    مفاتيحها يرجع تلقائيًا للتحويل المحلي — بحيث يبقى التحويل شغّالًا دومًا
    حتى بدون أي اتصال خارجي."""
    if CLOUDCONVERT_KEY:
        try:
            convert_via_cloudconvert(src_path, src_ext, target_ext, out_path)
            logger.info(f"[convert] تم التحويل عبر CloudConvert ✅")
            return
        except Exception as e:
            logger.warning(f"[convert] فشل CloudConvert: {e}. نجرب Convertio...")
    if CONVERTIO_KEY:
        try:
            convert_via_convertio(src_path, src_ext, target_ext, out_path)
            logger.info(f"[convert] تم التحويل عبر Convertio ✅")
            return
        except Exception as e:
            logger.warning(f"[convert] فشل Convertio: {e}. نجرب ConvertAPI...")
    if CONVERTAPI_SECRET:
        try:
            convert_via_convertapi(src_path, src_ext, target_ext, out_path)
            logger.info(f"[convert] تم التحويل عبر ConvertAPI ✅")
            return
        except Exception as e:
            logger.warning(f"[convert] فشل ConvertAPI: {e}. نرجع للتحويل المحلي...")
    convert_locally(src_path, src_ext, target_ext, out_path, title)
# ==============================================================================
# 6ه‍. نظام تقطيع الملفات الكبيرة والاستئناف (Chunking & Resume System)
# ==============================================================================
# يُستخدم لدمج/تحويل ملفات ضخمة (مئات الفصول) بدون تحميلها كاملة بالرام، مع
# قدرة استئناف حقيقية لو انطفى السيرفر بمنتصف العملية. مفعّل اختياريًا فقط
# (CHUNKING_ENABLED من قاعدة البيانات، افتراضيًا معطّل)، ولا يمس مسار التحويل
# العادي (convert_file) الحالي — أي ملف صغير يستمر يتحول بنفس الطريقة السابقة
# فورًا، والتقطيع يتفعّل فقط للملفات الكبيرة فعليًا أو لما يطلبه المطور صراحة.

# ── استيراد كسول لمكتبات Upstash Redis وpsutil (المراقبة/التخزين اختياريان) ──
_REDIS_LIBS_READY = False
_REDIS_IMPORT_ERROR = None
_redis_client = None
def _ensure_redis_client():
    """يحاول الاتصال بـ Upstash Redis عند أول استخدام فعلي. يرجع True/False،
    ولا يرمي استثناء أبدًا — أي فشل هنا يعني ببساطة إن RedisJobManager سيستخدم
    PostgreSQL كبديل تلقائي (انظر get_job_backend أدناه)."""
    global _REDIS_LIBS_READY, _REDIS_IMPORT_ERROR, _redis_client
    if _REDIS_LIBS_READY: return True
    if not (UPSTASH_REDIS_URL and UPSTASH_REDIS_TOKEN):
        _REDIS_IMPORT_ERROR = "متغيرات UPSTASH_REDIS_REST_URL / UPSTASH_REDIS_REST_TOKEN غير مضبوطة"
        return False
    try:
        from upstash_redis import Redis
        _redis_client = Redis(url=UPSTASH_REDIS_URL, token=UPSTASH_REDIS_TOKEN)
        _redis_client.ping()  # تأكيد فعلي إن الاتصال شغّال، مو بس إن المكتبة موجودة
        _REDIS_LIBS_READY = True
        return True
    except Exception as e:
        _REDIS_IMPORT_ERROR = str(e)
        _redis_client = None
        return False

_PSUTIL_READY = False
_PSUTIL_IMPORT_ERROR = None
def _ensure_psutil():
    global _PSUTIL_READY, _PSUTIL_IMPORT_ERROR
    if _PSUTIL_READY: return True
    try:
        global psutil
        import psutil
        _PSUTIL_READY = True
        return True
    except Exception as e:
        _PSUTIL_IMPORT_ERROR = str(e)
        return False


class TempFileManager:
    """مدير مركزي لكل الملفات المؤقتة المُنشأة أثناء التقطيع/الدمج: يسجّل كل
    ملف عند إنشائه، يحذفه بأمان (بمحاولات متكررة لو محجوز)، وعنده مهمة خلفية
    تنظّف أي ملف نسي أحد حذفه (أقدم من ساعتين) — ضمان مزدوج ضد تراكم القمامة
    على القرص حتى لو صار استثناء بمنتصف عملية ونسينا ننظّف يدويًا."""

    def __init__(self):
        self._registry = {}  # {path: registered_at_timestamp}
        self._lock = threading.Lock()

    def register(self, path):
        with self._lock:
            self._registry[path] = time.time()
        return path

    def safe_delete(self, path, retries=3, delay=0.5):
        """يحاول حذف ملف عدة مرات (لو محجوز مؤقتًا من عملية أخرى)، ويزيله من
        السجل بغض النظر عن نجاح الحذف الفعلي حتى ما يتكرر السجل بلا داعٍ."""
        with self._lock:
            self._registry.pop(path, None)
        for attempt in range(retries):
            try:
                if os.path.exists(path):
                    os.remove(path)
                return True
            except Exception as e:
                if attempt < retries - 1:
                    time.sleep(delay)
                else:
                    logger.warning(f"⚠️ فشل حذف ملف مؤقت {path} بعد {retries} محاولات: {e}")
        return False

    def cleanup_old(self, max_age_seconds=7200):
        """يحذف كل الملفات المسجّلة الأقدم من max_age_seconds (ساعتين افتراضيًا)."""
        with self._lock:
            now = time.time()
            stale = [p for p, t in self._registry.items() if now - t > max_age_seconds]
        for p in stale:
            self.safe_delete(p)

    def cleanup_all(self):
        """يحذف كل الملفات المسجّلة بغض النظر عن عمرها — يُستدعى عند إقلاع
        البوت لتنظيف أي بقايا من جلسة سابقة انتهت بشكل غير طبيعي."""
        with self._lock:
            paths = list(self._registry.keys())
        for p in paths:
            self.safe_delete(p)

    def stats(self):
        with self._lock:
            paths = list(self._registry.keys())
        total_size = 0
        existing = 0
        for p in paths:
            try:
                if os.path.exists(p):
                    total_size += os.path.getsize(p)
                    existing += 1
            except Exception:
                pass
        return {"registered": len(paths), "existing_on_disk": existing, "total_size_mb": round(total_size / (1024 * 1024), 2)}

    def cleanup_loop(self):
        """مهمة خلفية دائمة: تفحص كل ساعة وتنظّف أي ملف نُسي حذفه."""
        while True:
            try:
                self.cleanup_old(max_age_seconds=7200)
            except Exception as e:
                logger.error(f"❌ خطأ في تنظيف الملفات المؤقتة: {e}")
            time.sleep(3600)


temp_file_manager = TempFileManager()


class MemoryMonitor:
    """يراقب استهلاك الرام الفعلي للعملية (عبر psutil لو متاحة)، ويسمح لدوال
    التقطيع تنتظر حتى تنخفض الذاكرة قبل معالجة الجزء التالي — يمنع تراكم
    استهلاك الرام أثناء معالجة ملف ضخم بمئات الأجزاء المتتالية."""

    def __init__(self, max_mb=200):
        self.max_mb = max_mb

    def current_usage_mb(self):
        if not _ensure_psutil():
            return None
        try:
            return psutil.Process(os.getpid()).memory_info().rss / (1024 * 1024)
        except Exception:
            return None

    def wait_if_needed(self, max_wait_seconds=30):
        """لو الرام الحالي فوق الحد الأقصى، ينتظر (مع gc.collect() دوري) حتى
        تنخفض أو حتى تنتهي المهلة القصوى — لا يعلّق البوت للأبد حتى لو ظل
        الاستهلاك مرتفعًا لسبب خارج عن السيطرة."""
        usage = self.current_usage_mb()
        if usage is None or usage <= self.max_mb:
            return
        waited = 0
        while usage and usage > self.max_mb and waited < max_wait_seconds:
            gc.collect()
            time.sleep(1)
            waited += 1
            usage = self.current_usage_mb()
        if usage and usage > self.max_mb:
            logger.warning(f"⚠️ الرام لا تزال مرتفعة ({usage:.0f}MB) بعد {waited}ث انتظار، نتابع المعالجة على أي حال.")


memory_monitor = MemoryMonitor(max_mb=int(os.environ.get("MAX_RAM_MB", "200")))


class RedisJobManager:
    """مدير حالة "مهام التقطيع" (chunk jobs). يخزّن بـ Upstash Redis لو متاح
    (خفيف وسريع، لا يثقّل قاعدة البيانات الرئيسية بتحديثات متكررة كل جزء)، أو
    يرجع تلقائيًا لتخزين نفس البيانات داخل db["chunk_jobs"] (PostgreSQL) لو
    Redis غير مضبوط أو فشل الاتصال — بحيث تبقى ميزة الاستئناف تعمل بأي الحالتين.
    الواجهة (create_job/update_part/complete_job/fail_job/get_job/delete_job)
    موحّدة بغض النظر عن الخلفية الفعلية المستخدمة تحتها."""

    PREFIX = "chunkjob:"

    def __init__(self):
        self.use_redis = _ensure_redis_client()

    def _pg_jobs(self):
        return db.setdefault("chunk_jobs", {})

    def create_job(self, input_path, target_ext, total_parts, meta=None):
        job_id = str(uuid.uuid4())[:12]
        job_data = {
            "input_path": input_path, "target_ext": target_ext,
            "total_parts": total_parts, "completed_parts": 0,
            "parts_outputs": [], "status": "pending",
            "created_at": time.time(), "meta": meta or {},
        }
        if self.use_redis:
            try:
                _redis_client.set(f"{self.PREFIX}{job_id}", json.dumps(job_data))
                return job_id
            except Exception as e:
                logger.warning(f"⚠️ فشل إنشاء مهمة بـ Redis ({e})، نستخدم PostgreSQL بدلاً منه.")
                self.use_redis = False
        self._pg_jobs()[job_id] = job_data
        sync_db()
        return job_id

    def get_job(self, job_id):
        if self.use_redis:
            try:
                raw = _redis_client.get(f"{self.PREFIX}{job_id}")
                return json.loads(raw) if raw else None
            except Exception as e:
                logger.warning(f"⚠️ فشل قراءة مهمة من Redis ({e})، نبحث بـ PostgreSQL.")
        return self._pg_jobs().get(job_id)

    def _save_job(self, job_id, job_data):
        if self.use_redis:
            try:
                _redis_client.set(f"{self.PREFIX}{job_id}", json.dumps(job_data))
                return
            except Exception as e:
                logger.warning(f"⚠️ فشل حفظ مهمة بـ Redis ({e})، نستخدم PostgreSQL.")
                self.use_redis = False
        self._pg_jobs()[job_id] = job_data
        sync_db()

    def update_part(self, job_id, part_output_path):
        job = self.get_job(job_id)
        if not job:
            return None
        job["completed_parts"] = job.get("completed_parts", 0) + 1
        job.setdefault("parts_outputs", []).append(part_output_path)
        job["status"] = "processing"
        self._save_job(job_id, job)
        return job

    def complete_job(self, job_id, final_output):
        job = self.get_job(job_id)
        if not job:
            return None
        job["status"] = "completed"
        job["final_output"] = final_output
        self._save_job(job_id, job)
        return job

    def fail_job(self, job_id, error):
        job = self.get_job(job_id)
        if not job:
            return None
        job["status"] = "failed"
        job["error"] = str(error)[:300]
        self._save_job(job_id, job)
        return job

    def delete_job(self, job_id):
        if self.use_redis:
            try:
                _redis_client.delete(f"{self.PREFIX}{job_id}")
            except Exception as e:
                logger.warning(f"⚠️ فشل حذف مهمة من Redis: {e}")
        self._pg_jobs().pop(job_id, None)
        sync_db()

    def list_incomplete_jobs(self):
        """يرجّع كل المهام بحالة pending/processing — تُستخدم عند إقلاع البوت
        لعرض المهام المتوقفة على المطور واقتراح استئنافها."""
        results = []
        if self.use_redis:
            try:
                for key in _redis_client.keys(f"{self.PREFIX}*"):
                    raw = _redis_client.get(key)
                    if raw:
                        job = json.loads(raw)
                        if job.get("status") in ("pending", "processing"):
                            job["job_id"] = key[len(self.PREFIX):]
                            results.append(job)
                return results
            except Exception as e:
                logger.warning(f"⚠️ فشل مسح مهام Redis ({e})، نستخدم PostgreSQL.")
        for jid, job in self._pg_jobs().items():
            if job.get("status") in ("pending", "processing"):
                j = dict(job); j["job_id"] = jid
                results.append(j)
        return results

    def cleanup_old_jobs(self, days=7):
        """يحذف المهام المكتملة/الفاشلة الأقدم من days أيام — لا يمس المهام
        النشطة (pending/processing) مهما كان عمرها، حتى لا نفقد إمكانية استئنافها."""
        cutoff = time.time() - (days * 86400)
        if self.use_redis:
            try:
                for key in _redis_client.keys(f"{self.PREFIX}*"):
                    raw = _redis_client.get(key)
                    if raw:
                        job = json.loads(raw)
                        if job.get("status") in ("completed", "failed") and job.get("created_at", 0) < cutoff:
                            _redis_client.delete(key)
            except Exception as e:
                logger.warning(f"⚠️ فشل تنظيف مهام Redis القديمة: {e}")
        pg_jobs = self._pg_jobs()
        stale_ids = [jid for jid, j in pg_jobs.items() if j.get("status") in ("completed", "failed") and j.get("created_at", 0) < cutoff]
        for jid in stale_ids:
            pg_jobs.pop(jid, None)
        if stale_ids:
            sync_db()

    def backend_name(self):
        return "Upstash Redis" if self.use_redis else "PostgreSQL (fallback)"


job_manager = RedisJobManager()


# ── محرك التقطيع والدمج (PDF فقط حاليًا — الصيغة الوحيدة اللي تدعم تقسيم/دمج
# جزئي حقيقي على مستوى الصفحات عبر pypdf بدون فقدان بيانات) ──────────────────
def split_pdf_to_chunks(input_path, pages_per_chunk=50):
    """يقسّم PDF لأجزاء صغيرة (كل جزء ملف PDF مستقل)، كل جزء يُسجَّل فورًا
    بـ TempFileManager حتى ينضمن حذفه لاحقًا حتى لو صار خطأ بمنتصف العملية."""
    if not _ensure_convert_libs():
        raise RuntimeError(f"مكتبات التحويل غير مثبتة: {_CONVERT_IMPORT_ERROR}")
    reader = pypdf.PdfReader(input_path)
    total_pages = len(reader.pages)
    chunk_paths = []
    for start in range(0, total_pages, pages_per_chunk):
        writer = pypdf.PdfWriter()
        end = min(start + pages_per_chunk, total_pages)
        for page_num in range(start, end):
            writer.add_page(reader.pages[page_num])
        chunk_path = f"/tmp/chunk_{uuid.uuid4().hex[:8]}_{start}_{end}.pdf"
        with open(chunk_path, "wb") as f:
            writer.write(f)
        temp_file_manager.register(chunk_path)
        chunk_paths.append(chunk_path)
        del writer
        gc.collect()
    del reader
    gc.collect()
    return chunk_paths, total_pages

def merge_pdf_chunks(chunk_paths, out_path):
    """يدمج أجزاء PDF مكتملة التحويل لملف نهائي واحد، ثم يحذف كل الأجزاء
    فورًا بعد نجاح الدمج (عبر TempFileManager.safe_delete)."""
    if not _ensure_convert_libs():
        raise RuntimeError(f"مكتبات التحويل غير مثبتة: {_CONVERT_IMPORT_ERROR}")
    writer = pypdf.PdfWriter()
    for cp in chunk_paths:
        reader = pypdf.PdfReader(cp)
        for page in reader.pages:
            writer.add_page(page)
        del reader
    with open(out_path, "wb") as f:
        writer.write(f)
    del writer
    gc.collect()
    for cp in chunk_paths:
        temp_file_manager.safe_delete(cp)

def convert_with_chunking(src_path, src_ext, target_ext, out_path, title="Document", pages_per_chunk=None):
    """نقطة الدخول الذكية للتحويل الكبير: تتحقق أولاً من CHUNKING_ENABLED
    ونوع الملف (PDF فقط مدعوم للتقطيع الفعلي حاليًا). لو التقطيع غير مناسب
    أو معطّل، تستدعي convert_file العادية مباشرة بدون أي تغيير بسلوكها
    الحالي. لو التقطيع مناسب، تنشئ مهمة بـ job_manager، تقسّم الملف، تحوّل
    كل جزء عبر convert_file نفسها (فتستفيد من سلسلة CloudConvert/Convertio/
    ConvertAPI/محلي تلقائيًا)، تراقب الرام بين الأجزاء، وتدمج النتيجة أخيرًا."""
    cfg = db["config"].get("chunking_settings", {})
    enabled = cfg.get("enabled", False)
    chunk_size = pages_per_chunk or cfg.get("pages_per_chunk", 50)
    pause_between = cfg.get("pause_between_chunks", 3)

    if not enabled or src_ext.lower().lstrip(".") != "pdf":
        # التقطيع غير مفعّل أو الصيغة غير مدعومة للتقطيع الفعلي — نسار المسار
        # العادي بدون أي تغيير، بالضبط زي ما كان يشتغل قبل هالنظام بالكامل.
        convert_file(src_path, src_ext, target_ext, out_path, title)
        return None  # None = تم بدون إنشاء job (تحويل مباشر عادي)

    try:
        reader_check = None
        if _ensure_convert_libs():
            reader_check = pypdf.PdfReader(src_path)
        total_pages = len(reader_check.pages) if reader_check else 0
        del reader_check
    except Exception:
        total_pages = 0

    if total_pages <= chunk_size:
        # الملف أصلاً أصغر من حجم الجزء الواحد — ما فيه داعي للتقطيع، نحوّله مباشرة.
        convert_file(src_path, src_ext, target_ext, out_path, title)
        return None

    chunk_paths, total_pages = split_pdf_to_chunks(src_path, chunk_size)
    total_parts = len(chunk_paths)
    job_id = job_manager.create_job(src_path, target_ext, total_parts,
                                     meta={"title": title, "out_path": out_path, "total_pages": total_pages})
    logger.info(f"[chunking] بدأت مهمة {job_id}: {total_parts} جزء ({total_pages} صفحة)")

    converted_chunks = []
    try:
        for i, chunk_path in enumerate(chunk_paths):
            memory_monitor.wait_if_needed()
            chunk_out = f"/tmp/chunkout_{uuid.uuid4().hex[:8]}.{target_ext.lstrip('.')}"
            temp_file_manager.register(chunk_out)
            convert_file(chunk_path, "pdf", target_ext, chunk_out, f"{title} (جزء {i+1}/{total_parts})")
            converted_chunks.append(chunk_out)
            job_manager.update_part(job_id, chunk_out)
            temp_file_manager.safe_delete(chunk_path)
            gc.collect()
            if i < total_parts - 1:
                time.sleep(pause_between)

        if target_ext.lower().lstrip(".") == "pdf":
            merge_pdf_chunks(converted_chunks, out_path)
        else:
            # صيغ غير PDF: ما فيه دمج ثنائي موحّد بسيط وآمن لكل الصيغ، فندمج
            # كنص متسلسل عبر merge_files_to_one الموجودة أصلاً (تدعم docx/epub/txt).
            sections = [(f"جزء {i+1}", extract_text_from_file(cp, target_ext.lstrip("."))) for i, cp in enumerate(converted_chunks)]
            merge_files_to_one(sections, target_ext.lstrip("."), out_path, title)
            for cp in converted_chunks:
                temp_file_manager.safe_delete(cp)

        job_manager.complete_job(job_id, out_path)
        logger.info(f"[chunking] اكتملت مهمة {job_id} ✅")
        return job_id
    except Exception as e:
        job_manager.fail_job(job_id, str(e))
        logger.error(f"[chunking] فشلت مهمة {job_id}: {e}")
        for cp in chunk_paths + converted_chunks:
            temp_file_manager.safe_delete(cp)
        raise

def resume_chunk_job(job_id):
    """يستأنف مهمة تقطيع متوقفة (pending/processing) من حيث توقفت بالضبط —
    يتخطى الأجزاء المكتملة (المسجَّلة بـ parts_outputs) ويكمل الباقي فقط.
    ملاحظة مهمة: يتطلب إن ملف المصدر الأصلي (input_path) ما زال موجودًا على
    القرص؛ لو انحذف (مثلاً بعد إعادة إقلاع نظّفت /tmp)، الاستئناف يفشل برسالة
    واضحة بدل خطأ غامض، والمطور يقدر يرفع الملف من جديد ويبدأ مهمة جديدة."""
    job = job_manager.get_job(job_id)
    if not job:
        return None, "المهمة غير موجودة."
    if job["status"] not in ("pending", "processing", "failed"):
        return None, f"المهمة بحالة '{job['status']}' ولا تحتاج استئناف."
    src_path = job["input_path"]
    if not os.path.exists(src_path):
        return None, "ملف المصدر الأصلي لم يعد موجودًا على القرص، لا يمكن الاستئناف. ابدأ مهمة جديدة."

    meta = job.get("meta", {})
    target_ext = job["target_ext"]
    title = meta.get("title", "Document")
    out_path = meta.get("out_path", f"/tmp/resumed_{job_id}.{target_ext.lstrip('.')}")
    cfg = db["config"].get("chunking_settings", {})
    chunk_size = cfg.get("pages_per_chunk", 50)
    pause_between = cfg.get("pause_between_chunks", 3)

    already_done = job.get("completed_parts", 0)
    converted_chunks = list(job.get("parts_outputs", []))
    # نتأكد إن الأجزاء المحفوظة سابقًا ما زالت موجودة فعليًا على القرص (لو
    # انحذفت بتنظيف /tmp، لازم نعيد إنتاجها من جديد بدل نعتمد على مسار وهمي)
    converted_chunks = [c for c in converted_chunks if os.path.exists(c)]
    already_done = len(converted_chunks)

    try:
        chunk_paths, total_pages = split_pdf_to_chunks(src_path, chunk_size)
        total_parts = len(chunk_paths)
        remaining = chunk_paths[already_done:]
        logger.info(f"[chunking] استئناف مهمة {job_id}: {already_done}/{total_parts} مكتمل سابقًا، متبقي {len(remaining)}")

        for i, chunk_path in enumerate(remaining, start=already_done):
            memory_monitor.wait_if_needed()
            chunk_out = f"/tmp/chunkout_{uuid.uuid4().hex[:8]}.{target_ext.lstrip('.')}"
            temp_file_manager.register(chunk_out)
            convert_file(chunk_path, "pdf", target_ext, chunk_out, f"{title} (جزء {i+1}/{total_parts})")
            converted_chunks.append(chunk_out)
            job_manager.update_part(job_id, chunk_out)
            temp_file_manager.safe_delete(chunk_path)
            gc.collect()
            if i < total_parts - 1:
                time.sleep(pause_between)

        # الأجزاء اللي كانت مكتملة سابقًا (قبل الاستئناف) لسا PDF أصلي مو
        # محوَّل بعد لو كانت من split_pdf_to_chunks الجديدة — نحذفها بأمان
        for cp in chunk_paths[:already_done]:
            temp_file_manager.safe_delete(cp)

        if target_ext.lower().lstrip(".") == "pdf":
            merge_pdf_chunks(converted_chunks, out_path)
        else:
            sections = [(f"جزء {i+1}", extract_text_from_file(cp, target_ext.lstrip("."))) for i, cp in enumerate(converted_chunks)]
            merge_files_to_one(sections, target_ext.lstrip("."), out_path, title)
            for cp in converted_chunks:
                temp_file_manager.safe_delete(cp)

        job_manager.complete_job(job_id, out_path)
        logger.info(f"[chunking] اكتمل استئناف مهمة {job_id} ✅")
        return out_path, None
    except Exception as e:
        job_manager.fail_job(job_id, str(e))
        logger.error(f"[chunking] فشل استئناف مهمة {job_id}: {e}")
        return None, str(e)[:200]

def _chunking_panel_content():
    """يبني نص وأزرار لوحة تحكم نظام التقطيع والاستئناف."""
    cfg = db["config"]["chunking_settings"]
    status = "🟢 مفعّل" if cfg.get("enabled") else "🔴 معطّل"
    incomplete_count = len(job_manager.list_incomplete_jobs())
    text = (f"🧩 *نظام التقطيع والاستئناف*\n\nالحالة: {status}\n"
            f"خلفية التخزين: {job_manager.backend_name()}\n\n"
            f"صفحات/جزء: {cfg.get('pages_per_chunk', 50)}\n"
            f"انتظار بين الأجزاء: {cfg.get('pause_between_chunks', 3)}ث\n"
            f"حد الرام: {cfg.get('max_ram_mb', 200)} MB\n"
            f"الاحتفاظ بالمهام: {cfg.get('job_retention_days', 7)} أيام\n\n"
            f"⏳ مهام متوقفة: {incomplete_count}\n\n"
            f"_يعمل فقط على ملفات PDF أكبر من حجم الجزء الواحد. أي ملف أصغر أو "
            f"بصيغة أخرى يُحوَّل مباشرة بالطريقة العادية بدون تقطيع._")
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if cfg.get('enabled') else '🟢 تفعيل'}", callback_data="chunk_toggle"))
    m.add(types.InlineKeyboardButton("📄 صفحات/جزء", callback_data="chunk_set_pages"),
          types.InlineKeyboardButton("⏱️ الانتظار", callback_data="chunk_set_pause"))
    m.add(types.InlineKeyboardButton("🧠 حد الرام", callback_data="chunk_set_ram"))
    m.add(types.InlineKeyboardButton("📊 إحصائيات وتنظيف", callback_data="chunk_view_stats"))
    if incomplete_count:
        m.add(types.InlineKeyboardButton(f"⏳ المهام المتوقفة ({incomplete_count})", callback_data="chunk_view_incomplete"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_services_status"))
    return text, m

def send_chapter_file(cid, nid, chapter_num, fmt):
    novel = get_novel(nid)
    if not novel: return False
    ch = novel.get("chapters", {}).get(str(chapter_num))
    if not ch or not ch.get(fmt):
        bot.send_message(cid, f"❌ لا يوجد ملف بصيغة {fmt.upper()} لهذا الفصل بعد.")
        return False
    try:
        bot.send_document(cid, ch[fmt], caption=f"📖 *{md_safe(novel['title'])}* — الفصل {chapter_num} ({fmt.upper()})", parse_mode="Markdown", protect_content=should_protect_content(cid))
        track_event("novel_download", nid)
        return True
    except Exception as e:
        bot.send_message(cid, f"❌ خطأ بإرسال الملف: {e}")
        return False

def send_merged_file(cid, nid, group, fmt):
    novel = get_novel(nid)
    if not novel: return False
    mg = novel.get("merged", {}).get(group)
    if not mg or not mg.get(fmt):
        bot.send_message(cid, f"❌ لا يوجد ملف مدمج بصيغة {fmt.upper()} لهذي المجموعة بعد.")
        return False
    try:
        bot.send_document(cid, mg[fmt], caption=f"📦 *{md_safe(novel['title'])}* — الفصول {group} ({fmt.upper()})", parse_mode="Markdown", protect_content=should_protect_content(cid))
        track_event("novel_download", nid)
        return True
    except Exception as e:
        bot.send_message(cid, f"❌ خطأ بإرسال الملف: {e}")
        return False

# ==============================================================================
# 8. ADMIN KEYBOARDS
# ==============================================================================
def stats_menu_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton("📊 نظرة عامة", callback_data="stats_overview"))
    m.add(types.InlineKeyboardButton("📚 أكثر الروايات طلبًا", callback_data="stats_top_novels"))
    m.add(types.InlineKeyboardButton("🎨 أكثر أعمال المانجا طلبًا", callback_data="stats_top_manga"))
    m.add(types.InlineKeyboardButton("🎬 أكثر المسلسلات طلبًا", callback_data="stats_top_series"))
    m.add(types.InlineKeyboardButton("📄 أكثر الملفات طلبًا", callback_data="stats_top_items"))
    m.add(types.InlineKeyboardButton("🔍 أكثر كلمات البحث", callback_data="stats_top_searches"))
    m.add(types.InlineKeyboardButton("📈 نمو المستخدمين", callback_data="stats_growth"))
    visible = db["config"].get("stats_visible_to_users", False)
    m.add(types.InlineKeyboardButton(f"👁️ إظهار للمستخدمين: {'✅ مفعّل' if visible else '❌ معطّل'}", callback_data="stats_toggle_visible"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

def quick_toggle_kb():
    """سويتش سريع لتعطيل/تفعيل كل الميزات (أزرار رئيسية + ميزات مخصصة) بضغطة وحدة،
    تعطيل حقيقي يمنع الميزة من الشغل نهائيًا (حتى لو استُدعيت بدمج ميزة أو رابط مباشر)،
    مو مجرد إخفاء عن القائمة."""
    m = types.InlineKeyboardMarkup(row_width=1)
    for btn in sorted(db["config"].get("menu_buttons", []), key=lambda x: x.get("order", 99)):
        vis = "🟢" if btn.get("active", True) else "🔴"
        m.add(types.InlineKeyboardButton(f"{vis} {btn['label']}", callback_data=f"qtoggle_btn_{btn['id']}"))
    for fid, f in db["config"].get("custom_features", {}).items():
        vis = "🟢" if f.get("active", True) else "🔴"
        m.add(types.InlineKeyboardButton(f"{vis} ✨ {f['label']}", callback_data=f"qtoggle_feat_{fid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

def admin_main_kb():
    m = types.InlineKeyboardMarkup(row_width=2)
    m.add(
        types.InlineKeyboardButton("🎨 تخصيص القائمة",    callback_data="adm_menu_editor"),
        types.InlineKeyboardButton("✨ الميزات المخصصة",  callback_data="adm_features")
    )
    m.add(
        types.InlineKeyboardButton("📢 إذاعة",             callback_data="adm_broadcast"),
        types.InlineKeyboardButton("📊 إحصائيات",          callback_data="adm_stats")
    )
    m.add(
        types.InlineKeyboardButton("🔌 حالة الخدمات الخارجية", callback_data="adm_services_status")
    )
    m.add(
        types.InlineKeyboardButton("👤 المستخدمون",        callback_data="adm_users"),
        types.InlineKeyboardButton("🎟️ أكواد النقاط",     callback_data="adm_codes")
    )
    m.add(
        types.InlineKeyboardButton("📢 قنوات إجبارية",    callback_data="adm_channels"),
        types.InlineKeyboardButton("💰 مالية",              callback_data="adm_money")
    )
    m.add(types.InlineKeyboardButton("🔌 تشغيل/تعطيل الميزات (سويتش سريع)", callback_data="adm_quick_toggle"))
    m.add(types.InlineKeyboardButton("📦 حدود التنزيل المدمج", callback_data="adm_merge_limits"))
    m.add(types.InlineKeyboardButton("🔄 تحويل صيغة ملف", callback_data="adm_convert_file"))
    m.add(types.InlineKeyboardButton("🔒 حماية المحتوى", callback_data="adm_content_protection"))
    m.add(types.InlineKeyboardButton("📣 قناة إعلان التحديثات", callback_data="adm_announce_channel"))
    m.add(types.InlineKeyboardButton("💾 نسخ احتياطية واسترجاع", callback_data="adm_snapshots"))
    m.add(
        types.InlineKeyboardButton("🔧 الصيانة",           callback_data="adm_maintenance"),
        types.InlineKeyboardButton("✏️ رسالة ترحيب",      callback_data="adm_welcome")
    )
    m.add(types.InlineKeyboardButton("🏷️ إدارة التصنيفات", callback_data="adm_tags"))
    m.add(types.InlineKeyboardButton("📂 إدارة أقسام القائمة", callback_data="adm_list_sections"))
    m.add(types.InlineKeyboardButton("🎛️ تفعيل/تعطيل أنواع المحتوى", callback_data="adm_content_types"))
    m.add(types.InlineKeyboardButton("🗄️ قنوات الأرشيف والنسخ", callback_data="adm_archive_channels"))
    m.add(types.InlineKeyboardButton("🆕 استقبال الأعضاء الجدد", callback_data="adm_new_members"))
    m.add(types.InlineKeyboardButton("📰 نظام الأخبار", callback_data="adm_news"))
    m.add(types.InlineKeyboardButton("🕷️ سحب الفصول التلقائي", callback_data="adm_scraper"))
    m.add(types.InlineKeyboardButton("🌐 صلاحيات الترجمة", callback_data="adm_translate_access"))
    m.add(types.InlineKeyboardButton("💎 إدارة الاشتراكات", callback_data="adm_subs"))
    m.add(types.InlineKeyboardButton("👁️ تسمية زر شاهد الآن", callback_data="adm_watch_label"))
    m.add(types.InlineKeyboardButton("🤖 إنشاء ميزة بالذكاء الاصطناعي", callback_data="adm_ai_feature"))
    m.add(types.InlineKeyboardButton("🧠 تعديل منطق الميزات بالذكاء", callback_data="adm_ai_logic"))
    m.add(types.InlineKeyboardButton("🎮 المسابقات", callback_data="adm_contests"))
    m.add(types.InlineKeyboardButton("🤖 المحادثة الذكية", callback_data="adm_ai_talk"))
    m.add(types.InlineKeyboardButton("🔗 الدعوة التلقائية (موافقة)", callback_data="adm_auto_invite"))
    m.add(types.InlineKeyboardButton("🏠 خروج", callback_data="go_home"))
    return m

def menu_editor_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    buttons = sorted(db["config"].get("menu_buttons", []), key=lambda x: x.get("order", 99))
    for btn in buttons:
        vis = "👁️" if btn.get("visible", True) else "🙈"
        m.add(types.InlineKeyboardButton(
            f"{vis} {btn['label']} (ترتيب:{btn.get('order',0)})",
            callback_data=f"mebtn_{btn['id']}"
        ))
    m.add(types.InlineKeyboardButton("➕ إضافة زر جديد", callback_data="me_add_btn"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

def btn_editor_kb(btn_id):
    m = types.InlineKeyboardMarkup(row_width=2)
    btn = next((b for b in db["config"]["menu_buttons"] if b["id"] == btn_id), None)
    if not btn: return m
    vis_label = "🙈 إخفاء" if btn.get("visible", True) else "👁️ إظهار"
    en_label = "⛔ تعطيل الميزة بالكامل" if btn.get("active", True) else "✅ تفعيل الميزة"
    m.add(
        types.InlineKeyboardButton("✏️ تغيير الاسم",    callback_data=f"me_rename_{btn_id}"),
        types.InlineKeyboardButton("🔢 تغيير الترتيب",  callback_data=f"me_reorder_{btn_id}")
    )
    m.add(
        types.InlineKeyboardButton(vis_label,            callback_data=f"me_toggle_{btn_id}"),
        types.InlineKeyboardButton("🔗 تغيير الأمر (Action)", callback_data=f"me_action_{btn_id}")
    )
    m.add(types.InlineKeyboardButton(en_label, callback_data=f"me_enable_toggle_{btn_id}"))
    m.add(types.InlineKeyboardButton("🗑️ حذف الزر",       callback_data=f"me_delete_{btn_id}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_menu_editor"))
    return m

def features_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    for fid, f in db["config"].get("custom_features", {}).items():
        vis = "👁️" if f.get("visible", True) else "🙈"
        act = "🟢" if f.get("active", True) else "🔴"
        m.add(types.InlineKeyboardButton(
            f"{act}{vis} {f['label']}", callback_data=f"feat_edit_{fid}"))
    m.add(types.InlineKeyboardButton("🤖 إنشاء ميزة بالذكاء الاصطناعي", callback_data="adm_ai_feature"))
    m.add(types.InlineKeyboardButton("➕ إنشاء ميزة يدوياً", callback_data="feat_manual"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

def feat_editor_kb(fid):
    m = types.InlineKeyboardMarkup(row_width=2)
    f = db["config"]["custom_features"].get(fid, {})
    vis_label = "🙈 إخفاء" if f.get("visible", True) else "👁️ إظهار"
    act_label = "🔴 تعطيل" if f.get("active", True) else "🟢 تفعيل"
    dev_label = "🔓 عامة" if not f.get("dev_only", False) else "🔒 للمطور فقط"
    m.add(
        types.InlineKeyboardButton("✏️ تغيير الاسم",    callback_data=f"feat_rename_{fid}"),
        types.InlineKeyboardButton("💬 تغيير الرد",      callback_data=f"feat_resp_{fid}")
    )
    m.add(
        types.InlineKeyboardButton(vis_label,            callback_data=f"feat_vis_{fid}"),
        types.InlineKeyboardButton(act_label,            callback_data=f"feat_act_{fid}")
    )
    m.add(
        types.InlineKeyboardButton(dev_label,            callback_data=f"feat_priv_{fid}"),
        types.InlineKeyboardButton("🗑️ حذف الميزة",       callback_data=f"feat_del_{fid}")
    )
    m.add(
        types.InlineKeyboardButton("📋 الأزرار الفرعية", callback_data=f"feat_subs_{fid}")
    )
    m.add(types.InlineKeyboardButton("🤖 عدّل عبر الذكاء الاصطناعي", callback_data=f"feat_ai_edit_{fid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_features"))
    return m

def feat_sub_buttons_kb(parent_fid):
    m = types.InlineKeyboardMarkup(row_width=1)
    parent_feature = db["config"]["custom_features"].get(parent_fid)
    if not parent_feature: return m

    for sub_btn_data in parent_feature.get("sub_buttons", []):
        # Assuming sub_btn_data contains 'id' and 'label'
        sb_label = sub_btn_data['label']; sb_id = sub_btn_data['id']
        m.add(types.InlineKeyboardButton(f"➖ {sb_label}", callback_data=f"feat_del_sub_{parent_fid}_{sb_id}"))
    
    m.add(types.InlineKeyboardButton("➕ إضافة زر فرعي", callback_data=f"feat_add_sub_{parent_fid}"))
    m.add(types.InlineKeyboardButton("🔗 دمج ميزة أخرى هنا", callback_data=f"feat_nest_{parent_fid}"))
    m.add(types.InlineKeyboardButton("📐 شكل العرض", callback_data=f"feat_layout_{parent_fid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"feat_edit_{parent_fid}"))
    return m

def ai_logic_features_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    for btn in db["config"].get("menu_buttons", []):
        m.add(types.InlineKeyboardButton(f"⚙️ {btn['label']}", callback_data=f"ailogic_{btn['id']}"))
    for fid, f in db["config"].get("custom_features", {}).items():
        m.add(types.InlineKeyboardButton(f"✨ {f['label']}", callback_data=f"ailogic_feat_{fid}"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

# ── Personal List System (القائمة: مفضلة/مشاهدة حالياً/تمت المشاهدة/أرغب بمشاهدتها) ──

def list_sections():
    return db["config"].get("list_sections", [])

def get_section_name(sec_id):
    for s in list_sections():
        if s["id"] == sec_id: return s["name"]
    return sec_id

def user_list_items(ustr, sec_id):
    return db["users"].get(ustr, {}).get("list_items", {}).get(sec_id, [])

def add_to_list(ustr, sec_id, entry_type, entry_id):
    u = db["users"][ustr]
    u.setdefault("list_items", {}).setdefault(sec_id, [])
    if not any(e["type"] == entry_type and e["id"] == entry_id for e in u["list_items"][sec_id]):
        u["list_items"][sec_id].append({"type": entry_type, "id": entry_id})
    sync_db()

def remove_from_list(ustr, sec_id, entry_type, entry_id):
    u = db["users"][ustr]
    items = u.get("list_items", {}).get(sec_id, [])
    u["list_items"][sec_id] = [e for e in items if not (e["type"] == entry_type and e["id"] == entry_id)]
    sync_db()

def list_sections_kb():
    """يعرض أقسام القائمة (المفضلة، تمت المشاهدة...) وعدد العناصر بكل قسم."""
    m = types.InlineKeyboardMarkup(row_width=1)
    for s in list_sections():
        m.add(types.InlineKeyboardButton(s["name"], callback_data=f"listsec_{s['id']}"))
    m.add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
    return m

def list_section_entries_kb(ustr, sec_id):
    """يعرض محتويات قسم معين، مع رقم آخر فصل/حلقة انضافت وآخر تقدّم وصله المستخدم،
    وزر إزالة سريع لكل عنصر. يدعم روايات، مانجا، مسلسلات، وملفات عادية."""
    marker = db["config"].get("progress_marker", "🔴")
    m = types.InlineKeyboardMarkup(row_width=1)
    entries = user_list_items(ustr, sec_id)
    for e in entries:
        etype, eid = e["type"], e["id"]
        if etype == "novel":
            nv = get_novel(eid)
            if not nv: continue
            progress = get_user_novel_progress(ustr, eid)
            label = f"📚 {nv['title']}"
            if progress: label += f"  {marker}{progress}"
            m.add(
                types.InlineKeyboardButton(label, callback_data=f"open_novel_{eid}"),
                types.InlineKeyboardButton("❌", callback_data=f"removefromlist_{sec_id}_novel_{eid}")
            )
        elif etype == "manga":
            mv = get_manga(eid)
            if not mv: continue
            progress = get_user_manga_progress(ustr, eid)
            label = f"🎨 {mv['title']}"
            if progress: label += f"  {marker}{progress}"
            m.add(
                types.InlineKeyboardButton(label, callback_data=f"open_manga_{eid}"),
                types.InlineKeyboardButton("❌", callback_data=f"removefromlist_{sec_id}_manga_{eid}")
            )
        elif etype == "series":
            sv = get_series(eid)
            if not sv: continue
            progress = get_user_series_progress(ustr, eid)
            label = f"🎬 {sv['title']}"
            if progress: label += f"  {marker}{progress}"
            m.add(
                types.InlineKeyboardButton(label, callback_data=f"open_series_{eid}"),
                types.InlineKeyboardButton("❌", callback_data=f"removefromlist_{sec_id}_series_{eid}")
            )
        else:
            item = db["items"].get(eid)
            if item:
                m.add(
                    types.InlineKeyboardButton(f"📖 {item['title']}", callback_data=f"item_{eid}"),
                    types.InlineKeyboardButton("❌", callback_data=f"removefromlist_{sec_id}_item_{eid}")
                )
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="view_favs"))
    return m

def pick_section_kb(entry_type, entry_id):
    """يعرض للمستخدم الأقسام ليختار وين يضيف هذا العمل (رواية/مانجا/مسلسل/ملف) بقائمته."""
    m = types.InlineKeyboardMarkup(row_width=1)
    for s in list_sections():
        m.add(types.InlineKeyboardButton(s["name"], callback_data=f"addtolist_{s['id']}_{entry_type}_{entry_id}"))
    back_map = {"novel": f"open_novel_{entry_id}", "manga": f"open_manga_{entry_id}",
                "series": f"open_series_{entry_id}", "item": f"item_{entry_id}"}
    back_cb = back_map.get(entry_type, f"item_{entry_id}")
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=back_cb))
    return m

def admin_list_sections_kb():
    m = types.InlineKeyboardMarkup(row_width=1)
    for s in list_sections():
        m.add(types.InlineKeyboardButton(f"✏️ {s['name']}", callback_data=f"listsec_rename_{s['id']}"),)
    m.add(types.InlineKeyboardButton("➕ إضافة قسم جديد", callback_data="listsec_add"))
    m.add(types.InlineKeyboardButton(f"🏷️ تغيير علامة التقدم (حالياً: {db['config'].get('progress_marker','🔴')})", callback_data="listsec_marker"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

def open_feature(uid, cid, mid, fid, is_photo_msg=False):
    """يفتح ميزة (رئيسية أو مخصصة): يعرض نصها + أزرارها الفرعية بشكل العرض المختار.
    إذا كان أحد الأزرار الفرعية ميزة مدمجة (nested)، الضغط عليه يشغّل منطق الميزة
    الأصلية بالكامل (فعليًا، لا مجرد شكل) عبر استدعاء open_feature عليها بدورها."""
    def edit(text, reply_markup=None):
        try:
            if is_photo_msg:
                bot.edit_message_caption(caption=text, chat_id=cid, message_id=mid, reply_markup=reply_markup, parse_mode="Markdown")
            else:
                bot.edit_message_text(text, cid, mid, reply_markup=reply_markup, parse_mode="Markdown")
        except Exception as e:
            err = str(e)
            if "message is not modified" not in err:
                logger.warning(f"⚠️ open_feature edit() failed: {err}")
                try: bot.send_message(cid, text, reply_markup=reply_markup, parse_mode="Markdown")
                except Exception as e2: logger.error(f"❌ fallback send_message failed too: {e2}")

    ftype, fobj = get_feature_owner_type(fid)
    if not fobj:
        edit("❌ الميزة غير موجودة.", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home")))
        return
    if not fobj.get("active", True) and not is_admin(uid):
        edit(f"⚠️ ميزة «{fobj.get('label','')}» معطّلة حاليًا.", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home")))
        return

    label = fobj.get("label", "")
    body_text = fobj.get("response_text", f"*{label}*") if ftype == "custom" else f"*{label}*"
    sub_buttons = fobj.get("sub_buttons", [])
    layout = fobj.get("sub_layout", "grid2")

    buttons = []
    for sb in sub_buttons:
        if sb.get("nested_fid"):
            # زر يفتح ميزة أخرى مدمجة فعليًا
            buttons.append(types.InlineKeyboardButton(sb["label"], callback_data=f"open_feature_{sb['nested_fid']}"))
        elif sb.get("response_text") is not None:
            buttons.append(types.InlineKeyboardButton(sb["label"], callback_data=f"subfeat_{fid}_{sb['id']}"))
        elif sb.get("action"):
            buttons.append(types.InlineKeyboardButton(sb["label"], callback_data=sb["action"]))
        else:
            buttons.append(types.InlineKeyboardButton(sb["label"], callback_data=f"noop_{sb.get('id','')}"))

    m = apply_layout(None, buttons, layout)
    if is_admin(uid):
        m.add(types.InlineKeyboardButton("➕ زر فرعي", callback_data=f"feat_add_sub_{fid}"),
              types.InlineKeyboardButton("🔗 دمج ميزة", callback_data=f"feat_nest_{fid}"))
        m.add(types.InlineKeyboardButton("📐 شكل العرض", callback_data=f"feat_layout_{fid}"))
    pid_back = fobj.get("parent_id")
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"open_feature_{pid_back}" if pid_back else "go_home"))
    m.add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
    edit(body_text, m)

def new_member_panel_text():
    nm = db["config"]["new_member_settings"]
    status = "🟢 مفعّلة" if nm.get("enabled") else "🔴 معطّلة"
    groups = nm.get("watched_groups", [])
    dm_status = "🟢 مفعّلة" if nm.get("dm_enabled", True) else "🔴 معطّلة"
    ch = nm.get("auto_add_channel")
    return (f"🆕 *استقبال الأعضاء الجدد*\n\n"
            f"لما عضو جديد ينضم لأي مجموعة مراقَبة، البوت يرسله رسالة خاصة "
            f"(ويقدر يضيف رابط قناة للانضمام التلقائي).\n\n"
            f"الحالة العامة: {status}\n"
            f"عدد المجموعات المراقَبة: {len(groups)}\n"
            f"الرسالة الخاصة: {dm_status}\n"
            f"قناة الإضافة التلقائية: `{ch or 'غير مربوطة'}`\n\n"
            f"⚠️ ملاحظة: البوت لازم يكون *مشرف* بالمجموعة عشان يشوف الأعضاء الجدد، "
            f"وبالقناة عشان يسوي رابط دعوة تلقائي.")

def new_member_panel_kb():
    nm = db["config"]["new_member_settings"]
    ch = nm.get("auto_add_channel")
    m = types.InlineKeyboardMarkup(row_width=1)
    m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if nm.get('enabled') else '🟢 تفعيل'} الميزة", callback_data="nm_toggle_enabled"))
    m.add(types.InlineKeyboardButton("👥 إدارة المجموعات المراقَبة", callback_data="nm_manage_groups"))
    m.add(types.InlineKeyboardButton("✏️ تعديل نص الرسالة الخاصة", callback_data="nm_edit_dm_text"))
    m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if nm.get('dm_enabled', True) else '🟢 تفعيل'} الرسالة الخاصة", callback_data="nm_toggle_dm"))
    m.add(types.InlineKeyboardButton("🔘 زر اختياري بالرسالة", callback_data="nm_edit_dm_button"))
    m.add(types.InlineKeyboardButton("📡 ربط قناة الإضافة التلقائية", callback_data="nm_set_channel"))
    if ch:
        m.add(types.InlineKeyboardButton("❌ إلغاء ربط القناة", callback_data="nm_unset_channel"))
    m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
    return m

# ==============================================================================
# 9. CALLBACK HANDLER
# ==============================================================================
@bot.callback_query_handler(func=lambda call: True)
def on_callback(call):
    """غلاف حماية عام: أي خطأ غير متوقع بأي مكان بمنطق الأزرار ما عاد يصير سكوت
    تام بدون رد — المستخدم يشوف رسالة خطأ واضحة، والمطور يشوفها باللوق."""
    try:
        _on_callback_impl(call)
    except Exception as e:
        logger.error(f"❌ خطأ غير متوقع بمعالجة الزر (data={getattr(call, 'data', '?')}): {e}", exc_info=True)
        try:
            bot.answer_callback_query(call.id, "⚠️ صار خطأ غير متوقع، حاول مرة ثانية.", show_alert=True)
        except Exception:
            pass
        try:
            bot.send_message(call.message.chat.id, f"❌ صار خطأ غير متوقع أثناء تنفيذ العملية:\n`{str(e)[:300]}`", parse_mode="Markdown")
        except Exception:
            pass

def _on_callback_impl(call):
    uid = call.from_user.id; cid = call.message.chat.id; mid = call.message.message_id; ustr = str(uid); data = call.data
    is_photo_msg = bool(getattr(call.message, "photo", None))
    def edit(text, reply_markup=None):
        try:
            if is_photo_msg:
                # لا يمكن تحويل رسالة صورة (بوستر الرواية مثلاً) إلى رسالة نصية؛
                # نعدّل الكابشن بدل النص، وإذا فشل نرسل رسالة جديدة بدل ما تنكتم الحركة كليًا.
                bot.edit_message_caption(caption=text, chat_id=cid, message_id=mid, reply_markup=reply_markup, parse_mode="Markdown")
            else:
                bot.edit_message_text(text, cid, mid, reply_markup=reply_markup, parse_mode="Markdown")
        except Exception as e:
            err = str(e)
            if "message is not modified" in err:
                pass  # نفس المحتوى، ما يحتاج تحديث فعلي — هذا مو خطأ حقيقي
            else:
                logger.warning(f"⚠️ edit() failed (cid={cid}, mid={mid}, photo={is_photo_msg}): {err}")
                try:
                    bot.send_message(cid, text, reply_markup=reply_markup, parse_mode="Markdown")
                except Exception as e2:
                    logger.error(f"❌ fallback send_message failed too: {e2}")

    # حارس عام: لو الأمر (action) اللي بنشغّله يعود لزر رئيسي معطّل بالكامل من المطور،
    # نمنعه هنا قبل أي منطق تاني — بغض النظر جاء الاستدعاء من القائمة، دمج ميزة، أو استدعاء مباشر.
    _blocked_btn = next((b for b in db["config"]["menu_buttons"] if b.get("action") == data and not b.get("active", True)), None)
    if _blocked_btn and not is_admin(uid):
        bot.answer_callback_query(call.id, f"⚠️ ميزة «{_blocked_btn['label']}» معطّلة حاليًا.", show_alert=True)
        return

    if data == "do_start":
        ensure_user(call.message)
        db["users"][ustr]["seen_start_prompt"] = True; sync_db()
        unjoined = check_channels(uid)
        if unjoined:
            send_channels_msg(cid, unjoined)
        else:
            send_welcome_and_menu(cid, uid)
        return

    # ===== ميزة الدعوة التلقائية عند ضغط أي زر (باستثناء أزرار الإدارة والتحكم بالميزة نفسها) =====
    if db["config"]["auto_invite"].get("enabled") and "callback" in db["config"]["auto_invite"].get("trigger_on", []):
        if not data.startswith("noop_") and not data.startswith("admin_") and not data.startswith("adm_") and not data.startswith("invite_"):
            send_invite_to_user(uid, "ضغط زر في القناة/المجموعة")

    if data == "go_home": edit("🏠 *القائمة الرئيسية* — اختر من الأزرار أدناه 👇", build_main_menu(uid))
    elif data.startswith("noop_"):
        bot.answer_callback_query(call.id)
    elif data.startswith("open_feature_"):
        fid = data[13:]
        ftype_check, fobj_check = get_feature_owner_type(fid)
        real_action = fobj_check.get("action") if (ftype_check == "main" and fobj_check) else None
        if real_action:
            # الميزة المفتوحة هي زر رئيسي له منطق حقيقي (بحث، مفضلة، اشتراك...الخ)
            # حتى لو كانت مدمجة داخل ميزة ثانية، لازم تشتغل بمنطقها الأصلي كامل
            # بدل ما تُعرض كصفحة فارغة بدون وظيفة.
            call.data = real_action
            on_callback(call)
        else:
            open_feature(uid, cid, mid, fid, is_photo_msg=is_photo_msg)
    elif data.startswith("subfeat_"):
        parts = data.split("_", 2); parent_fid = parts[1]; sub_id = parts[2]
        parent_obj = db["config"]["custom_features"].get(parent_fid) or next((b for b in db["config"]["menu_buttons"] if b["id"] == parent_fid), None)
        sub = next((s for s in (parent_obj.get("sub_buttons", []) if parent_obj else []) if s["id"] == sub_id), None)
        if sub and sub.get("response_text"):
            m = types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"open_feature_{parent_fid}"), types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
            edit(sub["response_text"], m)
        else:
            bot.answer_callback_query(call.id, "❌ لا يوجد محتوى لهذا الزر.")
    elif data == "check_joined":
        chs = check_channels(uid)
        if chs: bot.answer_callback_query(call.id, "❌ ما زلت غير مشترك.", show_alert=True)
        else:
            pending = db["pending_actions"].pop(ustr, None)
            sync_db()
            if pending and pending.get("action") == "deep_link_after_join":
                payload = pending.get("payload", "")
                if payload.startswith("novel_") and get_novel(payload[6:]):
                    open_novel(uid, cid, mid, payload[6:], as_new_message=True); return
                if payload.startswith("manga_") and get_manga(payload[6:]):
                    open_manga(uid, cid, mid, payload[6:], as_new_message=True); return
                if payload.startswith("series_") and get_series(payload[7:]):
                    open_series(uid, cid, mid, payload[7:], as_new_message=True); return
            edit("✅ تم التحقق من اشتراكك. إليك القائمة الرئيسية 👇", build_main_menu(uid))
    elif data.startswith("nav_"):
        cat_id = data[4:]; cat = db["categories"].get(cat_id)
        if cat: edit(f"📁 *قسم: {cat['name']}*", build_category_menu(cat_id, uid))
    elif data.startswith("item_"):
        iid = data[5:]; item = db["items"].get(iid)
        if not item: return
        ensure_user(call.message)
        last_view = get_user_item_view(ustr, iid)  # نجيب آخر زيارة سابقة قبل ما نحدّثها للزيارة الحالية
        set_user_item_view(ustr, iid)
        track_event("item_view", iid)
        view_line = f"\n\n👁️ آخر مرة شاهدت/حمّلت هذا الملف: {last_view.split('.')[0]}" if last_view else ""
        m = types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("➕ نقل إلى القائمة", callback_data=f"addlist_item_{iid}"))
        if is_admin(uid): m.add(types.InlineKeyboardButton("🗑️ حذف", callback_data=f"del_item_{iid}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"nav_{item.get('category','root')}"))
        try:
            ftype = item.get("type", "document"); caption = f"*{item['title']}*{view_line}"
            protect = should_protect_content(uid)
            if ftype == "photo": bot.send_photo(cid, item["file_id"], caption=caption, parse_mode="Markdown", reply_markup=m, protect_content=protect)
            elif ftype == "video": bot.send_video(cid, item["file_id"], caption=caption, parse_mode="Markdown", reply_markup=m, protect_content=protect)
            elif ftype == "audio": bot.send_audio(cid, item["file_id"], caption=caption, parse_mode="Markdown", reply_markup=m, protect_content=protect)
            else: bot.send_document(cid, item["file_id"], caption=caption, parse_mode="Markdown", reply_markup=m, protect_content=protect)
            track_event("item_download", iid)
        except Exception as e: bot.send_message(cid, f"❌ خطأ: {e}")

    elif data.startswith("addtolist_"):
        # صيغة: addtolist_{sec_id}_{type}_{id}
        rest = data[10:]
        parts = rest.split("_", 2)
        sec_id, entry_type, entry_id = parts[0], parts[1], parts[2]
        add_to_list(ustr, sec_id, entry_type, entry_id)
        bot.answer_callback_query(call.id, f"✅ أُضيف إلى «{get_section_name(sec_id)}».", show_alert=True)
    elif data.startswith("removefromlist_"):
        rest = data[15:]
        parts = rest.split("_", 2)
        sec_id, entry_type, entry_id = parts[0], parts[1], parts[2]
        remove_from_list(ustr, sec_id, entry_type, entry_id)
        bot.answer_callback_query(call.id, "❌ أُزيل من القائمة.")
        edit(f"📂 *{get_section_name(sec_id)}:*", list_section_entries_kb(ustr, sec_id))
    elif data.startswith("addlist_"):
        # صيغة: addlist_{type}_{id} — يعرض للمستخدم اختيار القسم
        rest = data[8:]
        entry_type, entry_id = rest.split("_", 1)
        edit("📂 *اختر القسم اللي تبي تضيف له هذا العمل:*", pick_section_kb(entry_type, entry_id))
    elif data == "search_filter":
        db["pending_actions"][ustr] = {"action": "search_tag_select", "selected": []}
        sync_db()
        edit("🔍 *اختر تصنيف واحد أو أكثر (اضغط لتحديد/إلغاء)، ثم اضغط عرض النتائج — أو ابحث بالاسم مباشرة:*",
             build_tag_search_menu(selected=[], mode="search"))
    elif data == "search_by_name":
        db["pending_actions"][ustr] = {"action": "search_by_name"}; sync_db()
        bot.send_message(cid, "🔎 أرسل اسم الرواية أو الملف (ولو جزء منه)، وما يشترط تعرف تصنيفه:")
    elif data.startswith("searchtag_"):
        try:
            t_idx = int(data[10:])
            tags_list = db["config"].get("tags", [])
            tag = _tag_name(tags_list[t_idx]) if 0 <= t_idx < len(tags_list) else ""
        except ValueError:
            tag = data[10:]  # توافق مع بيانات قديمة كانت تحمل اسم التصنيف مباشرة
        pending = db["pending_actions"].get(ustr) or {}
        action = pending.get("action")
        if action in ("additem_tag", "addnovel_tag"):
            # وضع الإضافة: اختيار متعدد وغير محدود — الزر يبدّل التحديد (toggle) بدل الانتقال المباشر
            sel = list(pending.get("selected", []))
            if tag in sel: sel.remove(tag)
            else: sel.append(tag)
            pending["selected"] = sel
            db["pending_actions"][ustr] = pending; sync_db()
            edit("🏷️ *اختر تصنيف واحد أو أكثر (بدون حد أقصى)، ثم اضغط متابعة:*",
                 build_tag_search_menu(selected=sel, mode="assign"))
        elif action == "search_tag_select":
            sel = list(pending.get("selected", []))
            if tag in sel: sel.remove(tag)
            else: sel.append(tag)
            pending["selected"] = sel
            db["pending_actions"][ustr] = pending; sync_db()
            edit("🔍 *اختر تصنيف واحد أو أكثر (اضغط لتحديد/إلغاء)، ثم اضغط عرض النتائج — أو ابحث بالاسم مباشرة:*",
                 build_tag_search_menu(selected=sel, mode="search"))
        else:
            # احتياطي: تصرف قديم (تصنيف واحد مباشر) لأي حالة غير متوقعة
            edit(f"🔍 *نتائج تصنيف {tag}:*", types.InlineKeyboardMarkup().add(
                types.InlineKeyboardButton("⬅️ رجوع", callback_data="search_filter")))
    elif data == "confirm_tags_assign":
        pending = db["pending_actions"].get(ustr) or {}
        sel = pending.get("selected", [])
        if not sel:
            bot.answer_callback_query(call.id, "❌ اختر تصنيف واحد على الأقل.", show_alert=True)
        elif pending.get("action") == "additem_tag":
            db["pending_actions"][ustr] = {"action": "additem_title", "cat": pending["cat"], "tags": sel}
            sync_db(); bot.send_message(cid, "📖 أرسل اسم الملف:")
        elif pending.get("action") == "addnovel_tag":
            db["pending_actions"][ustr] = {"action": "addnovel_poster", "cat": pending["cat"], "title": pending["title"], "story": pending["story"], "tags": sel}
            sync_db(); bot.send_message(cid, "🖼️ أرسل صورة البوستر الآن:")
    elif data == "run_tag_search":
        pending = db["pending_actions"].get(ustr) or {}
        sel = pending.get("selected", [])
        if not sel:
            bot.answer_callback_query(call.id, "❌ اختر تصنيف واحد على الأقل.", show_alert=True)
        else:
            track_event("search", "، ".join(sel))
            # مطابقة OR: أي عنصر فيه تصنيف واحد على الأقل من المختارين يظهر بالنتائج
            item_results = [iid for iid, item in db["items"].items() if set(_item_tags(item)) & set(sel)]
            novel_results = [nid for nid, nv in db["novels"].items() if set(_item_tags(nv)) & set(sel)]
            manga_results = [mid for mid, mv in db["manga"].items() if set(_item_tags(mv)) & set(sel)]
            series_results = [sid for sid, sv in db["series"].items() if set(_item_tags(sv)) & set(sel)]
            m = types.InlineKeyboardMarkup()
            # لا نعرض اسم التصنيف بتاتًا بنتائج البحث — فقط عناوين الأعمال نفسها
            for iid in item_results[:15]:
                m.add(types.InlineKeyboardButton(f"📖 {db['items'][iid]['title']}", callback_data=f"item_{iid}"))
            for nid in novel_results[:15]:
                m.add(types.InlineKeyboardButton(f"📚 {db['novels'][nid]['title']}", callback_data=f"open_novel_{nid}"))
            for mid in manga_results[:15]:
                m.add(types.InlineKeyboardButton(f"🎨 {db['manga'][mid]['title']}", callback_data=f"open_manga_{mid}"))
            for sid in series_results[:15]:
                m.add(types.InlineKeyboardButton(f"🎬 {db['series'][sid]['title']}", callback_data=f"open_series_{sid}"))
            text = "🔍 *نتائج البحث:*" if (item_results or novel_results or manga_results or series_results) else "❌ لا توجد نتائج مطابقة."
            m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="search_filter"))
            edit(text, m)
    elif data == "view_notifications":
        ensure_user(call.message)
        for n in db["users"][ustr].get("notifications", []):
            n["read"] = True
        sync_db()
        notifs = db["users"][ustr].get("notifications", [])
        text = f"🔔 *التنبيهات ({len(notifs)}):*" if notifs else "🔔 *لا توجد تنبيهات بعد.*"
        edit(text, notifications_kb(ustr))
    elif data.startswith("opennotif_"):
        nid = data[10:]
        if get_novel(nid):
            open_novel(uid, cid, mid, nid, as_new_message=True)
        else:
            bot.answer_callback_query(call.id, "❌ العمل غير موجود حاليًا.", show_alert=True)
    elif data == "clear_notifications":
        ensure_user(call.message)
        db["users"][ustr]["notifications"] = []
        sync_db()
        edit("🔔 *لا توجد تنبيهات.*", notifications_kb(ustr))
    elif data == "view_favs":
        ensure_user(call.message)
        edit("📂 *قائمتي — اختر القسم:*", list_sections_kb())
    elif data.startswith("listsec_") and not data.startswith(("listsec_rename_", "listsec_add", "listsec_marker")):
        sec_id = data[8:]
        ensure_user(call.message)
        edit(f"📂 *{get_section_name(sec_id)}:*", list_section_entries_kb(ustr, sec_id))
    elif data == "my_account":
        u = ensure_user(call.message); sub = "✅ نشط" if check_sub(uid) else "❌ غير مشترك"
        edit(f"👤 *حسابي:*\n\nالاسم: {u.get('first_name','')}\nالنقاط: {u.get('points',0)} 💎\nالاشتراك: {sub}", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home")))
    elif data == "redeem_code":
        db["pending_actions"][ustr] = {"action": "redeem"}; sync_db(); bot.send_message(cid, "🎟️ أرسل الكود:")
    elif data == "contact_us":
        db["pending_actions"][ustr] = {"action": "inquiry"}; sync_db(); edit("💬 *الدعم:*\nأرسل استفسارك الآن وسيتم الرد قريباً.", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home")))
    
    # ── Admin Panel ──────────────────────────────────────────────────────────
    elif data == "admin_panel" and is_admin(uid): edit("⚙️ *لوحة التحكم:*", admin_main_kb())
    elif data == "adm_quick_toggle" and is_admin(uid):
        edit("🔌 *تشغيل/تعطيل الميزات*\n\n🟢 = مفعّلة وتشتغل عادي | 🔴 = معطّلة بالكامل (ما تشتغل حتى لو استُدعيت بأي طريقة)\n\nاضغط على أي ميزة عشان تبدّل حالتها فورًا.", quick_toggle_kb())
    elif data.startswith("qtoggle_btn_") and is_admin(uid):
        bid_ = data[12:]
        for btn in db["config"]["menu_buttons"]:
            if btn["id"] == bid_:
                btn["active"] = not btn.get("active", True)
                break
        sync_db()
        edit("🔌 *تشغيل/تعطيل الميزات*\n\n🟢 = مفعّلة وتشتغل عادي | 🔴 = معطّلة بالكامل (ما تشتغل حتى لو استُدعيت بأي طريقة)\n\nاضغط على أي ميزة عشان تبدّل حالتها فورًا.", quick_toggle_kb())
    elif data.startswith("qtoggle_feat_") and is_admin(uid):
        fid = data[13:]
        f = db["config"]["custom_features"].get(fid)
        if f:
            f["active"] = not f.get("active", True)
            sync_db()
        edit("🔌 *تشغيل/تعطيل الميزات*\n\n🟢 = مفعّلة وتشتغل عادي | 🔴 = معطّلة بالكامل (ما تشتغل حتى لو استُدعيت بأي طريقة)\n\nاضغط على أي ميزة عشان تبدّل حالتها فورًا.", quick_toggle_kb())
    elif data == "adm_announce_channel" and is_admin(uid):
        ac = db["config"]["announce_channel"]
        chat = ac.get("chat_id") or "غير مربوطة"
        mode_text = "كل الأعمال" if ac.get("mode") == "all" else f"أعمال مختارة ({len(ac.get('selected_novels', []))})"
        text = (f"📣 *قناة إعلان تحديثات الأعمال*\n\n"
                f"القناة الحالية: `{chat}`\n"
                f"الحالة: {'✅ مفعّلة' if ac.get('enabled') else '❌ معطّلة'}\n"
                f"وضع النشر: *{mode_text}*\n\n"
                f"كل ما ينضاف فصل جديد لعمل، ينشر إعلان بهذي القناة تلقائيًا مع زر "
                f"يودّي المستخدم للعمل بالبوت مباشرة.")
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("🔗 تعيين/تغيير القناة", callback_data="announce_set_channel"))
        m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if ac.get('enabled') else '🟢 تفعيل'} النشر التلقائي", callback_data="announce_toggle_enabled"))
        m.add(types.InlineKeyboardButton(f"{'📚 التبديل لكل الأعمال' if ac.get('mode')=='selected' else '🎯 التبديل لأعمال مختارة'}", callback_data="announce_toggle_mode"))
        if ac.get("mode") == "selected":
            m.add(types.InlineKeyboardButton("🎯 اختيار الأعمال المسموح لها", callback_data="announce_pick_novels"))
        m.add(types.InlineKeyboardButton("✏️ تعديل نص الإعلان", callback_data="announce_edit_template"))
        m.add(types.InlineKeyboardButton("✏️ تعديل اسم الزر", callback_data="announce_edit_button"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(text, m)
    elif data == "announce_set_channel" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "set_announce_channel"}; sync_db()
        bot.send_message(cid, "🔗 أرسل يوزرنيم القناة (مثل @MyChannel) أو آيدي القناة، والبوت لازم يكون مشرف فيها:")
    elif data == "announce_toggle_enabled" and is_admin(uid):
        db["config"]["announce_channel"]["enabled"] = not db["config"]["announce_channel"].get("enabled", True)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.")
        edit("📣 *قناة إعلان تحديثات الأعمال*", admin_main_kb())
    elif data == "announce_toggle_mode" and is_admin(uid):
        cur = db["config"]["announce_channel"].get("mode", "selected")
        db["config"]["announce_channel"]["mode"] = "all" if cur == "selected" else "selected"
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.")
        edit("📣 *قناة إعلان تحديثات الأعمال*", admin_main_kb())
    elif data == "announce_pick_novels" and is_admin(uid):
        selected = db["config"]["announce_channel"].get("selected_novels", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for nid, nv in db["novels"].items():
            mark = "✅ " if nid in selected else ""
            m.add(types.InlineKeyboardButton(f"{mark}{nv['title']}", callback_data=f"announce_toggle_novel_{nid}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_announce_channel"))
        edit("🎯 *اختر الأعمال المسموح لها تُعلن بالقناة (اضغط للتبديل):*", m)
    elif data.startswith("announce_toggle_novel_") and is_admin(uid):
        nid = data[22:]
        selected = db["config"]["announce_channel"].setdefault("selected_novels", [])
        if nid in selected: selected.remove(nid)
        else: selected.append(nid)
        sync_db()
        selected2 = db["config"]["announce_channel"].get("selected_novels", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for nid2, nv in db["novels"].items():
            mark = "✅ " if nid2 in selected2 else ""
            m.add(types.InlineKeyboardButton(f"{mark}{nv['title']}", callback_data=f"announce_toggle_novel_{nid2}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_announce_channel"))
        edit("🎯 *اختر الأعمال المسموح لها تُعلن بالقناة (اضغط للتبديل):*", m)
    elif data == "announce_edit_template" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "set_announce_template"}; sync_db()
        bot.send_message(cid, "✏️ أرسل نص الإعلان الجديد. استخدم `{title}` لاسم العمل و `{chapter}` لرقم الفصل:\n\nمثال: 🆕 تحديث جديد: *{title}* — الفصل {chapter}", parse_mode="Markdown")
    elif data == "announce_edit_button" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "set_announce_button"}; sync_db()
        bot.send_message(cid, "✏️ أرسل اسم الزر الجديد (مثل: 📖 اقرأ الآن):")
    elif data.startswith("archivepick_"):
        wid = data[12:]
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "archive_link_pick_work":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة.", show_alert=True); return
        content_type = pending["content_type"]
        pool = {"novels": db["novels"], "manga": db["manga"], "series": db["series"]}[content_type]
        w = pool.get(wid)
        if not w:
            bot.answer_callback_query(call.id, "❌ العمل غير موجود.", show_alert=True); return
        db["pending_actions"][ustr] = {**pending, "action": "archive_link_pick_chapter", "work_id": wid}
        sync_db()
        unit = "الحلقة" if content_type == "series" else "الفصل"
        edit(f"✅ العمل: *{w['title']}*\n\nأرسل رقم {unit}:", None)
    elif data == "adm_snapshots" and is_admin(uid):
        if not is_owner(uid):
            bot.answer_callback_query(call.id, "❌ هذي الميزة للمطور فقط، لأنها تقدر تغيّر كل بيانات البوت.", show_alert=True)
            return
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("💾 خذ نسخة احتياطية الآن", callback_data="snapshot_create"))
        m.add(types.InlineKeyboardButton("📋 عرض النسخ المتاحة", callback_data="snapshot_list"))
        m.add(types.InlineKeyboardButton("🔑 استرجاع بكود", callback_data="snapshot_restore_code"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit("💾 *النسخ الاحتياطية والاسترجاع*\n\nتقدر تأخذ نسخة كاملة من كل شي بالبوت (روايات، ملفات، ميزات، إعدادات، مستخدمين) بأي وقت، وترجعها لاحقًا بكود قصير — البوت يرجع بالضبط لنفس حالته وقت أخذ النسخة.", m)
    elif data == "snapshot_create" and is_owner(uid):
        bot.answer_callback_query(call.id, "⏳ جاري أخذ النسخة الاحتياطية...")
        code = create_snapshot(label=f"يدوية بواسطة المطور")
        if code:
            bot.send_message(cid, f"✅ *تم أخذ نسخة احتياطية كاملة بنجاح.*\n\nكود الاسترجاع الخاص بك:\n`{code}`\n\n"
                                   f"احتفظ بهذا الكود — أرسله لي بأي وقت (أو استخدم زر «استرجاع بكود») عشان "
                                   f"يرجع البوت بالكامل لنفس حالته الآن.", parse_mode="Markdown")
        else:
            bot.send_message(cid, "❌ فشل أخذ النسخة الاحتياطية. تأكد من الاتصال بقاعدة البيانات.")
    elif data == "snapshot_list" and is_owner(uid):
        snaps = list_snapshots()
        if not snaps:
            edit("📋 *لا توجد نسخ احتياطية بعد.*", admin_main_kb())
        else:
            lines = ["📋 *آخر 20 نسخة احتياطية:*\n"]
            for code, label, created_at in snaps:
                lines.append(f"`{code}` — {created_at.split('.')[0]} {f'({label})' if label else ''}")
            m = types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_snapshots"))
            edit("\n".join(lines), m)
    elif data == "snapshot_restore_code" and is_owner(uid):
        db["pending_actions"][ustr] = {"action": "restore_snapshot_code"}; sync_db()
        bot.send_message(cid, "🔑 أرسل كود الاسترجاع:")
    elif data.startswith("snapshot_confirm_restore_") and is_owner(uid):
        code = data[25:]
        bot.answer_callback_query(call.id, "⏳ جاري الاسترجاع...")
        ok = restore_snapshot(code)
        if ok:
            bot.send_message(cid, "✅ *تم استرجاع البوت بالكامل لنفس حالته وقت أخذ هذي النسخة.*", parse_mode="Markdown")
        else:
            bot.send_message(cid, "❌ فشل الاسترجاع. تأكد إن الكود صحيح.")
    elif data == "adm_content_protection" and is_admin(uid):
        cp = db["config"].get("content_protection", {})
        copy_on = cp.get("copy", True); save_on = cp.get("save", True)
        text = (f"🔒 *حماية المحتوى*\n\n"
                f"لما مفعّلة، المستخدمون العاديون ما يقدرون ينسخون/يعيدون توجيه/يحفظون "
                f"الملفات والفصول اللي يرسلها البوت. المطور والمشرفين مستثنون دائمًا "
                f"ويقدرون يتصرفون بالملفات بحرية للإدارة.\n\n"
                f"منع النسخ/إعادة التوجيه: {'✅ مفعّل' if copy_on else '❌ معطّل'}\n"
                f"منع الحفظ: {'✅ مفعّل' if save_on else '❌ معطّل'}")
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if copy_on else '🟢 تفعيل'} منع النسخ/التوجيه", callback_data="toggle_protect_copy"))
        m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if save_on else '🟢 تفعيل'} منع الحفظ", callback_data="toggle_protect_save"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(text, m)
    elif data == "toggle_protect_copy" and is_admin(uid):
        db["config"]["content_protection"]["copy"] = not db["config"]["content_protection"].get("copy", True)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.")
        edit("🔒 *حماية المحتوى*", admin_main_kb())
    elif data == "toggle_protect_save" and is_admin(uid):
        db["config"]["content_protection"]["save"] = not db["config"]["content_protection"].get("save", True)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.")
        edit("🔒 *حماية المحتوى*", admin_main_kb())
    elif data == "adm_convert_file" and can_use_convert(uid):
        if is_owner(uid):
            allowed = db["config"].get("convert_allowed_users", [])
            always_remove = db["config"].get("always_remove_phrases", [])
            phrasing_rules = db["config"].get("translation_phrasing_rules", {})
            pm = db["config"].get("public_merge_tool", {})
            m = types.InlineKeyboardMarkup(row_width=1)
            m.add(types.InlineKeyboardButton("📤 استخدم الأداة الآن", callback_data="convert_do_use"))
            m.add(types.InlineKeyboardButton("🧩 دمج عدة ملفات بملف واحد", callback_data="filemerge_start"))
            m.add(types.InlineKeyboardButton(f"👥 إدارة المستخدمين المسموح لهم ({len(allowed)})", callback_data="convert_manage_users"))
            m.add(types.InlineKeyboardButton(f"🗑️ عبارات الحذف الدائمة ({len(always_remove)})", callback_data="manage_always_remove"))
            m.add(types.InlineKeyboardButton(f"🌐 قواعد صياغة الترجمة ({len(phrasing_rules)})", callback_data="manage_phrasing_rules"))
            m.add(types.InlineKeyboardButton(f"🧩 إعدادات أداة الدمج العامة ({'✅ مفعّلة' if pm.get('enabled', True) else '❌ معطّلة'})", callback_data="manage_merge_tool_settings"))
            edit("🔄 *أدوات الملفات (تحويل صيغة / تعديل اسم / حذف نص / ترجمة / دمج)*", m)
        else:
            db["pending_actions"][ustr] = {"action": "convert_upload_wait"}; sync_db()
            bot.send_message(cid, "📤 أرسل الملف اللي تبي تشتغل عليه (PDF, DOCX, DOC, EPUB, TXT, HTML):")
    elif data == "manage_merge_tool_settings" and is_owner(uid):
        pm = db["config"].get("public_merge_tool", {})
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if pm.get('enabled', True) else '🟢 تفعيل'} أداة الدمج للمستخدمين", callback_data="toggle_merge_tool_enabled"))
        m.add(types.InlineKeyboardButton(f"✏️ تغيير الحد المجاني (حاليًا {pm.get('max_files_free', 5)} ملفات)", callback_data="set_merge_free_limit_files"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_convert_file"))
        edit("🧩 *إعدادات أداة الدمج العامة*\n\nالمطور والمستخدمين المسموح لهم بتحويل الصيغة دائمًا غير محدودين بعدد الملفات.", m)
    elif data == "toggle_merge_tool_enabled" and is_owner(uid):
        db["config"]["public_merge_tool"]["enabled"] = not db["config"]["public_merge_tool"].get("enabled", True)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.", show_alert=True)
        edit("🧩 *إعدادات أداة الدمج العامة*", admin_main_kb())
    elif data == "set_merge_free_limit_files" and is_owner(uid):
        db["pending_actions"][ustr] = {"action": "set_merge_free_limit_files_input"}; sync_db()
        bot.send_message(cid, "🔢 أرسل العدد الأقصى الجديد للملفات اللي يقدر المستخدم العادي يدمجها مجانًا بعملية وحدة:")
    elif data == "convert_do_use" and can_use_convert(uid):
        db["pending_actions"][ustr] = {"action": "convert_upload_wait"}; sync_db()
        bot.send_message(cid, "📤 أرسل الملف اللي تبي تشتغل عليه (PDF, DOCX, DOC, EPUB, TXT, HTML):")
    elif data == "manage_always_remove" and is_owner(uid):
        phrases = db["config"].get("always_remove_phrases", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for i, p in enumerate(phrases):
            m.add(types.InlineKeyboardButton(f"❌ {p[:40]}", callback_data=f"remove_alwaysphrase_{i}"))
        m.add(types.InlineKeyboardButton("➕ إضافة عبارة دائمة الحذف", callback_data="add_alwaysphrase"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_convert_file"))
        edit("🗑️ *عبارات تُحذف تلقائيًا من أي ملف* (بالإضافة لأي عبارة تكتبها وقت الاستخدام):", m)
    elif data == "add_alwaysphrase" and is_owner(uid):
        db["pending_actions"][ustr] = {"action": "add_always_remove_phrase"}; sync_db()
        bot.send_message(cid, "✏️ أرسل العبارة أو الرابط اللي تبي يُحذف دائمًا من كل ملف مستقبلًا:")
    elif data.startswith("remove_alwaysphrase_") and is_owner(uid):
        idx = int(data[21:])
        phrases = db["config"].get("always_remove_phrases", [])
        if 0 <= idx < len(phrases):
            phrases.pop(idx); sync_db()
        bot.answer_callback_query(call.id, "✅ تم الحذف.", show_alert=True)
        edit("🗑️ *عبارات تُحذف تلقائيًا*", admin_main_kb())
    elif data == "filemerge_start":
        if not db["config"].get("public_merge_tool", {}).get("enabled", True) and not can_use_convert(uid):
            bot.answer_callback_query(call.id, "❌ أداة الدمج غير مفعّلة حاليًا.", show_alert=True); return
        db["pending_actions"][ustr] = {"action": "filemerge_collect", "sections": [], "files_meta": []}
        sync_db()
        bot.send_message(cid, "🧩 *دمج عدة ملفات بملف واحد*\n\nأرسل أول ملف (PDF/DOCX/DOC/EPUB/TXT/HTML).", parse_mode="Markdown")
    elif data == "filemerge_add_more":
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "filemerge_ask_more":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة.", show_alert=True); return
        limit = db["config"].get("public_merge_tool", {}).get("max_files_free", 5)
        if not can_use_convert(uid) and len(pending["sections"]) >= limit:
            bot.answer_callback_query(call.id, f"❌ وصلت الحد الأقصى ({limit} ملفات) للدمج المجاني. اضغط «✅ خلصت» للدمج الآن.", show_alert=True)
            return
        db["pending_actions"][ustr] = {**pending, "action": "filemerge_collect"}
        sync_db()
        bot.send_message(cid, f"📤 أرسل الملف رقم {len(pending['sections'])+1}:")
    elif data == "filemerge_finish":
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "filemerge_ask_more" or not pending.get("sections"):
            bot.answer_callback_query(call.id, "❌ لازم ترفع ملف واحد على الأقل.", show_alert=True); return
        m = types.InlineKeyboardMarkup(row_width=2)
        all_formats = ["pdf", "docx", "epub", "txt", "html"]
        m.add(*[types.InlineKeyboardButton(f.upper(), callback_data=f"filemerge_fmt_{f}") for f in all_formats])
        db["pending_actions"][ustr] = {**pending, "action": "filemerge_choose_fmt"}
        sync_db()
        edit(f"✅ عندك {len(pending['sections'])} ملف جاهز للدمج.\n\nاختر صيغة الملف النهائي:", m)
    elif data.startswith("filemerge_fmt_"):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "filemerge_choose_fmt":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة.", show_alert=True); return
        target_fmt = data[14:]
        sections = pending["sections"]
        bot.answer_callback_query(call.id, "⏳ جاري الدمج بالخلفية... راح أرسل لك الملف هنا أول ما يخلص")
        out_path = f"/tmp/merged_{uuid.uuid4().hex[:8]}.{target_fmt}"
        db["pending_actions"].pop(ustr, None); sync_db()

        def _do_filemerge_job():
            t_start = time.time()
            try:
                merge_files_to_one(sections, target_fmt, out_path, title="Merged")
                logger.info(f"[filemerge] نجح دمج {len(sections)} ملف لـ {target_fmt} خلال {time.time()-t_start:.1f}ث")
                with open(out_path, "rb") as f:
                    bot.send_document(cid, f, visible_file_name=f"merged.{target_fmt}",
                                       caption=f"✅ تم دمج {len(sections)} ملف بملف {target_fmt.upper()} واحد.")
            except Exception as e:
                logger.error(f"[filemerge] فشل الدمج بعد {time.time()-t_start:.1f}ث: {e}")
                bot.send_message(cid, f"❌ فشل الدمج:\n`{str(e)[:300]}`", parse_mode="Markdown")
            finally:
                try: os.remove(out_path)
                except Exception: pass

        threading.Thread(target=_do_filemerge_job, daemon=True).start()
    elif data == "manage_phrasing_rules" and is_owner(uid):
        rules = db["config"].get("translation_phrasing_rules", {})
        m = types.InlineKeyboardMarkup(row_width=1)
        for k, v in rules.items():
            m.add(types.InlineKeyboardButton(f"❌ {k} → {v}", callback_data=f"remove_phrasing_{k}"))
        m.add(types.InlineKeyboardButton("➕ إضافة قاعدة صياغة", callback_data="add_phrasing_rule"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_convert_file"))
        edit("🌐 *قواعد صياغة ثابتة بالترجمة* (مثال: «إله» تُترجم دائمًا «حاكم»):", m)
    elif data == "add_phrasing_rule" and is_owner(uid):
        db["pending_actions"][ustr] = {"action": "add_phrasing_rule_original"}; sync_db()
        bot.send_message(cid, "✏️ أرسل الكلمة/العبارة الأصلية:")
    elif data.startswith("remove_phrasing_") and is_owner(uid):
        key = data[16:]
        rules = db["config"].get("translation_phrasing_rules", {})
        rules.pop(key, None); sync_db()
        bot.answer_callback_query(call.id, "✅ تم الحذف.", show_alert=True)
        edit("🌐 *قواعد صياغة ثابتة بالترجمة*", admin_main_kb())
    elif data == "convert_manage_users" and is_owner(uid):
        allowed = db["config"].get("convert_allowed_users", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for aid in allowed:
            m.add(types.InlineKeyboardButton(f"❌ إزالة {aid}", callback_data=f"convert_remove_user_{aid}"))
        m.add(types.InlineKeyboardButton("➕ إضافة مستخدم", callback_data="convert_add_user"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_convert_file"))
        edit("👥 *المستخدمون المسموح لهم استخدام تحويل الصيغة* (غير المطور):", m)
    elif data == "convert_add_user" and is_owner(uid):
        db["pending_actions"][ustr] = {"action": "convert_add_user_id"}; sync_db()
        bot.send_message(cid, "🆔 أرسل آيدي المستخدم اللي تبي تسمح له باستخدام تحويل الصيغة:")
    elif data.startswith("convert_remove_user_") and is_owner(uid):
        target = data[21:]
        allowed = db["config"].get("convert_allowed_users", [])
        if target in allowed:
            allowed.remove(target); sync_db()
        bot.answer_callback_query(call.id, "✅ تم إزالة المستخدم من القائمة.", show_alert=True)
        edit("👥 *المستخدمون المسموح لهم استخدام تحويل الصيغة* (غير المطور):", admin_main_kb())
    elif data == "filetool_convert" and can_use_convert(uid):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "file_tool_menu":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة، ابدأ من جديد.", show_alert=True); return
        src_ext = pending["src_ext"]
        all_formats = ["pdf", "docx", "doc", "epub", "txt", "html"]
        targets = [f for f in all_formats if f != src_ext]
        m = types.InlineKeyboardMarkup(row_width=2)
        m.add(*[types.InlineKeyboardButton(f.upper(), callback_data=f"convfmt_{f}") for f in targets])
        db["pending_actions"][ustr] = {**pending, "action": "convert_choose_fmt"}
        sync_db()
        edit("🔄 *اختر الصيغة اللي تبي تحوّل لها:*", m)
    elif data == "filetool_rename" and can_use_convert(uid):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "file_tool_menu":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة، ابدأ من جديد.", show_alert=True); return
        db["pending_actions"][ustr] = {**pending, "action": "rename_file_input"}
        sync_db()
        bot.send_message(cid, f"✏️ الاسم الحالي: *{pending['title']}*\n\nأرسل الاسم الجديد (بدون امتداد الصيغة):", parse_mode="Markdown")
    elif data == "filetool_removetext" and can_use_convert(uid):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "file_tool_menu":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة، ابدأ من جديد.", show_alert=True); return
        db["pending_actions"][ustr] = {**pending, "action": "removetext_input"}
        sync_db()
        bot.send_message(cid, "🗑️ أرسل العبارة أو الرابط اللي تبي تحذفه من الملف كامل (كل مرة يتكرر فيها بينحذف).\n\n"
                               "لو تبي تحذف أكثر من عبارة، أرسلهم كل وحدة بسطر منفصل بنفس الرسالة.")
    elif data == "filetool_translate" and can_use_translate(uid):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "file_tool_menu":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة، ابدأ من جديد.", show_alert=True); return
        m = types.InlineKeyboardMarkup(row_width=2)
        langs = [("ar", "🇸🇦 العربية"), ("en", "🇬🇧 English"), ("ko", "🇰🇷 한국어"), ("zh", "🇨🇳 中文")]
        m.add(*[types.InlineKeyboardButton(lbl, callback_data=f"translatelang_{code}") for code, lbl in langs])
        db["pending_actions"][ustr] = {**pending, "action": "translate_choose_lang"}
        sync_db()
        edit("🌐 *اختر اللغة اللي تبي تترجم الملف لها:*", m)
    elif data.startswith("translatelang_") and can_use_translate(uid):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "translate_choose_lang":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة، ابدأ من جديد.", show_alert=True); return
        target_lang = data[14:]
        src_path = pending["src_path"]; src_ext = pending["src_ext"]; title = pending.get("title", "Document")
        bot.answer_callback_query(call.id, "⏳ جاري الترجمة (قد تأخذ دقيقة للنصوص الطويلة)...")
        try:
            if not _ensure_convert_libs():
                raise RuntimeError(f"مكتبات التحويل غير مثبتة: {_CONVERT_IMPORT_ERROR}")
            text = extract_text_from_file(src_path, src_ext)
            res = translate_text(text, target_lang)
            if isinstance(res, dict) and "__ai_error__" in res:
                bot.send_message(cid, ai_error_message(res, uid), parse_mode="Markdown")
            else:
                translated = res["translated_text"]
                out_path = f"/tmp/translated_{uuid.uuid4().hex[:8]}.{src_ext}"
                if src_ext == "pdf": build_pdf_from_text(translated, out_path, title)
                elif src_ext in ("docx", "doc"): build_docx_from_text(translated, out_path)
                elif src_ext == "epub": build_epub_from_text(translated, out_path, title)
                elif src_ext == "txt": build_txt_from_text(translated, out_path)
                elif src_ext in ("html", "htm"): build_html_from_text(translated, out_path, title)
                with open(out_path, "rb") as f:
                    bot.send_document(cid, f, visible_file_name=f"{title}_{target_lang}.{src_ext}",
                                       caption=f"✅ تمت الترجمة بنجاح.")
                try: os.remove(out_path)
                except Exception: pass
        except Exception as e:
            bot.send_message(cid, f"❌ فشلت الترجمة:\n`{str(e)[:300]}`", parse_mode="Markdown")
        finally:
            try: os.remove(src_path)
            except Exception: pass
            db["pending_actions"].pop(ustr, None); sync_db()
    elif data.startswith("convfmt_") and can_use_convert(uid):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "convert_choose_fmt":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة، ابدأ من جديد.", show_alert=True); return
        target_fmt = data[8:]
        src_path = pending["src_path"]; src_ext = pending["src_ext"]; title = pending.get("title", "Document")
        if not os.path.exists(src_path):
            # الملف المؤقت بـ /tmp اختفى — غالبًا لأن البوت أعاد التشغيل أو
            # الاستضافة نظّفت /tmp بين وقت رفع الملف ووقت اختيار الصيغة
            # (الملفات المؤقتة مو دائمة ولا تُحفظ بقاعدة البيانات، بعكس
            # pending_actions نفسها اللي فضلت محفوظة بـ DB وتشاور لملف مفقود).
            logger.warning(f"[convert] الملف المصدر مفقود من /tmp للمستخدم {uid}: {src_path}")
            bot.answer_callback_query(call.id, "❌ انتهت صلاحية الملف المؤقت", show_alert=True)
            bot.send_message(cid, "⚠️ الملف الأصلي ماعاد موجود مؤقتًا (على الأغلب بسبب إعادة تشغيل الاستضافة). "
                                    "أعد رفع الملف من جديد وحاول التحويل مرة ثانية.")
            db["pending_actions"].pop(ustr, None); sync_db()
            return
        bot.answer_callback_query(call.id, "⏳ جاري التحويل بالخلفية... راح أرسل لك الملف هنا أول ما يخلص")
        out_path = f"/tmp/converted_{uuid.uuid4().hex[:8]}.{target_fmt}"
        logger.info(f"[convert] بدء تحويل {src_ext}->{target_fmt} للمستخدم {uid}")
        # ملاحظة مهمة: ما نمسح pending_actions ولا ملف المصدر هنا. نخلي الحالة
        # "convert_choose_fmt" باقية بعد التحويل، عشان المستخدم يقدر يحوّل نفس
        # الملف لأي عدد صيغ يبيه بدون ما يعيد رفعه كل مرة ("انتهت الجلسة").
        # الملف المصدر ينحذف فقط لما يبدأ عملية جديدة (رفع ملف ثاني) أو تنتهي
        # صلاحيته لأي سبب آخر.

        def _do_conversion_job():
            # نشغّل التحويل بخيط خلفية حقيقي بدون قتل قسري بمهلة ثابتة — التحويل
            # الفعلي (خصوصًا PDF مع تشكيل عربي لكل سطر) قد ياخذ وقت متفاوت حسب
            # حجم الملف وسرعة السيرفر، والقتل القسري عند 45 ثانية كان يفشّل حتى
            # التحويلات السليمة اللي بس بطيئة شوي بدل ما تكون معلّقة فعلاً.
            # بدل هذا، نخلي التحويل يكمل بالخلفية ونرسل النتيجة (نجاح أو خطأ)
            # للمستخدم أول ما يخلص، مهما طال الوقت.
            t_start = time.time()
            try:
                convert_file(src_path, src_ext, target_fmt, out_path, title)
                logger.info(f"[convert] نجح تحويل {src_ext}->{target_fmt} للمستخدم {uid} خلال {time.time()-t_start:.1f}ث")
                with open(out_path, "rb") as f:
                    bot.send_document(cid, f, visible_file_name=f"{title}.{target_fmt}",
                                       caption=f"✅ تم التحويل من {src_ext.upper()} إلى {target_fmt.upper()}.")
                # نعيد عرض قائمة الصيغ فورًا حتى يقدر يحوّل لصيغة ثانية بضغطة وحدة
                all_formats = ["pdf", "docx", "doc", "epub", "txt", "html"]
                targets = [f for f in all_formats if f != src_ext]
                m = types.InlineKeyboardMarkup(row_width=2)
                m.add(*[types.InlineKeyboardButton(f.upper(), callback_data=f"convfmt_{f}") for f in targets])
                m.add(types.InlineKeyboardButton("✅ خلصت من هذا الملف", callback_data="convert_done"))
                bot.send_message(cid, "🔄 تبي تحوّل نفس الملف لصيغة ثانية؟", reply_markup=m)
            except Exception as e:
                logger.error(f"[convert] فشل تحويل {src_ext}->{target_fmt} للمستخدم {uid} بعد {time.time()-t_start:.1f}ث: {e}")
                if isinstance(e, FileNotFoundError):
                    bot.send_message(cid, "⚠️ الملف الأصلي ماعاد موجود مؤقتًا (على الأغلب بسبب إعادة تشغيل الاستضافة). "
                                            "أعد رفع الملف من جديد وحاول التحويل مرة ثانية.")
                    db["pending_actions"].pop(ustr, None); sync_db()
                else:
                    bot.send_message(cid, f"❌ فشل التحويل:\n`{str(e)[:300]}`", parse_mode="Markdown")
            finally:
                try: os.remove(out_path)
                except Exception: pass

        threading.Thread(target=_do_conversion_job, daemon=True).start()
    elif data == "convert_done" and can_use_convert(uid):
        pending = db["pending_actions"].get(ustr, {})
        src_path = pending.get("src_path")
        if src_path:
            try: os.remove(src_path)
            except Exception: pass
        db["pending_actions"].pop(ustr, None); sync_db()
        bot.answer_callback_query(call.id, "✅ تم")
        bot.send_message(cid, "تم إنهاء الجلسة. أرسل لي ملف جديد أي وقت تبي تحوّله.")
    elif data == "adm_merge_limits" and is_admin(uid):
        ml = db["config"].get("merge_limits", {})
        free_v = ml.get("free_daily", 10)
        sub_v = ml.get("sub_daily")
        sub_text = "غير محدود ♾️" if sub_v is None else str(sub_v)
        text = (f"📦 *حدود التنزيلات المدمجة اليومية*\n\n"
                f"👤 المستخدم العادي: *{free_v}* ملف مدمج/يوم\n"
                f"💎 المشترك: *{sub_text}*\n\n"
                f"تقدر تخصص عدد مختلف لأي مستخدم بعينه من زر «تخصيص لمستخدم».")
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("✏️ تغيير حد المستخدم العادي", callback_data="mlim_edit_free"))
        m.add(types.InlineKeyboardButton("✏️ تغيير حد المشترك", callback_data="mlim_edit_sub"))
        m.add(types.InlineKeyboardButton("♾️ اجعل حد المشترك غير محدود", callback_data="mlim_sub_unlimited"))
        m.add(types.InlineKeyboardButton("👤 تخصيص لمستخدم معيّن", callback_data="mlim_edit_user"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(text, m)
    elif data == "mlim_edit_free" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "set_merge_free_limit"}; sync_db()
        bot.send_message(cid, "🔢 أرسل العدد الجديد لحد التنزيل المدمج اليومي للمستخدم العادي:")
    elif data == "mlim_edit_sub" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "set_merge_sub_limit"}; sync_db()
        bot.send_message(cid, "🔢 أرسل العدد الجديد لحد التنزيل المدمج اليومي للمشترك:")
    elif data == "mlim_sub_unlimited" and is_admin(uid):
        db["config"]["merge_limits"]["sub_daily"] = None; sync_db()
        bot.answer_callback_query(call.id, "✅ صار حد المشترك غير محدود.", show_alert=True)
    elif data == "mlim_edit_user" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "set_merge_user_override_id"}; sync_db()
        bot.send_message(cid, "🆔 أرسل آيدي المستخدم اللي تبي تخصص له حد مختلف:")
    elif data == "adm_stats" and is_admin(uid):
        edit("📊 *لوحة الإحصائيات المتطورة*\n\nاختر نوع الإحصائية:", stats_menu_kb())
    elif data == "adm_services_status" and is_admin(uid):
        def _key_status(val):
            return "🟢 مضبوط" if val else "⚪ غير مضبوط"
        ts = db["config"]["translation_settings"]
        engine = ts.get("engine", "ai")
        text = (
            "🔌 *حالة الخدمات الخارجية*\n\n"
            "*التحويل — سلسلة الأولوية:*\n"
            f"1️⃣ CloudConvert: {_key_status(CLOUDCONVERT_KEY)}\n"
            f"2️⃣ Convertio: {_key_status(CONVERTIO_KEY)}\n"
            f"3️⃣ ConvertAPI: {_key_status(CONVERTAPI_SECRET)}\n"
            f"4️⃣ محلي (fallback دائم): 🟢 يعمل دومًا\n\n"
            "*الترجمة:*\n"
            f"المحرك النشط الآن: {'🤖 ذكاء اصطناعي (OpenRouter)' if engine == 'ai' else '⚡ محلي (deep-translator)'}\n"
            f"ذكاء اصطناعي (OpenRouter): {_key_status(OR_KEY)}\n"
            f"محلي (deep-translator): {'🟢 مثبت' if _TRANSLATE_LIBS_READY else '⚪ لم يُختبر بعد'}\n\n"
            "*الأمان:*\n"
            f"فحص نوع الملف الحقيقي (python-magic): {'🟢 مثبت' if _SECURITY_LIBS_READY else '⚪ لم يُختبر بعد'}\n\n"
            "*خدمات أخرى مقروءة (بدون تكامل مبرمج بعد):*\n"
            f"HTML2PDF: {_key_status(HTML2PDF_KEY)} — OCF: {_key_status(OCF_KEY)} — ChangeThisFile: {_key_status(CHANGETHISFILE_KEY)}\n\n"
            "_ملاحظة: \"لم يُختبر بعد\" تعني إن المكتبة لسا ما استُدعيت مرة، مو "
            "بالضرورة غير مثبتة — تتحدث لـ 🟢/❌ عند أول استخدام فعلي._"
        )
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("🧩 نظام التقطيع والاستئناف", callback_data="adm_chunking"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(text, m)
    elif data == "adm_chunking" and is_admin(uid):
        text, markup = _chunking_panel_content(); edit(text, markup)
    elif data == "chunk_toggle" and is_admin(uid):
        cfg = db["config"]["chunking_settings"]
        cfg["enabled"] = not cfg.get("enabled", False)
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم {'تفعيل' if cfg['enabled'] else 'تعطيل'} التقطيع")
        text, markup = _chunking_panel_content(); edit(text, markup)
    elif data == "chunk_set_pages" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "chunk_set_pages"}; sync_db()
        bot.send_message(cid, "📄 أرسل عدد الصفحات بكل جزء (مثال: 50):")
    elif data == "chunk_set_pause" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "chunk_set_pause"}; sync_db()
        bot.send_message(cid, "⏱️ أرسل عدد ثواني الانتظار بين كل جزء (مثال: 3):")
    elif data == "chunk_set_ram" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "chunk_set_ram"}; sync_db()
        bot.send_message(cid, "🧠 أرسل الحد الأقصى للرام بالميجابايت قبل الانتظار (مثال: 200):")
    elif data == "chunk_view_stats" and is_admin(uid):
        stats = temp_file_manager.stats()
        ram = memory_monitor.current_usage_mb()
        ram_text = f"{ram:.1f} MB" if ram is not None else "غير متاح (psutil غير مثبتة)"
        text = (f"📊 *إحصائيات التقطيع*\n\nخلفية تخزين المهام: {job_manager.backend_name()}\n"
                f"الرام الحالي للبوت: {ram_text}\n\n"
                f"*الملفات المؤقتة المسجّلة:*\nالعدد: {stats['registered']}\n"
                f"موجودة فعليًا بالقرص: {stats['existing_on_disk']}\nالحجم الإجمالي: {stats['total_size_mb']} MB")
        m2 = types.InlineKeyboardMarkup(row_width=1)
        m2.add(types.InlineKeyboardButton("🧹 تنظيف الآن", callback_data="chunk_cleanup_now"))
        m2.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_chunking"))
        edit(text, m2)
    elif data == "chunk_cleanup_now" and is_admin(uid):
        temp_file_manager.cleanup_old(max_age_seconds=0)  # 0 = ينظف كل شيء مسجَّل فورًا بغض النظر عن العمر
        job_manager.cleanup_old_jobs(days=db["config"]["chunking_settings"].get("job_retention_days", 7))
        bot.answer_callback_query(call.id, "✅ تم تنظيف الملفات المؤقتة والمهام القديمة.", show_alert=True)
        text, markup = _chunking_panel_content(); edit(text, markup)
    elif data == "chunk_view_incomplete" and is_admin(uid):
        jobs = job_manager.list_incomplete_jobs()
        if not jobs:
            text = "✅ *لا توجد مهام تقطيع متوقفة حاليًا.*"
        else:
            text = f"⏳ *مهام متوقفة ({len(jobs)}):*\n\n"
            for j in jobs[:10]:
                meta = j.get("meta", {})
                text += f"• `{j['job_id']}` — {meta.get('title','?')} ({j.get('completed_parts',0)}/{j.get('total_parts','?')} جزء)\n"
        m2 = types.InlineKeyboardMarkup(row_width=1)
        for j in jobs[:10]:
            m2.add(types.InlineKeyboardButton(f"▶️ استئناف {j['job_id']}", callback_data=f"chunk_resume_{j['job_id']}"))
        m2.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_chunking"))
        edit(text, m2)
    elif data.startswith("chunk_resume_") and is_admin(uid):
        job_id = data[len("chunk_resume_"):]
        bot.answer_callback_query(call.id, "⏳ جاري استئناف المهمة بالخلفية...")
        def _do_resume():
            out_path, err = resume_chunk_job(job_id)
            if err:
                bot.send_message(cid, f"❌ فشل استئناف المهمة `{job_id}`:\n{err}", parse_mode="Markdown")
            else:
                try:
                    with open(out_path, "rb") as f:
                        bot.send_document(cid, f, caption=f"✅ اكتمل استئناف المهمة `{job_id}`", parse_mode="Markdown")
                except Exception as e:
                    bot.send_message(cid, f"✅ اكتملت المهمة لكن فشل إرسال الملف: {e}\nالمسار: {out_path}")
        threading.Thread(target=_do_resume, daemon=True).start()
    elif data == "stats_overview" and is_admin(uid):
        s = db["config"]["stats"]
        today = str(datetime.now().date())
        text = (f"📊 *نظرة عامة:*\n\n"
                f"👥 إجمالي المستخدمين: *{s.get('total_users',0)}*\n"
                f"🆕 مستخدمين جدد اليوم: *{s.get('daily_new_users',{}).get(today,0)}*\n"
                f"📥 إجمالي التنزيلات: *{s.get('total_downloads',0)}*\n"
                f"📥 تنزيلات اليوم: *{s.get('daily_downloads',{}).get(today,0)}*\n"
                f"🔍 إجمالي عمليات البحث: *{s.get('total_searches',0)}*\n"
                f"📁 عدد الأقسام: *{len(db['categories'])}*\n"
                f"📄 عدد الملفات: *{len(db['items'])}*\n"
                f"📚 عدد الروايات: *{len(db['novels'])}*\n"
                f"🎨 عدد أعمال المانجا: *{len(db['manga'])}*\n"
                f"🎬 عدد المسلسلات/الأفلام: *{len(db['series'])}*\n"
                f"🚫 محظورين: *{len(db['config'].get('banned_users',[]))}*\n\n"
                f"👁️ ظهور الإحصائيات للمستخدمين: {'✅ مفعّل' if db['config'].get('stats_visible_to_users') else '❌ معطّل'}")
        edit(text, stats_menu_kb())
    elif data == "stats_top_novels" and is_admin(uid):
        s = db["config"]["stats"]
        views = s.get("novel_views", {}); downloads = s.get("novel_downloads", {})
        combined = sorted(set(list(views.keys()) + list(downloads.keys())),
                           key=lambda n: downloads.get(n, 0) + views.get(n, 0), reverse=True)[:15]
        lines = ["📚 *أكثر الروايات طلبًا:*\n"]
        for nid in combined:
            nv = db["novels"].get(nid)
            if not nv: continue
            lines.append(f"• {nv['title']} — 👁️ {views.get(nid,0)} | 📥 {downloads.get(nid,0)}")
        edit("\n".join(lines) if len(lines) > 1 else "❌ ما فيه بيانات كافية بعد.", stats_menu_kb())
    elif data == "stats_top_manga" and is_admin(uid):
        s = db["config"]["stats"]
        views = s.get("manga_views", {}); downloads = s.get("manga_downloads", {})
        combined = sorted(set(list(views.keys()) + list(downloads.keys())),
                           key=lambda n: downloads.get(n, 0) + views.get(n, 0), reverse=True)[:15]
        lines = ["🎨 *أكثر أعمال المانجا/المانهوا طلبًا:*\n"]
        for mid in combined:
            mv = db["manga"].get(mid)
            if not mv: continue
            lines.append(f"• {mv['title']} — 👁️ {views.get(mid,0)} | 📥 {downloads.get(mid,0)}")
        edit("\n".join(lines) if len(lines) > 1 else "❌ ما فيه بيانات كافية بعد.", stats_menu_kb())
    elif data == "stats_top_series" and is_admin(uid):
        s = db["config"]["stats"]
        views = s.get("series_views", {}); downloads = s.get("series_downloads", {})
        combined = sorted(set(list(views.keys()) + list(downloads.keys())),
                           key=lambda n: downloads.get(n, 0) + views.get(n, 0), reverse=True)[:15]
        lines = ["🎬 *أكثر المسلسلات/الأفلام طلبًا:*\n"]
        for sid in combined:
            sv = db["series"].get(sid)
            if not sv: continue
            lines.append(f"• {sv['title']} — 👁️ {views.get(sid,0)} | 📥 {downloads.get(sid,0)}")
        edit("\n".join(lines) if len(lines) > 1 else "❌ ما فيه بيانات كافية بعد.", stats_menu_kb())
    elif data == "stats_top_items" and is_admin(uid):
        s = db["config"]["stats"]
        views = s.get("item_views", {}); downloads = s.get("item_downloads", {})
        combined = sorted(set(list(views.keys()) + list(downloads.keys())),
                           key=lambda i: downloads.get(i, 0) + views.get(i, 0), reverse=True)[:15]
        lines = ["📄 *أكثر الملفات طلبًا:*\n"]
        for iid in combined:
            it = db["items"].get(iid)
            if not it: continue
            lines.append(f"• {it['title']} — 👁️ {views.get(iid,0)} | 📥 {downloads.get(iid,0)}")
        edit("\n".join(lines) if len(lines) > 1 else "❌ ما فيه بيانات كافية بعد.", stats_menu_kb())
    elif data == "stats_top_searches" and is_admin(uid):
        s = db["config"]["stats"]
        q = sorted(s.get("search_queries", {}).items(), key=lambda x: x[1], reverse=True)[:20]
        lines = ["🔍 *أكثر الكلمات/التصنيفات بحثًا (يفيدك تعرف وش يطلبه الناس أكثر):*\n"]
        for query, count in q:
            lines.append(f"• {query} — {count} مرة")
        edit("\n".join(lines) if len(lines) > 1 else "❌ ما فيه بيانات كافية بعد.", stats_menu_kb())
    elif data == "stats_growth" and is_admin(uid):
        s = db["config"]["stats"]
        days = sorted(s.get("daily_new_users", {}).items())[-14:]
        lines = ["📈 *نمو المستخدمين آخر 14 يوم:*\n"]
        for day, count in days:
            lines.append(f"• {day}: +{count}")
        edit("\n".join(lines) if len(lines) > 1 else "❌ ما فيه بيانات كافية بعد.", stats_menu_kb())
    elif data == "stats_toggle_visible" and is_admin(uid):
        db["config"]["stats_visible_to_users"] = not db["config"].get("stats_visible_to_users", False)
        sync_db()
        bot.answer_callback_query(call.id, f"👁️ ظهور الإحصائيات للمستخدمين: {'✅ مفعّل الآن' if db['config']['stats_visible_to_users'] else '❌ معطّل الآن'}", show_alert=True)
        edit("📊 *لوحة الإحصائيات المتطورة*\n\nاختر نوع الإحصائية:", stats_menu_kb())
    elif data == "view_public_stats":
        if not db["config"].get("stats_visible_to_users") and not is_admin(uid):
            bot.answer_callback_query(call.id, "❌ الإحصائيات غير متاحة حاليًا.", show_alert=True)
        else:
            s = db["config"]["stats"]
            text = (f"📊 *إحصائيات البوت:*\n\n"
                    f"👥 المستخدمين: *{s.get('total_users',0)}*\n"
                    f"📥 التنزيلات: *{s.get('total_downloads',0)}*\n"
                    f"📚 الروايات: *{len(db['novels'])}*\n"
                    f"📄 الملفات: *{len(db['items'])}*")
            m = types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home"))
            edit(text, m)
    elif data == "adm_maintenance" and is_admin(uid):
        db["config"]["maintenance"] = not db["config"]["maintenance"]; sync_db(); bot.answer_callback_query(call.id, f"الصيانة: {'مفعّل ✅' if db['config']['maintenance'] else 'معطّل ❌'}", show_alert=True)
    elif data == "adm_broadcast" and is_admin(uid):
        if not has_permission(uid, "broadcast"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية الإذاعة.", show_alert=True); return
        db["pending_actions"][ustr] = {"action": "broadcast"}; sync_db(); bot.send_message(cid, "📢 أرسل رسالة الإذاعة:")
    elif data == "adm_welcome" and is_admin(uid):
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("✏️ تعديل الترحيب", callback_data="welcome_edit_start"))
        if db["config"].get("welcome_rich") or db["config"].get("welcome_msg"):
            m.add(types.InlineKeyboardButton("🗑️ حذف الترحيب", callback_data="welcome_delete"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit("👋 *إدارة رسالة الترحيب*\n\nترسل بشكل مستقل تمامًا قبل رسالة القائمة، وما تأثر على شكل الأزرار.", m)
    elif data == "welcome_edit_start" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "edit_welcome"}; sync_db()
        bot.send_message(cid, "✏️ أرسل الآن رسالة الترحيب:\n\n- نص فقط، أو\n- صورة (مع نص كابشن اختياري)، أو\n- فيديو (مع نص كابشن اختياري)\n\nما ترسله الآن هو اللي راح يظهر بالضبط، بشكل مستقل عن باقي الميزات.")
    elif data == "welcome_delete" and is_admin(uid):
        db["config"]["welcome_rich"] = None
        db["config"]["welcome_msg"] = ""
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم حذف رسالة الترحيب.", show_alert=True)
    elif data == "adm_watch_label" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "edit_watch_label"}; sync_db()
        current = db["config"].get("novel_watch_label", "👁️ شاهد الآن")
        bot.send_message(cid, f"👁️ النص الحالي: {current}\n\nأرسل النص الجديد لزر \"شاهد الآن\":")
    elif data == "adm_money" and is_admin(uid):
        s = db["config"]["sub_settings"]
        m = types.InlineKeyboardMarkup().add(
            types.InlineKeyboardButton("💳 إدارة خطط الاشتراك والأسعار", callback_data="adm_subs"),
            types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(f"💰 *المالية:*\n\nإجمالي إيرادات الاشتراكات: *{s.get('total_sub_revenue',0)}{s.get('currency','$')}*\n\n"
             f"لإدارة الأسعار والمدد، استخدم «إدارة خطط الاشتراك».", m)
    
    elif data == "adm_users" and is_admin(uid):
        m = types.InlineKeyboardMarkup(row_width=2).add(types.InlineKeyboardButton("🚫 حظر", callback_data="adm_ban"), types.InlineKeyboardButton("✅ رفع حظر", callback_data="adm_unban"), types.InlineKeyboardButton("👁️ عرض مستخدم", callback_data="adm_view_user"), types.InlineKeyboardButton("➕ مشرف", callback_data="adm_add_admin"), types.InlineKeyboardButton("✏️ تعديل صلاحيات مشرف", callback_data="adm_edit_admin"), types.InlineKeyboardButton("➖ إزالة مشرف", callback_data="adm_rem_admin"), types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit("👤 *إدارة المستخدمين:*", m)
    elif data == "adm_ban" and is_admin(uid):
        if not has_permission(uid, "ban_users"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية حظر المستخدمين.", show_alert=True); return
        db["pending_actions"][ustr] = {"action": "ban"}; sync_db(); bot.send_message(cid, "🚫 أرسل ID المستخدم للحظر:")
    elif data == "adm_unban" and is_admin(uid):
        if not has_permission(uid, "ban_users"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية رفع الحظر.", show_alert=True); return
        db["pending_actions"][ustr] = {"action": "unban"}; sync_db(); bot.send_message(cid, "✅ أرسل ID المستخدم لرفع الحظر:")
    elif data == "adm_view_user" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "view_user"}; sync_db(); bot.send_message(cid, "👤 أرسل ID المستخدم:")
    elif data == "adm_add_admin" and is_owner(uid):
        db["pending_actions"][ustr] = {"action": "add_admin_id"}; sync_db(); bot.send_message(cid, "➕ أرسل ID المشرف الجديد:")
    elif data.startswith("setexp_none_") and is_owner(uid):
        target = data[12:]
        if target in db["config"]["admins"] and isinstance(db["config"]["admins"][target], dict):
            db["config"]["admins"][target]["expires_at"] = None
            sync_db()
        bot.answer_callback_query(call.id, "✅ صار غير محدود.", show_alert=True)
    elif data.startswith("setexp_custom_") and is_owner(uid):
        target = data[14:]
        db["pending_actions"][ustr] = {"action": "set_admin_expiry_days", "target": target}
        sync_db()
        bot.send_message(cid, "🔢 أرسل عدد الأيام الجديد (تبدأ من الآن):")
    elif data == "adm_edit_admin" and is_owner(uid):
        m = types.InlineKeyboardMarkup()
        for aid, arec in db["config"]["admins"].items():
            aname = arec if isinstance(arec, str) else arec.get("title", aid)
            m.add(types.InlineKeyboardButton(f"✏️ {aname}", callback_data=f"editadm_{aid}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_users")); edit("✏️ *اختر مشرف لتعديل صلاحياته:*", m)
    elif data.startswith("editadm_") and is_owner(uid):
        target = data[8:]
        arec = db["config"]["admins"].get(target)
        if not arec:
            bot.answer_callback_query(call.id, "❌ غير موجود.", show_alert=True); return
        title = arec if isinstance(arec, str) else arec.get("title", target)
        current_perms = [] if isinstance(arec, str) else [k for k, v in arec.get("permissions", {}).items() if v]
        db["pending_actions"][ustr] = {"action": "add_admin_perms", "target": target, "title": title, "selected_perms": current_perms, "is_edit": True}
        sync_db()
        edit(f"🔐 *تعديل صلاحيات: {title}*\n\nاختر/ألغِ الصلاحيات، ثم اضغط متابعة:", admin_perm_picker_kb(current_perms))
    elif data == "adm_rem_admin" and is_owner(uid):
        m = types.InlineKeyboardMarkup()
        for aid, arec in db["config"]["admins"].items():
            aname = arec if isinstance(arec, str) else arec.get("title", aid)
            m.add(types.InlineKeyboardButton(f"❌ {aname}", callback_data=f"rmadm_{aid}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_users")); edit("➖ *إزالة مشرف:*", m)
    elif data.startswith("rmadm_") and is_owner(uid):
        db["config"]["admins"].pop(data[6:], None); sync_db(); bot.answer_callback_query(call.id, "✅ تم إزالة المشرف."); edit("➖ *إزالة مشرف:*", admin_main_kb())

    elif data.startswith("newadm_perm_") and is_owner(uid):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "add_admin_perms": return
        perm = data[12:]
        sel = list(pending.get("selected_perms", []))
        if perm in sel: sel.remove(perm)
        else: sel.append(perm)
        pending["selected_perms"] = sel
        db["pending_actions"][ustr] = pending; sync_db()
        edit("🔐 اختر صلاحيات هذا المشرف (تقدر تختار أكثر من وحدة، ثم اضغط متابعة):", admin_perm_picker_kb(sel))
    elif data == "newadm_perms_done" and is_owner(uid):
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "add_admin_perms" or not pending.get("selected_perms"):
            bot.answer_callback_query(call.id, "❌ اختر صلاحية واحدة على الأقل.", show_alert=True)
            return
        if pending.get("is_edit"):
            target, title, perms = pending["target"], pending["title"], pending["selected_perms"]
            arec = db["config"]["admins"].get(target, {})
            old_expiry = arec.get("expires_at") if isinstance(arec, dict) else None
            db["config"]["admins"][target] = {"title": title, "permissions": {p: True for p in perms}, "expires_at": old_expiry}
            db["pending_actions"].pop(ustr, None)
            sync_db()
            edit(f"✅ تم تحديث صلاحيات «{title}».\n\nتقدر تغيّر مدة انتهائه من نفس القائمة لو تبي.", admin_expiry_edit_kb(target))
            return
        pending["action"] = "add_admin_expiry"
        db["pending_actions"][ustr] = pending; sync_db()
        edit("⏳ *مدة صلاحية هذا المشرف:*\n\nلو حددت مدة، تنتهي صلاحياته تلقائيًا بعدها بدون أي إشعار له بذلك — وما راح يظهر له إنه مؤقت أصلًا.", admin_expiry_picker_kb())
    elif data == "newadm_expiry_none" and is_owner(uid):
        pending = db["pending_actions"].pop(ustr, {})
        sync_db()
        target, title, perms = pending["target"], pending["title"], pending["selected_perms"]
        perms_dict = {p: True for p in perms}
        db["config"]["admins"][target] = {"title": title, "permissions": perms_dict, "expires_at": None}
        sync_db()
        bot.send_message(cid, f"✅ تم إضافة المشرف «{title}» (ID: {target}) بدون مدة انتهاء.")
    elif data == "newadm_expiry_custom" and is_owner(uid):
        pending = db["pending_actions"].get(ustr, {})
        pending["action"] = "add_admin_expiry_days"
        db["pending_actions"][ustr] = pending; sync_db()
        bot.send_message(cid, "🔢 أرسل عدد الأيام اللي تبي صلاحياته تسري خلالها (مثال: 30):")

    elif data == "adm_codes" and is_admin(uid):
        if not has_permission(uid, "manage_codes"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية إدارة الأكواد.", show_alert=True); return
        m = types.InlineKeyboardMarkup().add(
            types.InlineKeyboardButton("➕ كود جديد (عشوائي)", callback_data="adm_gen_code"),
            types.InlineKeyboardButton("✏️ كود مخصص (تسميه بنفسك)", callback_data="adm_custom_code")
        )
        m.add(types.InlineKeyboardButton("📋 عرض الأكواد", callback_data="adm_list_codes"), types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit("🎟️ *أكواد النقاط:*", m)
    elif data == "adm_gen_code" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "gen_code"}; sync_db(); bot.send_message(cid, "🎟️ أرسل عدد النقاط للكود:")
    elif data == "adm_custom_code" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "custom_code_name"}; sync_db()
        bot.send_message(cid, "✏️ أرسل نص/اسم الكود اللي تبيه (مثال: WELCOME2026):")
    elif data == "adm_list_codes" and is_admin(uid):
        codes = db.get("codes", {}); text = "🎟️ *الأكواد:*\n\n" + "\n".join([f"{'✅' if i.get('used') else '🟢'} `{c}` — {i['points']} نقطة" for c, i in list(codes.items())[:20]]) if codes else "🎟️ لا توجد أكواد."
        edit(text, types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_codes")))

    elif data == "adm_channels" and is_admin(uid):
        chs = db["config"]["mandatory_channels"]; m = types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("➕ إضافة قناة", callback_data="adm_add_ch"))
        for ch in chs: m.add(types.InlineKeyboardButton(f"❌ حذف {ch}", callback_data=f"adm_delch_{ch}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit("📢 *القنوات الإجبارية:*\n" + ("\n".join(chs) if chs else "لا توجد قنوات"), m)
    elif data == "adm_add_ch" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "add_channel"}; sync_db(); bot.send_message(cid, "📢 أرسل معرف القناة (مثال: @mychannel):")
    elif data.startswith("adm_delch_") and is_admin(uid):
        ch = data[10:]; db["config"]["mandatory_channels"].remove(ch) if ch in db["config"]["mandatory_channels"] else None; sync_db(); bot.answer_callback_query(call.id, f"✅ تم حذف {ch}"); edit("📢 *القنوات الإجبارية:*", admin_main_kb())

    elif data.startswith("add_cat_") and is_admin(uid):
        if not has_permission(uid, "manage_categories"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية إدارة الأقسام.", show_alert=True); return
        db["pending_actions"][ustr] = {"action": "addcat", "parent": data[8:]}; sync_db(); bot.send_message(cid, "📁 أرسل اسم القسم:")
    elif data.startswith("add_item_") and is_admin(uid):
        if not has_permission(uid, "add_item"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية إضافة ملفات.", show_alert=True); return
        db["pending_actions"][ustr] = {"action": "additem_tag", "cat": data[9:], "selected": []}; sync_db(); bot.send_message(cid, "🏷️ اختر تصنيف واحد أو أكثر (بدون حد أقصى)، ثم اضغط متابعة:", reply_markup=build_tag_search_menu(selected=[], mode="assign"))
    elif data.startswith("del_cat_") and is_admin(uid):
        edit("⚠️ هل تريد حذف هذا القسم نهائياً؟", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("✅ تأكيد", callback_data=f"confirm_delcat_{data[8:]}"), types.InlineKeyboardButton("❌ إلغاء", callback_data=f"nav_{data[8:]}")))
    elif data.startswith("confirm_delcat_") and is_admin(uid):
        cat_id = data[15:]; cat = db["categories"].get(cat_id)
        if cat and cat.get("parent"):
            p = db["categories"].get(cat["parent"])
            if p and cat_id in p.get("children",[]): p["children"].remove(cat_id)
            # تنظيف الملفات والروايات التابعة للقسم قبل حذفه (تفادي بيانات يتيمة)
            for iid in cat.get("items", []): db["items"].pop(iid, None)
            for nid in cat.get("novels", []):
                db["novels"].pop(nid, None)
                for u in db["users"].values(): u.get("novel_progress", {}).pop(nid, None)
            db["config"].get("category_layouts", {}).pop(cat_id, None)
            db["categories"].pop(cat_id, None); sync_db()
        edit("✅ تم حذف القسم.", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home")))
    elif data.startswith("del_item_") and is_admin(uid):
        iid = data[9:]; [c.get("items",[]).remove(iid) for c in db["categories"].values() if iid in c.get("items",[])]; db["items"].pop(iid, None); sync_db(); bot.answer_callback_query(call.id, "✅ تم حذف الملف.")

    # ── Category Layout (شكل عرض الأقسام الفرعية) ───────────────────────────
    elif data.startswith("catlayout_") and is_admin(uid):
        cat_id = data[10:]
        edit("📐 *اختر شكل عرض الأقسام الفرعية:*", category_layout_picker_kb(cat_id))
    elif data.startswith("setcatlayout_") and is_admin(uid):
        parts = data.split("_", 2); cat_id = parts[1]; layout = parts[2]
        db["config"].setdefault("category_layouts", {})[cat_id] = layout
        sync_db(); bot.answer_callback_query(call.id, f"✅ تم اختيار: {layout}")
        cat = db["categories"].get(cat_id)
        edit(f"📁 *قسم: {cat['name']}*" if cat else "📁", build_category_menu(cat_id, uid))

    # ── Novel Creation Flow (إضافة رواية) ───────────────────────────────────
    elif data.startswith("add_novel_") and is_admin(uid):
        if not has_permission(uid, "add_novel"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية إضافة روايات.", show_alert=True)
            return
        cat_id = data[10:]
        db["pending_actions"][ustr] = {"action": "addnovel_title", "cat": cat_id}
        sync_db(); bot.send_message(cid, "📖 أرسل اسم الرواية:")
    elif data.startswith("add_manga_") and is_admin(uid):
        if not has_permission(uid, "add_novel"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية إضافة أعمال.", show_alert=True)
            return
        cat_id = data[10:]
        db["pending_actions"][ustr] = {"action": "addmanga_title", "cat": cat_id}
        sync_db(); bot.send_message(cid, "🎨 أرسل اسم العمل (مانهوا/مانجا):")
    elif data.startswith("add_series_") and is_admin(uid):
        if not has_permission(uid, "add_novel"):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية إضافة أعمال.", show_alert=True)
            return
        cat_id = data[11:]
        db["pending_actions"][ustr] = {"action": "addseries_title", "cat": cat_id}
        sync_db(); bot.send_message(cid, "🎬 أرسل اسم العمل (مسلسل/فيلم):")

    # ── Open / View Novel ────────────────────────────────────────────────────
    elif data.startswith("open_novel_"):
        nid = data[11:]
        open_novel(uid, cid, mid, nid)
    elif data.startswith("novel_admin_") and is_admin(uid):
        nid = data[12:]
        novel = get_novel(nid)
        if not can_manage_novel(uid, novel):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية إدارة هذي الرواية.", show_alert=True)
            return
        label = md_safe(novel["title"]) if novel else nid
        edit(f"⚙️ *إدارة الرواية: {label}*", novel_admin_kb(nid))
    elif data.startswith("novel_watch_"):
        nid = data[12:]
        ensure_user(call.message)
        edit("👁️ *اختر من فين تبدأ:*", novel_watch_kb(nid, ustr))
    elif data.startswith("comments_"):
        key = data[9:]
        back_cb = f"open_novel_{key[6:]}" if "_ch_" not in key else f"novel_ch_{key.split('_ch_')[0][6:]}_{key.split('_ch_')[1]}"
        comments = get_comments(key)
        text = f"💬 *التعليقات ({len(comments)}):*" if comments else "💬 *لا توجد تعليقات بعد، كن أول من يعلّق!*"
        edit(text, comments_view_kb(key, back_cb, uid))
    elif data.startswith("addcomment_"):
        key = data[11:]
        db["pending_actions"][ustr] = {"action": "write_comment", "key": key}; sync_db()
        bot.send_message(cid, "✍️ اكتب تعليقك:")
    elif data.startswith("delcomment_") and is_admin(uid):
        rest = data[11:]
        key, comment_id = rest.rsplit("_", 1)
        delete_comment(key, comment_id)
        bot.answer_callback_query(call.id, "🗑️ تم حذف التعليق.")
        back_cb = f"open_novel_{key[6:]}" if "_ch_" not in key else f"novel_ch_{key.split('_ch_')[0][6:]}_{key.split('_ch_')[1]}"
        comments = get_comments(key)
        text = f"💬 *التعليقات ({len(comments)}):*" if comments else "💬 *لا توجد تعليقات بعد، كن أول من يعلّق!*"
        edit(text, comments_view_kb(key, back_cb, uid))
    elif data.startswith("chgrid_"):
        rest = data[7:]
        nid, page_str = rest.rsplit("_", 1)
        page = int(page_str)
        novel = get_novel(nid)
        title = md_safe(novel["title"]) if novel else ""
        edit(f"🔢 *{title} — كل الفصول:*\n\nاضغط على أي رقم فصل للانتقال له مباشرة.", novel_chapter_grid_kb(nid, page))
    elif data.startswith("novel_ch_"):
        rest = data[9:]  # nid_chapternum
        parts = rest.rsplit("_", 1)
        nid = parts[0]; ch_str = parts[1]
        if ch_str == "None":
            bot.answer_callback_query(call.id, "❌ لا يوجد فصل.")
        else:
            ch_num = int(ch_str)
            set_user_novel_progress(ustr, nid, ch_num)
            novel = get_novel(nid)
            title = md_safe(novel["title"]) if novel else ""
            edit(f"📖 *{title}* — الفصل {ch_num}\n\nاختر صيغة التحميل أو تصفح فصل آخر:", novel_chapter_nav_kb(nid, ch_num))
    elif data.startswith("novel_merged_"):
        nid = data[13:]
        novel = get_novel(nid)
        if novel and novel_merge_groups(novel):
            edit("📦 *اختر مجموعة الفصول المدمجة:*", novel_merged_list_kb(nid))
        else:
            bot.answer_callback_query(call.id, "❌ لا توجد مجموعات مدمجة بعد.", show_alert=True)
    elif data.startswith("novel_mopen_"):
        rest = data[12:]; nid, group = rest.split("_", 1)
        edit(f"📦 *الفصول {group}* — اختر الصيغة:", novel_merged_format_kb(nid, group))
    elif data.startswith("novel_get_"):
        # صيغة البيانات: novel_get_{nid}_ch_{num}_{fmt}  أو  novel_get_{nid}_mg_{group}_{fmt}
        rest = data[10:]
        if "_ch_" in rest:
            marker_idx = rest.index("_ch_")
            nid = rest[:marker_idx]
            remainder = rest[marker_idx + 4:]  # "{num}_{fmt}"
            ch_num_str, fmt = remainder.rsplit("_", 1)
            ok = send_chapter_file(cid, nid, int(ch_num_str), fmt)
            if ok: bot.answer_callback_query(call.id, "✅ جاري الإرسال...")
        elif "_mg_" in rest:
            marker_idx = rest.index("_mg_")
            nid = rest[:marker_idx]
            remainder = rest[marker_idx + 4:]  # "{group}_{fmt}"
            group, fmt = remainder.rsplit("_", 1)
            allowed, used, limit = can_download_merge(uid)
            if not allowed:
                bot.answer_callback_query(
                    call.id,
                    f"❌ وصلت الحد اليومي للتنزيلات المدمجة ({used}/{limit}). "
                    f"{'اشترك للحصول على حد أعلى أو غير محدود.' if not check_sub(uid) else 'حاول مرة ثانية غدًا.'}",
                    show_alert=True)
            else:
                ok = send_merged_file(cid, nid, group, fmt)
                if ok:
                    record_merge_download(ustr)
                    bot.answer_callback_query(call.id, "✅ جاري الإرسال...")
        else:
            bot.answer_callback_query(call.id, "❌ طلب غير صالح.")

    # ── Novel Admin Actions ──────────────────────────────────────────────────
    elif data.startswith("novel_addch_") and is_admin(uid):
        nid = data[12:]
        novel = get_novel(nid)
        if not can_manage_novel(uid, novel):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية إضافة فصول لهذي الرواية.", show_alert=True)
            return
        db["pending_actions"][ustr] = {"action": "addnovel_ch_number", "nid": nid}
        sync_db(); bot.send_message(cid, "🔢 أرسل رقم الفصل:")
    elif data.startswith("novel_delch_") and is_admin(uid):
        nid = data[12:]
        novel = get_novel(nid)
        if not can_delete_chapter(uid, novel):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية حذف فصول من هذي الرواية.", show_alert=True)
            return
        chs = novel_chapter_numbers(novel) if novel else []
        if not chs:
            bot.answer_callback_query(call.id, "❌ لا توجد فصول لحذفها.", show_alert=True)
        else:
            db["pending_actions"][ustr] = {"action": "delnovel_ch_number", "nid": nid}
            sync_db()
            nums = "، ".join(str(n) for n in chs)
            bot.send_message(cid, f"🔢 أرسل رقم الفصل يلي تبي تحذفه.\n\nالفصول الموجودة: {nums}")
    elif data.startswith("confirm_delch_") and is_admin(uid):
        rest = data[14:]
        nid, ch_str = rest.rsplit("_", 1)
        novel = get_novel(nid)
        if not can_delete_chapter(uid, novel):
            bot.answer_callback_query(call.id, "❌ ما عندك صلاحية حذف فصول من هذي الرواية.", show_alert=True)
            return
        if novel and str(ch_str) in novel.get("chapters", {}):
            novel["chapters"].pop(str(ch_str), None)
            for u in db["users"].values():
                if u.get("novel_progress", {}).get(nid) == int(ch_str):
                    u["novel_progress"].pop(nid, None)
            sync_db()
            edit(f"✅ تم حذف الفصل {ch_str}.", novel_admin_kb(nid))
        else:
            bot.answer_callback_query(call.id, "❌ الفصل غير موجود.", show_alert=True)
    elif data.startswith("novel_mergepick_") and is_admin(uid):
        nid = data[16:]
        edit("📦 *دمج فصول موجودة*\n\nاختر طريقة تحديد الفصول اللي تبي تدمجها بملف واحد:", novel_merge_pick_kb(nid))
    elif data.startswith("novel_mergerange_") and is_admin(uid):
        nid = data[18:]
        novel = get_novel(nid)
        suggested = next_merge_range(novel) if novel else None
        db["pending_actions"][ustr] = {"action": "addnovel_merge_range", "nid": nid, "suggested": suggested}
        sync_db()
        hint = f"\n\n💡 المدى المقترح: {suggested}" if suggested else ""
        bot.send_message(cid, f"🔢 أرسل مدى الفصول (مثال: 1-25) — بيتم تجهيزها تلقائيًا من الفصول المرفوعة أصلاً:{hint}")
    elif data.startswith("novel_mergemanual_") and is_admin(uid):
        rest = data[19:]
        nid, page_str = rest.rsplit("_", 1)
        page = int(page_str)
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") == "addnovel_merge_manual" and pending.get("nid") == nid:
            selected = pending.get("selected", [])
        else:
            selected = []
            db["pending_actions"][ustr] = {"action": "addnovel_merge_manual", "nid": nid, "selected": selected}
            sync_db()
        edit("✅ *اختيار فصول محددة للدمج*\n\nاضغط على أرقام الفصول اللي تبيها (تقدر تختار أكثر من رقم، مو لازم متسلسلة)، وبعدها اضغط «تم، ابدأ الدمج».",
             build_novel_merge_manual_kb(nid, selected, page))
    elif data.startswith("novel_mergetoggle_") and is_admin(uid):
        rest = data[19:]
        nid, page_str, ch_str = rest.split("_", 2)
        page = int(page_str); ch_num = int(ch_str)
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "addnovel_merge_manual" or pending.get("nid") != nid:
            pending = {"action": "addnovel_merge_manual", "nid": nid, "selected": []}
        selected = pending.get("selected", [])
        if ch_num in selected:
            selected.remove(ch_num)
        else:
            selected.append(ch_num)
        pending["selected"] = selected
        db["pending_actions"][ustr] = pending
        sync_db()
        edit("✅ *اختيار فصول محددة للدمج*\n\nاضغط على أرقام الفصول اللي تبيها (تقدر تختار أكثر من رقم، مو لازم متسلسلة)، وبعدها اضغط «تم، ابدأ الدمج».",
             build_novel_merge_manual_kb(nid, selected, page))
    elif data.startswith("novel_mergeconfirm_") and is_admin(uid):
        nid = data[20:]
        pending = db["pending_actions"].get(ustr, {})
        selected = pending.get("selected", []) if pending.get("action") == "addnovel_merge_manual" and pending.get("nid") == nid else []
        db["pending_actions"].pop(ustr, None); sync_db()
        if not selected:
            bot.send_message(cid, "❌ ما فيه فصول محددة."); return
        bot.send_message(cid, f"⏳ جاري تجهيز ودمج {len(selected)} فصل بالخلفية، راح أرسل لك النتيجة هنا أول ما يخلص...")
        def _do_merge_job():
            ok, msg = merge_existing_chapters(nid, selected, cid)
            bot.send_message(cid, msg, parse_mode="Markdown" if ok else None)
        threading.Thread(target=_do_merge_job, daemon=True).start()
    elif data.startswith("novel_link_") and is_admin(uid):
        nid = data[11:]
        db["pending_actions"][ustr] = {"action": "addnovel_link", "nid": nid}
        sync_db(); bot.send_message(cid, "🔗 أرسل معرف القناة/المجموعة (مثال: @mychannel) أو أرسل 'إلغاء' لإلغاء الربط:")
    elif data.startswith("novel_del_") and is_admin(uid):
        nid = data[10:]
        edit("⚠️ هل تريد حذف هذه الرواية نهائياً؟ (كل الفصول ستُحذف)",
             types.InlineKeyboardMarkup().add(
                 types.InlineKeyboardButton("✅ تأكيد", callback_data=f"confirm_delnovel_{nid}"),
                 types.InlineKeyboardButton("❌ إلغاء", callback_data=f"open_novel_{nid}")))
    elif data.startswith("confirm_delnovel_") and is_admin(uid):
        nid = data[17:]
        novel = db["novels"].pop(nid, None)
        if novel:
            cat = db["categories"].get(novel.get("category"))
            if cat and nid in cat.get("novels", []): cat["novels"].remove(nid)
            for u in db["users"].values(): u.get("novel_progress", {}).pop(nid, None)
        sync_db()
        edit("✅ تم حذف الرواية.", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home")))

    # ── Manga/Manhwa: view & navigate ────────────────────────────────────────
    elif data.startswith("open_manga_"):
        mid = data[11:]
        open_manga(uid, cid, mid, mid)
    elif data.startswith("manga_admin_") and is_admin(uid):
        mid = data[12:]
        manga = get_manga(mid)
        label = md_safe(manga["title"]) if manga else mid
        edit(f"⚙️ *إدارة العمل: {label}*", manga_admin_kb(mid))
    elif data.startswith("manga_watch_"):
        mid = data[12:]
        ensure_user(call.message)
        edit("👁️ *اختر من فين تبدأ:*", manga_watch_kb(mid, ustr))
    elif data.startswith("mangagrid_"):
        rest = data[10:]
        mid, page_str = rest.rsplit("_", 1)
        manga = get_manga(mid)
        title = md_safe(manga["title"]) if manga else ""
        edit(f"🔢 *{title} — كل الفصول:*\n\nاضغط على أي رقم فصل للانتقال له مباشرة.", manga_chapter_grid_kb(mid, int(page_str)))
    elif data.startswith("manga_ch_"):
        rest = data[9:]
        parts = rest.rsplit("_", 1)
        mid = parts[0]; ch_str = parts[1]
        if ch_str == "None":
            bot.answer_callback_query(call.id, "❌ لا يوجد فصل.")
        else:
            ch_num = int(ch_str)
            set_user_manga_progress(ustr, mid, ch_num)
            manga = get_manga(mid)
            title = md_safe(manga["title"]) if manga else ""
            edit(f"🎨 *{title}* — الفصل {ch_num}\n\nاختر صيغة التحميل أو تصفح فصل آخر:", manga_chapter_nav_kb(mid, ch_num))
    elif data.startswith("manga_get_"):
        rest = data[10:]
        if "_ch_" in rest:
            marker_idx = rest.index("_ch_")
            mid = rest[:marker_idx]
            remainder = rest[marker_idx + 4:]
            ch_num_str, fmt = remainder.rsplit("_", 1)
            ok = send_manga_chapter_file(cid, mid, int(ch_num_str), fmt)
            if ok: bot.answer_callback_query(call.id, "✅ جاري الإرسال...")
        elif "_mg_" in rest:
            marker_idx = rest.index("_mg_")
            mid = rest[:marker_idx]
            remainder = rest[marker_idx + 4:]
            group, fmt = remainder.rsplit("_", 1)
            allowed, used, limit = can_download_merge(uid)
            if not allowed:
                bot.answer_callback_query(call.id, f"❌ وصلت الحد اليومي للتنزيلات المدمجة ({used}/{limit}).", show_alert=True)
            else:
                ok = send_manga_merged_file(cid, mid, group, fmt)
                if ok:
                    record_merge_download(ustr)
                    bot.answer_callback_query(call.id, "✅ جاري الإرسال...")
        else:
            bot.answer_callback_query(call.id, "❌ طلب غير صالح.")

    # ── Manga Admin Actions ──────────────────────────────────────────────────
    elif data.startswith("manga_addch_") and is_admin(uid):
        mid = data[12:]
        db["pending_actions"][ustr] = {"action": "addmanga_ch_number", "mid": mid}
        sync_db(); bot.send_message(cid, "🔢 أرسل رقم الفصل:")
    elif data.startswith("manga_delch_") and is_admin(uid):
        mid = data[12:]
        manga = get_manga(mid)
        chs = manga_chapter_numbers(manga) if manga else []
        if not chs:
            bot.answer_callback_query(call.id, "❌ لا توجد فصول لحذفها.", show_alert=True)
        else:
            db["pending_actions"][ustr] = {"action": "delmanga_ch_number", "mid": mid}
            sync_db()
            nums = "، ".join(str(n) for n in chs)
            bot.send_message(cid, f"🔢 أرسل رقم الفصل يلي تبي تحذفه.\n\nالفصول الموجودة: {nums}")
    elif data.startswith("confirm_delmangach_") and is_admin(uid):
        rest = data[19:]
        mid, ch_str = rest.rsplit("_", 1)
        manga = get_manga(mid)
        if manga and str(ch_str) in manga.get("chapters", {}):
            manga["chapters"].pop(str(ch_str), None)
            for u in db["users"].values():
                if u.get("manga_progress", {}).get(mid) == int(ch_str):
                    u["manga_progress"].pop(mid, None)
            sync_db()
            edit(f"✅ تم حذف الفصل {ch_str}.", manga_admin_kb(mid))
        else:
            bot.answer_callback_query(call.id, "❌ الفصل غير موجود.", show_alert=True)
    elif data.startswith("manga_addmerge_") and is_admin(uid):
        mid = data[15:]
        manga = get_manga(mid)
        suggested = manga_next_merge_range(manga) if manga else None
        db["pending_actions"][ustr] = {"action": "addmanga_merge_range", "mid": mid, "suggested": suggested}
        sync_db()
        hint = f"\n\n💡 المدى المقترح: {suggested}" if suggested else ""
        bot.send_message(cid, f"🔢 أرسل مدى الفصول للمجموعة (مثال: 1-25):{hint}")
    elif data.startswith("manga_link_") and is_admin(uid):
        mid = data[11:]
        db["pending_actions"][ustr] = {"action": "addmanga_link", "mid": mid}
        sync_db(); bot.send_message(cid, "🔗 أرسل معرف القناة/المجموعة (مثال: @mychannel) أو أرسل 'إلغاء' لإلغاء الربط:")
    elif data.startswith("manga_del_") and is_admin(uid):
        mid = data[10:]
        edit("⚠️ هل تريد حذف هذا العمل نهائياً؟ (كل الفصول ستُحذف)",
             types.InlineKeyboardMarkup().add(
                 types.InlineKeyboardButton("✅ تأكيد", callback_data=f"confirm_delmanga_{mid}"),
                 types.InlineKeyboardButton("❌ إلغاء", callback_data=f"open_manga_{mid}")))
    elif data.startswith("confirm_delmanga_") and is_admin(uid):
        mid = data[18:]
        manga = db["manga"].pop(mid, None)
        if manga:
            cat = db["categories"].get(manga.get("category"))
            if cat and mid in cat.get("manga", []): cat["manga"].remove(mid)
            for u in db["users"].values(): u.get("manga_progress", {}).pop(mid, None)
        sync_db()
        edit("✅ تم حذف العمل.", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home")))

    # ── Series/Movies: view & navigate ───────────────────────────────────────
    elif data.startswith("open_series_"):
        sid = data[12:]
        open_series(uid, cid, sid, sid)
    elif data.startswith("series_admin_") and is_admin(uid):
        sid = data[13:]
        series = get_series(sid)
        label = md_safe(series["title"]) if series else sid
        edit(f"⚙️ *إدارة العمل: {label}*", series_admin_kb(sid))
    elif data.startswith("series_watch_"):
        sid = data[13:]
        ensure_user(call.message)
        edit("👁️ *اختر من فين تبدأ:*", series_watch_kb(sid, ustr))
    elif data.startswith("seriesgrid_"):
        rest = data[11:]
        sid, page_str = rest.rsplit("_", 1)
        series = get_series(sid)
        title = md_safe(series["title"]) if series else ""
        edit(f"🔢 *{title} — كل الحلقات:*\n\nاضغط على أي رقم حلقة للانتقال لها مباشرة.", series_episode_grid_kb(sid, int(page_str)))
    elif data.startswith("series_ep_"):
        rest = data[10:]
        parts = rest.rsplit("_", 1)
        sid = parts[0]; ep_str = parts[1]
        if ep_str == "None":
            bot.answer_callback_query(call.id, "❌ لا توجد حلقة.")
        else:
            ep_num = int(ep_str)
            set_user_series_progress(ustr, sid, ep_num)
            series = get_series(sid)
            title = md_safe(series["title"]) if series else ""
            edit(f"🎬 *{title}* — الحلقة {ep_num}\n\nاختر الجودة:", series_quality_kb(sid, ep_num, uid))
    elif data.startswith("series_get_"):
        rest = data[11:]
        sid, ep_str, quality = rest.split("_", 2)
        ok = send_series_episode(cid, sid, int(ep_str), quality, uid)
        if ok: bot.answer_callback_query(call.id, "✅ جاري الإرسال...")

    # ── Series Admin Actions ─────────────────────────────────────────────────
    elif data.startswith("series_addep_") and is_admin(uid):
        sid = data[13:]
        db["pending_actions"][ustr] = {"action": "addseries_ep_number", "sid": sid}
        sync_db(); bot.send_message(cid, "🔢 أرسل رقم الحلقة:")
    elif data.startswith("epquality_") and is_admin(uid):
        quality = data[10:]
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "addseries_ep_quality":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة، ابدأ من جديد.", show_alert=True); return
        m = types.InlineKeyboardMarkup(row_width=2)
        m.add(
            types.InlineKeyboardButton("🆓 مجانية", callback_data=f"epaccess_{quality}_free"),
            types.InlineKeyboardButton("🔒 مشتركين فقط", callback_data=f"epaccess_{quality}_sub")
        )
        db["pending_actions"][ustr] = {**pending, "quality": quality}
        sync_db()
        edit(f"🎞️ الجودة: *{quality}*\n\nهل هذي الجودة مجانية لكل المستخدمين، أو مقصورة على المشتركين؟", m)
    elif data.startswith("epaccess_") and is_admin(uid):
        rest = data[9:]
        quality, access = rest.rsplit("_", 1)
        pending = db["pending_actions"].get(ustr, {})
        if pending.get("action") != "addseries_ep_quality":
            bot.answer_callback_query(call.id, "❌ انتهت الجلسة، ابدأ من جديد.", show_alert=True); return
        db["pending_actions"][ustr] = {"action": "addseries_ep_file", "sid": pending["sid"], "ep_num": pending["ep_num"],
                                        "quality": quality, "sub_only": (access == "sub")}
        sync_db()
        bot.send_message(cid, f"📥 أرسل ملف الفيديو الآن (جودة {quality}):")
    elif data.startswith("series_delep_") and is_admin(uid):
        sid = data[13:]
        series = get_series(sid)
        eps = series_episode_numbers(series) if series else []
        if not eps:
            bot.answer_callback_query(call.id, "❌ لا توجد حلقات لحذفها.", show_alert=True)
        else:
            db["pending_actions"][ustr] = {"action": "delseries_ep_number", "sid": sid}
            sync_db()
            nums = "، ".join(str(n) for n in eps)
            bot.send_message(cid, f"🔢 أرسل رقم الحلقة يلي تبي تحذفها.\n\nالحلقات الموجودة: {nums}")
    elif data.startswith("confirm_delep_") and is_admin(uid):
        rest = data[14:]
        sid, ep_str = rest.rsplit("_", 1)
        series = get_series(sid)
        if series and str(ep_str) in series.get("episodes", {}):
            series["episodes"].pop(str(ep_str), None)
            for u in db["users"].values():
                if u.get("series_progress", {}).get(sid) == int(ep_str):
                    u["series_progress"].pop(sid, None)
            sync_db()
            edit(f"✅ تم حذف الحلقة {ep_str}.", series_admin_kb(sid))
        else:
            bot.answer_callback_query(call.id, "❌ الحلقة غير موجودة.", show_alert=True)
    elif data.startswith("series_link_") and is_admin(uid):
        sid = data[12:]
        db["pending_actions"][ustr] = {"action": "addseries_link", "sid": sid}
        sync_db(); bot.send_message(cid, "🔗 أرسل معرف القناة/المجموعة (مثال: @mychannel) أو أرسل 'إلغاء' لإلغاء الربط:")
    elif data.startswith("series_del_") and is_admin(uid):
        sid = data[11:]
        edit("⚠️ هل تريد حذف هذا العمل نهائياً؟ (كل الحلقات ستُحذف)",
             types.InlineKeyboardMarkup().add(
                 types.InlineKeyboardButton("✅ تأكيد", callback_data=f"confirm_delseries_{sid}"),
                 types.InlineKeyboardButton("❌ إلغاء", callback_data=f"open_series_{sid}")))
    elif data.startswith("confirm_delseries_") and is_admin(uid):
        sid = data[19:]
        series = db["series"].pop(sid, None)
        if series:
            cat = db["categories"].get(series.get("category"))
            if cat and sid in cat.get("series", []): cat["series"].remove(sid)
            for u in db["users"].values(): u.get("series_progress", {}).pop(sid, None)
        sync_db()
        edit("✅ تم حذف العمل.", types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🏠 رئيسية", callback_data="go_home")))

    # ── Menu & Features Editor ──────────────────────────────────────────────
    elif data == "adm_tags" and is_admin(uid): edit("🏷️ *إدارة التصنيفات:*", tags_admin_kb())
    elif data == "adm_archive_channels" and is_admin(uid):
        ac = db["config"]["archive_channels"]
        allowed = db["config"].get("archive_upload_allowed_users", [])
        text = (f"🗄️ *قنوات الأرشيف والنسخ الاحتياطية*\n\n"
                f"كل ملف يُرفع بالبوت ينسخ تلقائيًا لقناة الأرشيف المناسبة، وأي ملف "
                f"يُرفع مباشرة بقناة الأرشيف يقدر يُربط بفصل/حلقة بالبوت.\n\n"
                f"📚 روايات: `{ac.get('novels') or 'غير مربوطة'}`\n"
                f"🎨 مانجا/مانهوا: `{ac.get('manga') or 'غير مربوطة'}`\n"
                f"🎬 مسلسلات/أفلام: `{ac.get('series') or 'غير مربوطة'}`\n\n"
                f"👥 المسموح لهم بالرفع المباشر بالقناة (غير المطور): {len(allowed)}")
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("🔗 ربط قناة الروايات", callback_data="archch_set_novels"))
        m.add(types.InlineKeyboardButton("🔗 ربط قناة المانجا", callback_data="archch_set_manga"))
        m.add(types.InlineKeyboardButton("🔗 ربط قناة المسلسلات", callback_data="archch_set_series"))
        m.add(types.InlineKeyboardButton("👥 إدارة المسموح لهم بالرفع المباشر", callback_data="archch_manage_users"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(text, m)
    elif data.startswith("archch_set_") and is_admin(uid):
        ctype = data[11:]
        db["pending_actions"][ustr] = {"action": "set_archive_channel", "ctype": ctype}
        sync_db()
        bot.send_message(cid, "🔗 أرسل يوزرنيم القناة (مثل @MyArchive) أو آيدي القناة، والبوت لازم يكون مشرف فيها:")
    elif data == "archch_manage_users" and is_admin(uid):
        allowed = db["config"].get("archive_upload_allowed_users", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for aid in allowed:
            m.add(types.InlineKeyboardButton(f"❌ إزالة {aid}", callback_data=f"archch_remove_user_{aid}"))
        m.add(types.InlineKeyboardButton("➕ إضافة مستخدم", callback_data="archch_add_user"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_archive_channels"))
        edit("👥 *المسموح لهم بالرفع المباشر بقنوات الأرشيف* (غير المطور):", m)
    elif data == "archch_add_user" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "archch_add_user_id"}; sync_db()
        bot.send_message(cid, "🆔 أرسل آيدي المستخدم اللي تبي تسمح له بالرفع المباشر لقنوات الأرشيف:")
    elif data.startswith("archch_remove_user_") and is_admin(uid):
        target = data[20:]
        allowed = db["config"].get("archive_upload_allowed_users", [])
        if target in allowed: allowed.remove(target); sync_db()
        bot.answer_callback_query(call.id, "✅ تم الإزالة.", show_alert=True)
        edit("👥 *المسموح لهم بالرفع المباشر*", admin_main_kb())
    elif data == "adm_translate_access" and is_admin(uid):
        ts = db["config"]["translation_settings"]
        access = "🌍 عام للكل" if ts.get("public") else f"🔒 خاص — {len(ts.get('allowed_users', []))} مستخدم مسموح لهم"
        engine = ts.get("engine", "ai")
        engine_label = "🤖 ذكاء اصطناعي (دقيق، أدبي، يستهلك رصيد OpenRouter)" if engine == "ai" else "⚡ محلي (Google عبر deep-translator، فوري ومجاني)"
        text = (f"🌐 *صلاحيات الترجمة*\n\n"
                f"الحالة: {access}\n"
                f"المحرك الحالي: {engine_label}\n\n"
                f"المطور دايمًا يقدر يستخدم الترجمة بغض النظر عن هالإعدادات. لو خليتها "
                f"«خاص» بدون إضافة أي مستخدم، تظهر لك أنت بس.")
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton(f"{'🔒 خاص' if ts.get('public') else '🌍 عام'} — تبديل الوصول", callback_data="tr_toggle_public"))
        m.add(types.InlineKeyboardButton(f"🔄 تبديل المحرك ({'إلى محلي ⚡' if engine == 'ai' else 'إلى ذكاء اصطناعي 🤖'})", callback_data="tr_toggle_engine"))
        m.add(types.InlineKeyboardButton("👥 إدارة المسموح لهم", callback_data="tr_manage_allowed"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(text, m)
    elif data == "tr_toggle_engine" and is_admin(uid):
        ts = db["config"]["translation_settings"]
        new_engine = "local" if ts.get("engine", "ai") == "ai" else "ai"
        if new_engine == "local" and not _ensure_translation_libs():
            bot.answer_callback_query(call.id, f"❌ مكتبة الترجمة المحلية غير مثبتة على الاستضافة: {_TRANSLATE_IMPORT_ERROR}", show_alert=True)
        else:
            ts["engine"] = new_engine
            sync_db()
            bot.answer_callback_query(call.id, f"✅ تم التبديل إلى: {'⚡ محلي' if new_engine == 'local' else '🤖 ذكاء اصطناعي'}", show_alert=True)
        call.data = "adm_translate_access"; on_callback(call)
    elif data == "tr_toggle_public" and is_admin(uid):
        ts = db["config"]["translation_settings"]
        ts["public"] = not ts.get("public", False)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.", show_alert=True)
        edit("🌐 *صلاحيات الترجمة*", admin_main_kb())
    elif data == "tr_manage_allowed" and is_admin(uid):
        allowed = db["config"]["translation_settings"].get("allowed_users", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for a in allowed:
            m.add(types.InlineKeyboardButton(f"❌ إزالة {a}", callback_data=f"tr_remove_allowed_{a}"))
        m.add(types.InlineKeyboardButton("➕ إضافة مستخدم", callback_data="tr_add_allowed"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_translate_access"))
        edit(f"👥 *المسموح لهم استخدام الترجمة* ({len(allowed)}):", m)
    elif data == "tr_add_allowed" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "tr_add_allowed_id"}; sync_db()
        bot.send_message(cid, "🆔 أرسل آيدي المستخدم اللي تبي تسمح له باستخدام الترجمة:")
    elif data.startswith("tr_remove_allowed_") and is_admin(uid):
        target = data[18:]
        allowed = db["config"]["translation_settings"].get("allowed_users", [])
        if target in allowed: allowed.remove(target); sync_db()
        bot.answer_callback_query(call.id, "✅ تم الإزالة.", show_alert=True)
        m = types.InlineKeyboardMarkup(row_width=1)
        for a in allowed:
            m.add(types.InlineKeyboardButton(f"❌ إزالة {a}", callback_data=f"tr_remove_allowed_{a}"))
        m.add(types.InlineKeyboardButton("➕ إضافة مستخدم", callback_data="tr_add_allowed"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_translate_access"))
        edit("👥 *المسموح لهم استخدام الترجمة*:", m)
    elif data == "adm_scraper" and is_admin(uid):
        ss = db["config"]["scraper_settings"]
        profiles = ss.get("site_profiles", {})
        jobs = db["scrape_jobs"]
        active = [j for j in jobs.values() if j["status"] == "running"]
        status = "🟢 مفعّلة" if ss.get("enabled") else "🔴 معطّلة"
        access = "عام للكل" if ss.get("public") else f"{len(ss.get('allowed_users', []))} مستخدم مسموح لهم"
        text = (f"🕷️ *سحب الفصول التلقائي*\n\n"
                f"الحالة العامة: {status}\n"
                f"الوصول: {access}\n"
                f"ملفات تعريف المواقع المحفوظة: {len(profiles)}\n"
                f"مهام سحب نشطة الآن: {len(active)}\n\n"
                f"⚠️ استخدم هذي الميزة فقط لسحب محتوى تملك حق نشره أو مصادر مفتوحة/بإذن.")
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton(f"{'🔴 تعطيل' if ss.get('enabled') else '🟢 تفعيل'} الميزة", callback_data="scr_toggle_enabled"))
        m.add(types.InlineKeyboardButton(f"{'🔒 خاص' if ss.get('public') else '🌍 عام'} — تبديل الوصول", callback_data="scr_toggle_public"))
        m.add(types.InlineKeyboardButton("👥 إدارة المسموح لهم", callback_data="scr_manage_allowed"))
        m.add(types.InlineKeyboardButton("🌐 ملفات تعريف المواقع", callback_data="scr_manage_profiles"))
        m.add(types.InlineKeyboardButton("▶️ بدء مهمة سحب جديدة", callback_data="scr_new_job"))
        if jobs:
            m.add(types.InlineKeyboardButton(f"📋 المهام ({len(jobs)})", callback_data="scr_list_jobs"))
        m.add(types.InlineKeyboardButton("🚀 السحب التلقائي (قائمة انتظار)", callback_data="scr_advanced"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(text, m)
    elif data == "scr_advanced" and is_admin(uid):
        text, markup = _scraper_advanced_panel_content(); edit(text, markup)
    elif data == "scr_toggle_auto" and is_admin(uid):
        ss = db["config"]["scraper_settings"]
        ss["auto_schedule_enabled"] = not ss.get("auto_schedule_enabled", False)
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم {'تفعيل' if ss['auto_schedule_enabled'] else 'تعطيل'} الجدولة")
        text, markup = _scraper_advanced_panel_content(); edit(text, markup)
    elif data == "scr_add_link" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "scr_add_link_wait"}; sync_db()
        bot.send_message(cid, "🔗 أرسل رابط الرواية (يجب وجود ملف تعريف مسجّل لدومين هذا الموقع "
                               "بقسم ⚙️ ملفات تعريف المواقع، وإلا سيُرفض الرابط تلقائيًا وقت التشغيل).\n\n"
                               "أرسل الرابط فقط، والبوت يحدد ملف التعريف ونوع العمل (رواية) تلقائيًا:")
    elif data == "scr_view_pending" and is_admin(uid):
        pending = db["config"]["scraper_settings"].get("pending_links", [])
        if not pending:
            text = "📋 *لا توجد روابط بقائمة الانتظار*"
        else:
            text = f"📋 *قائمة الانتظار ({len(pending)})*\n\n" + "\n".join(f"{i+1}. {l['url'][:50]}" for i, l in enumerate(pending[:15]))
            if len(pending) > 15:
                text += f"\n\n... و{len(pending)-15} رابط آخر"
        m2 = types.InlineKeyboardMarkup(row_width=1)
        if pending:
            m2.add(types.InlineKeyboardButton("🗑️ مسح القائمة بالكامل", callback_data="scr_clear_pending"))
        m2.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="scr_advanced"))
        edit(text, m2)
    elif data == "scr_clear_pending" and is_admin(uid):
        db["config"]["scraper_settings"]["pending_links"] = []; sync_db()
        bot.answer_callback_query(call.id, "✅ تم مسح قائمة الانتظار")
        text, markup = _scraper_advanced_panel_content(); edit(text, markup)
    elif data == "scr_view_history" and is_admin(uid):
        ss = db["config"]["scraper_settings"]
        processed = ss.get("processed_links", [])[-10:]
        failed = ss.get("failed_links", [])[-10:]
        text = "📊 *سجل السحب التلقائي*\n\n✅ *آخر مكتملة:*\n"
        text += "\n".join(f"• {l['url'][:45]}" for l in processed) if processed else "لا يوجد"
        text += "\n\n❌ *آخر فاشلة:*\n"
        text += "\n".join(f"• {l['url'][:45]} — {l.get('error','')[:40]}" for l in failed) if failed else "لا يوجد"
        m2 = types.InlineKeyboardMarkup(row_width=1)
        m2.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="scr_advanced"))
        edit(text, m2)
    elif data == "scr_edit_time" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "scr_edit_time"}; sync_db()
        bot.send_message(cid, "⏰ أرسل وقت التشغيل اليومي (مثال: 06:00):")
    elif data == "scr_edit_count" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "scr_edit_count"}; sync_db()
        bot.send_message(cid, "🔢 أرسل عدد الروابط المسحوبة بكل دورة جدولة (مثال: 1):")
    elif data == "scr_run_now" and is_admin(uid):
        if not db["config"]["scraper_settings"].get("pending_links"):
            bot.answer_callback_query(call.id, "⚠️ قائمة الانتظار فاضية.", show_alert=True)
        else:
            bot.answer_callback_query(call.id, "⏳ جاري تشغيل دورة السحب...")
            threading.Thread(target=run_auto_scrape_schedule, daemon=True).start()
            bot.send_message(cid, "✅ بدأت دورة السحب بالخلفية (قد تأخذ وقتًا حسب عدد الفصول).")
    elif data == "scr_toggle_enabled" and is_admin(uid):
        ss = db["config"]["scraper_settings"]
        ss["enabled"] = not ss.get("enabled", False)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.", show_alert=True)
        edit("🕷️ *سحب الفصول التلقائي*", admin_main_kb())
    elif data == "scr_toggle_public" and is_admin(uid):
        ss = db["config"]["scraper_settings"]
        ss["public"] = not ss.get("public", False)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.", show_alert=True)
        edit("🕷️ *سحب الفصول التلقائي*", admin_main_kb())
    elif data == "scr_manage_allowed" and is_admin(uid):
        allowed = db["config"]["scraper_settings"].get("allowed_users", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for a in allowed:
            m.add(types.InlineKeyboardButton(f"❌ إزالة {a}", callback_data=f"scr_remove_allowed_{a}"))
        m.add(types.InlineKeyboardButton("➕ إضافة مستخدم", callback_data="scr_add_allowed"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_scraper"))
        edit(f"👥 *المسموح لهم باستخدام السحب* ({len(allowed)}):", m)
    elif data == "scr_add_allowed" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "scr_add_allowed_id"}; sync_db()
        bot.send_message(cid, "🆔 أرسل آيدي المستخدم اللي تبي تسمح له باستخدام السحب:")
    elif data.startswith("scr_remove_allowed_") and is_admin(uid):
        target = data[19:]
        allowed = db["config"]["scraper_settings"].get("allowed_users", [])
        if target in allowed: allowed.remove(target); sync_db()
        bot.answer_callback_query(call.id, "✅ تم الإزالة.", show_alert=True)
        m = types.InlineKeyboardMarkup(row_width=1)
        for a in allowed:
            m.add(types.InlineKeyboardButton(f"❌ إزالة {a}", callback_data=f"scr_remove_allowed_{a}"))
        m.add(types.InlineKeyboardButton("➕ إضافة مستخدم", callback_data="scr_add_allowed"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_scraper"))
        edit("👥 *المسموح لهم باستخدام السحب*:", m)
    elif data == "scr_manage_profiles" and is_admin(uid):
        profiles = db["config"]["scraper_settings"].get("site_profiles", {})
        m = types.InlineKeyboardMarkup(row_width=1)
        for pid, p in profiles.items():
            m.add(types.InlineKeyboardButton(f"🌐 {p.get('name', pid)}", callback_data=f"scr_profile_view_{pid}"))
        m.add(types.InlineKeyboardButton("➕ إضافة موقع جديد", callback_data="scr_profile_add"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_scraper"))
        edit(f"🌐 *ملفات تعريف المواقع* ({len(profiles)}):\n\nكل موقع يحتاج إعداد selectors مخصصة له (CSS)، لأن كل موقع مبني ببنية HTML مختلفة.", m)
    elif data == "scr_profile_add" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "scr_profile_name"}; sync_db()
        bot.send_message(cid, "✏️ أرسل اسم مميز لملف التعريف هذا (مثلاً: اسم الموقع):")
    elif data.startswith("scr_profile_view_") and is_admin(uid):
        pid = data[17:]
        p = db["config"]["scraper_settings"]["site_profiles"].get(pid)
        if not p:
            bot.answer_callback_query(call.id, "❌ ملف التعريف غير موجود.", show_alert=True)
        else:
            text = (f"🌐 *{p.get('name', pid)}*\n\n"
                    f"عنوان العمل: `{p.get('title_selector') or '—'}`\n"
                    f"وصف العمل: `{p.get('description_selector') or '—'}`\n"
                    f"صورة الغلاف: `{p.get('poster_selector') or '—'}`\n"
                    f"محتوى الفصل: `{p.get('chapter_content_selector') or '—'}`\n"
                    f"عنوان الفصل: `{p.get('chapter_title_selector') or '—'}`\n"
                    f"رابط الفصل التالي: `{p.get('next_page_selector') or '—'}`\n\n"
                    f"🧪 رابط عينة للاختبار: `{p.get('test_url') or '— غير محدد —'}`")
            m = types.InlineKeyboardMarkup(row_width=2)
            m.add(types.InlineKeyboardButton("✏️ عنوان العمل", callback_data=f"scr_edit_sel_title_{pid}"),
                  types.InlineKeyboardButton("✏️ وصف العمل", callback_data=f"scr_edit_sel_desc_{pid}"))
            m.add(types.InlineKeyboardButton("✏️ صورة الغلاف", callback_data=f"scr_edit_sel_poster_{pid}"),
                  types.InlineKeyboardButton("✏️ محتوى الفصل", callback_data=f"scr_edit_sel_content_{pid}"))
            m.add(types.InlineKeyboardButton("✏️ عنوان الفصل", callback_data=f"scr_edit_sel_chtitle_{pid}"),
                  types.InlineKeyboardButton("✏️ الفصل التالي", callback_data=f"scr_edit_sel_next_{pid}"))
            m.add(types.InlineKeyboardButton("🧪 تعيين رابط العينة", callback_data=f"scr_set_test_url_{pid}"))
            m.add(types.InlineKeyboardButton("🗑️ حذف ملف التعريف", callback_data=f"scr_profile_delete_{pid}"))
            m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="scr_manage_profiles"))
            edit(text, m)
    elif data.startswith("scr_set_test_url_") and is_admin(uid):
        pid = data[len("scr_set_test_url_"):]
        db["pending_actions"][ustr] = {"action": "scr_set_test_url", "profile_id": pid}; sync_db()
        bot.send_message(cid, "🔗 أرسل رابط فصل حقيقي من هذا الموقع (يُستخدم لاختبار كل selector تعدّله تلقائيًا):")
    elif data.startswith("scr_edit_sel_") and is_admin(uid):
        # التنسيق: scr_edit_sel_<field>_<profile_id> — الحقل نفسه ما فيه "_"
        rest = data[13:]
        field, pid = rest.split("_", 1)
        field_labels = {
            "title": "عنوان العمل", "desc": "وصف العمل", "poster": "صورة الغلاف",
            "content": "محتوى الفصل", "chtitle": "عنوان الفصل", "next": "رابط الفصل التالي",
        }
        db["pending_actions"][ustr] = {"action": "scr_set_selector", "field": field, "profile_id": pid}; sync_db()
        bot.send_message(cid, f"✏️ أرسل CSS selector لـ «{field_labels.get(field, field)}» (مثلاً: `div.chapter-content`):", parse_mode="Markdown")
    elif data.startswith("scr_profile_delete_") and is_admin(uid):
        pid = data[19:]
        profiles = db["config"]["scraper_settings"]["site_profiles"]
        if pid in profiles: del profiles[pid]; sync_db()
        bot.answer_callback_query(call.id, "🗑️ تم الحذف.", show_alert=True)
        edit("🌐 *ملفات تعريف المواقع*", admin_main_kb())
    elif data == "scr_new_job" and is_admin(uid):
        profiles = db["config"]["scraper_settings"].get("site_profiles", {})
        if not profiles:
            bot.answer_callback_query(call.id, "❌ أضف ملف تعريف موقع أول من «ملفات تعريف المواقع».", show_alert=True)
        else:
            m = types.InlineKeyboardMarkup(row_width=1)
            for pid, p in profiles.items():
                m.add(types.InlineKeyboardButton(f"🌐 {p.get('name', pid)}", callback_data=f"scr_job_pick_profile_{pid}"))
            m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_scraper"))
            edit("🌐 *اختر ملف تعريف الموقع اللي تبي تسحب منه:*", m)
    elif data.startswith("scr_job_pick_profile_") and is_admin(uid):
        pid = data[21:]
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("📚 رواية", callback_data=f"scr_job_pick_type_novel_{pid}"))
        m.add(types.InlineKeyboardButton("🎨 مانجا/مانهوا", callback_data=f"scr_job_pick_type_manga_{pid}"))
        m.add(types.InlineKeyboardButton("🎬 مسلسل/فيلم", callback_data=f"scr_job_pick_type_series_{pid}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="scr_new_job"))
        edit("📚 *اختر نوع العمل اللي بتسحبه:*", m)
    elif data.startswith("scr_job_pick_type_") and is_admin(uid):
        rest = data[18:]
        wtype, pid = rest.split("_", 1)
        db["pending_actions"][ustr] = {"action": "scr_job_start_url", "profile_id": pid, "work_type": wtype}; sync_db()
        bot.send_message(cid, "🔗 أرسل رابط أول فصل بالعمل (نقطة البداية للسحب):")
    elif data == "scr_list_jobs" and is_admin(uid):
        jobs = db["scrape_jobs"]
        m = types.InlineKeyboardMarkup(row_width=1)
        status_emoji = {"running": "🟢", "paused": "🟡", "done": "✅", "failed": "🔴"}
        for jid, j in jobs.items():
            emoji = status_emoji.get(j["status"], "⚪")
            m.add(types.InlineKeyboardButton(f"{emoji} {jid} — {j['chapters_done']} فصل", callback_data=f"scr_job_view_{jid}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_scraper"))
        edit(f"📋 *مهام السحب* ({len(jobs)}):", m)
    elif data.startswith("scr_job_view_") and is_admin(uid):
        jid = data[13:]
        j = get_scrape_job(jid)
        if not j:
            bot.answer_callback_query(call.id, "❌ المهمة غير موجودة.", show_alert=True)
        else:
            text = (f"📋 *مهمة سحب {jid}*\n\n"
                    f"النوع: {j['work_type']}\n"
                    f"الحالة: {j['status']}\n"
                    f"الفصول المسحوبة: {j['chapters_done']}\n"
                    f"آخر رابط: `{j.get('last_chapter_url') or '—'}`\n"
                    + (f"الخطأ: {j['error']}\n" if j.get('error') else ""))
            m = types.InlineKeyboardMarkup(row_width=1)
            if j["status"] in ("running", "paused"):
                m.add(types.InlineKeyboardButton("▶️ متابعة السحب (دفعة جديدة)", callback_data=f"scr_job_continue_{jid}"))
            if j["status"] == "running":
                m.add(types.InlineKeyboardButton("⏸️ إيقاف مؤقت", callback_data=f"scr_job_pause_{jid}"))
            m.add(types.InlineKeyboardButton("🗑️ حذف المهمة", callback_data=f"scr_job_delete_{jid}"))
            m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="scr_list_jobs"))
            edit(text, m)
    elif data.startswith("scr_job_continue_") and is_admin(uid):
        jid = data[17:]
        j = get_scrape_job(jid)
        if j: j["status"] = "running"; sync_db()
        bot.answer_callback_query(call.id, "⏳ جاري سحب دفعة جديدة بالخلفية (قد يأخذ وقت)...")
        def _do_continue():
            updated = run_scrape_job_step(jid, max_chapters=5)
            if updated:
                bot.send_message(cid, f"✅ *تحديث المهمة {jid}*\n\nالحالة: {updated['status']}\nالفصول: {updated['chapters_done']}"
                                  + (f"\nالخطأ: {updated['error']}" if updated.get('error') else ""), parse_mode="Markdown")
        threading.Thread(target=_do_continue, daemon=True).start()
    elif data.startswith("scr_job_pause_") and is_admin(uid):
        jid = data[14:]
        j = get_scrape_job(jid)
        if j: j["status"] = "paused"; sync_db()
        bot.answer_callback_query(call.id, "⏸️ تم الإيقاف المؤقت.", show_alert=True)
        edit(f"📋 مهمة {jid}", admin_main_kb())
    elif data.startswith("scr_job_delete_") and is_admin(uid):
        jid = data[15:]
        if jid in db["scrape_jobs"]: del db["scrape_jobs"][jid]; sync_db()
        bot.answer_callback_query(call.id, "🗑️ تم الحذف.", show_alert=True)
        edit("📋 *مهام السحب*", admin_main_kb())
    elif data == "adm_news" and is_admin(uid):
        ns = db["config"]["news_settings"]
        drafts = [n for n in db["news_items"].values() if n["status"] == "draft"]
        text = (f"📰 *نظام الأخبار*\n\n"
                f"قناة النشر: `{ns.get('publish_channel') or 'غير مربوطة'}`\n"
                f"عدد المسودات الحالية: {len(drafts)}\n\n"
                f"ألصق رابط أي خبر/صفحة، والبوت يجيب العنوان والوصف والصورة تلقائيًا "
                f"عشان تراجعها وتعدّلها قبل النشر.")
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("➕ إضافة خبر من رابط", callback_data="news_add_from_url"))
        if drafts:
            m.add(types.InlineKeyboardButton(f"📋 المسودات ({len(drafts)})", callback_data="news_list_drafts"))
        m.add(types.InlineKeyboardButton("📡 ربط قناة النشر", callback_data="news_set_channel"))
        m.add(types.InlineKeyboardButton("🚀 السحب التلقائي", callback_data="news_advanced"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit(text, m)
    elif data == "news_advanced" and is_admin(uid):
        text, markup = _news_advanced_panel_content(); edit(text, markup)
    elif data == "news_toggle_scrape" and is_admin(uid):
        settings = db["config"]["news_settings"]
        settings["auto_scrape_enabled"] = not settings.get("auto_scrape_enabled", False)
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم {'تفعيل' if settings['auto_scrape_enabled'] else 'تعطيل'} السحب التلقائي")
        text, markup = _news_advanced_panel_content(); edit(text, markup)
    elif data == "news_toggle_autopub" and is_admin(uid):
        settings = db["config"]["news_settings"]
        settings["auto_publish"] = not settings.get("auto_publish", False)
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم {'تفعيل' if settings['auto_publish'] else 'تعطيل'} النشر التلقائي")
        text, markup = _news_advanced_panel_content(); edit(text, markup)
    elif data == "news_add_source" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "news_add_source"}; sync_db()
        bot.send_message(cid, "🔗 أرسل رابط المصدر (RSS أو صفحة أخبار):")
    elif data == "news_list_sources" and is_admin(uid):
        sources = db["config"]["news_settings"].get("sources", [])
        if not sources:
            text = "📋 *لا توجد مصادر.*"
        else:
            text = "📋 *المصادر الحالية:*\n" + "\n".join([f"• {src}" for src in sources])
        m2 = types.InlineKeyboardMarkup(row_width=1)
        for _i, src in enumerate(sources):
            m2.add(types.InlineKeyboardButton(f"❌ حذف {src[:30]}...", callback_data=f"news_del_source_{_i}"))
        m2.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="news_advanced"))
        edit(text, m2)
    elif data.startswith("news_del_source_") and is_admin(uid):
        idx_s = data[len("news_del_source_"):]
        sources = db["config"]["news_settings"].get("sources", [])
        try:
            idx = int(idx_s)
            if 0 <= idx < len(sources):
                sources.pop(idx); sync_db()
        except ValueError:
            pass
        bot.answer_callback_query(call.id, "✅ تم حذف المصدر")
        text, markup = _news_advanced_panel_content(); edit(text, markup)
    elif data == "news_set_time" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "news_set_time"}; sync_db()
        bot.send_message(cid, "⏰ أرسل وقت النشر التلقائي (مثل 06:00 أو 18:30):")
    elif data == "news_set_max" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "news_set_max"}; sync_db()
        bot.send_message(cid, "🔢 أرسل عدد الأخبار القصوى لكل دورة سحب (مثال: 5):")
    elif data == "news_scrape_now" and is_admin(uid):
        bot.answer_callback_query(call.id, "⏳ جاري السحب...")
        threading.Thread(target=auto_scrape_news, daemon=True).start()
        bot.send_message(cid, "✅ تم بدء سحب الأخبار بالخلفية (ستظهر كمسودات جديدة أو تُنشر تلقائيًا حسب إعداداتك).")
    elif data == "news_add_from_url" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "news_url_wait"}; sync_db()
        bot.send_message(cid, "🔗 أرسل رابط الخبر/الصفحة:")
    elif data == "news_set_channel" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "news_set_channel_id"}; sync_db()
        bot.send_message(cid, "📡 أرسل يوزرنيم قناة النشر (مثل @MyNewsChannel) أو آيديها، والبوت لازم يكون مشرف فيها:")
    elif data == "news_list_drafts" and is_admin(uid):
        drafts = [n for n in db["news_items"].values() if n["status"] == "draft"]
        m = types.InlineKeyboardMarkup(row_width=1)
        for n in drafts:
            m.add(types.InlineKeyboardButton(f"📰 {n['title'][:40]}", callback_data=f"news_review_{n['id']}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_news"))
        edit(f"📋 *المسودات* ({len(drafts)}):", m)
    elif data.startswith("news_review_") and is_admin(uid):
        news_id = data[12:]
        edit(news_review_text(news_id), news_review_kb(news_id))
    elif data.startswith("news_edit_title_") and is_admin(uid):
        news_id = data[16:]
        db["pending_actions"][ustr] = {"action": "news_set_title", "news_id": news_id}; sync_db()
        bot.send_message(cid, "✏️ أرسل العنوان الجديد:")
    elif data.startswith("news_edit_desc_") and is_admin(uid):
        news_id = data[15:]
        db["pending_actions"][ustr] = {"action": "news_set_desc", "news_id": news_id}; sync_db()
        bot.send_message(cid, "📝 أرسل الوصف الجديد:")
    elif data.startswith("news_edit_cat_") and is_admin(uid):
        news_id = data[14:]
        db["pending_actions"][ustr] = {"action": "news_set_cat", "news_id": news_id}; sync_db()
        bot.send_message(cid, "🏷️ أرسل اسم التصنيف (مثلاً: رواية، أنمي، مانهوا):")
    elif data.startswith("news_edit_poster_") and is_admin(uid):
        news_id = data[17:]
        db["pending_actions"][ustr] = {"action": "news_set_poster", "news_id": news_id}; sync_db()
        bot.send_message(cid, "🖼️ أرسل صورة جديدة، أو رابط صورة مباشر:")
    elif data.startswith("news_link_pick_type_") and is_admin(uid):
        news_id = data[20:]
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("📚 رواية", callback_data=f"news_link_type_novel_{news_id}"))
        m.add(types.InlineKeyboardButton("🎨 مانجا/مانهوا", callback_data=f"news_link_type_manga_{news_id}"))
        m.add(types.InlineKeyboardButton("🎬 مسلسل/فيلم", callback_data=f"news_link_type_series_{news_id}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"news_review_{news_id}"))
        edit("🔗 *اختر نوع العمل اللي تبي تربط الخبر فيه:*", m)
    elif data.startswith("news_link_type_") and is_admin(uid):
        rest = data[15:]
        ltype, news_id = rest.split("_", 1)
        db["pending_actions"][ustr] = {"action": "news_link_search", "news_id": news_id, "ltype": ltype}; sync_db()
        bot.send_message(cid, "🔍 أرسل اسم العمل (أو جزء منه) للبحث عنه:")
    elif data.startswith("news_link_confirm_") and is_admin(uid):
        rest = data[18:]
        ltype, work_id, news_id = rest.split("_", 2)
        n = get_news_item(news_id)
        if n:
            n["linked_type"] = ltype
            n["linked_id"] = work_id
            sync_db()
            bot.answer_callback_query(call.id, "✅ تم الربط.", show_alert=True)
        bot.send_message(cid, news_review_text(news_id), parse_mode="Markdown", reply_markup=news_review_kb(news_id))
    elif data.startswith("news_unlink_") and is_admin(uid):
        news_id = data[12:]
        n = get_news_item(news_id)
        if n:
            n["linked_type"] = None; n["linked_id"] = None; sync_db()
        edit(news_review_text(news_id), news_review_kb(news_id))
    elif data.startswith("news_preview_") and is_admin(uid):
        news_id = data[13:]
        n = get_news_item(news_id)
        if not n:
            bot.answer_callback_query(call.id, "❌ الخبر غير موجود.", show_alert=True)
        else:
            caption = f"*{n['title']}*\n\n{n['description']}"
            if n.get("category"): caption += f"\n\n🏷️ {n['category']}"
            try:
                if n.get("poster_file_id"):
                    bot.send_photo(cid, n["poster_file_id"], caption=caption, parse_mode="Markdown", reply_markup=news_publish_kb(news_id))
                elif n.get("poster_url"):
                    bot.send_photo(cid, n["poster_url"], caption=caption, parse_mode="Markdown", reply_markup=news_publish_kb(news_id))
                else:
                    bot.send_message(cid, caption, parse_mode="Markdown", reply_markup=news_publish_kb(news_id))
            except Exception as e:
                bot.send_message(cid, f"{caption}\n\n⚠️ ما قدرت أعرض الصورة: {e}", parse_mode="Markdown", reply_markup=news_publish_kb(news_id))
    elif data.startswith("news_publish_") and is_admin(uid):
        news_id = data[13:]
        ok, err = publish_news_item(news_id)
        if ok:
            bot.answer_callback_query(call.id, "✅ تم النشر.", show_alert=True)
            edit("📰 *نظام الأخبار*", admin_main_kb())
        else:
            bot.answer_callback_query(call.id, f"❌ فشل النشر: {err}", show_alert=True)
    elif data.startswith("news_delete_") and is_admin(uid):
        news_id = data[12:]
        if news_id in db["news_items"]:
            del db["news_items"][news_id]; sync_db()
        bot.answer_callback_query(call.id, "🗑️ تم حذف المسودة.", show_alert=True)
        edit("📰 *نظام الأخبار*", admin_main_kb())
    elif data == "adm_new_members" and is_admin(uid):
        edit(new_member_panel_text(), new_member_panel_kb())
    elif data == "nm_toggle_enabled" and is_admin(uid):
        nm = db["config"]["new_member_settings"]
        nm["enabled"] = not nm.get("enabled", False)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.", show_alert=True)
        edit(new_member_panel_text(), new_member_panel_kb())
    elif data == "nm_manage_groups" and is_admin(uid):
        groups = db["config"]["new_member_settings"].get("watched_groups", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for g in groups:
            m.add(types.InlineKeyboardButton(f"❌ إزالة {g}", callback_data=f"nm_remove_group_{g}"))
        m.add(types.InlineKeyboardButton("➕ إضافة مجموعة", callback_data="nm_add_group"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_new_members"))
        edit("👥 *المجموعات المراقَبة* (البوت لازم يكون مشرف فيها):", m)
    elif data == "nm_add_group" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "nm_add_group_id"}; sync_db()
        bot.send_message(cid, "🆔 أرسل يوزرنيم المجموعة (مثل @MyGroup) أو آيديها، والبوت لازم يكون مشرف فيها:")
    elif data.startswith("nm_remove_group_") and is_admin(uid):
        target = data[16:]
        groups = db["config"]["new_member_settings"].get("watched_groups", [])
        if target in groups: groups.remove(target); sync_db()
        bot.answer_callback_query(call.id, "✅ تم الإزالة.", show_alert=True)
        m = types.InlineKeyboardMarkup(row_width=1)
        for g in groups:
            m.add(types.InlineKeyboardButton(f"❌ إزالة {g}", callback_data=f"nm_remove_group_{g}"))
        m.add(types.InlineKeyboardButton("➕ إضافة مجموعة", callback_data="nm_add_group"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_new_members"))
        edit("👥 *المجموعات المراقَبة* (البوت لازم يكون مشرف فيها):", m)
    elif data == "nm_edit_dm_text" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "nm_set_dm_text"}; sync_db()
        bot.send_message(cid, "✏️ أرسل نص الرسالة الخاصة اللي تنرسل للعضو الجديد:")
    elif data == "nm_toggle_dm" and is_admin(uid):
        nm = db["config"]["new_member_settings"]
        nm["dm_enabled"] = not nm.get("dm_enabled", True)
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم التحديث.", show_alert=True)
        edit(new_member_panel_text(), new_member_panel_kb())
    elif data == "nm_edit_dm_button" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "nm_set_dm_button_label"}; sync_db()
        bot.send_message(cid, "🔘 أرسل نص الزر (أو أرسل `بدون` لإلغاء الزر):")
    elif data == "nm_set_channel" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "nm_set_channel_id"}; sync_db()
        bot.send_message(cid, "📡 أرسل يوزرنيم قناة الإضافة التلقائية (مثل @MyChannel) أو آيديها، والبوت لازم يكون مشرف فيها بصلاحية دعوة أعضاء:")
    elif data == "nm_unset_channel" and is_admin(uid):
        db["config"]["new_member_settings"]["auto_add_channel"] = None
        sync_db()
        bot.answer_callback_query(call.id, "✅ تم إلغاء الربط.", show_alert=True)
        edit(new_member_panel_text(), new_member_panel_kb())
    elif data == "adm_content_types" and is_admin(uid):
        cf = db["config"]["content_features"]
        m = types.InlineKeyboardMarkup(row_width=1)
        for key, info in cf.items():
            status = "🟢 مفعّلة" if info.get("active", True) else "🔴 معطّلة"
            m.add(types.InlineKeyboardButton(f"{info['label']} — {status}", callback_data=f"toggle_ctype_{key}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit("🎛️ *تفعيل/تعطيل أنواع المحتوى*\n\nلما تعطّل نوع، يختفي بالكامل عن المستخدمين العاديين "
             "(بس يبقى ظاهر لك ولباقي المشرفين للإدارة). اضغط على أي نوع لتبديل حالته.", m)
    elif data.startswith("toggle_ctype_") and is_admin(uid):
        key = data[13:]
        cf = db["config"]["content_features"]
        if key in cf:
            cf[key]["active"] = not cf[key].get("active", True)
            sync_db()
        m = types.InlineKeyboardMarkup(row_width=1)
        for k, info in cf.items():
            status = "🟢 مفعّلة" if info.get("active", True) else "🔴 معطّلة"
            m.add(types.InlineKeyboardButton(f"{info['label']} — {status}", callback_data=f"toggle_ctype_{k}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="admin_panel"))
        edit("🎛️ *تفعيل/تعطيل أنواع المحتوى*", m)
    elif data == "adm_list_sections" and is_admin(uid):
        edit("📂 *إدارة أقسام القائمة* (المفضلة/تمت المشاهدة/إلخ)\n\nتقدر تغيّر أي اسم، تضيف قسم جديد، أو تغيّر علامة التقدم اللي تظهر جنب رقم آخر فصل.", admin_list_sections_kb())
    elif data.startswith("listsec_rename_") and is_admin(uid):
        sec_id = data[15:]
        db["pending_actions"][ustr] = {"action": "rename_list_section", "sec_id": sec_id}; sync_db()
        bot.send_message(cid, f"✏️ أرسل الاسم الجديد لقسم «{get_section_name(sec_id)}»:")
    elif data == "listsec_add" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "add_list_section"}; sync_db()
        bot.send_message(cid, "➕ أرسل اسم القسم الجديد (مثلاً: أشاهدها الآن):")
    elif data == "listsec_marker" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "edit_progress_marker"}; sync_db()
        bot.send_message(cid, "🏷️ أرسل الرمز أو النص اللي تبيه يظهر جنب رقم آخر فصل بالقائمة (مثال: 🔴 أو NEW أو 🆕):")
    elif data == "tags_layout" and is_admin(uid):
        edit("📐 *اختر شكل عرض التصنيفات للمستخدم:*", tags_layout_kb())
    elif data.startswith("settagslayout_") and is_admin(uid):
        layout = data[14:]
        db["config"]["tags_layout"] = layout; sync_db()
        edit("🏷️ *إدارة التصنيفات:*", tags_admin_kb())
    elif data.startswith("tag_edit_") and is_admin(uid):
        idx = int(data[9:]); tags_list = db["config"].get("tags", [])
        name = _tag_name(tags_list[idx]) if 0 <= idx < len(tags_list) else "؟"
        edit(f"🏷️ *تعديل التصنيف: {name}*", tag_editor_kb(idx))
    elif data == "tag_add" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "add_tag_name"}; sync_db(); bot.send_message(cid, "➕ أرسل اسم التصنيف الجديد:")
    elif data.startswith("tag_rename_") and is_admin(uid):
        idx = int(data[11:]); db["pending_actions"][ustr] = {"action": "rename_tag", "idx": idx}; sync_db()
        bot.send_message(cid, "✏️ أرسل الاسم الجديد للتصنيف:")
    elif data.startswith("tag_toggle_") and is_admin(uid):
        idx = int(data[11:]); tags_list = db["config"].get("tags", [])
        if 0 <= idx < len(tags_list):
            t = tags_list[idx]
            if isinstance(t, dict):
                t["active"] = not t.get("active", True)
            else:
                tags_list[idx] = {"name": t, "active": False}
            sync_db()
        edit("🏷️ *إدارة التصنيفات:*", tags_admin_kb())
    elif data.startswith("tag_up_") and is_admin(uid):
        idx = int(data[7:]); tags_list = db["config"].get("tags", [])
        if idx > 0:
            tags_list[idx-1], tags_list[idx] = tags_list[idx], tags_list[idx-1]; sync_db()
        edit("🏷️ *إدارة التصنيفات:*", tags_admin_kb())
    elif data.startswith("tag_down_") and is_admin(uid):
        idx = int(data[9:]); tags_list = db["config"].get("tags", [])
        if idx < len(tags_list) - 1:
            tags_list[idx+1], tags_list[idx] = tags_list[idx], tags_list[idx+1]; sync_db()
        edit("🏷️ *إدارة التصنيفات:*", tags_admin_kb())
    elif data.startswith("tag_del_") and is_admin(uid):
        idx = int(data[8:]); tags_list = db["config"].get("tags", [])
        if 0 <= idx < len(tags_list):
            tags_list.pop(idx); sync_db()
        edit("🏷️ *إدارة التصنيفات:*", tags_admin_kb())

    elif data == "adm_menu_editor" and is_admin(uid): edit("🎨 *تخصيص القائمة:*", menu_editor_kb())
    elif data.startswith("mebtn_") and is_admin(uid): edit(f"🎨 *تعديل الزر: {data[6:]}*", btn_editor_kb(data[6:]))
    elif data.startswith("me_rename_") and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "rename_btn", "btn_id": data[10:]}; sync_db(); bot.send_message(cid, "✏️ أرسل الاسم الجديد للزر:")
    elif data.startswith("me_reorder_") and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "reorder_btn", "btn_id": data[11:]}; sync_db(); bot.send_message(cid, "🔢 أرسل رقم الترتيب:")
    elif data.startswith("me_action_") and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "change_btn_action", "btn_id": data[10:]}; sync_db()
        bot.send_message(cid, "🔗 أرسل الأمر (Action) الجديد للزر:")
    elif data.startswith("me_toggle_") and is_admin(uid):
        bid = data[10:]; [setattr(type('obj', (object,), b), 'visible', not b.get('visible', True)) or b.update({'visible': not b.get('visible', True)}) for b in db["config"]["menu_buttons"] if b["id"] == bid]; sync_db(); edit("🎨 *تخصيص القائمة:*", menu_editor_kb())
    elif data.startswith("me_enable_toggle_") and is_admin(uid):
        bid = data[18:]
        for b in db["config"]["menu_buttons"]:
            if b["id"] == bid:
                b["active"] = not b.get("active", True)
                break
        sync_db(); edit("🎨 *تخصيص القائمة:*", menu_editor_kb())
    elif data.startswith("me_delete_") and is_admin(uid):
        bid = data[10:]; db["config"]["menu_buttons"] = [b for b in db["config"]["menu_buttons"] if b["id"] != bid]; sync_db(); edit("🎨 *تخصيص القائمة:*", menu_editor_kb())
    elif data == "me_add_btn" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "add_btn_label"}; sync_db(); bot.send_message(cid, "➕ أرسل اسم الزر الجديد:")

    elif data == "adm_features" and is_admin(uid): edit("✨ *الميزات المخصصة:*", features_kb())
    elif data == "feat_manual" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "manual_feat_label"}; sync_db(); bot.send_message(cid, "✏️ أرسل اسم الميزة:")
    elif data.startswith("feat_edit_") and is_admin(uid): edit(f"✨ *تعديل الميزة: {data[10:]}*", feat_editor_kb(data[10:]))
    elif data.startswith("feat_rename_") and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "rename_feat", "fid": data[12:]}; sync_db(); bot.send_message(cid, "✏️ أرسل الاسم الجديد:")
    elif data.startswith("feat_resp_") and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "resp_feat", "fid": data[10:]}; sync_db(); bot.send_message(cid, "💬 أرسل الرد الجديد:")
    elif data.startswith("feat_vis_") and is_admin(uid):
        f = db["config"]["custom_features"].get(data[9:]); f["visible"] = not f.get("visible", True) if f else True; sync_db(); edit("✨ *الميزات المخصصة:*", features_kb())
    elif data.startswith("feat_act_") and is_admin(uid):
        f = db["config"]["custom_features"].get(data[9:]); f["active"] = not f.get("active", True) if f else True; sync_db(); edit("✨ *الميزات المخصصة:*", features_kb())
    elif data.startswith("feat_priv_") and is_admin(uid):
        fid = data[10:]; f = db["config"]["custom_features"].get(fid); f["dev_only"] = not f.get("dev_only", False) if f else False; sync_db()
        flabel = f['label'] if f else fid
        edit(f"✨ *تعديل الميزة: {flabel}*", feat_editor_kb(fid))
    elif data.startswith("feat_del_") and is_admin(uid) and not data.startswith("feat_del_sub_"):
        fid_del = data[9:]
        db["config"]["custom_features"].pop(fid_del, None)
        db["config"]["feature_nesting_map"].pop(fid_del, None)
        for parent_id, children in list(db["config"]["feature_nesting_map"].items()):
            if fid_del in children: children.remove(fid_del)
        sync_db(); edit("✨ *الميزات المخصصة:*", features_kb())
    elif data.startswith("feat_subs_") and is_admin(uid):
        parent_fid = data[10:]
        pf_label = db["config"]["custom_features"][parent_fid]["label"]
        edit(f"📋 *الأزرار الفرعية للميزة: {pf_label}*", feat_sub_buttons_kb(parent_fid))
    elif data.startswith("feat_add_sub_") and is_admin(uid):
        parent_fid = data[13:]
        db["pending_actions"][ustr] = {"action": "add_sub_feature_label", "parent_fid": parent_fid};
        sync_db();
        bot.send_message(cid, "➕ أرسل اسم الزر الفرعي الجديد:")
    elif data.startswith("feat_nest_") and is_admin(uid):
        parent_fid = data[10:]
        edit("🔗 *اختر الميزة اللي تبي تدمجها هنا:*", nest_feature_picker_kb(parent_fid))
    elif data.startswith("donest_") and is_admin(uid):
        parts = data.split("_", 2); parent_fid = parts[1]; child_fid = parts[2]
        nest_feature(parent_fid, child_fid)
        sync_db(); bot.answer_callback_query(call.id, "✅ تم الدمج.")
        edit("📋 *الأزرار الفرعية:*", feat_sub_buttons_kb(parent_fid))
    elif data.startswith("feat_layout_") and is_admin(uid):
        fid = data[12:]
        edit("📐 *اختر شكل عرض الأزرار الفرعية:*", layout_picker_kb(fid))
    elif data.startswith("setlayout_") and is_admin(uid):
        parts = data.split("_", 2); fid = parts[1]; layout = parts[2]
        target = db["config"]["custom_features"].get(fid) or next((b for b in db["config"]["menu_buttons"] if b["id"] == fid), None)
        if target:
            target["sub_layout"] = layout; sync_db()
        bot.answer_callback_query(call.id, f"✅ تم اختيار: {layout}")
        edit("📋 *الأزرار الفرعية:*", feat_sub_buttons_kb(fid))
    elif data.startswith("feat_del_sub_") and is_admin(uid):
        parts = data.split("_"); parent_fid = parts[3]; sub_btn_id = parts[4]
        parent_feature = db["config"]["custom_features"].get(parent_fid)
        if parent_feature:
            parent_feature["sub_buttons"] = [b for b in parent_feature["sub_buttons"] if b["id"] != sub_btn_id]
            sync_db()
            pf_label2 = parent_feature['label']
            edit(f"📋 *الأزرار الفرعية للميزة: {pf_label2}*", feat_sub_buttons_kb(parent_fid))
        else:
            bot.answer_callback_query(call.id, "❌ الميزة الأم غير موجودة.", show_alert=True)

    elif data == "adm_ai_logic" and is_admin(uid): edit("🧠 *تعديل منطق الميزات بالذكاء:*", ai_logic_features_kb())
    elif data.startswith("ailogic_") and is_admin(uid):
        bid = data[8:]; db["pending_actions"][ustr] = {"action": "ai_logic_request", "btn_id": bid}; sync_db(); bot.send_message(cid, f"🧠 أرسل طلبك لتعديل منطق '{bid}' (نص و/أو صورة):")

    # ── نظام المسابقات المتطور (Contests) ───────────────────────────────────
    elif data == "adm_contests" and is_admin(uid):
        text, markup = _contests_panel_content(); edit(text, markup)
    elif data == "contest_toggle_enabled" and is_admin(uid):
        cfg = db["config"]["contests"]
        cfg["enabled"] = not cfg.get("enabled", False)
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم {'تفعيل' if cfg['enabled'] else 'تعطيل'}")
        text, markup = _contests_panel_content(); edit(text, markup)
    elif data == "contest_toggle_mode" and is_admin(uid):
        cfg = db["config"]["contests"]
        cfg["mode"] = "manual" if cfg.get("mode") == "auto" else "auto"
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم {'تشغيل الجدولة' if cfg['mode']=='auto' else 'إيقاف الجدولة'}")
        text, markup = _contests_panel_content(); edit(text, markup)
    elif data == "contest_start_now" and is_admin(uid):
        if db["config"]["contests"]["active_contests"]:
            bot.answer_callback_query(call.id, "⚠️ توجد مسابقة نشطة حاليًا", show_alert=True)
        else:
            bot.answer_callback_query(call.id, "⏳ جاري بدء المسابقة...")
            threading.Thread(target=start_contest, daemon=True).start()
            bot.send_message(cid, "✅ تم بدء المسابقة (ستُنشر بالهدف المحدد خلال لحظات).")
    elif data == "contest_toggle_answer" and is_admin(uid):
        cfg = db["config"]["contests"]
        cfg["answer_mode"] = "comment" if cfg.get("answer_mode", "button") == "button" else "button"
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم التبديل إلى: {'📝 تعليق' if cfg['answer_mode']=='comment' else '🔘 أزرار'}")
        text, markup = _contests_panel_content(); edit(text, markup)
    elif data == "contest_set_target" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "contest_set_target"}
        sync_db()
        bot.send_message(cid, "🎯 أرسل آيدي القناة/المجموعة المستهدفة (مثال: @MyChannel أو -100123456):")
    elif data == "contest_settings" and is_admin(uid):
        text, markup = _contests_settings_panel_content(); edit(text, markup)
    elif data == "contest_set_type" and is_admin(uid):
        m = types.InlineKeyboardMarkup(row_width=2)
        for t in ["mixed", "novel", "manga", "series", "anime", "general"]:
            m.add(types.InlineKeyboardButton(t.capitalize(), callback_data=f"contest_set_type_{t}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="contest_settings"))
        edit("📝 *اختر نوع الأسئلة:*", m)
    elif data.startswith("contest_set_type_") and is_admin(uid):
        t = data[len("contest_set_type_"):]
        db["config"]["contests"]["question_type"] = t
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم تعيين النوع: {t}")
        text, markup = _contests_settings_panel_content(); edit(text, markup)
    elif data == "contest_set_count" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "contest_set_count"}
        sync_db()
        bot.send_message(cid, "🔢 أرسل عدد الأسئلة (مثال: 5):")
    elif data == "contest_set_choices" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "contest_set_choices"}
        sync_db()
        bot.send_message(cid, "🔢 أرسل عدد الخيارات (2 أو 4):")
    elif data == "contest_set_source" and is_admin(uid):
        m = types.InlineKeyboardMarkup(row_width=2)
        m.add(types.InlineKeyboardButton("🤖 ذكاء اصطناعي", callback_data="contest_set_source_ai"))
        m.add(types.InlineKeyboardButton("📚 قاعدة البيانات", callback_data="contest_set_source_database"))
        m.add(types.InlineKeyboardButton("🔄 مختلط", callback_data="contest_set_source_mixed"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="contest_settings"))
        edit("📚 *اختر مصدر الأسئلة:*", m)
    elif data.startswith("contest_set_source_") and is_admin(uid):
        src = data[len("contest_set_source_"):]
        db["config"]["contests"]["question_source"] = src
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم تعيين المصدر: {src}")
        text, markup = _contests_settings_panel_content(); edit(text, markup)
    elif data == "contest_set_time" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "contest_set_time"}
        sync_db()
        bot.send_message(cid, "⏰ أرسل وقت الجدولة (مثال: 18:00):")
    elif data == "contest_set_days" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "contest_set_days"}
        sync_db()
        bot.send_message(cid, "📅 أرسل أيام الجدولة (مثال: sat,sun,mon أو all):")
    elif data == "contest_set_prizes" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "contest_set_prizes"}
        sync_db()
        bot.send_message(cid, "🎁 أرسل الجوائز بالتنسيق:\n`نقاط: 50`\n`أيام_vip: 3`\n`رتبة: بطل`\n\nأرسل `بدون` لإلغاء الجوائز.", parse_mode="Markdown")
    elif data == "contest_past_results" and is_admin(uid):
        past = db["config"]["contests"].get("past_contests", [])
        if not past:
            text = "🏆 *لا توجد مسابقات سابقة*"
        else:
            text = f"🏆 *آخر {min(10, len(past))} مسابقة*\n\n"
            for c in past[-10:]:
                text += f"• {c['name']} - {c['start_time'][:10]} - {len(c.get('participants', {}))} مشارك\n"
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_contests"))
        edit(text, m)
    elif data.startswith("contest_answer_") and not is_banned(uid):
        # الصيغة: contest_answer_<contest_id>_<q_index>_<option_index>  (contest_id بلا شرطات سفلية)
        parts = data[len("contest_answer_"):].split("_")
        if len(parts) == 3:
            c_id, q_idx_s, opt_idx_s = parts
            try:
                q_idx, opt_idx = int(q_idx_s), int(opt_idx_s)
                ok, msg = record_contest_answer(c_id, uid, q_idx, opt_idx)
                bot.answer_callback_query(call.id, msg, show_alert=True)
            except ValueError:
                bot.answer_callback_query(call.id, "❌ بيانات غير صالحة.", show_alert=True)
        else:
            bot.answer_callback_query(call.id, "❌ بيانات غير صالحة.", show_alert=True)

    # ── نظام "المحادثة الذكية" (AI Talk) ───────────────────────────────────
    elif data == "adm_ai_talk" and is_admin(uid):
        text, markup = _ai_talk_panel_content(); edit(text, markup)
    elif data == "talk_toggle" and is_admin(uid):
        cfg = db["config"]["ai_talk"]
        cfg["enabled"] = not cfg["enabled"]
        sync_db()
        text, markup = _ai_talk_panel_content(); edit(text, markup)
    elif data == "talk_toggle_auto" and is_admin(uid):
        cfg = db["config"]["ai_talk"]
        cfg["auto_post"] = not cfg["auto_post"]
        sync_db()
        text, markup = _ai_talk_panel_content(); edit(text, markup)
    elif data == "talk_personality" and is_admin(uid):
        m = types.InlineKeyboardMarkup(row_width=2)
        personalities = [
            ("friendly", "😊 ودود"), ("professional", "🎩 محترف"), ("sarcastic", "😏 ساخر"),
            ("enthusiastic", "🔥 حماسي"), ("calm", "😌 هادئ"), ("mysterious", "🕵️ غامض"),
        ]
        for key, label in personalities:
            m.add(types.InlineKeyboardButton(label, callback_data=f"talkp_{key}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_ai_talk"))
        edit("🎭 *اختر شخصية المحادثة الذكية:*", m)
    elif data.startswith("talkp_") and is_admin(uid):
        personality = data[len("talkp_"):]
        db["config"]["ai_talk"]["personality"] = personality
        sync_db()
        text, markup = _ai_talk_panel_content(); edit(text, markup)
    elif data == "talk_create_now" and is_admin(uid):
        bot.answer_callback_query(call.id, "⏳ جاري إنشاء المنشور...")
        def _do_create_talk_post():
            try:
                cfg = db["config"]["ai_talk"]
                post = create_smart_talk_post(cfg)
                channel = cfg.get("post_channel") or cid
                bot.send_message(channel, post, parse_mode="Markdown")
                cfg.setdefault("posts_history", []).append({"timestamp": datetime.now().isoformat(), "text": post})
                sync_db()
            except Exception as e:
                bot.send_message(cid, f"❌ فشل إنشاء المنشور: {e}")
        threading.Thread(target=_do_create_talk_post, daemon=True).start()

    # ── لوحة الدعوة التلقائية (Auto Invite) ─────────────────────────────────
    elif data == "adm_auto_invite" and is_admin(uid):
        text, markup = _auto_invite_panel_content(); edit(text, markup)
    elif data == "invite_toggle_enabled" and is_admin(uid):
        cfg = db["config"]["auto_invite"]
        cfg["enabled"] = not cfg.get("enabled", False)
        sync_db()
        bot.answer_callback_query(call.id, f"✅ تم {'تفعيل' if cfg['enabled'] else 'تعطيل'}")
        text, markup = _auto_invite_panel_content(); edit(text, markup)
    elif data == "invite_set_target" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "invite_set_target"}
        sync_db()
        bot.send_message(cid, "🔗 أرسل آيدي القناة/المجموعة المستهدفة (مثال: @MyChannel أو -100123456):")
    elif data == "invite_edit_text" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "invite_edit_text"}
        sync_db()
        bot.send_message(cid, "📝 أرسل النص الجديد (استخدم `{link}` لإدراج رابط الدعوة):", parse_mode="Markdown")
    elif data == "invite_edit_button" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "invite_edit_button"}
        sync_db()
        bot.send_message(cid, "🔘 أرسل النص الجديد للزر:")
    elif data == "invite_triggers" and is_admin(uid):
        cfg = db["config"]["auto_invite"]
        current = cfg.get("trigger_on", [])
        m = types.InlineKeyboardMarkup(row_width=1)
        for t in ["callback", "message", "new_member"]:
            mark = "✅ " if t in current else ""
            label = {"callback": "ضغط أزرار", "message": "رسائل في المجموعة", "new_member": "دخول أعضاء جدد"}[t]
            m.add(types.InlineKeyboardButton(f"{mark}{label}", callback_data=f"invite_toggle_trigger_{t}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_auto_invite"))
        edit("⚙️ *اختر أنواع التفاعل التي تُرسل الدعوة:*", m)
    elif data.startswith("invite_toggle_trigger_") and is_admin(uid):
        trigger = data[len("invite_toggle_trigger_"):]
        cfg = db["config"]["auto_invite"]
        triggers = cfg.setdefault("trigger_on", [])
        if trigger in triggers:
            triggers.remove(trigger)
        else:
            triggers.append(trigger)
        sync_db()
        m = types.InlineKeyboardMarkup(row_width=1)
        for t in ["callback", "message", "new_member"]:
            mark = "✅ " if t in triggers else ""
            label = {"callback": "ضغط أزرار", "message": "رسائل في المجموعة", "new_member": "دخول أعضاء جدد"}[t]
            m.add(types.InlineKeyboardButton(f"{mark}{label}", callback_data=f"invite_toggle_trigger_{t}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_auto_invite"))
        edit("⚙️ *اختر أنواع التفاعل التي تُرسل الدعوة:*", m)

    # ── Subscription System ─────────────────────────────────────────────────
    elif data == "my_subscription":
        u = ensure_user(call.message)
        edit(sub_status_text(ustr), sub_menu_kb(uid))
    elif data == "sub_plans":
        edit("💳 *خطط الاشتراك المتاحة:*\n" + sub_plans_text(), sub_plans_kb())
    elif data == "sub_redeem":
        db["pending_actions"][ustr] = {"action": "sub_redeem_code"}; sync_db(); bot.send_message(cid, "🎟️ أرسل كود الاشتراك:")
    elif data.startswith("sub_request_"):
        days = data[12:]
        plans = db["config"].get("sub_plans", {})
        price = plans.get(days, "؟")
        u_info = call.from_user
        uname = f"@{u_info.username}" if u_info.username else "(بدون يوزر)"
        fullname = f"{u_info.first_name or ''} {u_info.last_name or ''}".strip()
        owner_msg = (
            f"💎 *طلب اشتراك جديد*\n\n"
            f"👤 الاسم: {fullname}\n"
            f"🔗 اليوزر: {uname}\n"
            f"🆔 المعرف: `{ustr}`\n"
            f"📦 الخطة المطلوبة: {days} يوم — {price}$\n\n"
            f"تواصل مع المستخدم للاتفاق على الدفع، ثم فعّل اشتراكه من الزر تحت."
        )
        owner_kb = types.InlineKeyboardMarkup().add(
            types.InlineKeyboardButton("✅ تفعيل هذا الاشتراك الآن", callback_data=f"quickact_{ustr}_{days}")
        )
        try:
            bot.send_message(OWNER_ID, owner_msg, reply_markup=owner_kb, parse_mode="Markdown")
            edit(f"✅ تم إرسال طلبك (خطة {days} يوم) إلى الإدارة.\nراح يتواصلون معك قريبًا لإتمام العملية.",
                 types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="my_subscription")))
        except Exception as e:
            bot.answer_callback_query(call.id, f"❌ تعذر إرسال الطلب: {e}", show_alert=True)
    elif data.startswith("quickact_") and is_admin(uid):
        parts = data.split("_", 2); target_ustr = parts[1]; days = int(parts[2])
        new_expiry = activate_subscription(target_ustr, days)
        edit(f"✅ تم تفعيل اشتراك {days} يوم للمستخدم `{target_ustr}`.\nينتهي في: {new_expiry.strftime('%Y-%m-%d')}")
        try:
            bot.send_message(int(target_ustr), f"🎉 تم تفعيل اشتراكك ({days} يوم)!\nينتهي في: {new_expiry.strftime('%Y-%m-%d')}")
        except: pass
    elif data == "adm_subs" and is_admin(uid):
        edit("💎 *إدارة الاشتراكات:*", admin_subs_kb())
    elif data == "adm_sub_addplan" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "new_plan_days"}; sync_db()
        bot.send_message(cid, "🔢 كم عدد أيام الخطة الجديدة؟ (اكتب أي رقم تبيه، مثلاً 14 أو 365)")
    elif data.startswith("editplan_") and is_admin(uid):
        days = data[9:]
        price = db["config"].get("sub_plans", {}).get(days, 0)
        m = types.InlineKeyboardMarkup(row_width=1)
        m.add(types.InlineKeyboardButton("✏️ تغيير السعر", callback_data=f"editplan_price_{days}"))
        m.add(types.InlineKeyboardButton("🗑️ حذف هذي الخطة", callback_data=f"editplan_del_{days}"))
        m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_subs"))
        edit(f"📅 *خطة {days} يوم*\n\nالسعر الحالي: {price}$", m)
    elif data.startswith("editplan_price_") and is_admin(uid):
        days = data[15:]
        db["pending_actions"][ustr] = {"action": "add_sub_price", "days": days}; sync_db()
        bot.send_message(cid, f"💰 أرسل السعر الجديد لخطة {days} يوم بالدولار:")
    elif data.startswith("editplan_del_") and is_admin(uid):
        days = data[13:]
        db["config"]["sub_plans"].pop(days, None); sync_db()
        bot.answer_callback_query(call.id, f"✅ تم حذف خطة {days} يوم.", show_alert=True)
        edit("💳 *إدارة خطط الاشتراك:*", admin_subs_kb())
    elif data == "adm_sub_gencode" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "gen_sub_code"}; sync_db(); bot.send_message(cid, "🎟️ كم عدد أيام الاشتراك لهذا الكود؟")
    elif data == "adm_sub_manual" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "manual_activate_sub"}; sync_db()
        bot.send_message(cid, "✍️ أرسل: ID_المستخدم عدد_الأيام\nمثال: 123456789 30")
    elif data == "adm_sub_list" and is_admin(uid):
        subbed = [(u_id, u.get("sub_expiry")) for u_id, u in db["users"].items() if u.get("sub_expiry") and check_sub(u_id)]
        text = "💎 *المشتركون الحاليون:*\n\n" + ("\n".join([f"`{u_id}` — حتى {exp[:10]}" for u_id, exp in subbed]) if subbed else "لا يوجد مشتركون حاليًا.")
        edit(text, types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="adm_subs")))

    # ── AI Studio (Interactive AI feature/logic builder) ───────────────────
    elif data == "adm_ai_feature" and is_admin(uid):
        db["pending_actions"][ustr] = {"action": "ai_studio_chat", "mode": "new", "history": []}
        sync_db()
        bot.send_message(cid, "🤖 *استوديو الذكاء الاصطناعي*\n\nصف لي وش تبي (بالنص و/أو بصورة)، وممكن نتحاور لين توضح الفكرة أكثر.\nمثال: \"أبغى ميزة تعرض روايات مقترحة حسب تصنيف يختاره المستخدم\"", parse_mode="Markdown")
    elif data.startswith("feat_ai_edit_") and is_admin(uid):
        target_fid = data[13:]
        db["pending_actions"][ustr] = {"action": "ai_studio_chat", "mode": "edit", "history": [], "target_fid": target_fid}
        sync_db()
        bot.send_message(cid, "🤖 *تعديل الميزة عبر الذكاء الاصطناعي*\n\nصف لي وش تبي تغيّره في هذي الميزة (نص و/أو صورة):", parse_mode="Markdown")

# ==============================================================================
# 9.5 SLASH COMMANDS (تظهر بقائمة "/" بتلغرام)
# ==============================================================================
@bot.message_handler(commands=["start"])
def cmd_start(msg):
    uid = msg.from_user.id
    ensure_user(msg)
    db["users"][str(uid)]["seen_start_prompt"] = True; sync_db()
    if is_banned(str(uid)):
        return
    if db["config"].get("maintenance") and not is_admin(uid):
        bot.send_message(msg.chat.id, "🛠️ البوت تحت الصيانة حاليًا، حاول لاحقًا.")
        return
    unjoined = check_channels(uid)
    if unjoined:
        # نحفظ رابط العمق (لو فيه) عشان نفتحه تلقائيًا بعد ما يشترك بالقنوات المطلوبة
        parts = msg.text.split(maxsplit=1)
        if len(parts) > 1 and parts[1].strip():
            db["pending_actions"][str(uid)] = {"action": "deep_link_after_join", "payload": parts[1].strip()}
            sync_db()
        send_channels_msg(msg.chat.id, unjoined)
        return
    # رابط عمق من قناة (Deep Link): t.me/BotUsername?start=novel_<id>|manga_<id>|series_<id> يفتح العمل مباشرة
    parts = msg.text.split(maxsplit=1)
    if len(parts) > 1:
        payload = parts[1].strip()
        if payload.startswith("novel_") and get_novel(payload[6:]):
            open_novel(uid, msg.chat.id, None, payload[6:], as_new_message=True); return
        if payload.startswith("manga_") and get_manga(payload[6:]):
            open_manga(uid, msg.chat.id, None, payload[6:], as_new_message=True); return
        if payload.startswith("series_") and get_series(payload[7:]):
            open_series(uid, msg.chat.id, None, payload[7:], as_new_message=True); return
    send_welcome_and_menu(msg.chat.id, uid)

@bot.message_handler(commands=["search"])
def cmd_search(msg):
    uid = msg.from_user.id; ustr = str(uid)
    if is_banned(ustr): return
    ensure_user(msg)
    db["pending_actions"][ustr] = {"action": "search_tag_select", "selected": []}
    sync_db()
    bot.send_message(msg.chat.id, "🔍 *اختر تصنيف واحد أو أكثر (اضغط لتحديد/إلغاء)، ثم اضغط عرض النتائج — أو ابحث بالاسم مباشرة:*",
                      reply_markup=build_tag_search_menu(selected=[], mode="search"), parse_mode="Markdown")

@bot.message_handler(commands=["list", "mylist"])
def cmd_list(msg):
    uid = msg.from_user.id; ustr = str(uid)
    if is_banned(ustr): return
    ensure_user(msg)
    bot.send_message(msg.chat.id, "📂 *قائمتي — اختر القسم:*", reply_markup=list_sections_kb(), parse_mode="Markdown")

@bot.message_handler(commands=["help"])
def cmd_help(msg):
    bot.send_message(msg.chat.id,
        "ℹ️ *الأوامر المتاحة:*\n\n"
        "/start — القائمة الرئيسية\n"
        "/search — البحث بالتصنيف أو بالاسم\n"
        "/list — قائمتي (المفضلة، أشاهدها الآن، إلخ)\n"
        "/help — عرض هذي الرسالة",
        parse_mode="Markdown")

@bot.message_handler(commands=["admin"])
def cmd_admin(msg):
    uid = msg.from_user.id
    if not is_admin(uid): return
    bot.send_message(msg.chat.id, "⚙️ *لوحة التحكم:*", reply_markup=admin_main_kb(), parse_mode="Markdown")

@bot.message_handler(commands=["contests"])
def cmd_contests(msg):
    """اختصار مباشر للوحة تحكم المسابقات، بدل المرور بلوحة التحكم الرئيسية كاملة."""
    uid = msg.from_user.id
    if not is_admin(uid): return
    text, markup = _contests_panel_content()
    bot.send_message(msg.chat.id, text, reply_markup=markup, parse_mode="Markdown")

@bot.message_handler(commands=["talk"])
def cmd_talk(msg):
    """اختصار مباشر للوحة تحكم المحادثة الذكية، بدل المرور بلوحة التحكم الرئيسية كاملة."""
    uid = msg.from_user.id
    if not is_admin(uid): return
    text, markup = _ai_talk_panel_content()
    bot.send_message(msg.chat.id, text, reply_markup=markup, parse_mode="Markdown")

@bot.message_handler(commands=["invite"])
def cmd_auto_invite(msg):
    """اختصار مباشر للوحة تحكم الدعوة التلقائية، بدل المرور بلوحة التحكم الرئيسية كاملة."""
    uid = msg.from_user.id
    if not is_admin(uid): return
    text, markup = _auto_invite_panel_content()
    bot.send_message(msg.chat.id, text, reply_markup=markup, parse_mode="Markdown")

@bot.message_handler(commands=["set_chunking"])
def cmd_set_chunking(msg):
    """أمر سريع لتفعيل/تعطيل التقطيع مباشرة بالكتابة: /set_chunking on أو off."""
    uid = msg.from_user.id
    if not is_admin(uid): return
    parts = msg.text.strip().split()
    if len(parts) != 2 or parts[1].lower() not in ("on", "off"):
        bot.send_message(msg.chat.id, "الاستخدام: `/set_chunking on` أو `/set_chunking off`", parse_mode="Markdown")
        return
    db["config"]["chunking_settings"]["enabled"] = (parts[1].lower() == "on")
    sync_db()
    bot.send_message(msg.chat.id, f"✅ تم {'تفعيل' if parts[1].lower() == 'on' else 'تعطيل'} نظام التقطيع.")

@bot.message_handler(commands=["chunk_status"])
def cmd_chunk_status(msg):
    """اختصار مباشر للوحة تحكم التقطيع الكاملة."""
    uid = msg.from_user.id
    if not is_admin(uid): return
    text, markup = _chunking_panel_content()
    bot.send_message(msg.chat.id, text, reply_markup=markup, parse_mode="Markdown")

@bot.message_handler(commands=["chunk_stats"])
def cmd_chunk_stats(msg):
    """يعرض إحصائيات الملفات المؤقتة واستهلاك الرام الحالي مباشرة."""
    uid = msg.from_user.id
    if not is_admin(uid): return
    stats = temp_file_manager.stats()
    ram = memory_monitor.current_usage_mb()
    ram_text = f"{ram:.1f} MB" if ram is not None else "غير متاح (psutil غير مثبتة)"
    text = (f"📊 *إحصائيات التقطيع*\n\nخلفية تخزين المهام: {job_manager.backend_name()}\n"
            f"الرام الحالي للبوت: {ram_text}\n\n"
            f"الملفات المؤقتة المسجّلة: {stats['registered']}\n"
            f"موجودة فعليًا بالقرص: {stats['existing_on_disk']}\nالحجم الإجمالي: {stats['total_size_mb']} MB")
    bot.send_message(msg.chat.id, text, parse_mode="Markdown")

@bot.message_handler(commands=["convert"])
def cmd_convert(msg):
    """يفتح أداة تحويل صيغ الملفات لأي مستخدم مسموح له صراحة، حتى لو مو مشرف أصلًا."""
    uid = msg.from_user.id; ustr = str(uid)
    if is_banned(ustr): return
    if not can_use_convert(uid):
        return  # صامت تمامًا لغير المصرّح لهم، حتى ما يعرفون إن الميزة موجودة أصلًا
    ensure_user(msg)
    db["pending_actions"][ustr] = {"action": "convert_upload_wait"}; sync_db()
    bot.send_message(msg.chat.id, "📤 أرسل الملف اللي تبي تحوّل صيغته (PDF أو DOCX/DOC أو EPUB):")

def setup_bot_commands():
    """يسجل قائمة الأوامر اللي تظهر لما يكتب المستخدم / بمحادثة البوت."""
    try:
        bot.set_my_commands([
            types.BotCommand("start", "🏠 القائمة الرئيسية"),
            types.BotCommand("search", "🔍 بحث"),
            types.BotCommand("list", "📂 قائمتي"),
            types.BotCommand("help", "ℹ️ المساعدة"),
        ])
        logger.info("✅ تم تسجيل قائمة أوامر / بتلغرام.")
    except Exception as e:
        logger.error(f"❌ فشل تسجيل قائمة الأوامر: {e}")

# ==============================================================================
# 10. MESSAGES HANDLER
# ==============================================================================
def can_use_archive_upload(uid):
    """المطور دايمًا مسموح له. أي مستخدم ثاني لازم يكون بقائمة المسموح لهم صراحة
    (زي نظام السماح بتحويل الصيغة تمامًا)."""
    if is_owner(uid): return True
    return str(uid) in db["config"].get("archive_upload_allowed_users", [])

def get_or_create_group_invite_link(channel_ref):
    """يسوي رابط دعوة دائم لقناة الإضافة التلقائية ويخزّنه بالـ config عشان ما نطلب
    رابط جديد من تليجرام كل مرة (تجنب حدود الاستخدام)، إلا لو صار خطأ فيه فعليًا."""
    cache = db["config"]["new_member_settings"].setdefault("_invite_link_cache", {})
    cached = cache.get(str(channel_ref))
    if cached:
        return cached
    try:
        link_obj = bot.create_chat_invite_link(channel_ref, name="ترحيب أعضاء جدد")
        cache[str(channel_ref)] = link_obj.invite_link
        sync_db()
        return link_obj.invite_link
    except Exception as e:
        logger.error(f"❌ فشل إنشاء رابط دعوة للقناة {channel_ref}: {e}")
        return None

@bot.message_handler(content_types=["new_chat_members"])
def on_new_chat_members(msg):
    """يكتشف انضمام عضو جديد لأي مجموعة مراقَبة، يرسله رسالة خاصة ترحيبية (مع زر
    اختياري)، وإذا كانت قناة الإضافة التلقائية مربوطة يرسل له رابط دعوتها كمان.
    ما يشتغل إلا على المجموعات اللي المطور ضافها صراحة لقائمة المراقبة."""
    # ===== ميزة الدعوة التلقائية عند انضمام عضو جديد (مستقلة عن new_member_settings،
    # تعمل بأي مجموعة بغض النظر عن قائمة المراقبة أدناه) =====
    if db["config"]["auto_invite"].get("enabled") and "new_member" in db["config"]["auto_invite"].get("trigger_on", []):
        for _im in msg.new_chat_members:
            if not _im.is_bot:
                send_invite_to_user(_im.id, "انضمام جديد للمجموعة")

    nm = db["config"].get("new_member_settings", {})
    if not nm.get("enabled"):
        return
    chat_id = str(msg.chat.id)
    chat_username = f"@{msg.chat.username}" if msg.chat.username else None
    watched = nm.get("watched_groups", [])
    if not (chat_id in watched or (chat_username and chat_username in watched)):
        return  # المجموعة هذي مو بقائمة المراقبة، نتجاهل
    for member in msg.new_chat_members:
        if member.is_bot:
            continue  # نتجاهل دخول بوتات ثانية للمجموعة
        target_uid = member.id
        if nm.get("dm_enabled", True):
            m = None
            btn_label = nm.get("dm_button_label")
            btn_url = nm.get("dm_button_url")
            if btn_label and btn_url:
                m = types.InlineKeyboardMarkup()
                m.add(types.InlineKeyboardButton(btn_label, url=btn_url))
            try:
                bot.send_message(target_uid, nm.get("dm_message", "👋 أهلاً فيك!"), reply_markup=m)
            except Exception:
                # الأغلب لأن العضو ما بدأ محادثة مع البوت من قبل (خصوصية تليجرام) —
                # نسجله عشان المطور يقدر يتابعه يدويًا لو احتاج
                failed = nm.setdefault("dm_failed_users", [])
                if str(target_uid) not in failed:
                    failed.append(str(target_uid))
                    if len(failed) > 500:  # نحافظ على حجم القائمة معقول
                        del failed[:len(failed) - 500]
                    sync_db()
                continue  # ما نكمل لمحاولة إرسال رابط القناة إذا فشلت الرسالة الأساسية أصلًا
        auto_channel = nm.get("auto_add_channel")
        if auto_channel:
            invite_link = get_or_create_group_invite_link(auto_channel)
            if invite_link:
                m2 = types.InlineKeyboardMarkup()
                m2.add(types.InlineKeyboardButton("📡 انضم الآن", url=invite_link))
                try:
                    bot.send_message(target_uid, nm.get("auto_add_channel_message", "📢 انضم لقناتنا أيضًا:"), reply_markup=m2)
                except Exception:
                    pass  # لو فشلت هذي بالذات نتجاهلها بصمت، الرسالة الأساسية أهم

@bot.channel_post_handler(content_types=["document", "video"])
def on_channel_post(msg):
    """يكتشف أي ملف يترفع مباشرة بأي قناة أرشيف مربوطة، ويسأل الشخص اللي رفعه (إذا
    كان مسموح له) وش الرواية/العمل والفصل اللي يخص هذا الملف، عشان يربطهم تلقائيًا."""
    chat_id = str(msg.chat.id)
    chat_username = f"@{msg.chat.username}" if msg.chat.username else None
    archive_channels = db["config"].get("archive_channels", {})
    content_type = None
    for ctype, ch in archive_channels.items():
        if ch and (str(ch) == chat_id or str(ch) == chat_username):
            content_type = ctype
            break
    if not content_type:
        return  # القناة هذي مو مربوطة كأرشيف لأي نوع محتوى، نتجاهل الرسالة
    poster_id = msg.from_user.id if msg.from_user else None
    if poster_id is None or not can_use_archive_upload(poster_id):
        return  # ما نرد إطلاقًا على شخص غير مصرّح له، حتى ما يعرف الميزة موجودة
    file_id = msg.document.file_id if msg.content_type == "document" else msg.video.file_id
    db["pending_actions"][str(poster_id)] = {
        "action": "archive_link_pick_work", "content_type": content_type,
        "file_id": file_id, "file_kind": msg.content_type, "archive_chat": chat_id, "archive_msg_id": msg.message_id
    }
    sync_db()
    label = {"novels": "رواية", "manga": "عمل مانجا/مانهوا", "series": "مسلسل/فيلم"}[content_type]
    bot.send_message(poster_id, f"📥 استلمت ملف بقناة أرشيف {label}. أرسل اسم {label} اللي يخص هذا الملف (أو معرّفه لو تعرفه):")

@bot.message_handler(content_types=["text","document","photo","video","audio"])
def on_message(msg):
    """غلاف حماية عام: نفس فكرة on_callback — أي خطأ غير متوقع بمعالجة رسالة
    (رفع ملف، إدخال نص، إلخ) ما عاد يصير سكوت تام بدون رد للمستخدم."""
    try:
        _on_message_impl(msg)
    except Exception as e:
        logger.error(f"❌ خطأ غير متوقع بمعالجة رسالة: {e}", exc_info=True)
        try:
            bot.send_message(msg.chat.id, f"❌ صار خطأ غير متوقع أثناء تنفيذ العملية:\n`{str(e)[:300]}`", parse_mode="Markdown")
        except Exception:
            pass

def _on_message_impl(msg):
    uid = msg.from_user.id; ustr = str(uid)
    if is_banned(ustr): return
    pending = db["pending_actions"].pop(ustr, None)
    if pending:
        sync_db(); action = pending.get("action",""); txt = msg.text.strip() if msg.content_type == "text" else ""
        if action == "broadcast":
            sent = 0; [bot.send_message(int(t), txt) or time.sleep(0.05) for t in db["users"]] if msg.content_type == "text" else None; bot.send_message(msg.chat.id, "📢 تم الإرسال.")
        elif action == "rename_list_section" and txt:
            for s in db["config"]["list_sections"]:
                if s["id"] == pending["sec_id"]:
                    s["name"] = txt; break
            sync_db(); bot.send_message(msg.chat.id, "✅ تم تغيير اسم القسم.")
        elif action == "add_list_section" and txt:
            new_id = "sec_" + str(uuid.uuid4())[:8]
            db["config"]["list_sections"].append({"id": new_id, "name": txt})
            sync_db(); bot.send_message(msg.chat.id, f"✅ تم إضافة قسم «{txt}».")
        elif action == "set_archive_channel" and txt:
            target = txt.strip()
            try:
                test_chat = bot.get_chat(target)
                db["config"]["archive_channels"][pending["ctype"]] = target
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم ربط قناة الأرشيف: {test_chat.title or target}")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ ما قدرت أوصل للقناة. تأكد إن البوت مشرف فيها.\n`{e}`", parse_mode="Markdown")
        elif action == "archch_add_user_id" and txt:
            target = txt.strip()
            allowed = db["config"].setdefault("archive_upload_allowed_users", [])
            if target in allowed:
                bot.send_message(msg.chat.id, "⚠️ هذا المستخدم مسموح له أصلًا.")
            else:
                allowed.append(target); sync_db()
                bot.send_message(msg.chat.id, f"✅ تم السماح للمستخدم {target} بالرفع المباشر لقنوات الأرشيف.")
        elif action == "nm_add_group_id" and txt:
            target = txt.strip()
            try:
                test_chat = bot.get_chat(target)
                groups = db["config"]["new_member_settings"].setdefault("watched_groups", [])
                if target in groups:
                    bot.send_message(msg.chat.id, "⚠️ هذي المجموعة مضافة أصلًا.")
                else:
                    groups.append(target); sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم إضافة مجموعة للمراقبة: {test_chat.title or target}")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ ما قدرت أوصل للمجموعة. تأكد إن البوت مشرف فيها.\n`{e}`", parse_mode="Markdown")
        elif action == "nm_set_dm_text" and txt:
            db["config"]["new_member_settings"]["dm_message"] = txt
            sync_db()
            bot.send_message(msg.chat.id, "✅ تم تحديث نص الرسالة الخاصة.")
        elif action == "nm_set_dm_button_label" and txt:
            nm = db["config"]["new_member_settings"]
            if txt.strip() in ("بدون", "بلا", "إلغاء", "الغاء"):
                nm["dm_button_label"] = None; nm["dm_button_url"] = None
                sync_db()
                bot.send_message(msg.chat.id, "✅ تم إلغاء الزر.")
            else:
                nm["dm_button_label"] = txt.strip()
                db["pending_actions"][ustr] = {"action": "nm_set_dm_button_url"}; sync_db()
                bot.send_message(msg.chat.id, "🔗 طيب، الحين أرسل رابط الزر (لازم يبدأ بـ https://):")
        elif action == "nm_set_dm_button_url" and txt:
            url = txt.strip()
            if not url.startswith("http"):
                bot.send_message(msg.chat.id, "❌ الرابط لازم يبدأ بـ http:// أو https://. حاول مرة ثانية:")
                db["pending_actions"][ustr] = {"action": "nm_set_dm_button_url"}; sync_db()
            else:
                db["config"]["new_member_settings"]["dm_button_url"] = url
                sync_db()
                bot.send_message(msg.chat.id, "✅ تم ضبط زر الرسالة الخاصة.")
        elif action == "nm_set_channel_id" and txt:
            target = txt.strip()
            try:
                test_chat = bot.get_chat(target)
                db["config"]["new_member_settings"]["auto_add_channel"] = target
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم ربط قناة الإضافة التلقائية: {test_chat.title or target}")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ ما قدرت أوصل للقناة. تأكد إن البوت مشرف فيها بصلاحية دعوة أعضاء.\n`{e}`", parse_mode="Markdown")
        elif action == "news_url_wait" and txt:
            url = txt.strip()
            if not url.startswith("http"):
                bot.send_message(msg.chat.id, "❌ الرابط لازم يبدأ بـ http:// أو https://. حاول مرة ثانية:")
                db["pending_actions"][ustr] = {"action": "news_url_wait"}; sync_db()
            else:
                wait_msg = bot.send_message(msg.chat.id, "⏳ جاري جلب بيانات الرابط...")
                meta = extract_url_metadata(url)
                try: bot.delete_message(msg.chat.id, wait_msg.message_id)
                except Exception: pass
                if meta.get("error"):
                    bot.send_message(msg.chat.id, f"❌ ما قدرت أجيب بيانات الرابط:\n`{meta['error']}`", parse_mode="Markdown")
                else:
                    news_id = create_news_item(meta, uid)
                    bot.send_message(msg.chat.id, "✅ تم جلب البيانات. راجعها وعدّلها قبل النشر:")
                    bot.send_message(msg.chat.id, news_review_text(news_id), parse_mode="Markdown", reply_markup=news_review_kb(news_id))
        elif action == "news_set_channel_id" and txt:
            target = txt.strip()
            try:
                test_chat = bot.get_chat(target)
                db["config"]["news_settings"]["publish_channel"] = target
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم ربط قناة النشر: {test_chat.title or target}")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ ما قدرت أوصل للقناة. تأكد إن البوت مشرف فيها.\n`{e}`", parse_mode="Markdown")
        elif action == "news_add_source" and txt:
            url = txt.strip()
            if url.startswith("http"):
                sources = db["config"]["news_settings"].setdefault("sources", [])
                if url not in sources:
                    sources.append(url); sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم إضافة المصدر: {url}")
                else:
                    bot.send_message(msg.chat.id, "⚠️ هذا المصدر مضاف مسبقًا.")
            else:
                bot.send_message(msg.chat.id, "❌ رابط غير صالح.")
        elif action == "news_set_time" and txt:
            if re.match(r"^\d{2}:\d{2}$", txt.strip()):
                db["config"]["news_settings"]["publish_time"] = txt.strip(); sync_db()
                bot.send_message(msg.chat.id, f"✅ تم ضبط وقت النشر: {txt.strip()}")
            else:
                bot.send_message(msg.chat.id, "❌ صيغة غير صحيحة. استخدم HH:MM (مثل 06:00)")
        elif action == "news_set_max" and txt:
            try:
                n = int(txt.strip())
                if n > 0:
                    db["config"]["news_settings"]["max_news_per_day"] = n; sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم ضبط الحد الأقصى: {n} خبر لكل دورة")
                else:
                    bot.send_message(msg.chat.id, "❌ أدخل رقم موجب.")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أدخل رقم صحيح.")
        elif action == "chunk_set_pages" and txt:
            try:
                n = int(txt.strip())
                if n >= 5:
                    db["config"]["chunking_settings"]["pages_per_chunk"] = n; sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم ضبط صفحات كل جزء: {n}")
                else:
                    bot.send_message(msg.chat.id, "❌ أدخل رقم 5 أو أكثر (أجزاء أصغر من كذا غير عملية).")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أدخل رقم صحيح.")
        elif action == "chunk_set_pause" and txt:
            try:
                n = float(txt.strip())
                if n >= 0:
                    db["config"]["chunking_settings"]["pause_between_chunks"] = n; sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم ضبط الانتظار بين الأجزاء: {n}ث")
                else:
                    bot.send_message(msg.chat.id, "❌ أدخل رقم موجب أو صفر.")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أدخل رقم صحيح.")
        elif action == "chunk_set_ram" and txt:
            try:
                n = int(txt.strip())
                if n >= 50:
                    db["config"]["chunking_settings"]["max_ram_mb"] = n
                    memory_monitor.max_mb = n
                    sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم ضبط حد الرام: {n} MB")
                else:
                    bot.send_message(msg.chat.id, "❌ أدخل رقم 50 أو أكثر (حد أقل من كذا غير واقعي).")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أدخل رقم صحيح.")
        elif action == "scr_add_allowed_id" and txt:
            target = txt.strip()
            allowed = db["config"]["scraper_settings"].setdefault("allowed_users", [])
            if target in allowed:
                bot.send_message(msg.chat.id, "⚠️ هذا المستخدم مسموح له أصلًا.")
            else:
                allowed.append(target); sync_db()
                bot.send_message(msg.chat.id, f"✅ تم السماح للمستخدم {target} باستخدام السحب.")
        elif action == "tr_add_allowed_id" and txt:
            target = txt.strip()
            allowed = db["config"]["translation_settings"].setdefault("allowed_users", [])
            if target in allowed:
                bot.send_message(msg.chat.id, "⚠️ هذا المستخدم مسموح له أصلًا.")
            else:
                allowed.append(target); sync_db()
                bot.send_message(msg.chat.id, f"✅ تم السماح للمستخدم {target} باستخدام الترجمة.")
        elif action == "scr_profile_name" and txt:
            pid = str(uuid.uuid4())[:8]
            db["config"]["scraper_settings"]["site_profiles"][pid] = {
                "name": txt.strip(), "domain": "", "test_url": "",
                "title_selector": "", "description_selector": "", "poster_selector": "",
                "chapter_content_selector": "", "chapter_title_selector": "", "next_page_selector": "",
            }
            sync_db()
            bot.send_message(msg.chat.id, f"✅ تم إنشاء ملف تعريف «{txt.strip()}». الحين حدد الـ selectors من قائمة إدارة ملفات التعريف.")
        elif action == "scr_set_test_url" and txt:
            pid = pending.get("profile_id", "")
            profile = db["config"]["scraper_settings"]["site_profiles"].get(pid)
            url = txt.strip()
            if not profile:
                bot.send_message(msg.chat.id, "❌ ملف التعريف غير موجود.")
            elif not url.startswith("http"):
                bot.send_message(msg.chat.id, "❌ الرابط لازم يبدأ بـ http:// أو https://")
            else:
                profile["test_url"] = url
                sync_db()
                bot.send_message(msg.chat.id, "✅ تم حفظ رابط العينة. كل selector تعدّله الحين راح يُختبر عليه تلقائيًا.")
        elif action == "scr_set_selector" and txt:
            field = pending.get("field", ""); pid = pending.get("profile_id", "")
            field_map = {
                "title": "title_selector", "desc": "description_selector", "poster": "poster_selector",
                "content": "chapter_content_selector", "chtitle": "chapter_title_selector", "next": "next_page_selector",
            }
            key = field_map.get(field)
            profile = db["config"]["scraper_settings"]["site_profiles"].get(pid)
            if not profile or not key:
                bot.send_message(msg.chat.id, "❌ ملف التعريف غير موجود.")
            else:
                profile[key] = txt.strip()
                sync_db()
                bot.send_message(msg.chat.id, "✅ تم تحديث الـ selector.")
                # اختبار فوري على رابط عينة (لو محفوظ بالبروفايل) — يكشف selector
                # خاطئ الآن بدل ما يُكتشف لاحقًا بمنتصف مهمة سحب طويلة (paused).
                test_url = profile.get("test_url")
                if test_url:
                    def _test_selector():
                        try:
                            if field in ("content", "chtitle", "next"):
                                result = scrape_chapter(test_url, profile)
                                ok = bool(result and result.get("text"))
                                preview = (result.get("text", "")[:120] + "…") if ok else "لا يوجد نص مستخرَج"
                            else:
                                result = scrape_work_metadata(test_url, profile)
                                ok = bool(result and result.get("title"))
                                preview = result.get("title", "—") if ok else "لا يوجد عنوان مستخرَج"
                            status = "✅ الاختبار نجح" if ok else "⚠️ الاختبار فشل — راجع الـ selector"
                            bot.send_message(msg.chat.id, f"🧪 *اختبار تلقائي على رابط العينة:*\n{status}\n`{preview}`", parse_mode="Markdown")
                        except Exception as e:
                            bot.send_message(msg.chat.id, f"🧪 فشل الاختبار التلقائي: {e}")
                    threading.Thread(target=_test_selector, daemon=True).start()
                else:
                    bot.send_message(msg.chat.id, "💡 لو تضيف «رابط عينة» لهذا الملف (من قائمة إدارته)، بقدر أختبر كل selector تلقائيًا فور تعديله.")
        elif action == "scr_job_start_url" and txt:
            url = txt.strip()
            if not url.startswith("http"):
                bot.send_message(msg.chat.id, "❌ الرابط لازم يبدأ بـ http:// أو https://. حاول مرة ثانية:")
                db["pending_actions"][ustr] = pending; sync_db()
            else:
                profile_id = pending.get("profile_id", ""); work_type = pending.get("work_type", "novel")
                ss = db["config"]["scraper_settings"]
                job_id = create_scrape_job(profile_id, url, work_type, uid, translate_to=ss.get("default_translate_to"))
                bot.send_message(msg.chat.id, f"⏳ بدأت مهمة السحب `{job_id}`. جاري سحب أول دفعة بالخلفية (قد يأخذ وقت)...", parse_mode="Markdown")
                target_chat = msg.chat.id
                def _do_first_batch():
                    updated = run_scrape_job_step(job_id, max_chapters=5)
                    if updated:
                        bot.send_message(target_chat, f"✅ *تحديث المهمة {job_id}*\n\nالحالة: {updated['status']}\nالفصول: {updated['chapters_done']}"
                                          + (f"\nالخطأ: {updated['error']}" if updated.get('error') else ""), parse_mode="Markdown")
                threading.Thread(target=_do_first_batch, daemon=True).start()
        elif action == "scr_add_link_wait" and txt:
            url = txt.strip()
            if not url.startswith("http"):
                bot.send_message(msg.chat.id, "❌ الرابط لازم يبدأ بـ http:// أو https://")
            else:
                profile_id = get_matching_profile_id(url)
                if not profile_id:
                    bot.send_message(msg.chat.id, "❌ ما فيه ملف تعريف مسجّل لهذا الموقع. سجّله أولاً من "
                                                    "⚙️ ملفات تعريف المواقع، ثم أعد إرسال الرابط.")
                else:
                    add_pending_scrape_link(url, profile_id, "novel")
                    bot.send_message(msg.chat.id, f"✅ تم إضافة الرابط لقائمة الانتظار (سيُسحب تلقائيًا حسب الجدولة، "
                                                    f"أو فورًا لو ضغطت 'تشغيل دورة الآن').")
        elif action == "scr_edit_time" and txt:
            if re.match(r"^\d{2}:\d{2}$", txt.strip()):
                db["config"]["scraper_settings"]["schedule_time"] = txt.strip(); sync_db()
                bot.send_message(msg.chat.id, f"✅ تم ضبط وقت التشغيل: {txt.strip()}")
            else:
                bot.send_message(msg.chat.id, "❌ صيغة غير صحيحة. استخدم HH:MM (مثل 06:00)")
        elif action == "scr_edit_count" and txt:
            try:
                n = int(txt.strip())
                if n > 0:
                    db["config"]["scraper_settings"]["novels_per_schedule"] = n; sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم ضبط عدد الروابط لكل دورة: {n}")
                else:
                    bot.send_message(msg.chat.id, "❌ أدخل رقم موجب.")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أدخل رقم صحيح.")
        elif action == "news_set_title" and txt:
            n = get_news_item(pending.get("news_id", ""))
            if n:
                n["title"] = txt; sync_db()
                bot.send_message(msg.chat.id, "✅ تم تحديث العنوان.")
                bot.send_message(msg.chat.id, news_review_text(n["id"]), parse_mode="Markdown", reply_markup=news_review_kb(n["id"]))
        elif action == "news_set_desc" and txt:
            n = get_news_item(pending.get("news_id", ""))
            if n:
                n["description"] = txt; sync_db()
                bot.send_message(msg.chat.id, "✅ تم تحديث الوصف.")
                bot.send_message(msg.chat.id, news_review_text(n["id"]), parse_mode="Markdown", reply_markup=news_review_kb(n["id"]))
        elif action == "news_set_cat" and txt:
            n = get_news_item(pending.get("news_id", ""))
            if n:
                n["category"] = txt; sync_db()
                bot.send_message(msg.chat.id, "✅ تم تحديث التصنيف.")
                bot.send_message(msg.chat.id, news_review_text(n["id"]), parse_mode="Markdown", reply_markup=news_review_kb(n["id"]))
        elif action == "news_set_poster":
            n = get_news_item(pending.get("news_id", ""))
            if n:
                if msg.content_type == "photo":
                    n["poster_file_id"] = msg.photo[-1].file_id
                    n["poster_url"] = None
                    sync_db()
                    bot.send_message(msg.chat.id, "✅ تم تحديث الصورة.")
                elif txt.strip().startswith("http"):
                    n["poster_url"] = txt.strip()
                    n["poster_file_id"] = None
                    sync_db()
                    bot.send_message(msg.chat.id, "✅ تم تحديث الصورة.")
                else:
                    bot.send_message(msg.chat.id, "❌ أرسل صورة أو رابط صورة مباشر يبدأ بـ http.")
                    db["pending_actions"][ustr] = pending; sync_db()
                    return
                bot.send_message(msg.chat.id, news_review_text(n["id"]), parse_mode="Markdown", reply_markup=news_review_kb(n["id"]))
        elif action == "news_link_search" and txt:
            news_id = pending.get("news_id", ""); ltype = pending.get("ltype", "novel")
            n = get_news_item(news_id)
            coll_key = {"novel": "novels", "manga": "manga", "series": "series"}.get(ltype, "novels")
            collection = db.get(coll_key, {})
            query = txt.strip().lower()
            matches = [(k, v) for k, v in collection.items() if query in v.get("title", "").lower()][:10]
            if not matches:
                bot.send_message(msg.chat.id, "❌ ما لقيت أي نتيجة. حاول باسم مختلف:")
                db["pending_actions"][ustr] = pending; sync_db()
            else:
                m = types.InlineKeyboardMarkup(row_width=1)
                for k, v in matches:
                    m.add(types.InlineKeyboardButton(v["title"], callback_data=f"news_link_confirm_{ltype}_{k}_{news_id}"))
                m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data=f"news_review_{news_id}"))
                bot.send_message(msg.chat.id, "🔍 *النتائج:*", parse_mode="Markdown", reply_markup=m)
        elif action == "set_announce_channel" and txt:
            target = txt.strip()
            try:
                test_chat = bot.get_chat(target)
                db["config"]["announce_channel"]["chat_id"] = target
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم ربط قناة الإعلانات: {test_chat.title or target}")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ ما قدرت أوصل للقناة. تأكد إن البوت مشرف فيها والمعرّف صحيح.\n`{e}`", parse_mode="Markdown")
        elif action == "set_announce_template" and txt:
            db["config"]["announce_channel"]["message_template"] = txt
            sync_db(); bot.send_message(msg.chat.id, "✅ تم تحديث نص الإعلان.")
        elif action == "restore_snapshot_code" and txt:
            code = txt.strip().upper()
            snaps = list_snapshots()
            match = next((s for s in snaps if s[0] == code), None)
            if not match:
                # قد تكون النسخة أقدم من آخر 20، نحاول التحقق منها مباشرة
                conn = None
                try:
                    conn = get_conn(); cur = conn.cursor()
                    cur.execute("SELECT code, label, created_at FROM bot_snapshots WHERE code=%s", (code,))
                    match = cur.fetchone(); cur.close()
                except Exception: match = None
                finally:
                    if conn: release_conn(conn)
            if not match:
                bot.send_message(msg.chat.id, "❌ الكود غير صحيح أو غير موجود.")
            else:
                m = types.InlineKeyboardMarkup().add(
                    types.InlineKeyboardButton("✅ نعم، استرجع الآن", callback_data=f"snapshot_confirm_restore_{code}"),
                    types.InlineKeyboardButton("❌ إلغاء", callback_data="adm_snapshots"))
                bot.send_message(msg.chat.id, f"⚠️ *تأكيد الاسترجاع*\n\nراح يرجع البوت بالكامل لحالته وقت `{match[2].split('.')[0]}`، "
                                               f"وأي تغيير صار بعد هذا التاريخ (روايات، فصول، مستخدمين، ميزات) بينحذف. "
                                               f"متأكد؟", reply_markup=m, parse_mode="Markdown")
        elif action == "write_comment" and txt:
            ensure_user(msg)
            add_comment(pending["key"], uid, txt.strip()[:300])
            bot.send_message(msg.chat.id, "✅ تم إضافة تعليقك.")
        elif action == "set_announce_button" and txt:
            db["config"]["announce_channel"]["button_label"] = txt
            sync_db(); bot.send_message(msg.chat.id, "✅ تم تحديث اسم الزر.")
        elif action == "add_always_remove_phrase" and txt:
            phrase = txt.strip()
            db["config"].setdefault("always_remove_phrases", [])
            if phrase in db["config"]["always_remove_phrases"]:
                bot.send_message(msg.chat.id, "⚠️ هذي العبارة مسجّلة أصلًا.")
            else:
                db["config"]["always_remove_phrases"].append(phrase); sync_db()
                bot.send_message(msg.chat.id, f"✅ راح تُحذف «{phrase}» تلقائيًا من كل ملف مستقبلًا.")
        elif action == "add_phrasing_rule_original" and txt:
            db["pending_actions"][ustr] = {"action": "add_phrasing_rule_replacement", "original": txt.strip()}
            sync_db(); bot.send_message(msg.chat.id, f"✏️ أرسل الصياغة اللي تبي «{txt.strip()}» تترجم لها دائمًا:")
        elif action == "set_merge_free_limit_files_input" and txt:
            try:
                n = int(txt.strip())
                db["config"]["public_merge_tool"]["max_files_free"] = n
                sync_db(); bot.send_message(msg.chat.id, f"✅ صار الحد المجاني {n} ملف لكل عملية دمج.")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح.")
        elif action == "filemerge_collect":
            if msg.content_type != "document":
                bot.send_message(msg.chat.id, "❌ أرسل الملف كمستند (Document).")
                db["pending_actions"][ustr] = pending; sync_db(); return
            fname = msg.document.file_name or "file"
            f_ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
            if f_ext not in ("pdf", "docx", "doc", "epub", "txt", "html", "htm"):
                bot.send_message(msg.chat.id, "❌ صيغة غير مدعومة.")
                db["pending_actions"][ustr] = pending; sync_db(); return
            try:
                file_info = bot.get_file(msg.document.file_id)
                downloaded = bot.download_file(file_info.file_path)
                tmp_path = f"/tmp/mergein_{uuid.uuid4().hex[:8]}.{f_ext}"
                with open(tmp_path, "wb") as f:
                    f.write(downloaded)
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ فشل تحميل الملف: {e}"); return
            default_label = f"الفصل {len(pending['sections']) + 1}"
            db["pending_actions"][ustr] = {**pending, "action": "filemerge_label",
                                            "tmp_path": tmp_path, "tmp_ext": f_ext, "default_label": default_label}
            sync_db()
            bot.send_message(msg.chat.id, f"✏️ أرسل عنوان هذا الجزء (أو أرسل «تخطي» ليستخدم الاسم الافتراضي: {default_label}):")
        elif action == "filemerge_label" and txt:
            label = txt.strip() if txt.strip().lower() not in ("تخطي", "skip") else pending["default_label"]
            tmp_path = pending["tmp_path"]; tmp_ext = pending["tmp_ext"]
            try:
                if not _ensure_convert_libs():
                    raise RuntimeError(f"مكتبات التحويل غير مثبتة: {_CONVERT_IMPORT_ERROR}")
                text = extract_text_from_file(tmp_path, tmp_ext)
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ فشل قراءة الملف: {e}")
                try: os.remove(tmp_path)
                except Exception: pass
                db["pending_actions"][ustr] = {**pending, "action": "filemerge_collect"}; sync_db(); return
            try: os.remove(tmp_path)
            except Exception: pass
            sections = pending["sections"] + [(label, text)]
            db["pending_actions"][ustr] = {"action": "filemerge_ask_more", "sections": sections}
            sync_db()
            m = types.InlineKeyboardMarkup(row_width=1)
            m.add(types.InlineKeyboardButton("➕ أضف ملف آخر", callback_data="filemerge_add_more"))
            m.add(types.InlineKeyboardButton("✅ خلصت، ادمجهم الآن", callback_data="filemerge_finish"))
            bot.send_message(msg.chat.id, f"✅ تمت إضافة «{label}» ({len(sections)} ملف حتى الآن).", reply_markup=m)
        elif action == "add_phrasing_rule_replacement" and txt:
            original = pending["original"]; replacement = txt.strip()
            db["config"].setdefault("translation_phrasing_rules", {})
            db["config"]["translation_phrasing_rules"][original] = replacement
            sync_db()
            bot.send_message(msg.chat.id, f"✅ صار «{original}» تترجم دائمًا كـ «{replacement}».")
        elif action == "convert_add_user_id" and txt:
            target = txt.strip()
            allowed = db["config"].setdefault("convert_allowed_users", [])
            if target in allowed:
                bot.send_message(msg.chat.id, "⚠️ هذا المستخدم مسموح له أصلًا.")
            else:
                allowed.append(target); sync_db()
                bot.send_message(msg.chat.id, f"✅ تم السماح للمستخدم {target} باستخدام تحويل الصيغة (يقدر يستخدم /convert).")
        elif action == "archive_link_pick_work" and txt:
            content_type = pending["content_type"]
            query = txt.strip().lower()
            pool = {"novels": db["novels"], "manga": db["manga"], "series": db["series"]}[content_type]
            matches = [(wid, w) for wid, w in pool.items() if query in (w.get("title") or "").lower() or query == wid]
            if not matches:
                bot.send_message(msg.chat.id, "❌ ما لقيت عمل بهذا الاسم. أرسل الاسم كامل أو جزء منه:")
                db["pending_actions"][ustr] = pending; sync_db(); return
            if len(matches) == 1:
                wid, w = matches[0]
                db["pending_actions"][ustr] = {**pending, "action": "archive_link_pick_chapter", "work_id": wid}
                sync_db()
                unit = "الحلقة" if content_type == "series" else "الفصل"
                bot.send_message(msg.chat.id, f"✅ العمل: *{w['title']}*\n\nأرسل رقم {unit}:", parse_mode="Markdown")
            else:
                m = types.InlineKeyboardMarkup(row_width=1)
                for wid, w in matches[:10]:
                    m.add(types.InlineKeyboardButton(w["title"], callback_data=f"archivepick_{wid}"))
                bot.send_message(msg.chat.id, "🔍 لقيت أكثر من عمل مطابق، اختر الصحيح:", reply_markup=m)
        elif action == "archive_link_pick_chapter" and txt:
            try:
                ch_num = int(txt.strip())
            except:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح."); return
            content_type = pending["content_type"]; work_id = pending["work_id"]
            file_id = pending["file_id"]; file_kind = pending["file_kind"]
            if content_type == "novels":
                work = get_novel(work_id)
                fmt = "epub" if file_kind == "document" and False else "pdf"  # يُحدد لاحقًا بامتداد الملف الفعلي لو توفر
                work.setdefault("chapters", {}).setdefault(str(ch_num), {"pdf": None, "epub": None})
                work["chapters"][str(ch_num)]["pdf"] = file_id
                push_notification_all(f"🆕 فصل جديد: {work['title']} — الفصل {ch_num}", target_novel=work_id)
            elif content_type == "manga":
                work = get_manga(work_id)
                work.setdefault("chapters", {}).setdefault(str(ch_num), {"pdf": None, "epub": None})
                work["chapters"][str(ch_num)]["pdf"] = file_id
                push_notification_all(f"🆕 فصل جديد: {work['title']} — الفصل {ch_num}", target_novel=None)
            else:
                work = get_series(work_id)
                work.setdefault("episodes", {}).setdefault(str(ch_num), {})
                work["episodes"][str(ch_num)]["720p"] = {"file_id": file_id, "sub_only": False}
                push_notification_all(f"🆕 حلقة جديدة: {work['title']} — الحلقة {ch_num}", target_novel=None)
            db["archive_index"].setdefault(content_type, {})
            db["archive_index"][content_type][str(pending["archive_msg_id"])] = {
                "work_id": work_id, "chapter": str(ch_num), "fmt": "auto", "file_id": file_id
            }
            sync_db()
            bot.send_message(msg.chat.id, f"✅ تم ربط الملف بـ «{work['title']}» — الفصل/الحلقة {ch_num} بنجاح.")
        elif action == "rename_file_input" and txt:
            new_title = txt.strip()
            src_path = pending["src_path"]; src_ext = pending["src_ext"]
            out_path = f"/tmp/renamed_{uuid.uuid4().hex[:8]}.{src_ext}"
            try:
                import shutil
                shutil.copy(src_path, out_path)
                with open(out_path, "rb") as f:
                    bot.send_document(msg.chat.id, f, visible_file_name=f"{new_title}.{src_ext}",
                                       caption=f"✅ تم تغيير الاسم إلى: *{new_title}*", parse_mode="Markdown")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ فشل تغيير الاسم: {e}")
            finally:
                for p in (src_path, out_path):
                    try: os.remove(p)
                    except Exception: pass
                db["pending_actions"].pop(ustr, None); sync_db()
        elif action == "removetext_input" and txt:
            phrases = [p for p in txt.split("\n") if p.strip()]
            src_path = pending["src_path"]; src_ext = pending["src_ext"]; title = pending.get("title", "Document")
            out_path = f"/tmp/cleaned_{uuid.uuid4().hex[:8]}.{src_ext}"
            try:
                if not _ensure_convert_libs():
                    raise RuntimeError(f"مكتبات التحويل غير مثبتة: {_CONVERT_IMPORT_ERROR}")
                text = extract_text_from_file(src_path, src_ext)
                cleaned = remove_text_from_content(text, phrases)
                if src_ext == "pdf": build_pdf_from_text(cleaned, out_path, title)
                elif src_ext in ("docx", "doc"): build_docx_from_text(cleaned, out_path)
                elif src_ext == "epub": build_epub_from_text(cleaned, out_path, title)
                elif src_ext == "txt": build_txt_from_text(cleaned, out_path)
                elif src_ext in ("html", "htm"): build_html_from_text(cleaned, out_path, title)
                with open(out_path, "rb") as f:
                    bot.send_document(msg.chat.id, f, visible_file_name=f"{title}.{src_ext}",
                                       caption=f"✅ تم حذف {len(phrases)} عبارة من الملف وتنظيف الفراغات الناتجة.")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ فشلت العملية:\n`{str(e)[:300]}`", parse_mode="Markdown")
            finally:
                for p in (src_path, out_path):
                    try: os.remove(p)
                    except Exception: pass
                db["pending_actions"].pop(ustr, None); sync_db()
        elif action == "convert_upload_wait":
            if msg.content_type != "document":
                bot.send_message(msg.chat.id, "❌ أرسل الملف كـ document (ملف)، مو صورة أو نص.")
                db["pending_actions"][ustr] = pending; sync_db(); return
            fname = msg.document.file_name or "file"
            src_ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
            if src_ext not in ("pdf", "docx", "doc", "epub", "txt", "html", "htm"):
                bot.send_message(msg.chat.id, "❌ صيغة غير مدعومة. الصيغ المدعومة: PDF, DOCX, DOC, EPUB, TXT, HTML.")
                db["pending_actions"][ustr] = pending; sync_db(); return
            # لو فيه ملف من جلسة سابقة ما انحذف (المستخدم رفع ملف جديد بدون ما
            # يضغط "خلصت")، ننظفه الحين قبل ما نستبدله بالملف الجديد.
            old_src = pending.get("src_path")
            if old_src:
                try: os.remove(old_src)
                except Exception: pass
            try:
                file_info = bot.get_file(msg.document.file_id)
                downloaded = bot.download_file(file_info.file_path)
                src_path = f"/tmp/convert_src_{uuid.uuid4().hex[:8]}.{src_ext}"
                with open(src_path, "wb") as f:
                    f.write(downloaded)
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ فشل تحميل الملف: {e}")
                return
            # فحص أمني: نتأكد إن نوع الملف الحقيقي (من محتواه) يطابق فعليًا
            # الامتداد المُرسَل، لمنع ملفات متنكرة بامتداد بريء (مثال: ملف
            # تنفيذي باسم "chapter.pdf"). fail-open لو المكتبة غير مثبتة.
            is_valid, detected_mime = verify_file_type(src_path, [src_ext])
            if not is_valid:
                try: os.remove(src_path)
                except Exception: pass
                bot.send_message(msg.chat.id, f"❌ الملف المُرسَل لا يطابق فعليًا صيغة *{src_ext.upper()}* "
                                                f"(النوع الحقيقي المكتشف: `{detected_mime}`). تأكد من الملف وأعد المحاولة.",
                                  parse_mode="Markdown")
                return
            title = fname.rsplit(".", 1)[0]
            db["pending_actions"][ustr] = {"action": "file_tool_menu", "src_path": src_path, "src_ext": src_ext, "title": title}
            sync_db()
            m = types.InlineKeyboardMarkup(row_width=1)
            m.add(types.InlineKeyboardButton("🔄 تحويل الصيغة", callback_data="filetool_convert"))
            m.add(types.InlineKeyboardButton("✏️ تغيير اسم الملف", callback_data="filetool_rename"))
            m.add(types.InlineKeyboardButton("🗑️ حذف نص/روابط معيّنة من الملف", callback_data="filetool_removetext"))
            m.add(types.InlineKeyboardButton("🌐 ترجمة الملف", callback_data="filetool_translate"))
            bot.send_message(msg.chat.id, f"📄 الملف: *{fname}*\nالصيغة الحالية: *{src_ext.upper()}*\n\nاختر العملية اللي تبيها:", reply_markup=m, parse_mode="Markdown")
        elif action == "edit_progress_marker" and txt:
            db["config"]["progress_marker"] = txt
            sync_db(); bot.send_message(msg.chat.id, f"✅ تم تغيير علامة التقدم إلى: {txt}")
        elif action == "set_merge_free_limit" and txt:
            try:
                n = int(txt.strip())
                db["config"]["merge_limits"]["free_daily"] = n
                sync_db(); bot.send_message(msg.chat.id, f"✅ صار حد المستخدم العادي {n} ملف مدمج/يوم.")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح.")
        elif action == "set_merge_sub_limit" and txt:
            try:
                n = int(txt.strip())
                db["config"]["merge_limits"]["sub_daily"] = n
                sync_db(); bot.send_message(msg.chat.id, f"✅ صار حد المشترك {n} ملف مدمج/يوم.")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح.")
        elif action == "set_merge_user_override_id" and txt:
            target = txt.strip()
            db["pending_actions"][ustr] = {"action": "set_merge_user_override_value", "target": target}
            sync_db(); bot.send_message(msg.chat.id, "🔢 أرسل العدد المسموح لهذا المستخدم يوميًا (أو اكتب 0 للإلغاء والرجوع للحد العام، أو اكتب -1 لجعله غير محدود):")
        elif action == "set_merge_user_override_value" and txt:
            try:
                n = int(txt.strip())
                target = pending["target"]
                if n == 0:
                    db["config"]["merge_limits"]["user_overrides"].pop(target, None)
                    bot.send_message(msg.chat.id, f"✅ تم إلغاء التخصيص للمستخدم {target}، رجع للحد العام.")
                else:
                    db["config"]["merge_limits"]["user_overrides"][target] = n
                    display = "غير محدود ♾️" if n == -1 else str(n)
                    bot.send_message(msg.chat.id, f"✅ صار حد المستخدم {target} = {display}.")
                sync_db()
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح.")
        elif action == "edit_welcome":
            caption = msg.caption.strip() if msg.content_type in ("photo", "video") and msg.caption else (txt or None)
            rich = {"text": None, "photo_id": None, "video_id": None}
            if msg.content_type == "photo":
                rich["photo_id"] = msg.photo[-1].file_id
                rich["text"] = caption
            elif msg.content_type == "video":
                rich["video_id"] = msg.video.file_id
                rich["text"] = caption
            elif msg.content_type == "text" and txt:
                rich["text"] = txt
            else:
                bot.send_message(msg.chat.id, "❌ أرسل نص أو صورة أو فيديو صالح.")
                db["pending_actions"][ustr] = pending; sync_db(); return
            db["config"]["welcome_rich"] = rich
            db["config"]["welcome_msg"] = rich["text"] or ""  # للتوافق مع أي كود قديم يقرأ welcome_msg مباشرة
            sync_db(); bot.send_message(msg.chat.id, "✅ تم تحديث رسالة الترحيب.")
        elif action == "edit_watch_label" and txt:
            db["config"]["novel_watch_label"] = txt; sync_db(); bot.send_message(msg.chat.id, f"✅ تم تحديث نص الزر إلى: {txt}")
        elif action == "search_by_name" and txt:
            query = txt.strip().lower()
            track_event("search", txt.strip())
            m = types.InlineKeyboardMarkup(row_width=1)
            matches = []
            for iid, item in db["items"].items():
                if query in (item.get("title") or "").lower():
                    matches.append(("📄", item["title"], f"item_{iid}"))
            for nid, novel in db["novels"].items():
                if query in (novel.get("title") or "").lower():
                    matches.append(("📖", novel["title"], f"open_novel_{nid}"))
            for mid, manga in db["manga"].items():
                if query in (manga.get("title") or "").lower():
                    matches.append(("🎨", manga["title"], f"open_manga_{mid}"))
            for sid, series in db["series"].items():
                if query in (series.get("title") or "").lower():
                    matches.append(("🎬", series["title"], f"open_series_{sid}"))
            for icon, title, cb in matches[:20]:
                m.add(types.InlineKeyboardButton(f"{icon} {title}", callback_data=cb))
            m.add(types.InlineKeyboardButton("⬅️ رجوع", callback_data="search_filter"))
            if matches:
                extra = f"\n(عرض أول 20 من أصل {len(matches)})" if len(matches) > 20 else ""
                text = f"🔎 *نتائج البحث عن \"{txt.strip()}\":*{extra}"
            else:
                text = f"❌ ما لقيت شي يطابق \"{txt.strip()}\"."
            bot.send_message(msg.chat.id, text, reply_markup=m, parse_mode="Markdown")
        elif action == "ban" and txt:
            db["config"]["banned_users"].append(txt); sync_db(); bot.send_message(msg.chat.id, "🚫 تم الحظر.")
        elif action == "unban" and txt:
            if txt in db["config"]["banned_users"]: db["config"]["banned_users"].remove(txt); sync_db()
            bot.send_message(msg.chat.id, "✅ تم رفع الحظر.")
        elif action == "view_user" and txt:
            u = db["users"].get(txt); bot.send_message(msg.chat.id, f"👤 *بيانات المستخدم {txt}:*\n{json.dumps(u, indent=2)}" if u else "❌ غير موجود.")
        elif action == "set_admin_expiry_days" and txt:
            try:
                days = int(txt.strip())
                target = pending["target"]
                if target in db["config"]["admins"] and isinstance(db["config"]["admins"][target], dict):
                    db["config"]["admins"][target]["expires_at"] = (datetime.now() + timedelta(days=days)).isoformat()
                    sync_db()
                    bot.send_message(msg.chat.id, f"✅ صارت مدته {days} يوم من الآن.")
                else:
                    bot.send_message(msg.chat.id, "❌ المشرف غير موجود.")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح.")
        elif action == "add_admin_expiry_days" and txt:
            try:
                days = int(txt.strip())
                target, title, perms = pending["target"], pending["title"], pending["selected_perms"]
                perms_dict = {p: True for p in perms}
                expires = (datetime.now() + timedelta(days=days)).isoformat()
                db["config"]["admins"][target] = {"title": title, "permissions": perms_dict, "expires_at": expires}
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم إضافة المشرف «{title}» (ID: {target}) لمدة {days} يوم. المشرف نفسه ما راح يشوف تاريخ الانتهاء.")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أرسل رقم أيام صحيح.")
        elif action == "add_admin_id" and txt:
            target = txt.strip()
            if target in db["config"]["admins"] or (target.isdigit() and int(target) == OWNER_ID):
                bot.send_message(msg.chat.id, "⚠️ هذا المستخدم مشرف أصلًا أو هو المطور نفسه.")
            else:
                db["pending_actions"][ustr] = {"action": "add_admin_title", "target": target}
                sync_db(); bot.send_message(msg.chat.id, "🏷️ أرسل لقب هذا المشرف (مثال: مشرف الروايات):")
        elif action == "add_admin_title" and txt:
            db["pending_actions"][ustr] = {"action": "add_admin_perms", "target": pending["target"], "title": txt.strip(), "selected_perms": []}
            sync_db()
            bot.send_message(msg.chat.id, "🔐 اختر صلاحيات هذا المشرف (تقدر تختار أكثر من وحدة، ثم اضغط متابعة):",
                              reply_markup=admin_perm_picker_kb([]))
        elif action == "gen_code" and txt:
            try:
                code = str(uuid.uuid4()).split("-")[0].upper(); db["codes"][code] = {"points": int(txt), "used": False}; sync_db()
                bot.send_message(msg.chat.id, f"🎟️ تم إنشاء كود: `{code}`", parse_mode="Markdown")
            except: bot.send_message(msg.chat.id, "❌ خطأ.")
        elif action == "custom_code_name" and txt:
            code = txt.strip().upper()
            if code in db.get("codes", {}):
                bot.send_message(msg.chat.id, "❌ هذا الكود مستخدم بالفعل بالنظام، جرب اسم ثاني.")
            else:
                db["pending_actions"][ustr] = {"action": "custom_code_points", "code": code}; sync_db()
                bot.send_message(msg.chat.id, "💎 كم عدد النقاط لهذا الكود؟")
        elif action == "custom_code_points" and txt:
            try:
                pts = int(txt)
                db["pending_actions"][ustr] = {"action": "custom_code_expiry", "code": pending["code"], "points": pts}; sync_db()
                bot.send_message(msg.chat.id, "⏳ كم مدة صلاحية الكود بالأيام؟ (أرسل رقم، أو أرسل 'بدون' لعدم وضع مدة):")
            except: bot.send_message(msg.chat.id, "❌ رقم غير صحيح.")
        elif action == "custom_code_expiry" and txt:
            code = pending["code"]; pts = pending["points"]
            entry = {"points": pts, "used": False}
            if txt.strip() not in ("بدون", "بلا", "no", "none"):
                try:
                    days = int(txt.strip())
                    entry["expires_at"] = (datetime.now() + timedelta(days=days)).isoformat()
                except:
                    bot.send_message(msg.chat.id, "❌ مدة غير صحيحة، أرسل رقم أيام أو 'بدون'."); return
            db["codes"][code] = entry; sync_db()
            exp_txt = f"\nينتهي بعد: {txt.strip()} يوم" if "expires_at" in entry else "\nبدون مدة انتهاء"
            bot.send_message(msg.chat.id, f"✅ تم إنشاء الكود: `{code}`\nالنقاط: {pts}{exp_txt}\n(يعمل لشخص واحد فقط، يتعطل بعد أول استخدام)", parse_mode="Markdown")
        elif action == "add_channel" and txt:
            db["config"]["mandatory_channels"].append(txt); sync_db(); bot.send_message(msg.chat.id, "✅ تمت إضافة القناة.")
        elif action == "addcat" and txt:
            nid = str(uuid.uuid4())[:8]; p = pending.get("parent","root")
            db["categories"][nid] = {"id": nid, "name": txt, "parent": p, "children": [], "items": []}
            db["categories"].get(p, db["categories"]["root"])["children"].append(nid); sync_db(); bot.send_message(msg.chat.id, f"✅ تم إنشاء: {txt}")
        elif action == "additem_tag":
            db["pending_actions"][ustr] = {"action": "additem_title", "cat": pending["cat"], "tags": [txt] if txt else ["عام"]}; bot.send_message(msg.chat.id, "📖 أرسل اسم الملف:")
        elif action == "additem_title" and txt:
            db["pending_actions"][ustr] = {"action": "additem_file", "cat": pending["cat"], "tags": pending.get("tags", ["عام"]), "title": txt}; bot.send_message(msg.chat.id, "📥 أرسل الملف الآن:")
        elif action == "additem_file":
            fid = msg.photo[-1].file_id if msg.content_type == "photo" else msg.video.file_id if msg.content_type == "video" else msg.document.file_id if msg.content_type == "document" else msg.audio.file_id if msg.content_type == "audio" else None
            if fid:
                nid = str(uuid.uuid4())[:8]; db["items"][nid] = {"id": nid, "title": pending["title"], "file_id": fid, "type": msg.content_type, "tags": pending.get("tags", ["عام"]), "category": pending["cat"]}
                if pending["cat"] in db["categories"]: db["categories"][pending["cat"]]["items"].append(nid)
                sync_db(); bot.send_message(msg.chat.id, "✅ تمت الإضافة.")
        elif action == "add_tag_name" and txt:
            db["config"].setdefault("tags", []).append({"name": txt, "active": True}); sync_db()
            bot.send_message(msg.chat.id, f"✅ تم إضافة التصنيف: {txt}")
        elif action == "rename_tag" and txt:
            tags_list = db["config"].get("tags", []); idx = pending.get("idx", -1)
            if 0 <= idx < len(tags_list):
                if isinstance(tags_list[idx], dict): tags_list[idx]["name"] = txt
                else: tags_list[idx] = {"name": txt, "active": True}
                sync_db(); bot.send_message(msg.chat.id, "✅ تم تغيير اسم التصنيف.")
            else:
                bot.send_message(msg.chat.id, "❌ التصنيف غير موجود.")
        elif action == "rename_btn" and txt:
            for b in db["config"]["menu_buttons"]:
                if b["id"] == pending["btn_id"]: b["label"] = txt; break
            sync_db(); bot.send_message(msg.chat.id, "✅ تم تغيير الاسم.")
        elif action == "reorder_btn" and txt:
            try:
                for b in db["config"]["menu_buttons"]:
                    if b["id"] == pending["btn_id"]: b["order"] = int(txt); break
                sync_db(); bot.send_message(msg.chat.id, "✅ تم تغيير الترتيب.")
            except: bot.send_message(msg.chat.id, "❌ رقم غير صحيح.")
        elif action == "change_btn_action" and txt:
            for b in db["config"]["menu_buttons"]:
                if b["id"] == pending["btn_id"]: b["action"] = txt; break
            sync_db(); bot.send_message(msg.chat.id, "✅ تم تغيير أمر الزر.")
        elif action == "add_btn_label" and txt:
            db["pending_actions"][ustr] = {"action": "add_btn_action", "label": txt}; bot.send_message(msg.chat.id, "🔗 أرسل الأمر (Action):")
        elif action == "add_btn_action" and txt:
            db["config"]["menu_buttons"].append({"id": str(uuid.uuid4())[:8], "label": pending["label"], "action": txt, "visible": True, "order": 99}); sync_db(); bot.send_message(msg.chat.id, "✅ تم إضافة الزر.")
        elif action == "manual_feat_label" and txt:
            db["pending_actions"][ustr] = {"action": "manual_feat_resp", "label": txt}; bot.send_message(msg.chat.id, "💬 أرسل الرد التلقائي:")
        elif action == "manual_feat_resp" and txt:
            fid = str(uuid.uuid4())[:8]; db["config"]["custom_features"][fid] = {"label": pending["label"], "response_text": txt, "active": True, "visible": True, "sub_buttons": []}; sync_db(); bot.send_message(msg.chat.id, "✅ تم إنشاء الميزة.")
        elif action == "rename_feat" and txt:
            f = db["config"]["custom_features"].get(pending["fid"]); f["label"] = txt if f else ""; sync_db(); bot.send_message(msg.chat.id, "✅ تم التغيير.")
        elif action == "resp_feat" and txt:
            f = db["config"]["custom_features"].get(pending["fid"]); f["response_text"] = txt if f else ""; sync_db(); bot.send_message(msg.chat.id, "✅ تم التغيير.")
        elif action == "add_sub_feature_label" and txt:
            db["pending_actions"][ustr] = {"action": "add_sub_feature_action", "parent_fid": pending["parent_fid"], "label": txt};
            sync_db();
            bot.send_message(msg.chat.id, "🔗 أرسل الأمر (Action) للزر الفرعي:")
        elif action == "add_sub_feature_action" and txt:
            parent_feature = db["config"]["custom_features"].get(pending["parent_fid"])
            if parent_feature:
                sub_btn_id = str(uuid.uuid4())[:8]
                parent_feature.setdefault("sub_buttons", []).append({"id": sub_btn_id, "label": pending["label"], "action": txt})
                sync_db()
                bot.send_message(msg.chat.id, "✅ تم إضافة الزر الفرعي.")
            else:
                bot.send_message(msg.chat.id, "❌ الميزة الأم غير موجودة.")
        elif action == "addsub_feat_label" and txt:
            db["pending_actions"][ustr] = {"action": "addsub_feat_resp", "fid": pending["fid"], "label": txt}; bot.send_message(msg.chat.id, "💬 أرسل الرد للزر الفرعي:")
        elif action == "addsub_feat_resp" and txt:
            f = db["config"]["custom_features"].get(pending["fid"])
            if f: f.setdefault("sub_buttons", []).append({"id": str(uuid.uuid4())[:8], "label": pending["label"], "response_text": txt}); sync_db(); bot.send_message(msg.chat.id, "✅ تم إضافة الزر الفرعي.")
        elif action == "ai_logic_request":
            handle_ai_logic_request(msg, pending, txt)
        elif action == "ai_studio_chat":
            h_studio(msg, pending, txt)
        elif action == "sub_redeem_code" and txt:
            do_sub_redeem(msg.chat.id, ustr, txt)
        elif action == "new_plan_days" and txt:
            try:
                days = int(txt.strip())
                if str(days) in db["config"].get("sub_plans", {}):
                    bot.send_message(msg.chat.id, f"⚠️ فيه خطة بنفس المدة ({days} يوم) مسبقًا، أرسل السعر لتحديثها:")
                else:
                    bot.send_message(msg.chat.id, f"💰 كم سعر خطة {days} يوم بالدولار؟")
                db["pending_actions"][ustr] = {"action": "new_plan_price", "days": days}; sync_db()
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أرسل رقم أيام صحيح.")
        elif action == "new_plan_price" and txt:
            try:
                days = pending["days"]; price = float(txt.strip())
                db["config"]["sub_plans"][str(days)] = price; sync_db()
                bot.send_message(msg.chat.id, f"✅ تمت إضافة/تحديث خطة {days} يوم بسعر {price}$")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أرسل سعر صحيح.")
        elif action == "add_sub_price" and txt:
            try:
                days = int(pending.get("days", 30)); price = float(txt)
                db["config"]["sub_plans"][str(days)] = price; sync_db()
                bot.send_message(msg.chat.id, f"✅ تم تحديد سعر خطة {days} يوم = {price}$")
            except: bot.send_message(msg.chat.id, "❌ قيمة غير صحيحة.")
        elif action == "gen_sub_code" and txt:
            try:
                days = int(txt)
                code = "SUB-" + str(uuid.uuid4()).split("-")[0].upper()
                db["codes"][code] = {"sub_days": days, "used": False}; sync_db()
                bot.send_message(msg.chat.id, f"🎟️ كود اشتراك جديد ({days} يوم):\n`{code}`", parse_mode="Markdown")
            except: bot.send_message(msg.chat.id, "❌ رقم غير صحيح.")
        elif action == "manual_activate_sub" and txt:
            try:
                target, days = txt.split()
                activate_subscription(target, int(days))
                bot.send_message(msg.chat.id, f"✅ تم تفعيل اشتراك {days} يوم للمستخدم {target}.")
            except: bot.send_message(msg.chat.id, "❌ الصيغة: ID عدد_الأيام (مثال: 123456789 30)")
        elif action == "set_sub_layout" and txt:
            fid = pending.get("fid")
            layout = txt.strip().lower()
            if layout not in ("vertical", "horizontal", "grid2", "grid3"):
                bot.send_message(msg.chat.id, "❌ اختر: vertical أو horizontal أو grid2 أو grid3"); return
            target = db["config"]["custom_features"].get(fid) or next((b for b in db["config"]["menu_buttons"] if b["id"] == fid), None)
            if target:
                target["sub_layout"] = layout; sync_db()
                bot.send_message(msg.chat.id, f"✅ تم تعيين طريقة العرض: {layout}")

        # ── Novel creation flow (إنشاء رواية جديدة) ─────────────────────────
        elif action == "addnovel_title" and txt:
            db["pending_actions"][ustr] = {"action": "addnovel_story", "cat": pending["cat"], "title": txt}
            sync_db(); bot.send_message(msg.chat.id, "📝 أرسل قصة/وصف الرواية:")
        elif action == "addnovel_story" and txt:
            db["pending_actions"][ustr] = {"action": "addnovel_tag", "cat": pending["cat"], "title": pending["title"], "story": txt, "selected": []}
            sync_db(); bot.send_message(msg.chat.id, "🏷️ اختر تصنيف واحد أو أكثر (بدون حد أقصى)، ثم اضغط متابعة:", reply_markup=build_tag_search_menu(selected=[], mode="assign"))
        elif action == "addnovel_tag":
            # النص هنا يأتي من ضغط زر تصنيف (searchtag_) وليس رسالة نصية، لذا هذا احتياطي فقط
            db["pending_actions"][ustr] = {"action": "addnovel_poster", "cat": pending["cat"], "title": pending["title"], "story": pending["story"], "tags": [txt] if txt else ["عام"]}
            sync_db(); bot.send_message(msg.chat.id, "🖼️ أرسل صورة البوستر الآن:")
        elif action == "addnovel_poster":
            poster_id = msg.photo[-1].file_id if msg.content_type == "photo" else None
            if not poster_id:
                bot.send_message(msg.chat.id, "❌ لازم ترسل صورة. حاول مرة أخرى:")
                db["pending_actions"][ustr] = pending; sync_db(); return
            nid = create_novel(pending["title"], pending["story"], pending.get("tags", ["عام"]), poster_id, pending["cat"], created_by=uid)
            push_notification_all(f"📚 عمل جديد: {pending['title']}", target_novel=nid)
            bot.send_message(msg.chat.id, f"✅ تم إنشاء الرواية *{pending['title']}* بنجاح!\nالآن ارفع الفصول من قائمة إدارة الرواية.", parse_mode="Markdown")
            open_novel(uid, msg.chat.id, None, nid)

        # ── Add chapter flow (رفع فصل) ───────────────────────────────────────
        elif action == "addnovel_ch_number" and txt:
            try:
                ch_num = int(txt.strip())
            except:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح للفصل."); return
            db["pending_actions"][ustr] = {"action": "addnovel_ch_file", "nid": pending["nid"], "ch_num": ch_num}
            sync_db(); bot.send_message(msg.chat.id, f"📥 أرسل ملف الفصل {ch_num} (PDF أو EPUB كمستند):")
        elif action == "addnovel_ch_file":
            if msg.content_type != "document":
                bot.send_message(msg.chat.id, "❌ أرسل الملف كمستند (Document)."); return
            fname = (msg.document.file_name or "").lower()
            fmt = "epub" if fname.endswith(".epub") else "pdf"
            nid = pending["nid"]; ch_num = pending["ch_num"]
            novel = get_novel(nid)
            if not novel:
                bot.send_message(msg.chat.id, "❌ الرواية غير موجودة."); return
            novel.setdefault("chapters", {}).setdefault(str(ch_num), {"pdf": None, "epub": None})
            novel["chapters"][str(ch_num)][fmt] = msg.document.file_id
            sync_db()
            announce_new_chapter(nid, ch_num)
            mirror_to_archive("novels", nid, str(ch_num), fmt, msg.document.file_id)
            push_notification_all(f"🆕 فصل جديد: {novel['title']} — الفصل {ch_num}", target_novel=nid)
            bot.send_message(msg.chat.id, f"✅ تم رفع الفصل {ch_num} بصيغة {fmt.upper()}.\nإذا تبي ترفع الصيغة الثانية لنفس الفصل، اضغط 'رفع فصل جديد' وأدخل نفس الرقم.")

        # ── Delete chapter flow (حذف فصل) ───────────────────────────────────
        elif action == "delnovel_ch_number" and txt:
            nid = pending["nid"]
            novel = get_novel(nid)
            try:
                ch_num = int(txt.strip())
            except:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح للفصل."); return
            if not novel or str(ch_num) not in novel.get("chapters", {}):
                bot.send_message(msg.chat.id, f"❌ لا يوجد فصل بالرقم {ch_num} في هذه الرواية.")
                return
            db["pending_actions"].pop(ustr, None)
            sync_db()
            bot.send_message(
                msg.chat.id,
                f"⚠️ هل تريد حذف الفصل {ch_num} نهائياً؟ (كل صيغ الملف PDF/EPUB لهذا الفصل ستُحذف)",
                reply_markup=types.InlineKeyboardMarkup().add(
                    types.InlineKeyboardButton("✅ تأكيد الحذف", callback_data=f"confirm_delch_{nid}_{ch_num}"),
                    types.InlineKeyboardButton("❌ إلغاء", callback_data=f"open_novel_{nid}")
                )
            )

        # ── Add merged group flow (رفع فصول مدمجة) ──────────────────────────
        elif action == "addnovel_merge_range" and txt:
            rng = txt.strip()
            if "-" not in rng or not all(p.isdigit() for p in rng.split("-", 1)):
                bot.send_message(msg.chat.id, "❌ الصيغة يجب أن تكون مثل: 1-25"); return
            start_s, end_s = rng.split("-", 1)
            start, end = int(start_s), int(end_s)
            if start > end:
                bot.send_message(msg.chat.id, "❌ رقم البداية لازم يكون أصغر من أو يساوي النهاية."); return
            nid = pending["nid"]
            db["pending_actions"].pop(ustr, None); sync_db()
            chapter_list = list(range(start, end + 1))
            bot.send_message(msg.chat.id, f"⏳ جاري تجهيز ودمج الفصول {rng} بالخلفية، راح أرسل لك النتيجة هنا أول ما يخلص...")
            def _do_merge_job2():
                ok, result_msg = merge_existing_chapters(nid, chapter_list, msg.chat.id)
                bot.send_message(msg.chat.id, result_msg, parse_mode="Markdown" if ok else None)
            threading.Thread(target=_do_merge_job2, daemon=True).start()

        # ── Link channel/group to novel ──────────────────────────────────────
        elif action == "addnovel_link" and txt:
            nid = pending["nid"]; novel = get_novel(nid)
            if not novel:
                bot.send_message(msg.chat.id, "❌ الرواية غير موجودة."); return
            if txt.strip() in ("إلغاء", "الغاء", "cancel"):
                novel["linked_chat"] = None
                bot.send_message(msg.chat.id, "✅ تم إلغاء الربط.")
            else:
                novel["linked_chat"] = txt.strip()
                bot.send_message(msg.chat.id, f"✅ تم ربط الرواية بـ {txt.strip()}")
            sync_db()

        # ══════════════════ MANGA/MANHWA FLOWS ══════════════════════════════
        elif action == "addmanga_title" and txt:
            db["pending_actions"][ustr] = {"action": "addmanga_story", "cat": pending["cat"], "title": txt}
            sync_db(); bot.send_message(msg.chat.id, "📝 أرسل قصة/وصف العمل:")
        elif action == "addmanga_story" and txt:
            db["pending_actions"][ustr] = {"action": "addmanga_tag", "cat": pending["cat"], "title": pending["title"], "story": txt, "selected": []}
            sync_db(); bot.send_message(msg.chat.id, "🏷️ اختر تصنيف واحد أو أكثر، ثم اضغط متابعة:", reply_markup=build_tag_search_menu(selected=[], mode="assign"))
        elif action == "addmanga_tag":
            db["pending_actions"][ustr] = {"action": "addmanga_poster", "cat": pending["cat"], "title": pending["title"], "story": pending["story"], "tags": [txt] if txt else ["عام"]}
            sync_db(); bot.send_message(msg.chat.id, "🖼️ أرسل صورة البوستر الآن:")
        elif action == "addmanga_poster":
            poster_id = msg.photo[-1].file_id if msg.content_type == "photo" else None
            if not poster_id:
                bot.send_message(msg.chat.id, "❌ لازم ترسل صورة. حاول مرة أخرى:")
                db["pending_actions"][ustr] = pending; sync_db(); return
            mid = create_manga(pending["title"], pending["story"], pending.get("tags", ["عام"]), poster_id, pending["cat"], created_by=uid)
            push_notification_all(f"🎨 عمل جديد: {pending['title']}", target_novel=None)
            bot.send_message(msg.chat.id, f"✅ تم إنشاء العمل *{pending['title']}* بنجاح!\nالآن ارفع الفصول من قائمة إدارة العمل.", parse_mode="Markdown")
            open_manga(uid, msg.chat.id, None, mid)
        elif action == "addmanga_ch_number" and txt:
            try:
                ch_num = int(txt.strip())
            except:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح للفصل."); return
            db["pending_actions"][ustr] = {"action": "addmanga_ch_file", "mid": pending["mid"], "ch_num": ch_num}
            sync_db(); bot.send_message(msg.chat.id, f"📥 أرسل ملف الفصل {ch_num} (PDF أو EPUB كمستند):")
        elif action == "addmanga_ch_file":
            if msg.content_type != "document":
                bot.send_message(msg.chat.id, "❌ أرسل الملف كمستند (Document)."); return
            fname = (msg.document.file_name or "").lower()
            fmt = "epub" if fname.endswith(".epub") else "pdf"
            mid = pending["mid"]; ch_num = pending["ch_num"]
            manga = get_manga(mid)
            if not manga:
                bot.send_message(msg.chat.id, "❌ العمل غير موجود."); return
            manga.setdefault("chapters", {}).setdefault(str(ch_num), {"pdf": None, "epub": None})
            manga["chapters"][str(ch_num)][fmt] = msg.document.file_id
            sync_db()
            mirror_to_archive("manga", mid, str(ch_num), fmt, msg.document.file_id)
            push_notification_all(f"🆕 فصل جديد: {manga['title']} — الفصل {ch_num}", target_novel=None)
            bot.send_message(msg.chat.id, f"✅ تم رفع الفصل {ch_num} بصيغة {fmt.upper()}.")
        elif action == "delmanga_ch_number" and txt:
            mid = pending["mid"]; manga = get_manga(mid)
            try:
                ch_num = int(txt.strip())
            except:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح للفصل."); return
            if not manga or str(ch_num) not in manga.get("chapters", {}):
                bot.send_message(msg.chat.id, f"❌ لا يوجد فصل بالرقم {ch_num} في هذا العمل.")
                return
            db["pending_actions"].pop(ustr, None); sync_db()
            bot.send_message(msg.chat.id, f"⚠️ هل تريد حذف الفصل {ch_num} نهائياً؟",
                reply_markup=types.InlineKeyboardMarkup().add(
                    types.InlineKeyboardButton("✅ تأكيد الحذف", callback_data=f"confirm_delmangach_{mid}_{ch_num}"),
                    types.InlineKeyboardButton("❌ إلغاء", callback_data=f"open_manga_{mid}")))
        elif action == "addmanga_merge_range" and txt:
            rng = txt.strip()
            if "-" not in rng or not all(p.isdigit() for p in rng.split("-", 1)):
                bot.send_message(msg.chat.id, "❌ الصيغة يجب أن تكون مثل: 1-25"); return
            db["pending_actions"][ustr] = {"action": "addmanga_merge_file", "mid": pending["mid"], "range": rng}
            sync_db(); bot.send_message(msg.chat.id, f"📥 أرسل ملف المجموعة {rng} (PDF أو EPUB كمستند):")
        elif action == "addmanga_merge_file":
            if msg.content_type != "document":
                bot.send_message(msg.chat.id, "❌ أرسل الملف كمستند (Document)."); return
            fname = (msg.document.file_name or "").lower()
            fmt = "epub" if fname.endswith(".epub") else "pdf"
            mid = pending["mid"]; rng = pending["range"]
            manga = get_manga(mid)
            if not manga:
                bot.send_message(msg.chat.id, "❌ العمل غير موجود."); return
            manga.setdefault("merged", {}).setdefault(rng, {"pdf": None, "epub": None})
            manga["merged"][rng][fmt] = msg.document.file_id
            sync_db()
            bot.send_message(msg.chat.id, f"✅ تم رفع مجموعة الفصول {rng} بصيغة {fmt.upper()}.")
        elif action == "addmanga_link" and txt:
            mid = pending["mid"]; manga = get_manga(mid)
            if not manga:
                bot.send_message(msg.chat.id, "❌ العمل غير موجود."); return
            if txt.strip() in ("إلغاء", "الغاء", "cancel"):
                manga["linked_chat"] = None
                bot.send_message(msg.chat.id, "✅ تم إلغاء الربط.")
            else:
                manga["linked_chat"] = txt.strip()
                bot.send_message(msg.chat.id, f"✅ تم ربط العمل بـ {txt.strip()}")
            sync_db()

        # ══════════════════ SERIES/MOVIES FLOWS ═════════════════════════════
        elif action == "addseries_title" and txt:
            db["pending_actions"][ustr] = {"action": "addseries_story", "cat": pending["cat"], "title": txt}
            sync_db(); bot.send_message(msg.chat.id, "📝 أرسل قصة/وصف العمل:")
        elif action == "addseries_story" and txt:
            db["pending_actions"][ustr] = {"action": "addseries_tag", "cat": pending["cat"], "title": pending["title"], "story": txt, "selected": []}
            sync_db(); bot.send_message(msg.chat.id, "🏷️ اختر تصنيف واحد أو أكثر، ثم اضغط متابعة:", reply_markup=build_tag_search_menu(selected=[], mode="assign"))
        elif action == "addseries_tag":
            db["pending_actions"][ustr] = {"action": "addseries_poster", "cat": pending["cat"], "title": pending["title"], "story": pending["story"], "tags": [txt] if txt else ["عام"]}
            sync_db(); bot.send_message(msg.chat.id, "🖼️ أرسل صورة البوستر الآن:")
        elif action == "addseries_poster":
            poster_id = msg.photo[-1].file_id if msg.content_type == "photo" else None
            if not poster_id:
                bot.send_message(msg.chat.id, "❌ لازم ترسل صورة. حاول مرة أخرى:")
                db["pending_actions"][ustr] = pending; sync_db(); return
            sid = create_series(pending["title"], pending["story"], pending.get("tags", ["عام"]), poster_id, pending["cat"], created_by=uid)
            push_notification_all(f"🎬 عمل جديد: {pending['title']}", target_novel=None)
            bot.send_message(msg.chat.id, f"✅ تم إنشاء العمل *{pending['title']}* بنجاح!\nالآن ارفع الحلقات من قائمة إدارة العمل.", parse_mode="Markdown")
            open_series(uid, msg.chat.id, None, sid)
        elif action == "addseries_ep_number" and txt:
            try:
                ep_num = int(txt.strip())
            except:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح للحلقة."); return
            db["pending_actions"][ustr] = {"action": "addseries_ep_quality", "sid": pending["sid"], "ep_num": ep_num}
            sync_db()
            qm = types.InlineKeyboardMarkup(row_width=3)
            qm.add(*[types.InlineKeyboardButton(q, callback_data=f"epquality_{q}") for q in SERIES_QUALITIES])
            bot.send_message(msg.chat.id, f"🎞️ اختر جودة الفيديو اللي راح ترفعه الآن للحلقة {ep_num}:", reply_markup=qm)
        elif action == "addseries_ep_file":
            if msg.content_type != "video":
                bot.send_message(msg.chat.id, "❌ أرسل فيديو صحيح."); return
            sid = pending["sid"]; ep_num = pending["ep_num"]; quality = pending["quality"]; sub_only = pending["sub_only"]
            series = get_series(sid)
            if not series:
                bot.send_message(msg.chat.id, "❌ العمل غير موجود."); return
            series.setdefault("episodes", {}).setdefault(str(ep_num), {})
            series["episodes"][str(ep_num)][quality] = {"file_id": msg.video.file_id, "sub_only": sub_only}
            sync_db()
            mirror_to_archive("series", sid, str(ep_num), quality, msg.video.file_id, file_kind="video")
            push_notification_all(f"🆕 حلقة جديدة: {series['title']} — الحلقة {ep_num} ({quality})", target_novel=None)
            bot.send_message(msg.chat.id, f"✅ تم رفع الحلقة {ep_num} بجودة {quality} ({'مقصورة على المشتركين 🔒' if sub_only else 'مجانية 🆓'}).\nتقدر ترفع جودة ثانية لنفس الحلقة من 'رفع حلقة جديدة' بنفس الرقم.")
        elif action == "delseries_ep_number" and txt:
            sid = pending["sid"]; series = get_series(sid)
            try:
                ep_num = int(txt.strip())
            except:
                bot.send_message(msg.chat.id, "❌ أرسل رقم صحيح للحلقة."); return
            if not series or str(ep_num) not in series.get("episodes", {}):
                bot.send_message(msg.chat.id, f"❌ لا توجد حلقة بالرقم {ep_num} في هذا العمل.")
                return
            db["pending_actions"].pop(ustr, None); sync_db()
            bot.send_message(msg.chat.id, f"⚠️ هل تريد حذف الحلقة {ep_num} نهائياً (كل الجودات)؟",
                reply_markup=types.InlineKeyboardMarkup().add(
                    types.InlineKeyboardButton("✅ تأكيد الحذف", callback_data=f"confirm_delep_{sid}_{ep_num}"),
                    types.InlineKeyboardButton("❌ إلغاء", callback_data=f"open_series_{sid}")))
        elif action == "addseries_link" and txt:
            sid = pending["sid"]; series = get_series(sid)
            if not series:
                bot.send_message(msg.chat.id, "❌ العمل غير موجود."); return
            if txt.strip() in ("إلغاء", "الغاء", "cancel"):
                series["linked_chat"] = None
                bot.send_message(msg.chat.id, "✅ تم إلغاء الربط.")
            else:
                series["linked_chat"] = txt.strip()
                bot.send_message(msg.chat.id, f"✅ تم ربط العمل بـ {txt.strip()}")
            sync_db()

        # ── إعدادات نظام المسابقات المتطور ──────────────────────────────────
        elif action == "contest_set_target" and txt:
            target = txt.strip()
            try:
                test = bot.get_chat(target)
                db["config"]["contests"]["target_chat_id"] = target
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم تعيين الهدف: {test.title or target}")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ معرف غير صحيح:\n`{e}`", parse_mode="Markdown")
        elif action == "contest_set_count" and txt:
            try:
                count = int(txt.strip())
                if count > 0:
                    db["config"]["contests"]["question_count"] = count
                    sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم تعيين عدد الأسئلة: {count}")
                else:
                    bot.send_message(msg.chat.id, "❌ أدخل رقم موجب")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أدخل رقم صحيح")
        elif action == "contest_set_choices" and txt:
            try:
                choices = int(txt.strip())
                if choices in (2, 4):
                    db["config"]["contests"]["choices_count"] = choices
                    sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم تعيين عدد الخيارات: {choices}")
                else:
                    bot.send_message(msg.chat.id, "❌ اختر 2 أو 4 فقط")
            except ValueError:
                bot.send_message(msg.chat.id, "❌ أدخل رقم صحيح")
        elif action == "contest_set_time" and txt:
            if re.match(r"^\d{2}:\d{2}$", txt.strip()):
                db["config"]["contests"]["schedule_time"] = txt.strip()
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم تعيين وقت الجدولة: {txt.strip()}")
            else:
                bot.send_message(msg.chat.id, "❌ صيغة غير صحيحة (مثال: 18:00)")
        elif action == "contest_set_days" and txt:
            days_map = {
                "sat": "sat", "sun": "sun", "mon": "mon", "tue": "tue", "wed": "wed", "thu": "thu", "fri": "fri",
                "السبت": "sat", "الأحد": "sun", "الاثنين": "mon", "الثلاثاء": "tue",
                "الأربعاء": "wed", "الخميس": "thu", "الجمعة": "fri",
            }
            txt_lower = txt.strip().lower()
            if txt_lower == "all":
                days = ["sat", "sun", "mon", "tue", "wed", "thu", "fri"]
            else:
                days = []
                for part in txt_lower.split(","):
                    part = part.strip()
                    if part in days_map:
                        days.append(days_map[part])
                days = list(dict.fromkeys(days))
            if days:
                db["config"]["contests"]["schedule_days"] = days
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم تعيين الأيام: {', '.join(days)}")
            else:
                bot.send_message(msg.chat.id, "❌ أيام غير صالحة")
        elif action == "contest_set_prizes" and txt:
            if txt.strip().lower() in ("بدون", "لا", "none"):
                db["config"]["contests"]["prizes"] = {}
                sync_db()
                bot.send_message(msg.chat.id, "✅ تم إلغاء الجوائز")
            else:
                prizes = {}
                for line in txt.split("\n"):
                    line = line.strip()
                    if ":" in line:
                        key, value = line.split(":", 1)
                        key = key.strip().lower(); value = value.strip()
                        if key in ("نقاط", "points"):
                            try: prizes["points"] = int(value)
                            except ValueError: pass
                        elif key in ("أيام_vip", "vip_days"):
                            try: prizes["vip_days"] = int(value)
                            except ValueError: pass
                        elif key in ("رتبة", "role"):
                            prizes["role"] = value
                if prizes:
                    db["config"]["contests"]["prizes"] = prizes
                    sync_db()
                    bot.send_message(msg.chat.id, f"✅ تم تعيين الجوائز:\n{_format_prizes(prizes)}")
                else:
                    bot.send_message(msg.chat.id, "❌ لم يتم التعرف على الجوائز. استخدم التنسيق المطلوب.")

        # ── إعدادات ميزة الدعوة التلقائية (auto_invite) ─────────────────────
        elif action == "invite_set_target" and txt:
            target = txt.strip()
            try:
                test = bot.get_chat(target)
                db["config"]["auto_invite"]["target_chat_id"] = target
                db["config"]["auto_invite"]["invite_link"] = None  # إعادة إنشاء الرابط عند أول استخدام
                sync_db()
                bot.send_message(msg.chat.id, f"✅ تم تعيين الوجهة: {test.title or target}")
            except Exception as e:
                bot.send_message(msg.chat.id, f"❌ معرف غير صحيح:\n`{e}`", parse_mode="Markdown")
        elif action == "invite_edit_text" and txt:
            if "{link}" not in txt:
                bot.send_message(msg.chat.id, "⚠️ النص يجب أن يحتوي على `{link}` لإدراج الرابط.")
            else:
                db["config"]["auto_invite"]["confirmation_text"] = txt
                sync_db()
                bot.send_message(msg.chat.id, "✅ تم تحديث نص الموافقة.")
        elif action == "invite_edit_button" and txt:
            db["config"]["auto_invite"]["button_label"] = txt.strip()
            sync_db()
            bot.send_message(msg.chat.id, "✅ تم تحديث زر الموافقة.")
        return

    # ===== معالجة إجابات المسابقات في وضع التعليق (خارج نظام pending العادي) =====
    if db["config"]["contests"].get("enabled"):
        active = db["config"]["contests"].get("active_contests", {})
        for contest_id, contest in list(active.items()):
            if contest.get("answer_mode") == "comment" and contest.get("status") == "active":
                target = db["config"]["contests"].get("target_chat_id") or db["config"]["announce_channel"].get("chat_id")
                if target and str(msg.chat.id) == str(target):
                    if msg.content_type == "text" and not msg.from_user.is_bot:
                        handle_comment_answer(msg, contest_id, msg.from_user.id)
                        return

    # ===== ميزة الدعوة التلقائية عند استلام رسالة =====
    if db["config"]["auto_invite"].get("enabled") and "message" in db["config"]["auto_invite"].get("trigger_on", []):
        if msg.chat.type in ("group", "supergroup", "channel") and not msg.from_user.is_bot:
            send_invite_to_user(msg.from_user.id, "رسالة في المجموعة/القناة")

    if msg.content_type == "text":
        ensure_user(msg)
        u = db["users"][ustr]
        if not u.get("seen_start_prompt"):
            # أول رسالة عشوائية بس نعرض له زر "ابدأ" — بعدها نسكت تمامًا على أي نص عشوائي آخر.
            u["seen_start_prompt"] = True
            sync_db()
            m = types.InlineKeyboardMarkup().add(types.InlineKeyboardButton("🚀 ابدأ", callback_data="do_start"))
            bot.send_message(msg.chat.id, "👋 اضغط الزر تحت للبدء، أو استخدم /start بأي وقت.", reply_markup=m)
        # لو already شاف البرومبت قبل كذا، ما نرد بشي إطلاقًا على نص عشوائي.

if __name__ == '__main__':
    logger.info("🚀 TRAIKA V5 FINAL — STARTING...")
    # ملاحظة: init_db() اتنفذ أصلاً عند تحميل الملف (أعلاه)، ما نعيد استدعاءها هنا —
    # استدعاؤها مرتين كان يفتح اتصال قاعدة بيانات ثاني منفصل، ولو صار فيه أي تعثر
    # مؤقت بالاتصال الثاني (حتى لو الأول نجح تمامًا)، كان يقلب _db_load_failed=True
    # بالغلط ويوقف الحفظ والنسخ الاحتياطية بصمت تام بدون أي خطأ ظاهر للمستخدم.
    setup_bot_commands()
    # تسخين مسبق لخط PDF العربي بخيط منفصل (daemon) بدون ما يوقف إقلاع البوت.
    # كذا لو أول مستخدم يجرب "تحويل لـ PDF" ما يضطر ينتظر تحميل الخط من
    # الشبكة (حتى 20 ثانية) جوا مهلة التحويل نفسها (45 ثانية) — الخط يكون
    # جاهز بالكاش المحلي مسبقًا في أغلب الأحيان.
    threading.Thread(target=_ensure_arabic_font, daemon=True).start()
    # نظام المسابقات (فحص كل دقيقة) ونظام المحادثة الذكية (فحص كل 30 ثانية) —
    # كل واحد بخيط daemon منفصل حتى ما يوقف تشغيل البوت لو صار فيه استثناء داخلي.
    threading.Thread(target=contests_scheduler_loop, daemon=True).start()
    threading.Thread(target=ai_talk_scheduler_loop, daemon=True).start()
    # جدولة سحب الأخبار التلقائي (فحص كل 5 دقائق أو حسب وقت النشر المضبوط)
    threading.Thread(target=news_scheduler_loop, daemon=True).start()
    # جدولة السحب التلقائي للروايات من قائمة الانتظار (فحص كل دقيقة)
    threading.Thread(target=scraper_scheduler_loop, daemon=True).start()
    # نظام التقطيع: تنظيف فوري لأي بقايا ملفات من جلسة سابقة انتهت بشكل غير
    # طبيعي (كراش/انطفاء مفاجئ)، ثم مهمة خلفية دائمة تنظّف كل ساعة تلقائيًا.
    try:
        memory_monitor.max_mb = db["config"]["chunking_settings"].get("max_ram_mb", 200)
    except Exception:
        pass
    threading.Thread(target=temp_file_manager.cleanup_loop, daemon=True).start()
    if _db_load_failed:
        try:
            bot.send_message(OWNER_ID,
                "🚨 *تنبيه حرج:* فشل الاتصال بقاعدة البيانات عند إقلاع البوت.\n\n"
                "البوت يشتغل الآن بذاكرة فاضية مؤقتة، وتم *منع الحفظ نهائيًا* حماية "
                "لبياناتك الحقيقية من الاستبدال. راجع سجل الأخطاء (logs) وتأكد من "
                "`BOT_DATABASE_URL` واتصال قاعدة البيانات (Neon)، ثم أعد تشغيل البوت.",
                parse_mode="Markdown")
        except Exception as e:
            logger.error(f"❌ فشل إرسال تنبيه فشل القاعدة للمالك: {e}")
    while True:
        try: bot.infinity_polling(timeout=60, long_polling_timeout=60)
        except Exception as e:
            logger.error(f"Polling error: {e}")
            sync_db_now()  # نضمن حفظ آخر تغييرات قبل إعادة المحاولة
            time.sleep(5)
